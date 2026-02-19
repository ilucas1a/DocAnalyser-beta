"""
save_utils.py

Universal save/export functions for DocAnalyser.
Handles TXT, DOCX, RTF, and PDF exports with consistent metadata.

Usage:
    from save_utils import save_document_to_file, get_clean_filename, get_document_metadata
"""

import os
import re
import datetime
from tkinter import filedialog, messagebox


def save_document_to_file(
    filepath: str,
    content_text: str,
    title: str,
    source: str = "Unknown",
    published_date: str = None,
    imported_date: str = "Unknown",
    doc_class: str = "source",
    export_format: str = "docx",
    provider: str = None,
    model: str = None
) -> bool:
    """
    Universal save function for all document exports.
    
    Args:
        filepath: Full path to save file
        content_text: The document content to save
        title: Document title (used in header)
        source: Source URL or path
        published_date: Original publication date (optional)
        imported_date: When document was imported
        doc_class: Document class (source, processed_output, thread, etc.)
        export_format: One of 'txt', 'docx', 'rtf', 'pdf'
        provider: AI provider used (for processed outputs)
        model: AI model used (for processed outputs)
    
    Returns:
        True if successful, False otherwise
    """
    from document_export import get_export_date
    export_date = get_export_date()
    
    try:
        if export_format == 'txt':
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write("=" * 80 + "\n")
                f.write(f"{title}\n")
                f.write("=" * 80 + "\n\n")
                f.write("DOCUMENT INFORMATION:\n")
                f.write(f"  Title: {title}\n")
                f.write(f"  Source: {source}\n")
                if published_date:
                    f.write(f"  Published: {published_date}\n")
                f.write(f"  Imported: {imported_date}\n")
                f.write(f"  Exported: {export_date}\n")
                f.write(f"  Type: {doc_class}\n")
                if provider and model:
                    f.write(f"  Processed By: {provider} / {model}\n")
                f.write("\n" + "=" * 80 + "\n\n")
                f.write(content_text)
            return True
        
        elif export_format == 'docx' or export_format == 'rtf':
            from docx import Document
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            
            doc_obj = Document()
            
            # Title
            title_para = doc_obj.add_heading(title[:100], 0)
            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Metadata section
            doc_obj.add_heading('Document Information', level=2)
            meta_para = doc_obj.add_paragraph()
            meta_para.add_run("Title: ").bold = True
            meta_para.add_run(f"{title}\n")
            meta_para.add_run("Source: ").bold = True
            meta_para.add_run(f"{source}\n")
            if published_date:
                meta_para.add_run("Published: ").bold = True
                meta_para.add_run(f"{published_date}\n")
            meta_para.add_run("Imported: ").bold = True
            meta_para.add_run(f"{imported_date}\n")
            meta_para.add_run("Exported: ").bold = True
            meta_para.add_run(f"{export_date}\n")
            meta_para.add_run("Type: ").bold = True
            meta_para.add_run(f"{doc_class}\n")
            if provider and model:
                meta_para.add_run("Processed By: ").bold = True
                meta_para.add_run(f"{provider} / {model}\n")
            
            doc_obj.add_paragraph()  # Spacing
            doc_obj.add_heading('Content', level=2)
            
            # Add content paragraphs
            for line in content_text.split('\n'):
                if line.strip():
                    doc_obj.add_paragraph(line)
            
            # For RTF, save as docx with adjusted path
            if export_format == 'rtf':
                filepath = filepath.replace('.rtf', '.docx')
            
            doc_obj.save(filepath)
            return True
        
        elif export_format == 'pdf':
            try:
                from reportlab.lib.pagesizes import letter
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from reportlab.lib.units import inch
                from reportlab.lib.enums import TA_CENTER
                
                pdf_doc = SimpleDocTemplate(filepath, pagesize=letter)
                styles = getSampleStyleSheet()
                story = []
                
                # Title style
                title_style = ParagraphStyle(
                    'CustomTitle',
                    parent=styles['Heading1'],
                    fontSize=20,
                    textColor='#2E4053',
                    spaceAfter=20,
                    alignment=TA_CENTER
                )
                story.append(Paragraph(title[:100], title_style))
                story.append(Spacer(1, 0.2 * inch))
                
                # Metadata section
                meta_style = ParagraphStyle(
                    'Meta',
                    parent=styles['Normal'],
                    fontSize=10,
                    textColor='#555555'
                )
                story.append(Paragraph("<b>Document Information</b>", styles['Heading2']))
                story.append(Paragraph(f"<b>Title:</b> {title}", meta_style))
                story.append(Paragraph(f"<b>Source:</b> {source}", meta_style))
                if published_date:
                    story.append(Paragraph(f"<b>Published:</b> {published_date}", meta_style))
                story.append(Paragraph(f"<b>Imported:</b> {imported_date}", meta_style))
                story.append(Paragraph(f"<b>Exported:</b> {export_date}", meta_style))
                story.append(Paragraph(f"<b>Type:</b> {doc_class}", meta_style))
                if provider and model:
                    story.append(Paragraph(f"<b>Processed By:</b> {provider} / {model}", meta_style))
                story.append(Spacer(1, 0.3 * inch))
                
                # Content section
                story.append(Paragraph("<b>Content</b>", styles['Heading2']))
                story.append(Spacer(1, 0.1 * inch))
                
                content_style = ParagraphStyle(
                    'Content',
                    parent=styles['Normal'],
                    fontSize=10
                )
                
                for line in content_text.split('\n'):
                    if line.strip():
                        # Escape special characters for PDF
                        safe_line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        story.append(Paragraph(safe_line, content_style))
                        story.append(Spacer(1, 0.1 * inch))
                
                pdf_doc.build(story)
                return True
                
            except ImportError:
                print("❌ reportlab library not installed for PDF export")
                return False
        
        else:
            print(f"❌ Unknown export format: {export_format}")
            return False
            
    except Exception as e:
        print(f"❌ Error saving document: {str(e)}")
        return False


def get_clean_filename(title: str, max_length: int = 100) -> str:
    """
    Convert a title to a safe filename.
    
    Args:
        title: The document title
        max_length: Maximum filename length
    
    Returns:
        Clean filename without extension
    """
    # Remove invalid characters for Windows filenames
    clean_name = re.sub(r'[<>:"/\\|?*]', '-', title)
    # Remove non-printable characters
    clean_name = ''.join(char for char in clean_name if ord(char) >= 32)
    # Collapse multiple spaces/dashes
    clean_name = re.sub(r'[-\s]+', ' ', clean_name).strip()
    # Truncate if needed
    if len(clean_name) > max_length:
        clean_name = clean_name[:max_length].rsplit(' ', 1)[0]
    return clean_name if clean_name else 'document'


def get_document_metadata(app_instance, get_document_by_id_func=None) -> dict:
    """
    Extract document metadata from the app instance.
    
    Args:
        app_instance: The DocAnalyser app instance (self)
        get_document_by_id_func: Optional function to get document from database
    
    Returns:
        Dictionary with title, source, imported_date, doc_class
    """
    # Get title - try multiple sources
    title = "Untitled Document"
    
    # 1. Try current_document_metadata
    if hasattr(app_instance, 'current_document_metadata') and app_instance.current_document_metadata:
        title = app_instance.current_document_metadata.get('title', title)
    
    # 2. If still untitled, try to get from database
    if title == "Untitled Document" and hasattr(app_instance, 'current_document_id') and app_instance.current_document_id:
        if get_document_by_id_func:
            try:
                doc = get_document_by_id_func(app_instance.current_document_id)
                if doc:
                    title = doc.get('title', title)
            except:
                pass
    
    # 3. Final fallback to source
    if title == "Untitled Document":
        title = getattr(app_instance, 'current_document_source', 'document')
    
    # Get source
    source = getattr(app_instance, 'current_document_source', 'Unknown')
    
    # Get doc class
    doc_class = getattr(app_instance, 'current_document_class', 'source')
    
    # Get imported date
    imported_date = "Unknown"
    if hasattr(app_instance, 'current_document_metadata') and app_instance.current_document_metadata:
        fetched = app_instance.current_document_metadata.get('fetched', 'Unknown')
        if fetched != 'Unknown':
            try:
                from utils import format_display_date
                imported_date = format_display_date(fetched)
            except:
                imported_date = fetched
    
    # Get published date
    published_date = None
    if hasattr(app_instance, 'current_document_metadata') and app_instance.current_document_metadata:
        raw_pub_date = app_instance.current_document_metadata.get('published_date')
        if raw_pub_date:
            try:
                from utils import format_display_date
                published_date = format_display_date(raw_pub_date)
            except:
                published_date = raw_pub_date
    
    # Get AI provider and model (for processed outputs)
    provider = None
    model = None
    if hasattr(app_instance, 'current_document_metadata') and app_instance.current_document_metadata:
        provider = app_instance.current_document_metadata.get('provider')
        model = app_instance.current_document_metadata.get('model')
    
    # Fallback to current settings if not in metadata
    if not provider and hasattr(app_instance, 'provider_var'):
        provider = app_instance.provider_var.get()
    if not model and hasattr(app_instance, 'model_var'):
        model = app_instance.model_var.get()
    
    return {
        'title': title,
        'source': source,
        'published_date': published_date,
        'imported_date': imported_date,
        'doc_class': doc_class,
        'provider': provider,
        'model': model
    }


def prompt_and_save_document(
    app_instance,
    content_text: str,
    title: str = None,
    source: str = None,
    imported_date: str = None,
    doc_class: str = None,
    export_format: str = "docx",
    parent_window=None,
    get_document_by_id_func=None
) -> bool:
    """
    Show save dialog and save document using universal save function.
    
    Args:
        app_instance: The DocAnalyser app instance (self)
        content_text: The content to save
        title, source, imported_date, doc_class: Optional overrides
        export_format: 'txt', 'docx', 'rtf', 'pdf'
        parent_window: Parent window for dialog
        get_document_by_id_func: Optional function to get document from database
    
    Returns:
        True if saved successfully
    """
    # Get metadata, using overrides if provided
    meta = get_document_metadata(app_instance, get_document_by_id_func)
    title = title or meta['title']
    source = source or meta['source']
    published_date = meta.get('published_date')  # Extract published date from metadata
    imported_date = imported_date or meta['imported_date']
    doc_class = doc_class or meta['doc_class']
    provider = meta.get('provider')
    model = meta.get('model')
    
    # Get clean filename
    clean_name = get_clean_filename(title)
    default_name = f"{clean_name}.{export_format}"
    
    # File type descriptions
    filetypes = {
        'txt': [("Text files", "*.txt"), ("All files", "*.*")],
        'docx': [("Word documents", "*.docx"), ("All files", "*.*")],
        'rtf': [("RTF documents", "*.rtf"), ("All files", "*.*")],
        'pdf': [("PDF files", "*.pdf"), ("All files", "*.*")]
    }
    
    # Show save dialog
    filepath = filedialog.asksaveasfilename(
        title="Save Document As",
        defaultextension=f".{export_format}",
        initialfile=default_name,
        filetypes=filetypes.get(export_format, [("All files", "*.*")])
    )
    
    if not filepath:
        return False  # User cancelled
    
    # Save using universal function
    success = save_document_to_file(
        filepath=filepath,
        content_text=content_text,
        title=title,
        source=source,
        published_date=published_date,  # Pass published date to save function
        imported_date=imported_date,
        doc_class=doc_class,
        export_format=export_format,
        provider=provider,
        model=model
    )
    
    if success:
        app_instance.set_status(f"✅ Saved to: {os.path.basename(filepath)}")
        
        # Show appropriate message for RTF
        if export_format == 'rtf':
            messagebox.showinfo(
                "RTF Export",
                f"Document saved as DOCX at:\n{filepath.replace('.rtf', '.docx')}\n\n"
                "To convert to RTF, open in Word and Save As RTF."
            )
        else:
            messagebox.showinfo("Saved", f"Document saved to:\n{filepath}")
        return True
    else:
        messagebox.showerror("Export Error", "Failed to save document. Check console for details.")
        return False
