"""
document_export.py - Consolidated Document Export Module
Handles all document exports (txt, docx, rtf, pdf) in a single location.

This module consolidates export functionality that was previously duplicated
across Main.py and thread_viewer.py.
"""

import os
import re
import datetime
from typing import Dict, List, Optional, Tuple, Any


def get_export_date() -> str:
    """Get current date/time formatted for export metadata (DD-Mon-YYYY HH:MM:SS)"""
    now = datetime.datetime.now()
    return now.strftime('%d-%b-%Y %H:%M:%S')


def sanitize_filename(text: str, max_length: int = 100) -> str:
    """Convert text to a safe filename"""
    clean_name = re.sub(r'[<>:"/\\|?*]', '-', text)
    clean_name = ''.join(char for char in clean_name if ord(char) >= 32)
    clean_name = re.sub(r'[-\s]+', ' ', clean_name).strip()
    if len(clean_name) > max_length:
        clean_name = clean_name[:max_length].rsplit(' ', 1)[0]
    return clean_name if clean_name else 'document'


def get_file_extension_and_types(format: str) -> Tuple[str, List[Tuple[str, str]]]:
    """Get file extension and filetypes for save dialog"""
    format_map = {
        'txt': ('.txt', [("Text files", "*.txt"), ("All files", "*.*")]),
        'docx': ('.docx', [("Word Document", "*.docx"), ("All files", "*.*")]),
        'rtf': ('.rtf', [("RTF files", "*.rtf"), ("All files", "*.*")]),
        'pdf': ('.pdf', [("PDF files", "*.pdf"), ("All files", "*.*")]),
    }
    return format_map.get(format, ('.txt', [("All files", "*.*")]))


# =============================================================================
# DOCUMENT EXPORT - For source documents and AI products
# =============================================================================

def export_document(
    filepath: str,
    content: str,
    format: str,
    metadata: Dict[str, Any],
    show_messages: bool = True
) -> Tuple[bool, str]:
    """
    Universal document export function.
    
    Args:
        filepath: Full path to save the file
        content: The document content text
        format: Export format ('txt', 'docx', 'rtf', 'pdf')
        metadata: Dict with keys:
            - title: Document title
            - source: Source URL/path
            - published_date: Original publication date (optional)
            - imported_date: When document was imported
            - doc_class: 'source' or 'product'
            - provider: AI provider (optional)
            - model: AI model (optional)
        show_messages: Whether to show success/error message boxes
    
    Returns:
        Tuple of (success: bool, message: str)
    """
    format = format.lower()
    export_date = get_export_date()
    
    # Extract metadata with defaults
    title = metadata.get('title', 'Untitled Document')
    source = metadata.get('source', 'Unknown')
    published_date = metadata.get('published_date')
    imported_date = metadata.get('imported_date', 'Unknown')
    doc_class = metadata.get('doc_class', 'source')
    provider = metadata.get('provider')
    model = metadata.get('model')
    
    try:
        if format == 'txt':
            return _export_as_txt(filepath, content, title, source, published_date, 
                                  imported_date, export_date, doc_class, provider, model, show_messages)
        elif format == 'docx':
            return _export_as_docx(filepath, content, title, source, published_date,
                                   imported_date, export_date, doc_class, provider, model, show_messages)
        elif format == 'rtf':
            return _export_as_rtf(filepath, content, title, source, published_date,
                                  imported_date, export_date, doc_class, provider, model, show_messages)
        elif format == 'pdf':
            return _export_as_pdf(filepath, content, title, source, published_date,
                                  imported_date, export_date, doc_class, provider, model, show_messages)
        else:
            return False, f"Unsupported format: {format}"
    except Exception as e:
        error_msg = f"Export failed: {str(e)}"
        if show_messages:
            from tkinter import messagebox
            messagebox.showerror("Export Error", error_msg)
        return False, error_msg


def _export_as_txt(filepath, content, title, source, published_date, 
                   imported_date, export_date, doc_class, provider, model, show_messages) -> Tuple[bool, str]:
    """Export document as plain text"""
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write("=" * 80 + "\n")
        f.write(f"{title}\n")
        f.write("=" * 80 + "\n\n")
        f.write("DOCUMENT INFORMATION:\n")
        f.write(f"  Title: {title}\n")
        f.write(f"  Source: {source}\n")
        if published_date:
            f.write(f"  Published: {published_date}\n")
        f.write(f"  Imported: {imported_date}\n")
        f.write(f"  Exported: {export_date}\n")
        f.write(f"  Type: {doc_class}\n")
        if provider and model:
            f.write(f"  Processed By: {provider} / {model}\n")
        f.write("\n" + "=" * 80 + "\n\n")
        f.write(content)
    
    if show_messages:
        from tkinter import messagebox
        messagebox.showinfo("Saved", f"✅ Text saved to:\n{filepath}")
    return True, filepath


def _export_as_docx(filepath, content, title, source, published_date,
                    imported_date, export_date, doc_class, provider, model, show_messages) -> Tuple[bool, str]:
    """Export document as Word docx"""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    except ImportError:
        if show_messages:
            from tkinter import messagebox
            messagebox.showerror("Error", "python-docx library not installed.\n\nInstall with: pip install python-docx")
        return False, "python-docx not installed"
    
    doc = Document()
    
    # Add title heading
    title_para = doc.add_heading(title[:100], 0)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add metadata section
    doc.add_heading('Document Information', level=2)
    meta_para = doc.add_paragraph()
    meta_para.add_run("Title: ").bold = True
    meta_para.add_run(f"{title}\n")
    meta_para.add_run("Source: ").bold = True
    # When the source is a URL (e.g. a YouTube video for a subscription
    # doc, or a Substack article), render it as a real clickable
    # hyperlink so the reader can jump straight from the docx to the
    # source.  Plain non-URL strings (filenames, "Unknown", etc.) keep
    # their original plain-text rendering.
    _src_str = str(source) if source else ""
    if _src_str.startswith(("http://", "https://")):
        from docx_helpers import add_external_hyperlink
        add_external_hyperlink(meta_para, _src_str, _src_str)
        meta_para.add_run("\n")
    else:
        meta_para.add_run(f"{_src_str}\n")
    if published_date:
        meta_para.add_run("Published: ").bold = True
        meta_para.add_run(f"{published_date}\n")
    meta_para.add_run("Imported: ").bold = True
    meta_para.add_run(f"{imported_date}\n")
    meta_para.add_run("Exported: ").bold = True
    meta_para.add_run(f"{export_date}\n")
    meta_para.add_run("Type: ").bold = True
    meta_para.add_run(f"{doc_class}\n")
    if provider and model:
        meta_para.add_run("Processed By: ").bold = True
        meta_para.add_run(f"{provider} / {model}\n")
    
    doc.add_paragraph()  # Spacing
    
    # Add content heading
    doc.add_heading('Content', level=2)
    
    # Add content with markdown formatting support
    _add_markdown_content_to_docx(doc, content)
    
    doc.save(filepath)
    
    if show_messages:
        from tkinter import messagebox
        messagebox.showinfo("Saved", f"✅ Document saved to:\n{filepath}")
    return True, filepath


def _export_as_rtf(filepath, content, title, source, published_date,
                   imported_date, export_date, doc_class, provider, model, show_messages) -> Tuple[bool, str]:
    """Export document as RTF"""
    try:
        rtf_content = []
        rtf_content.append(r'{\rtf1\ansi\deff0')
        rtf_content.append(r'{\fonttbl{\f0 Times New Roman;}}')
        rtf_content.append(r'\f0\fs24')
        
        # Title (large, bold)
        safe_title = title.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')
        rtf_content.append(r'{\b\fs32 ' + safe_title + r'}\par')
        rtf_content.append(r'\par')
        
        # Metadata section
        rtf_content.append(r'{\b\fs28 Document Information}\par')
        rtf_content.append(r'{\b Title: }' + safe_title + r'\par')
        safe_source = source.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')
        rtf_content.append(r'{\b Source: }' + safe_source + r'\par')
        if published_date:
            rtf_content.append(r'{\b Published: }' + published_date + r'\par')
        rtf_content.append(r'{\b Imported: }' + imported_date + r'\par')
        rtf_content.append(r'{\b Exported: }' + export_date + r'\par')
        rtf_content.append(r'{\b Type: }' + doc_class + r'\par')
        if provider and model:
            rtf_content.append(r'{\b Processed By: }' + f"{provider} / {model}" + r'\par')
        rtf_content.append(r'\par')
        
        # Content section
        rtf_content.append(r'{\b\fs28 Content}\par')
        for line in content.split('\n'):
            if line.strip():
                safe_line = line.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')
                rtf_content.append(safe_line + r'\par')
        
        rtf_content.append(r'}')
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write('\n'.join(rtf_content))
        
        if show_messages:
            from tkinter import messagebox
            messagebox.showinfo("Saved", f"✅ Document saved as RTF:\n{filepath}")
        return True, filepath
        
    except Exception as e:
        # Fallback to TXT
        txt_path = filepath.replace('.rtf', '.txt')
        if show_messages:
            from tkinter import messagebox
            messagebox.showerror("RTF Error", f"Failed to create RTF:\n{str(e)}\n\nSaving as TXT instead...")
        return _export_as_txt(txt_path, content, title, source, published_date,
                              imported_date, export_date, doc_class, provider, model, show_messages)


def _export_as_pdf(filepath, content, title, source, published_date,
                   imported_date, export_date, doc_class, provider, model, show_messages) -> Tuple[bool, str]:
    """Export document as PDF"""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import inch
        from reportlab.lib.enums import TA_CENTER
    except ImportError:
        # Fallback to TXT
        txt_path = filepath.replace('.pdf', '.txt')
        if show_messages:
            from tkinter import messagebox
            messagebox.showerror("Error", 
                "reportlab library not installed.\n\n"
                "Install with: pip install reportlab\n\n"
                "Falling back to TXT format...")
        return _export_as_txt(txt_path, content, title, source, published_date,
                              imported_date, export_date, doc_class, provider, model, show_messages)
    
    pdf_doc = SimpleDocTemplate(filepath, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=20,
        textColor='#2E4053',
        spaceAfter=20,
        alignment=TA_CENTER
    )
    story.append(Paragraph(title[:100], title_style))
    story.append(Spacer(1, 0.2 * inch))
    
    # Metadata section
    meta_style = ParagraphStyle(
        'Meta',
        parent=styles['Normal'],
        fontSize=10,
        textColor='#555555'
    )
    story.append(Paragraph("<b>Document Information</b>", styles['Heading2']))
    story.append(Paragraph(f"<b>Title:</b> {_escape_pdf_text(title)}", meta_style))
    story.append(Paragraph(f"<b>Source:</b> {_escape_pdf_text(source)}", meta_style))
    if published_date:
        story.append(Paragraph(f"<b>Published:</b> {published_date}", meta_style))
    story.append(Paragraph(f"<b>Imported:</b> {imported_date}", meta_style))
    story.append(Paragraph(f"<b>Exported:</b> {export_date}", meta_style))
    story.append(Paragraph(f"<b>Type:</b> {doc_class}", meta_style))
    if provider and model:
        story.append(Paragraph(f"<b>Processed By:</b> {provider} / {model}", meta_style))
    story.append(Spacer(1, 0.3 * inch))
    
    # Content section
    story.append(Paragraph("<b>Content</b>", styles['Heading2']))
    
    # Process content with markdown formatting and paragraph preservation.
    # Phase 1c fix2: pre-process the content so that [Back](#key-points)
    # links inside digest detail sections return to the originating
    # Key Points bullet (no-op on non-digest content).
    _split_into_pdf_paragraphs(_rewire_digest_back_links(content), story, styles['Normal'])
    
    pdf_doc.build(story)
    
    if show_messages:
        from tkinter import messagebox
        messagebox.showinfo("Saved", f"✅ Document saved as PDF:\n{filepath}")
    return True, filepath


def _escape_pdf_text(text: str) -> str:
    """Escape special characters for PDF/reportlab"""
    text = text.replace('&', '&amp;')
    text = text.replace('<', '&lt;')
    text = text.replace('>', '&gt;')
    return text


def _markdown_to_pdf_html(text: str) -> str:
    """
    Convert markdown formatting to reportlab-compatible HTML.
    Converts **bold** to <b>bold</b> and *italic* to <i>italic</i>.
    Also escapes special characters.
    """
    import re
    
    # First escape special characters (but not * which we need for markdown)
    text = text.replace('&', '&amp;')
    text = text.replace('<', '&lt;')
    text = text.replace('>', '&gt;')
    
    # Convert **bold** to <b>bold</b> (do this before italic)
    text = re.sub(r'\*\*([^*]+)\*\*', r'<b>\1</b>', text)
    
    # Convert *italic* to <i>italic</i>
    text = re.sub(r'\*([^*]+)\*', r'<i>\1</i>', text)
    
    # Convert # headers to bold (simplified for PDF)
    lines = text.split('\n')
    converted_lines = []
    for line in lines:
        if line.strip().startswith('#'):
            # Remove # and make bold
            header_text = line.lstrip('#').strip()
            converted_lines.append(f'<b>{header_text}</b>')
        else:
            converted_lines.append(line)
    
    return '\n'.join(converted_lines)


def _rewire_digest_back_links(content: str) -> str:
    """
    Rewire digest [Back](#key-points) links to return to the specific
    Key Points bullet that referenced each detail section, rather than
    to the top of Key Points.

    Scans the content for Key Points bullets that contain a
    [Detail](#point-N) style reference.  For the first bullet that
    references each #point-N, appends a {#kp-point-N} anchor marker to
    that bullet, and rewrites the [Back](#key-points) link inside the
    matching ### ... {#point-N} detail section to target #kp-point-N
    instead.  The renderer's existing {#id}-stripping logic then emits
    the correct <a name="kp-point-N"/> destination on the bullet, so
    tapping Back in the PDF lands at the start of the exact originating
    bullet.

    Safe fallbacks:
      - If a point is referenced from multiple bullets, only the first
        receives the reverse anchor; later bullets keep their Detail
        link without a return hook.
      - If a detail section's #point-N has no matching Key Points
        bullet, its [Back](#key-points) is left untouched (same
        behaviour as before this helper existed).
      - If the content contains no digest patterns at all, the input
        is returned unchanged — the helper is a no-op on non-digest
        content and so is safe to apply to any AI response.
    """
    if not content:
        return content

    lines = content.split('\n')

    # ── Pass 1: locate the first bullet that references each #point-N ──
    # Match [text](#point-<id>) where id is alphanumeric with dashes.
    # Bullets are identified by a leading "- " or "* " after stripping.
    bullet_detail_re = re.compile(r'\[[^\[\]]+\]\(#(point-[\w-]+)\)')
    point_to_bullet_idx: Dict[str, int] = {}

    for idx, line in enumerate(lines):
        s = line.strip()
        # Recognise markdown bullets ("- " / "* ") and literal U+2022
        # bullets ("• "), which is what the digest generator actually
        # emits for Key Points items.
        if not (s.startswith('- ') or s.startswith('* ') or s.startswith('\u2022 ')):
            continue
        for m in bullet_detail_re.finditer(line):
            point_id = m.group(1)
            if point_id not in point_to_bullet_idx:
                point_to_bullet_idx[point_id] = idx

    if not point_to_bullet_idx:
        return content  # No digest patterns — return untouched

    # ── Pass 2: inject {#kp-point-N} at the end of each tagged bullet ──
    # Skip injection if the anchor already exists (idempotent; safe to
    # run twice on the same content).
    for point_id, idx in point_to_bullet_idx.items():
        line = lines[idx]
        anchor_marker = f'{{#kp-{point_id}}}'
        if anchor_marker in line:
            continue
        stripped = line.rstrip()
        trailing_ws = line[len(stripped):]
        lines[idx] = f'{stripped} {anchor_marker}{trailing_ws}'

    # ── Pass 3: rewrite [Back](#key-points) inside matching detail sections ──
    # Walk lines tracking the current section's rewrite target.  A ### heading
    # with a known {#point-N} anchor sets the target; any ## heading or a
    # ### heading without a known anchor clears it.
    heading_anchor_re = re.compile(r'^###\s+.*?\{#(point-[\w-]+)\}')
    back_link_re = re.compile(r'\[([^\[\]]+)\]\(#key-points\)')

    current_target = None  # e.g. "kp-point-1" when inside a mapped section

    for idx, line in enumerate(lines):
        s = line.strip()
        if s.startswith('### '):
            m = heading_anchor_re.match(s)
            if m and m.group(1) in point_to_bullet_idx:
                current_target = f'kp-{m.group(1)}'
            else:
                current_target = None
            continue
        if s.startswith('## '):
            # New top-level section — detail sections have ended
            current_target = None
            continue
        if current_target and '(#key-points)' in line:
            lines[idx] = back_link_re.sub(
                lambda m, t=current_target: f'[{m.group(1)}](#{t})',
                line,
            )

    return '\n'.join(lines)


def _split_into_pdf_paragraphs(content: str, story: list, style):
    """
    Split AI response content into reportlab Paragraph/Spacer objects,
    correctly handling markdown bullets, numbered lists, headings, bold, italic.
    """
    from reportlab.platypus import Paragraph, Spacer
    from reportlab.lib.styles import ParagraphStyle

    # Indented style for bullet / numbered items
    indent_style = ParagraphStyle(
        '_IndentStyle', parent=style,
        leftIndent=20, spaceAfter=2
    )
    heading_style = ParagraphStyle(
        '_HeadingStyle', parent=style,
        fontName='Helvetica-Bold', fontSize=12, spaceBefore=8, spaceAfter=3
    )

    # ── Phase 1c fix1 marker ───────────────────────────────────────────
    # Pre-scan the whole content for Pandoc-style {#anchor-id} definitions
    # so inline() can detect and soften orphan internal links.  reportlab
    # refuses to build a PDF that contains a <link href="#missing"> to an
    # undefined destination, so any [text](#x) where #x is not defined
    # somewhere in this content gets rendered as plain text rather than a
    # link.  External links (href not starting with '#') always pass
    # through as <link href="...">.
    defined_anchors = set(re.findall(r'\{#([^}]+)\}', content))

    def inline(t: str) -> str:
        """Convert inline **bold** / *italic* / [text](url) to reportlab XML.

        # ── Phase 1c: internal link support ─────────────────────────────
        The [text](url) pass converts markdown links to reportlab's
        native <link href="url">text</link> tag, which every modern PDF
        viewer renders as a clickable navigation element.  Runs before
        bold/italic so **bold** or *italic* inside a link's text still
        gets converted.  Non-greedy brackets handle the [[Detail](#...)]
        double-bracket pattern by matching only the inner pair.

        Phase 1c fix1: internal anchor links whose target is not defined
        anywhere in this content are rendered as plain text, not <link>
        tags, because reportlab errors out at build time on unresolved
        internal destinations.  External links (https:// etc.) always
        pass through.
        """
        t = t.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

        def _link_sub(m):
            link_text = m.group(1)
            href      = m.group(2)
            if href.startswith('#') and href[1:] not in defined_anchors:
                return link_text
            # Phase 1c fix2: render links in traditional web-link blue
            # (#0645AD — darker than pure #0000EE so it reads comfortably
            # on printed PDFs as well as on-screen).  Applies to both
            # internal (#anchor) and external (https://...) links.
            return f'<link href="{href}" color="#0645AD">{link_text}</link>'

        t = re.sub(r'\[([^\[\]]+)\]\(([^)]+)\)', _link_sub, t)
        t = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', t)
        t = re.sub(r'\*(.*?)\*', r'<i>\1</i>', t)
        return t

    def safe_para(html: str, st) -> None:
        try:
            story.append(Paragraph(html, st))
        except Exception:
            plain = re.sub(r'<[^>]+>', '', html)
            try:
                story.append(Paragraph(plain[:2000], st))
            except Exception:
                pass

    ol_counter = 0
    saved_ol = 0
    in_ul = False

    for line in content.split('\n'):
        s = line.strip()

        # Blank line — close bullet list, small spacer
        if not s:
            if in_ul:
                in_ul = False
            story.append(Spacer(1, 3))
            continue

        # Headings
        # Phase 1c: a Pandoc-style {#anchor-id} is stripped from the
        # visible text and prepended as a reportlab <a name="..."/>
        # destination so internal [text](#anchor-id) links target it.
        # Phase 1c fix1: the regex is no longer end-anchored, so it also
        # catches headings whose {#id} is followed by trailing markup
        # like `## Sources {#sources} [[Back](#introduction)]`.
        if s.startswith('### '):
            in_ul = False; ol_counter = 0; saved_ol = 0
            raw = s[4:]
            m_anchor = re.search(r'\s*\{#([^}]+)\}\s*', raw)
            anchor_tag = ''
            if m_anchor:
                anchor_tag = f'<a name="{m_anchor.group(1)}"/>'
                raw = (raw[:m_anchor.start()] + ' ' + raw[m_anchor.end():]).strip()
            safe_para(f'{anchor_tag}{inline(raw)}', heading_style)
            story.append(Spacer(1, 2))
            continue
        if s.startswith('## ') or s.startswith('# '):
            in_ul = False; ol_counter = 0; saved_ol = 0
            text = s.lstrip('#').strip()
            m_anchor = re.search(r'\s*\{#([^}]+)\}\s*', text)
            anchor_tag = ''
            if m_anchor:
                anchor_tag = f'<a name="{m_anchor.group(1)}"/>'
                text = (text[:m_anchor.start()] + ' ' + text[m_anchor.end():]).strip()
            safe_para(f'{anchor_tag}<b>{inline(text)}</b>', heading_style)
            story.append(Spacer(1, 2))
            continue

        # Horizontal rule
        if s == '---':
            in_ul = False; ol_counter = 0; saved_ol = 0
            story.append(Spacer(1, 4))
            continue

        # Bullet point
        # Phase 1c fix2: a Pandoc-style {#anchor-id} anywhere in the
        # bullet text is stripped and prepended as a reportlab
        # <a name="..."/> destination so that internal [text](#anchor-id)
        # links can target a specific bullet.  Same pattern as the
        # heading handlers above; harmless on any bullet without an
        # anchor marker.  Used by _rewire_digest_back_links to make
        # [Back] links return to the exact Key Points bullet that
        # referenced each detail section.
        if s.startswith('- ') or s.startswith('* '):
            if not in_ul:
                # Entering bullet list — save any active numbered list counter
                if ol_counter > 0:
                    saved_ol = ol_counter
                in_ul = True
            raw = s[2:]
            m_anchor = re.search(r'\s*\{#([^}]+)\}\s*', raw)
            anchor_tag = ''
            if m_anchor:
                anchor_tag = f'<a name="{m_anchor.group(1)}"/>'
                raw = (raw[:m_anchor.start()] + ' ' + raw[m_anchor.end():]).strip()
            safe_para(f'{anchor_tag}\u2022\u00a0\u00a0{inline(raw)}', indent_style)
            story.append(Spacer(1, 2))
            continue

        # Numbered item
        if re.match(r'^\d+\.\s+', s):
            if in_ul:
                # Returning from bullet sub-list — restore counter
                in_ul = False
                if saved_ol > 0:
                    ol_counter = saved_ol
                    saved_ol = 0
            if ol_counter == 0 and saved_ol == 0:
                ol_counter = 0  # Fresh list
            ol_counter += 1
            text = re.sub(r'^\d+\.\s+', '', s)
            safe_para(f'<b>{ol_counter}.</b>\u00a0{inline(text)}', indent_style)
            story.append(Spacer(1, 2))
            continue

        # Regular paragraph
        # Phase 1c fix3: strip a trailing Pandoc-style {#anchor-id} and
        # emit it as a reportlab <a name="..."/> destination, mirroring
        # the heading and bullet handlers above.  This covers digest
        # content that uses literal "• " (U+2022) bullets rather than
        # markdown "- " / "* ", which fall through to this branch but
        # still need anchor support so that [Back] links target the
        # correct originating line.  Harmless on any paragraph without
        # an anchor marker.
        if in_ul:
            in_ul = False
        m_anchor = re.search(r'\s*\{#([^}]+)\}\s*', s)
        anchor_tag = ''
        if m_anchor:
            anchor_tag = f'<a name="{m_anchor.group(1)}"/>'
            s = (s[:m_anchor.start()] + ' ' + s[m_anchor.end():]).strip()
        safe_para(f'{anchor_tag}{inline(s)}', style)
        story.append(Spacer(1, 4))


# =============================================================================
# CONVERSATION THREAD EXPORT - For conversation threads
# =============================================================================

# ── Phase 1b(i) marker ───────────────────────────────────────────────────────
# The helpers below and the apply_opening_rules parameter were added in
# Phase 1b(i) of the export redesign.  They standardise the metadata
# header across all four export formats (via MetadataBlock) and let
# callers request the "drop opening prompt / suppress first AI avatar /
# keep avatars on follow-ups" rules that the Copy-for-Word/Email path
# already uses.  Callers that pre-build a MetadataBlock should place it
# under thread_metadata["metadata_block"]; otherwise a simple block is
# synthesised from the legacy flat-string fields for backward compatibility.


def _get_metadata_block(thread_metadata: Dict[str, Any]):
    """Return a MetadataBlock derived from thread_metadata.

    If the caller provided a pre-built block under the "metadata_block"
    key (the Thread Viewer does this so the header matches the copy
    path exactly), it is used verbatim.  Otherwise a basic block is
    built from the legacy doc_title / source_info / provider / model
    fields.  Never raises - falls back to an empty block on any error.
    """
    try:
        block = thread_metadata.get('metadata_block')
        if block is not None:
            return block
        from thread_viewer_metadata import MetadataBlock, Source
        sources = []
        src = thread_metadata.get('source_info') or ''
        if src and src != 'N/A':
            sources.append(Source(name=str(src).strip()))
        published = thread_metadata.get('published_date')
        published_str = str(published) if published and str(published) != 'N/A' else ''
        fetched = thread_metadata.get('fetched_date')
        fetched_str = str(fetched) if fetched and str(fetched) != 'N/A' else ''
        provider = thread_metadata.get('provider') or ''
        model    = thread_metadata.get('model') or ''
        return MetadataBlock(
            title=str(thread_metadata.get('doc_title', '') or ''),
            sources=sources,
            ai_provider='' if provider == 'N/A' else str(provider),
            ai_model='' if model == 'N/A' else str(model),
            published_date=published_str,
            imported_date=fetched_str,
        )
    except Exception:
        from thread_viewer_metadata import MetadataBlock
        return MetadataBlock()


def _thread_render_plan(messages: List[Dict], apply_opening_rules: bool):
    """Pre-compute which messages to skip and which AI turn is 'first'.

    Returns a tuple (skip_first_user: bool, first_assistant_idx: Optional[int]).

    When apply_opening_rules is False (the default), skip_first_user is
    False and first_assistant_idx is None - every message renders with
    its avatar, matching the pre-Phase-1b behaviour.

    When apply_opening_rules is True:
      * If the thread opens with a user message, it is treated as the
        "please summarise…" prompt that produced the first AI response
        and is skipped entirely from the output.
      * The first assistant message (after that skip, if any) is marked
        so the renderer can omit its avatar label - its own title /
        headings already demarcate it.
    """
    if not apply_opening_rules or not messages:
        return False, None
    skip_first_user = messages[0].get('role') == 'user'
    first_assistant_idx = None
    for i, m in enumerate(messages):
        if skip_first_user and i == 0:
            continue
        if m.get('role') == 'assistant':
            first_assistant_idx = i
            break
    return skip_first_user, first_assistant_idx


# ── Phase 1b(i) polish marker ───────────────────────────────────────────────
# The block= parameter and the MetadataBlock fallback were added in the
# Phase 1b(i) polish pass.  DocAnalyser's own threads don't always carry
# per-message provider/model keys, so without the fallback the follow-up
# AI turn renders as a bare "AI" instead of using the provider/model that
# was actually recorded in the document's metadata.
def _ai_avatar_label(msg: Dict, block=None) -> str:
    """Build the AI avatar label using per-message provider/model.

    Prefers msg['provider'] / msg['model'] when available so each turn
    displays the actual model that ran it (helpful if the user switched
    models between turns).  When the per-message fields are absent -
    the common case in DocAnalyser's own threads, which don't currently
    populate them - falls back to the MetadataBlock's ai_provider /
    ai_model (the provider/model recorded when the document was produced)
    rather than a bare "AI" label.  Returns a generic "AI" label only
    when neither source has a value.
    """
    provider = (msg.get('provider') or '').strip()
    model    = (msg.get('model') or '').strip()
    if not provider and block is not None:
        provider = (getattr(block, 'ai_provider', '') or '').strip()
    if not model and block is not None:
        model = (getattr(block, 'ai_model', '') or '').strip()
    timestamp = (msg.get('timestamp') or '').strip()
    time_str  = f" [{timestamp}]" if timestamp else ""
    if provider and model and model != provider:
        return f"🤖 {provider} ({model}){time_str}"
    label = provider or "AI"
    return f"🤖 {label}{time_str}"


def _user_avatar_label(msg: Dict) -> str:
    """Build the YOU avatar label, with the message timestamp if present."""
    timestamp = (msg.get('timestamp') or '').strip()
    time_str  = f" [{timestamp}]" if timestamp else ""
    return f"🧑 YOU{time_str}"


def export_conversation_thread(
    filepath: str,
    format: str,
    thread_messages: List[Dict],
    thread_metadata: Dict[str, Any],
    show_messages: bool = True,
    apply_opening_rules: bool = False,
) -> Tuple[bool, str]:
    """
    Export a conversation thread to file.
    
    Args:
        filepath: Full path to save the file
        format: Export format ('txt', 'docx', 'rtf', 'pdf')
        thread_messages: List of message dicts with 'role' and 'content' keys
        thread_metadata: Dict with keys:
            - doc_title: Source document title
            - source_info: Source URL/path
            - published_date: Original publication date (optional)
            - fetched_date: When document was imported
            - provider: AI provider
            - model: AI model
            - message_count: Number of exchanges
        show_messages: Whether to show success/error message boxes
    
    Returns:
        Tuple of (success: bool, message: str)
    """
    format = format.lower()
    current_time = datetime.datetime.now()
    
    try:
        if format == 'txt':
            return _export_thread_as_txt(filepath, thread_messages, thread_metadata, current_time, show_messages, apply_opening_rules)
        elif format == 'docx':
            return _export_thread_as_docx(filepath, thread_messages, thread_metadata, current_time, show_messages, apply_opening_rules)
        elif format == 'rtf':
            return _export_thread_as_rtf(filepath, thread_messages, thread_metadata, current_time, show_messages, apply_opening_rules)
        elif format == 'pdf':
            return _export_thread_as_pdf(filepath, thread_messages, thread_metadata, current_time, show_messages, apply_opening_rules)
        else:
            return False, f"Unsupported format: {format}"
    except Exception as e:
        error_msg = f"Export failed: {str(e)}"
        if show_messages:
            from tkinter import messagebox
            messagebox.showerror("Export Error", error_msg)
        return False, error_msg


def _export_thread_as_txt(filepath, messages, metadata, current_time,
                          show_messages, apply_opening_rules=False) -> Tuple[bool, str]:
    """Export conversation thread as plain text.

    Header comes from MetadataBlock (same set of fields as the other
    three formats).  When apply_opening_rules is True the opening user
    prompt is dropped and the first AI response is rendered without an
    avatar label; follow-up turns always keep their labels.
    """
    block = _get_metadata_block(metadata)
    skip_first_user, first_assistant_idx = _thread_render_plan(
        messages, apply_opening_rules)

    with open(filepath, 'w', encoding='utf-8') as f:
        # Header
        for line in block.to_save_plain_lines():
            f.write(line + "\n")
        f.write("\n" + "-" * 60 + "\n\n")

        # Messages
        for idx, msg in enumerate(messages):
            if skip_first_user and idx == 0:
                continue

            role    = msg.get('role', 'unknown')
            content = msg.get('content', '')

            if role == "user":
                f.write(_user_avatar_label(msg) + "\n")
                f.write("-" * 40 + "\n")
                f.write(content + "\n\n")
            else:
                # Assistant: omit avatar only for the first AI turn when
                # apply_opening_rules is active.
                if idx != first_assistant_idx:
                    f.write(_ai_avatar_label(msg, block) + "\n")
                    f.write("-" * 40 + "\n")
                f.write(content + "\n\n")
                # End-of-AI-turn divider - only when there's another
                # message after this one, so a one-shot briefing doesn't
                # end with a trailing divider.
                if idx < len(messages) - 1:
                    f.write("\u2500" * 60 + "\n\n")

    if show_messages:
        from tkinter import messagebox
        messagebox.showinfo("Saved", f"✅ Thread saved to:\n{filepath}")
    return True, filepath


def _export_thread_as_docx(filepath, messages, metadata, current_time,
                           show_messages, apply_opening_rules=False) -> Tuple[bool, str]:
    """Export conversation thread as a Word .docx file.

    Header comes from MetadataBlock.to_docx_runs().  When
    apply_opening_rules is True the opening user prompt is dropped and
    the first AI turn renders without its avatar label; follow-up turns
    keep their avatar labels.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    except ImportError:
        if show_messages:
            from tkinter import messagebox
            messagebox.showerror("Error",
                "python-docx library not installed.\n\n"
                "Install with: pip install python-docx")
        return False, "python-docx not installed"

    block = _get_metadata_block(metadata)
    skip_first_user, first_assistant_idx = _thread_render_plan(
        messages, apply_opening_rules)

    doc = Document()

    # Metadata header (title + info block + dividers)
    block.to_docx_runs(doc)
    doc.add_paragraph()  # breathing room before the thread

    # Thread
    for idx, msg in enumerate(messages):
        if skip_first_user and idx == 0:
            continue

        role    = msg.get('role', 'unknown')
        content = msg.get('content', '')

        if role == "user":
            header_para = doc.add_paragraph()
            run = header_para.add_run(_user_avatar_label(msg))
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(46, 64, 83)  # Dark blue
        else:
            # Assistant: omit avatar only for the first AI turn when
            # apply_opening_rules is active.
            if idx != first_assistant_idx:
                header_para = doc.add_paragraph()
                run = header_para.add_run(_ai_avatar_label(msg, block))
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(22, 83, 126)  # Blue

        # Body: markdown-formatted content
        _add_markdown_content_to_docx(doc, content)

        # End-of-AI-turn divider - only when there's another message
        # after this one.  Subtle centred em-dash rule in light grey,
        # smaller than the MetadataBlock header dividers so the visual
        # hierarchy stays "header > turn separator".
        if role == "assistant" and idx < len(messages) - 1:
            rule_para = doc.add_paragraph()
            rule_run = rule_para.add_run('\u2500' * 30)
            rule_run.font.size = Pt(8)
            rule_run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            rule_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Breathing room between turns
        if idx < len(messages) - 1:
            doc.add_paragraph()

    doc.save(filepath)

    if show_messages:
        from tkinter import messagebox
        messagebox.showinfo("Saved", f"✅ Thread saved as Word document:\n{filepath}")
    return True, filepath


def _export_thread_as_rtf(filepath, messages, metadata, current_time,
                          show_messages, apply_opening_rules=False) -> Tuple[bool, str]:
    """Export conversation thread as RTF.

    Header comes from MetadataBlock.to_rtf_lines().  When
    apply_opening_rules is True the opening user prompt is dropped and
    the first AI turn renders without its avatar label; follow-up turns
    keep their avatar labels.
    """
    block = _get_metadata_block(metadata)
    skip_first_user, first_assistant_idx = _thread_render_plan(
        messages, apply_opening_rules)

    def _rtf_esc(t):
        if not t:
            return ""
        t = str(t)
        t = t.replace('\\', '\\\\')
        t = t.replace('{', '\\{')
        t = t.replace('}', '\\}')
        return t

    try:
        rtf_content = []
        rtf_content.append(r'{\rtf1\ansi\deff0')
        rtf_content.append(r'{\fonttbl{\f0 Times New Roman;}}')
        rtf_content.append(r'\f0\fs24')

        # Metadata header (title + info block)
        rtf_content.extend(block.to_rtf_lines())

        # Thin rule
        rtf_content.append(r'\brdrb\brdrs\brdrw10\par\par')

        # Thread
        for idx, msg in enumerate(messages):
            if skip_first_user and idx == 0:
                continue

            role    = msg.get('role', 'unknown')
            content = msg.get('content', '')

            if role == "user":
                rtf_content.append(r'{\b ' + _rtf_esc(_user_avatar_label(msg)) + r'}\par')
            else:
                # Assistant: omit avatar only for the first AI turn when
                # apply_opening_rules is active.
                if idx != first_assistant_idx:
                    rtf_content.append(r'{\b ' + _rtf_esc(_ai_avatar_label(msg, block)) + r'}\par')

            for line in _rtf_esc(content).split('\n'):
                rtf_content.append(line + r'\par')
            rtf_content.append(r'\par')

            # End-of-AI-turn divider - only when there's another
            # message after this one.  Uses RTF's native bottom-border
            # rule, same primitive the header uses.
            if role == "assistant" and idx < len(messages) - 1:
                rtf_content.append(r'\brdrb\brdrs\brdrw10\par\par')

        rtf_content.append(r'}')

        with open(filepath, 'w', encoding='utf-8') as f:
            f.write('\n'.join(rtf_content))

        if show_messages:
            from tkinter import messagebox
            messagebox.showinfo("Saved", f"✅ Thread saved as RTF:\n{filepath}")
        return True, filepath

    except Exception as e:
        # Fallback to TXT on RTF failure
        txt_path = filepath.replace('.rtf', '.txt')
        if show_messages:
            from tkinter import messagebox
            messagebox.showerror("RTF Error",
                f"Failed to create RTF:\n{str(e)}\n\nSaving as TXT instead...")
        return _export_thread_as_txt(txt_path, messages, metadata, current_time,
                                     show_messages, apply_opening_rules)


def _export_thread_as_pdf(filepath, messages, metadata, current_time,
                          show_messages, apply_opening_rules=False) -> Tuple[bool, str]:
    """Export conversation thread as PDF.

    Header comes from MetadataBlock.to_pdf_story().  When
    apply_opening_rules is True the opening user prompt is dropped and
    the first AI turn renders without its avatar label; follow-up turns
    keep their avatar labels.
    """
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.colors import HexColor
    except ImportError:
        txt_path = filepath.replace('.pdf', '.txt')
        if show_messages:
            from tkinter import messagebox
            messagebox.showerror("Error",
                "reportlab library not installed.\n\n"
                "Install with: pip install reportlab\n\n"
                "Falling back to TXT format...")
        return _export_thread_as_txt(txt_path, messages, metadata, current_time,
                                     show_messages, apply_opening_rules)

    block = _get_metadata_block(metadata)
    skip_first_user, first_assistant_idx = _thread_render_plan(
        messages, apply_opening_rules)

    pdf_doc = SimpleDocTemplate(filepath, pagesize=letter)
    styles = getSampleStyleSheet()

    user_style = ParagraphStyle(
        'UserStyle',
        parent=styles['Normal'],
        fontSize=10,
        textColor=HexColor('#2E4053'),
        fontName='Helvetica-Bold',
        spaceBefore=10,
    )
    ai_style = ParagraphStyle(
        'AIStyle',
        parent=styles['Normal'],
        fontSize=10,
        textColor=HexColor('#16537E'),
        fontName='Helvetica-Bold',
        spaceBefore=10,
    )

    story = []

    # Metadata header (title + info block + dividers)
    story.extend(block.to_pdf_story(styles))

    # Thread
    for idx, msg in enumerate(messages):
        if skip_first_user and idx == 0:
            continue

        role    = msg.get('role', 'unknown')
        content = msg.get('content', '')

        if role == "user":
            # Strip the emoji for PDF - reportlab's default fonts don't
            # render emoji, and the plain text label reads cleanly enough.
            label = _user_avatar_label(msg).replace('🧑 ', '').strip()
            story.append(Paragraph(_escape_pdf_text(label), user_style))
        else:
            if idx != first_assistant_idx:
                label = _ai_avatar_label(msg, block).replace('🤖 ', '').strip()
                story.append(Paragraph(_escape_pdf_text(label), ai_style))

        # Phase 1c fix4: apply the same digest back-link rewiring in the
        # thread PDF export path — digests delivered as AI responses in
        # the Thread Viewer are exported via this function, not via
        # _export_as_pdf.  Safe per-message: each AI turn is processed
        # in isolation and the helper is a no-op on non-digest content.
        _split_into_pdf_paragraphs(_rewire_digest_back_links(content), story, styles['Normal'])

        # End-of-AI-turn divider - only when there's another message
        # after this one.  Thin light-grey HRFlowable; lighter than the
        # header divider so the visual hierarchy stays "header > turn".
        if role == "assistant" and idx < len(messages) - 1:
            story.append(HRFlowable(width="100%", thickness=0.5,
                                    color=HexColor('#DDDDDD'),
                                    spaceBefore=8, spaceAfter=8))

        story.append(Spacer(1, 12))

    pdf_doc.build(story)

    if show_messages:
        from tkinter import messagebox
        messagebox.showinfo("Saved", f"✅ Thread saved as PDF:\n{filepath}")
    return True, filepath


# =============================================================================
# SHARED HELPERS - Markdown processing for Word documents
# =============================================================================

def _add_markdown_content_to_docx(doc, content: str):
    """
    Add markdown-formatted content to a Word document.

    Supports:
      - Headings: # / ## / ### / ... (capped at level 4)
      - Bullets: - item, * item, • item
      - Numbered lists: 1. item, 2. item, ...
      - Regular paragraphs
      - Inline formatting via _add_inline_markdown_to_paragraph:
          **bold**, *italic*, [text](https://...), [text](#anchor)
      - Pandoc-style {#anchor-id} markers at the end of any block,
        which become real Word bookmarks (the destinations of internal
        [text](#anchor-id) hyperlinks).

    Applies _rewire_digest_back_links() first, so digest [Back] links
    return to the originating Key Points bullet rather than the top of
    the Key Points section.  Safe / no-op for non-digest content.

    Collapses consecutive blank lines to avoid double paragraph markers.
    """
    import re
    from docx_helpers import add_bookmark

    # Preprocess digest [Back] links so each detail section's "Back"
    # button targets the specific Key Points bullet that referenced it.
    # Helper is a no-op on non-digest content.
    content = _rewire_digest_back_links(content)

    lines = content.split('\n')
    last_was_empty = False

    for line in lines:
        # Skip empty lines but track for spacing
        if not line.strip():
            if not last_was_empty:
                last_was_empty = True
            continue

        last_was_empty = False
        s = line.strip()

        # ── Strip and remember a Pandoc-style {#anchor} marker ────────
        # Anchors can appear at the end of a heading, bullet, numbered
        # item, or regular paragraph.  We strip the marker from the
        # visible text and attach a bookmark to whichever paragraph we
        # build below.  Same logic as the PDF path's anchor handling.
        anchor_name = None
        m_anchor = re.search(r'\s*\{#([^}]+)\}\s*', s)
        if m_anchor:
            anchor_name = m_anchor.group(1)
            s = (s[:m_anchor.start()] + ' ' + s[m_anchor.end():]).strip()

        # Headings: # / ## / ### / ...
        if s.startswith('#'):
            level = 0
            for char in s:
                if char == '#':
                    level += 1
                else:
                    break
            header_text = s.lstrip('#').strip()
            level = min(level, 4)
            heading_para = doc.add_heading('', level=level)
            _add_inline_markdown_to_paragraph(heading_para, header_text)
            if anchor_name:
                add_bookmark(heading_para, anchor_name)
            continue

        # Bullets: -, *, • (U+2022).  • is what the digest generator
        # actually emits for Key Points items, so we recognise it
        # alongside the markdown forms.
        if s.startswith(('- ', '* ', '\u2022 ')):
            bullet_text = s[2:].strip()
            para = doc.add_paragraph(style='List Bullet')
            _add_inline_markdown_to_paragraph(para, bullet_text)
            if anchor_name:
                add_bookmark(para, anchor_name)
            continue

        # Numbered list: 1. item / 2. item / ...
        if re.match(r'^\d+\.\s+', s):
            num_text = re.sub(r'^\d+\.\s+', '', s)
            para = doc.add_paragraph(style='List Number')
            _add_inline_markdown_to_paragraph(para, num_text)
            if anchor_name:
                add_bookmark(para, anchor_name)
            continue

        # Regular paragraph with inline formatting
        para = doc.add_paragraph()
        _add_inline_markdown_to_paragraph(para, s)
        if anchor_name:
            add_bookmark(para, anchor_name)


def _add_inline_markdown_to_paragraph(paragraph, text: str):
    """
    Add text to a paragraph with inline markdown formatting.

    Supports:
      - **bold**
      - *italic*
      - [text](https://...)   — external hyperlinks
      - [text](#anchor-id)    — internal hyperlinks (link to a bookmark)

    Links are processed first because they're the highest-precedence
    pattern (a link's text could itself contain * or **, but bold/italic
    inside link text is intentionally NOT supported — keep it simple).
    Bold and italic are then applied to the spans of plain text between
    links.

    Edge case: [[Detail](#point-1)] (digest-style outer brackets) parses
    as literal '[' + hyperlink 'Detail' + literal ']' because the link
    regex requires non-bracket chars inside the label.  That's the
    desired behaviour — the outer brackets render as plain text and only
    the inner [Detail](#point-1) becomes a clickable link.
    """
    import re
    from docx_helpers import add_external_hyperlink, add_internal_hyperlink

    if not text:
        return

    link_pattern = re.compile(r'\[([^\[\]]+)\]\(([^)]+)\)')

    pos = 0
    for m in link_pattern.finditer(text):
        # Render any plain text before this link with bold/italic applied.
        if m.start() > pos:
            _add_styled_runs(paragraph, text[pos:m.start()])

        label  = m.group(1)
        target = m.group(2)

        if target.startswith('#'):
            add_internal_hyperlink(paragraph, target[1:], label)
        elif target.startswith(('http://', 'https://', 'mailto:')):
            add_external_hyperlink(paragraph, target, label)
        else:
            # Unknown target type — render the original markdown verbatim
            # rather than producing a broken link.
            paragraph.add_run(f'[{label}]({target})')

        pos = m.end()

    # Render any remaining text after the last link.
    if pos < len(text):
        _add_styled_runs(paragraph, text[pos:])


def _add_styled_runs(paragraph, text: str):
    """Apply **bold** / *italic* substitutions to a span of plain text and
    append the resulting runs to `paragraph`.  Used by
    _add_inline_markdown_to_paragraph for the spans BETWEEN links.
    """
    import re
    if not text:
        return

    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*)'
    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**') and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)
