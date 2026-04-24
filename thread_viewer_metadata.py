"""
thread_viewer_metadata.py - Single source of truth for document metadata display

This module centralises how a document's metadata block is built and rendered,
removing the duplication that previously lived inline across many copy paths
in thread_viewer_copy.py.

Usage pattern
-------------
    block = MetadataBlock.from_document(doc, provider, model)
    plain_lines = block.to_plain_lines()
    html_parts  = block.to_html_parts()
    whatsapp    = block.to_whatsapp_lines()

Design
------
MetadataBlock is a plain data holder.  from_document() is the one place that
decides what goes into each field based on the document's metadata dict.  The
three render methods are each responsible for one output surface:

    to_plain_lines()     - returns list[str]  for plain-text copies
    to_html_parts()      - returns list[str]  of HTML snippets (inline-styled)
    to_whatsapp_lines()  - returns list[str]  for WhatsApp/Telegram copies

Keeping the structured representation separate from the rendering makes it
easy to add new fields or reshape the output in the future without hunting
through eight copy methods.
"""

from __future__ import annotations

import datetime
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional


# ─── Date formatting ─────────────────────────────────────────────────────────
#
# The app's shared utility format_display_date() produces "17-Apr-2026"
# (hyphen-separated).  The new metadata format uses space-separated months
# like "17 Apr 2026", which reads better in prose-style "Title: …" lines.
# We format locally here rather than touching format_display_date() because
# other parts of the app depend on the hyphenated form.

def _format_date_spaced(date_input: Any) -> str:
    """Return a date as '17 Apr 2026' or empty string on failure."""
    if not date_input:
        return ""

    if isinstance(date_input, datetime.datetime):
        return date_input.strftime("%d %b %Y").lstrip("0")

    date_str = str(date_input).strip()

    # Try common formats first — matches the logic in utils.format_display_date
    formats_to_try = [
        "%Y%m%d",
        "%Y-%m-%d",
        "%Y-%m-%dT%H:%M:%S.%f",
        "%Y-%m-%dT%H:%M:%S.%f%z",
        "%Y-%m-%dT%H:%M:%S%z",
        "%Y-%m-%dT%H:%M:%SZ",
        "%Y-%m-%dT%H:%M:%S",
    ]
    for fmt in formats_to_try:
        try:
            dt = datetime.datetime.strptime(date_str, fmt)
            # Use a cross-platform leading-zero strip
            formatted = dt.strftime("%d %b %Y")
            if formatted.startswith("0"):
                formatted = formatted[1:]
            return formatted
        except ValueError:
            continue

    try:
        from dateutil import parser
        dt = parser.parse(date_str)
        formatted = dt.strftime("%d %b %Y")
        if formatted.startswith("0"):
            formatted = formatted[1:]
        return formatted
    except Exception:
        pass

    # Already formatted as "17-Apr-2026" - convert to spaces
    if "-" in date_str and len(date_str) >= 9:
        try:
            dt = datetime.datetime.strptime(date_str, "%d-%b-%Y")
            formatted = dt.strftime("%d %b %Y")
            if formatted.startswith("0"):
                formatted = formatted[1:]
            return formatted
        except ValueError:
            pass

    return date_str


# ─── Data holders ────────────────────────────────────────────────────────────

@dataclass
class Source:
    """One source in the Source / Source(s) line.

    For a per-item subscription doc the block has a single Source with the
    analyst's name and optionally an interviewee.  For a digest the block has
    multiple Sources, one per subscription that contributed to the digest.
    """
    name: str
    interviewee: str = ""

    def to_display_string(self) -> str:
        """Render this single source as '<name> (interviewee: <x>)' or '<name>'."""
        if self.interviewee:
            return f"{self.name} (interviewee: {self.interviewee})"
        return self.name


@dataclass
class MetadataBlock:
    """Structured representation of a document's metadata header.

    Populate via MetadataBlock.from_document(doc, provider, model).
    Render via the to_plain_lines / to_html_parts / to_whatsapp_lines methods.
    """
    title: str = ""
    sources: List[Source] = field(default_factory=list)
    ai_provider: str = ""
    ai_model: str = ""
    published_date: str = ""   # Empty string if not available
    imported_date: str = ""    # Empty string if not available
    url: str = ""              # Source URL (per-item docs only; digests hold URLs per-source)
    is_digest: bool = False

    # ── Builder ──────────────────────────────────────────────────────────────

    @classmethod
    def from_document(
        cls,
        doc: Optional[Dict[str, Any]],
        fallback_provider: str = "",
        fallback_model: str = "",
        fallback_title: str = "",
        fallback_source_name: str = "",
    ) -> "MetadataBlock":
        """Build a MetadataBlock from a document dict.

        The AI provider/model come from the document's metadata if present
        (which is what was actually used to generate the response), falling
        back to the live provider/model values passed in.  This is more
        accurate than always using the live vars, which can drift if the
        user changes the dropdown after running a prompt.
        """
        if not doc:
            return cls(
                title=fallback_title,
                sources=[Source(name=fallback_source_name)] if fallback_source_name else [],
                ai_provider=fallback_provider,
                ai_model=fallback_model,
            )

        meta = doc.get("metadata") or {}
        if not isinstance(meta, dict):
            meta = {}

        title = doc.get("title") or fallback_title

        # AI provider / model: prefer metadata (what actually ran); fall back
        # to the live dropdown vars only if metadata is silent.
        ai_provider = meta.get("ai_provider") or fallback_provider
        ai_model    = meta.get("ai_model")    or fallback_model

        # Published date (per-item docs) and imported date (both).
        published_date = _format_date_spaced(meta.get("published_date", ""))
        imported_date  = _format_date_spaced(doc.get("fetched", ""))

        # ── Decide whether this is a digest and build the sources list ──────
        is_digest = bool(meta.get("digest"))
        sources: List[Source] = []

        if is_digest:
            # New-format digests carry a "sources" list of {name, interviewee}
            # dicts.  Older digests only have the flat "subscription_names"
            # list — fall back to that for backward compatibility.
            raw_sources = meta.get("sources")
            if isinstance(raw_sources, list) and raw_sources:
                for item in raw_sources:
                    if isinstance(item, dict):
                        sources.append(Source(
                            name=str(item.get("name", "")).strip(),
                            interviewee=str(item.get("interviewee", "")).strip(),
                        ))
                    elif isinstance(item, str):
                        # Tolerate a list of bare name strings.
                        sources.append(Source(name=item.strip()))
            else:
                legacy_names = meta.get("subscription_names") or []
                if isinstance(legacy_names, list):
                    for name in legacy_names:
                        if name:
                            sources.append(Source(name=str(name).strip()))

            # If we still have no sources (very old digest with neither
            # field), fall back to the document's source string.
            if not sources:
                legacy_source = doc.get("source") or fallback_source_name
                if legacy_source:
                    sources.append(Source(name=str(legacy_source)))

        else:
            # Per-item document: build a single Source.  For subscription
            # docs use the subscription_name + interviewee from metadata;
            # otherwise fall back to the document's source field.
            sub_name    = meta.get("subscription_name", "")
            interviewee = meta.get("interviewee", "")

            if sub_name:
                sources.append(Source(
                    name=str(sub_name).strip(),
                    interviewee=str(interviewee).strip() if interviewee else "",
                ))
            else:
                # Not a subscription doc — use the document's own source string
                fallback = doc.get("source") or fallback_source_name
                if fallback:
                    sources.append(Source(name=str(fallback)))

        # Per-item docs may carry a `url` in metadata (set when the source
        # was fetched from YouTube / Substack / RSS / podcast).  Digests
        # leave this blank because they aggregate many sources - each
        # source's URL is surfaced inline in the Sources section of the
        # digest output instead.
        url = "" if is_digest else str(meta.get("url", "") or "").strip()

        return cls(
            title=title,
            sources=sources,
            ai_provider=ai_provider,
            ai_model=ai_model,
            published_date=published_date,
            imported_date=imported_date,
            url=url,
            is_digest=is_digest,
        )

    # ── Helpers ──────────────────────────────────────────────────────────────

    def source_label(self) -> str:
        """Return 'Source(s)' for digests with multiple sources, else 'Source'."""
        if self.is_digest and len(self.sources) > 1:
            return "Source(s)"
        return "Source"

    def sources_display(self) -> str:
        """Render the sources list as a single comma-separated string.

        Used by all three renderers.  Returns empty string if no sources.
        """
        if not self.sources:
            return ""
        return ", ".join(s.to_display_string() for s in self.sources)

    def ai_display(self) -> str:
        """Render the AI line value, e.g. 'Anthropic (Claude) / claude-opus-4-6'."""
        if not self.ai_provider and not self.ai_model:
            return ""
        if self.ai_provider and self.ai_model:
            return f"{self.ai_provider} / {self.ai_model}"
        return self.ai_provider or self.ai_model

    # ── Plain text rendering ─────────────────────────────────────────────────

    def to_plain_lines(self) -> List[str]:
        """Return the metadata block as a list of plain-text lines.

        Produces, for example:
            Title: Subscription digest 17 Apr 2026
            Source(s): Alexander Mercouris (interviewee: Pepe Escobar), Glenn Diesen
            AI: Anthropic (Claude) / claude-opus-4-6
        """
        lines: List[str] = []
        if self.title:
            lines.append(f"Title: {self.title}")

        sources_str = self.sources_display()
        if sources_str:
            lines.append(f"{self.source_label()}: {sources_str}")

        ai_str = self.ai_display()
        if ai_str:
            lines.append(f"AI: {ai_str}")

        return lines

    # ── HTML rendering (Gmail-compatible inline styles) ──────────────────────

    def to_html_parts(self, escape_html_fn=None) -> List[str]:
        """Return the metadata block as a list of HTML snippets.

        Each snippet is a complete element (e.g. '<p …>…</p>').  The caller
        glues them together with '\n'.join(...) and wraps in a document.

        Args:
            escape_html_fn: optional callable that escapes HTML special chars.
                Pass in self._escape_html from CopyMixin to stay consistent
                with the rest of the file.  Falls back to a minimal built-in.
        """
        if escape_html_fn is None:
            def escape_html_fn(s):  # type: ignore
                if s is None:
                    return ""
                return (str(s)
                        .replace("&", "&amp;")
                        .replace("<", "&lt;")
                        .replace(">", "&gt;")
                        .replace('"', "&quot;"))

        parts: List[str] = []

        # Title — centred bold paragraph so it stands out at the top.
        if self.title:
            parts.append(
                f'<p style="color: #2C3E50; font-size: 16pt; font-weight: bold; '
                f'text-align: center; margin: 0 0 8pt 0;">'
                f'{escape_html_fn(self.title)}</p>'
            )

        # Thin rule separator under the title.
        parts.append(
            '<p style="border-bottom: 1px solid #ccc; margin: 0 0 8pt 0; '
            'padding: 0; line-height: 1px;">&nbsp;</p>'
        )

        # Source(s) / URL / Published / Imported / AI lines, rendered as a
        # single small grey paragraph.  Each line is emitted only when it
        # has a value, so per-item docs and digests both look tidy without
        # empty "URL:" or "Published:" rows.
        info_lines: List[str] = []

        sources_str = self.sources_display()
        if sources_str:
            info_lines.append(
                f'<b>{self.source_label()}:</b> {escape_html_fn(sources_str)}'
            )

        # URL is rendered as a real <a href> so Word / Outlook turn it
        # into a clickable link.  Digests skip this line - the AI
        # surfaces per-source URLs inline in the Sources section.
        if self.url and not self.is_digest:
            safe_url = escape_html_fn(self.url)
            info_lines.append(
                f'<b>URL:</b> <a href="{safe_url}">{safe_url}</a>'
            )

        if self.published_date:
            info_lines.append(
                f'<b>Published:</b> {escape_html_fn(self.published_date)}'
            )

        if self.imported_date:
            info_lines.append(
                f'<b>Imported:</b> {escape_html_fn(self.imported_date)}'
            )

        ai_str = self.ai_display()
        if ai_str:
            info_lines.append(f'<b>AI:</b> {escape_html_fn(ai_str)}')

        if info_lines:
            parts.append(
                '<p style="font-size: 10pt; color: #555; margin: 4pt 0;">'
                + '<br>'.join(info_lines) +
                '</p>'
            )

        # Closing rule.
        parts.append(
            '<p style="border-bottom: 1px solid #ccc; margin: 0 0 8pt 0; '
            'padding: 0; line-height: 1px;">&nbsp;</p>'
        )

        return parts

    # ── Save-file rendering (Phase 1b) ─────────────────────────────────────
    #
    # The three renderers below (txt / docx / rtf / pdf) are driven by a
    # single helper, _info_field_pairs(), so the set of fields that appear
    # in the header stays consistent across all save surfaces.  The set
    # matches to_html_parts() - Source(s), URL, Published, Imported, AI -
    # but is emitted as plain (label, value) tuples so each format can
    # decide how to lay them out.

    def _info_field_pairs(self):
        """Return (label, value) tuples for every non-empty info field.

        Order matches to_html_parts():
            Source(s)    always first when populated
            URL          per-item docs only (digests suppress this and
                         surface per-source URLs inline in the body)
            Published    per-item docs
            Imported     always if known
            AI           provider / model
        """
        pairs = []
        sources_str = self.sources_display()
        if sources_str:
            pairs.append((self.source_label(), sources_str))
        if self.url and not self.is_digest:
            pairs.append(("URL", self.url))
        if self.published_date:
            pairs.append(("Published", self.published_date))
        if self.imported_date:
            pairs.append(("Imported", self.imported_date))
        ai_str = self.ai_display()
        if ai_str:
            pairs.append(("AI", ai_str))
        return pairs

    def to_save_plain_lines(self) -> List[str]:
        """Return the metadata block as plain-text lines for a .txt save.

        Unlike to_plain_lines() - which is tuned for short plain-text
        clipboard copies and emits only Title / Source / AI - this version
        emits every header field (Title, Source(s), URL, Published,
        Imported, AI) so the saved file carries the same information as
        the .docx / .pdf surfaces.
        """
        lines: List[str] = []
        if self.title:
            lines.append(f"Title: {self.title}")
        for label, value in self._info_field_pairs():
            lines.append(f"{label}: {value}")
        return lines

    def to_docx_runs(self, doc) -> None:
        """Add the metadata header to a python-docx Document in place.

        Renders a centred bold title, a thin divider, one small grey line
        per info field, and a closing divider.  Leaves the document in a
        state where the caller can append body paragraphs immediately.
        """
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        if self.title:
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(self.title)
            title_run.bold = True
            title_run.font.size = Pt(16)
            title_run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        sep1 = doc.add_paragraph('\u2500' * 60)
        sep1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for label, value in self._info_field_pairs():
            para = doc.add_paragraph()
            label_run = para.add_run(f"{label}: ")
            label_run.bold = True
            label_run.font.size = Pt(10)
            label_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            value_run = para.add_run(str(value))
            value_run.font.size = Pt(10)
            value_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            para.paragraph_format.space_after = Pt(2)

        sep2 = doc.add_paragraph('\u2500' * 60)
        sep2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def to_rtf_lines(self) -> List[str]:
        """Return the metadata header as RTF code lines.

        Caller glues them in between the RTF header (``{\\rtf1...``)
        and the body.  Braces and backslashes are escaped inside field
        values to stop them from corrupting the RTF structure.
        """
        def _rtf_esc(t):
            if not t:
                return ""
            t = str(t)
            t = t.replace('\\', '\\\\')
            t = t.replace('{', '\\{')
            t = t.replace('}', '\\}')
            return t

        lines: List[str] = []
        if self.title:
            lines.append(r'\pard\qc\b\fs32 ' + _rtf_esc(self.title)
                         + r'\par\b0\fs22\ql\par')
        lines.append(r'\pard\fs20')
        for label, value in self._info_field_pairs():
            lines.append(r'{\b ' + label + r': }' + _rtf_esc(value) + r'\par')
        lines.append(r'\fs22\par')
        return lines

    def to_pdf_story(self, styles) -> list:
        """Return a list of reportlab Flowables for the metadata header.

        Args:
            styles: the result of reportlab.lib.styles.getSampleStyleSheet().
        """
        # ── Phase 1b(i) PDF divider fix ─────────────────────────────
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.colors import HexColor
        from reportlab.platypus import Paragraph, Spacer, HRFlowable

        def _esc(s):
            return (str(s)
                    .replace('&', '&amp;')
                    .replace('<', '&lt;')
                    .replace('>', '&gt;'))

        title_style = ParagraphStyle(
            '_MBTitle', parent=styles['Heading1'],
            fontSize=16, textColor=HexColor('#2C3E50'),
            alignment=TA_CENTER, spaceAfter=4,
        )
        info_style = ParagraphStyle(
            '_MBInfo', parent=styles['Normal'],
            fontSize=10, textColor=HexColor('#555555'),
            spaceAfter=2,
        )

        story = []
        if self.title:
            story.append(Paragraph(_esc(self.title), title_style))
        story.append(HRFlowable(width="100%", thickness=0.5,
                                color=HexColor('#CCCCCC'),
                                spaceBefore=2, spaceAfter=6))

        for label, value in self._info_field_pairs():
            if label == "URL":
                safe_url = _esc(value)
                story.append(Paragraph(
                    f'<b>{label}:</b> <link href="{safe_url}" color="blue">'
                    f'{safe_url}</link>',
                    info_style,
                ))
            else:
                story.append(Paragraph(f'<b>{label}:</b> {_esc(value)}', info_style))

        story.append(HRFlowable(width="100%", thickness=0.5,
                                color=HexColor('#CCCCCC'),
                                spaceBefore=6, spaceAfter=10))
        return story

    # ── WhatsApp rendering ───────────────────────────────────────────────────

    def to_whatsapp_lines(self) -> List[str]:
        """Return the metadata block as WhatsApp-friendly plain-text lines.

        Uses a row of em-dashes as a visual separator above and below the
        source/AI lines, and wraps the title in *…* to make it bold.
        """
        lines: List[str] = []
        if self.title:
            lines.append(f"*{self.title}*")
            lines.append("————————————————————")

        sources_str = self.sources_display()
        if sources_str:
            lines.append(f"{self.source_label()}: {sources_str}")

        ai_str = self.ai_display()
        if ai_str:
            lines.append(f"AI: {ai_str}")

        lines.append("————————————————————")
        lines.append("")
        return lines
