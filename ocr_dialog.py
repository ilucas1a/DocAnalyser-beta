"""
ocr_dialog.py - Multi-Image OCR Dialog
Modal dialog for processing multiple images/pages as a single document.
Follows the same UX pattern as dictation_dialog.py for consistency.

Supports both image files and PDFs. PDFs are converted to temporary images for processing.
All OCR processing logic is in ocr_handler.py - this module is purely UI.
"""

import os
import re
import tempfile
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from typing import List


class MultiImageOCRDialog:
    """
    Modal dialog for OCR processing multiple images into a single document.
    Mirrors the dictation dialog workflow: add pages, process all, get combined result.
    
    Accepts both image files and PDFs. PDFs are automatically converted to images.
    """
    
    # Supported image formats
    IMAGE_EXTENSIONS = ('.png', '.jpg', '.jpeg', '.tif', '.tiff', '.bmp', '.gif', '.webp')
    PDF_EXTENSIONS = ('.pdf',)
    ALL_EXTENSIONS = IMAGE_EXTENSIONS + PDF_EXTENSIONS
    
    def __init__(self, parent, app):
        self.parent = parent
        self.app = app
        self.image_files: List[str] = []  # List of image file paths (including converted PDF pages)
        self.temp_files: List[str] = []   # Temp files to clean up when dialog closes
        self.ocr_result = None
        self.processing = False
        self.cancelled = False
        
        # Get OCR settings from app config
        self.ocr_mode = app.config.get("ocr_mode", "local_first")
        self.ocr_language = app.config.get("ocr_language", "eng")
        self.ocr_quality = app.config.get("ocr_quality", "balanced")
        self.confidence_threshold = app.config.get("ocr_confidence_threshold", 60)
        
        # Get AI settings for cloud OCR
        self.provider = app.provider_var.get()
        self.model = app.model_var.get()
        self.api_key = app.config.get("keys", {}).get(self.provider, "")
        
        # Text type setting (printed vs handwriting)
        self.text_type = app.config.get("ocr_text_type", "printed")
        
        # Create dialog window
        self.dialog = tk.Toplevel(parent)
        self.dialog.title("📄 Multi-Page OCR")
        self.dialog.geometry("540x620")
        self.dialog.resizable(True, True)
        self.dialog.minsize(480, 500)
        self.dialog.transient(parent)
        # Note: Not using grab_set() as it interferes with drag-drop from Explorer
        # self.dialog.grab_set()
        
        # Center on parent
        self.dialog.update_idletasks()
        x = parent.winfo_x() + (parent.winfo_width() - 540) // 2
        y = parent.winfo_y() + (parent.winfo_height() - 620) // 2
        self.dialog.geometry(f"+{x}+{y}")
        
        # Ensure dialog is on top and has focus
        self.dialog.lift()
        self.dialog.focus_force()
        
        self._setup_ui()
        
        # Setup drag-drop
        self._setup_drag_drop()
        
        # Handle window close
        self.dialog.protocol("WM_DELETE_WINDOW", self._on_close)
    
    def _setup_ui(self):
        """Create the dialog UI."""
        main_frame = ttk.Frame(self.dialog, padding=15)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Title
        title_frame = ttk.Frame(main_frame)
        title_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(
            title_frame, 
            text="📄 Multi-Page OCR",
            font=('Arial', 14, 'bold')
        ).pack(side=tk.LEFT)
        
        # Page counter
        self.page_count_label = ttk.Label(
            title_frame,
            text="",
            font=('Arial', 10),
            foreground='#666666'
        )
        self.page_count_label.pack(side=tk.RIGHT)
        
        # Instructions with drag-drop hint
        ttk.Label(
            main_frame,
            text="Drag and drop file(s) from File Explorer onto the list below,\nor copy file(s) in File Explorer (Ctrl+C) and paste here (Ctrl+V).",
            font=('Arial', 9),
            foreground='gray',
            wraplength=480
        ).pack(fill=tk.X, pady=(0, 10))
        
        # Image list frame with drag-drop target styling
        self.list_frame = ttk.LabelFrame(main_frame, text="Pages to Process", padding=10)
        self.list_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        
        # Listbox with scrollbar
        list_container = ttk.Frame(self.list_frame)
        list_container.pack(fill=tk.BOTH, expand=True)
        
        self.image_listbox = tk.Listbox(
            list_container,
            selectmode=tk.EXTENDED,
            font=('Arial', 10),
            bg='#FFFDE6',  # Cream background
            height=8
        )
        self.image_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        scrollbar = ttk.Scrollbar(list_container, orient=tk.VERTICAL, command=self.image_listbox.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.image_listbox.config(yscrollcommand=scrollbar.set)
        
        # Placeholder text when empty
        self._show_placeholder()
        
        # Bind keyboard shortcuts
        self.image_listbox.bind('<Delete>', lambda e: self._remove_selected())
        self.image_listbox.bind('<BackSpace>', lambda e: self._remove_selected())
        self.image_listbox.bind('<<ListboxSelect>>', self._on_selection_change)
        
        # Bind Ctrl+V for paste (on both dialog and listbox)
        self.dialog.bind('<Control-v>', self._paste_files)
        self.dialog.bind('<Control-V>', self._paste_files)
        self.image_listbox.bind('<Control-v>', self._paste_files)
        self.image_listbox.bind('<Control-V>', self._paste_files)
        
        # List management buttons
        list_btn_frame = ttk.Frame(self.list_frame)
        list_btn_frame.pack(fill=tk.X, pady=(10, 0))
        
        self.add_btn = ttk.Button(
            list_btn_frame,
            text="📁 Open File Explorer",
            command=self._open_file_explorer,
            width=20
        )
        self.add_btn.pack(side=tk.LEFT, padx=(0, 5))
        
        self.remove_btn = ttk.Button(
            list_btn_frame,
            text="🗑️ Remove",
            command=self._remove_selected,
            width=14,
            state=tk.DISABLED
        )
        self.remove_btn.pack(side=tk.LEFT, padx=(0, 5))
        
        self.clear_btn = ttk.Button(
            list_btn_frame,
            text="🧹 Clear All",
            command=self._clear_all,
            width=10,
            state=tk.DISABLED
        )
        self.clear_btn.pack(side=tk.LEFT)
        
        # Move up/down buttons
        ttk.Separator(list_btn_frame, orient=tk.VERTICAL).pack(side=tk.LEFT, padx=10, fill=tk.Y)
        
        self.move_up_btn = ttk.Button(
            list_btn_frame,
            text="⬆️",
            command=self._move_up,
            width=3,
            state=tk.DISABLED
        )
        self.move_up_btn.pack(side=tk.LEFT, padx=(0, 2))
        
        self.move_down_btn = ttk.Button(
            list_btn_frame,
            text="⬇️",
            command=self._move_down,
            width=3,
            state=tk.DISABLED
        )
        self.move_down_btn.pack(side=tk.LEFT)
        
        # Status/progress area
        self.status_frame = ttk.Frame(main_frame)
        self.status_frame.pack(fill=tk.X, pady=(0, 10))
        
        self.status_label = ttk.Label(
            self.status_frame,
            text="",
            font=('Arial', 10),
            foreground='gray'
        )
        self.status_label.pack()
        
        self.progress_var = tk.DoubleVar(value=0)
        self.progress_bar = ttk.Progressbar(
            self.status_frame,
            variable=self.progress_var,
            maximum=100,
            mode='determinate'
        )
        # Progress bar hidden initially
        
        # Text Type selector (Printed vs Handwriting)
        text_type_frame = ttk.Frame(main_frame)
        text_type_frame.pack(fill=tk.X, pady=(0, 10))
        
        text_type_label = ttk.Label(
            text_type_frame,
            text="Text Type:",
            font=('Arial', 10)
        )
        text_type_label.pack(side=tk.LEFT, padx=(0, 10))
        
        # Add context help to label
        try:
            from context_help import add_help, HELP_TEXTS
            if HELP_TEXTS and 'ocr_text_type_dropdown' in HELP_TEXTS:
                add_help(text_type_label, **HELP_TEXTS['ocr_text_type_dropdown'])
        except ImportError:
            pass
        
        self.text_type_var = tk.StringVar(value=self.text_type)
        
        # Import text type options
        from ocr_handler import OCR_TEXT_TYPES
        
        text_type_combo = ttk.Combobox(
            text_type_frame,
            textvariable=self.text_type_var,
            state="readonly",
            width=60
        )
        text_type_combo['values'] = [
            f"{info['label']}" for info in OCR_TEXT_TYPES.values()
        ]
        # Set current selection
        for i, (key, info) in enumerate(OCR_TEXT_TYPES.items()):
            if key == self.text_type:
                text_type_combo.current(i)
                break
        text_type_combo.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        # Bind selection change to update mode description
        text_type_combo.bind('<<ComboboxSelected>>', self._on_text_type_change)
        
        # Store reference for later
        self.text_type_combo = text_type_combo
        
        # Add context help to text type dropdown
        try:
            from context_help import add_help, HELP_TEXTS
            if HELP_TEXTS and 'ocr_text_type_dropdown' in HELP_TEXTS:
                add_help(text_type_combo, **HELP_TEXTS['ocr_text_type_dropdown'])
        except ImportError:
            pass
        
        # Context hint frame (for handwriting mode)
        self.context_frame = ttk.LabelFrame(main_frame, text="Document Context (helps with difficult handwriting)", padding=5)
        # Initially hidden - shown when handwriting is selected
        
        context_hint_text = (
            "Describe the document to help the AI interpret unclear words.\n"
            "Example: \"Letter from father to daughter, ~1975. Names: John, Margaret. Topics: family news.\""
        )
        ttk.Label(
            self.context_frame,
            text=context_hint_text,
            font=('Arial', 8),
            foreground='gray',
            wraplength=460
        ).pack(fill=tk.X, pady=(0, 5))
        
        self.context_var = tk.StringVar(value="")
        self.context_entry = ttk.Entry(
            self.context_frame,
            textvariable=self.context_var,
            font=('Arial', 10)
        )
        self.context_entry.pack(fill=tk.X)
        
        # Show context frame if handwriting is already selected
        if self.text_type == "handwriting":
            self.context_frame.pack(fill=tk.X, pady=(0, 10))
        
        # OCR mode indicator
        mode_frame = ttk.Frame(main_frame)
        mode_frame.pack(fill=tk.X, pady=(0, 10))
        
        mode_text = self._get_mode_description()
        self.mode_label = ttk.Label(
            mode_frame,
            text=mode_text,
            font=('Arial', 9),
            foreground='gray',
            wraplength=480,
            justify=tk.CENTER
        )
        self.mode_label.pack()
        
        # Main action buttons
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(fill=tk.X, pady=(10, 0))
        
        self.process_btn = ttk.Button(
            btn_frame,
            text="📝 Process OCR",
            command=self._process_ocr,
            width=15,
            state=tk.DISABLED
        )
        self.process_btn.pack(side=tk.LEFT, padx=(0, 10))
        
        # Add context help to Process OCR button
        try:
            from context_help import add_help, HELP_TEXTS
            if HELP_TEXTS and 'process_ocr_button' in HELP_TEXTS:
                add_help(self.process_btn, **HELP_TEXTS['process_ocr_button'])
        except ImportError:
            pass
        
        self.cancel_btn = ttk.Button(
            btn_frame,
            text="Cancel",
            command=self._on_close,
            width=10
        )
        self.cancel_btn.pack(side=tk.RIGHT)
    
    def _setup_drag_drop(self):
        """Setup drag and drop support for the dialog."""
        
        # Method 1: Try windnd library (best for Windows)
        try:
            import windnd
            
            def handle_drop(files):
                """Handle dropped files from windnd."""
                # CRITICAL: windnd calls this from a non-main thread!
                # We must NOT touch any tkinter widgets directly here.
                # Only process data and schedule UI updates via after().
                try:
                    # Decode files (windnd may return bytes)
                    file_list = []
                    for f in files:
                        if isinstance(f, bytes):
                            f = f.decode('utf-8', errors='replace')
                        file_list.append(f)
                    
                    # Filter to only supported files (no tkinter here!)
                    valid_files = []
                    for f in file_list:
                        if os.path.isfile(f):
                            ext = os.path.splitext(f)[1].lower()
                            if ext in self.ALL_EXTENSIONS:
                                valid_files.append(f)
                    
                    if valid_files:
                        # Schedule UI update on main thread
                        # Use a copy of the list to avoid closure issues
                        files_to_add = list(valid_files)
                        self.dialog.after(10, lambda: self._handle_dropped_files(files_to_add))
                    
                except Exception as e:
                    print(f"Drop error: {e}")
            
            # Ensure dialog is fully realized before hooking
            self.dialog.update_idletasks()
            self.dialog.update()
            
            # Register the dialog window for drops
            windnd.hook_dropfiles(self.dialog, func=handle_drop)
            return
        except ImportError:
            pass  # windnd not available
        except Exception as e:
            print(f"windnd setup failed: {e}")
        
        # Method 2: Try tkinterdnd2 (cross-platform)
        try:
            from tkinterdnd2 import DND_FILES
            
            # For tkinterdnd2 to work on Toplevel, the root must be TkinterDnD.Tk
            # Check if our parent supports it
            if hasattr(self.dialog, 'drop_target_register'):
                self.dialog.drop_target_register(DND_FILES)
                self.dialog.dnd_bind('<<Drop>>', self._handle_dnd_drop)
            if hasattr(self.image_listbox, 'drop_target_register'):
                self.image_listbox.drop_target_register(DND_FILES)
                self.image_listbox.dnd_bind('<<Drop>>', self._handle_dnd_drop)
            print("DEBUG: tkinterdnd2 setup complete")
            return
        except ImportError:
            pass
        except Exception as e:
            print(f"DEBUG: tkinterdnd2 setup failed: {e}")
        
        # Method 3: If no drag-drop library available, update the placeholder text
        # to indicate drag-drop isn't available
        print("DEBUG: No drag-drop library available")
        self._show_placeholder_no_dnd()
    
    def _handle_dnd_drop(self, event):
        """Handle drag-and-drop file drop event."""
        # event.data contains the dropped file paths
        files = self._parse_drop_data(event.data)
        self._add_files_to_list(files)
    
    def _parse_drop_data(self, data: str) -> List[str]:
        """Parse dropped file data (handles different formats)."""
        files = []
        
        # Handle different drop data formats
        if '{' in data:
            # Tcl list format: {path1} {path2}
            import re
            files = re.findall(r'\{([^}]+)\}', data)
        else:
            # Space-separated or single file
            files = data.strip().split()
        
        # Filter to only supported files (images and PDFs)
        valid_files = []
        for f in files:
            f = f.strip().strip('"').strip("'")
            if os.path.isfile(f):
                ext = os.path.splitext(f)[1].lower()
                if ext in self.ALL_EXTENSIONS:
                    valid_files.append(f)
        
        return valid_files
    
    def _handle_dropped_files(self, files: List[str]):
        """Handle dropped files on the main thread."""
        try:
            for filepath in files:
                if filepath in self.image_files:
                    continue  # Skip duplicates
                
                ext = os.path.splitext(filepath)[1].lower()
                
                if ext in self.PDF_EXTENSIONS:
                    # Convert PDF pages to temp images
                    temp_images = self._convert_pdf_to_images(filepath)
                    for temp_img in temp_images:
                        if temp_img not in self.image_files:
                            self.image_files.append(temp_img)
                else:
                    # Regular image file
                    self.image_files.append(filepath)
            
            self._refresh_listbox()
            self.status_label.config(text=f"Ready to process {len(self.image_files)} page(s)")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to add dropped files: {e}", parent=self.dialog)
    
    def _add_files_to_list(self, files: List[str]):
        """Add files to the image list. PDFs are converted to temp images."""
        if not files:
            return
        
        try:
            # Sort naturally
            files = sorted(files, key=lambda f: self._natural_sort_key(os.path.basename(f)))
            
            added = 0
            for filepath in files:
                ext = os.path.splitext(filepath)[1].lower()
                
                if ext in self.PDF_EXTENSIONS:
                    # Convert PDF pages to temp images
                    temp_images = self._convert_pdf_to_images(filepath)
                    for temp_img in temp_images:
                        if temp_img not in self.image_files:
                            self.image_files.append(temp_img)
                            added += 1
                else:
                    # Regular image file
                    if filepath not in self.image_files:
                        self.image_files.append(filepath)
                        added += 1
            
            if added > 0:
                self._refresh_listbox()
                self.status_label.config(text=f"Ready to process {len(self.image_files)} page(s)")
        except Exception as e:
            import traceback
            error_msg = traceback.format_exc()
            messagebox.showerror(
                "Error Adding Files",
                f"Failed to add files:\n{str(e)}\n\nDetails:\n{error_msg}",
                parent=self.dialog
            )
    
    def _convert_pdf_to_images(self, pdf_path: str) -> List[str]:
        """
        Convert a PDF file to temporary image files.
        Returns list of temp image paths.
        """
        temp_images = []
        pdf_name = os.path.splitext(os.path.basename(pdf_path))[0]
        
        try:
            from pdf2image import convert_from_path
            
            self.status_label.config(text=f"Converting {os.path.basename(pdf_path)}...")
            self.dialog.update()
            
            # Convert PDF to images
            images = convert_from_path(pdf_path, dpi=300)
            
            # Save each page as a temp image
            for page_num, image in enumerate(images, start=1):
                # Create temp file with meaningful name
                temp_dir = tempfile.gettempdir()
                temp_filename = f"docanalyzer_ocr_{pdf_name}_page{page_num}.png"
                temp_path = os.path.join(temp_dir, temp_filename)
                
                # Save as PNG for best quality
                image.save(temp_path, 'PNG')
                temp_images.append(temp_path)
                self.temp_files.append(temp_path)  # Track for cleanup
            
            self.status_label.config(text=f"Converted {len(images)} page(s) from {os.path.basename(pdf_path)}")
            
        except ImportError:
            messagebox.showerror(
                "PDF Support Missing",
                "PDF conversion requires pdf2image and Poppler.\n\n"
                "Please install them or use image files instead.",
                parent=self.dialog
            )
        except Exception as e:
            messagebox.showerror(
                "PDF Conversion Error",
                f"Could not convert PDF:\n{str(e)}",
                parent=self.dialog
            )
        
        return temp_images
    
    def _on_listbox_click(self, event):
        """Handle listbox click - used to clear placeholder."""
        pass  # Placeholder handling is automatic
    
    def _show_placeholder(self):
        """Show placeholder text when listbox is empty."""
        if not self.image_files:
            self.image_listbox.config(fg='gray')
            self.image_listbox.delete(0, tk.END)
            self.image_listbox.insert(tk.END, "  Drag and drop file(s) here")
            self.image_listbox.insert(tk.END, "  or Ctrl+C in Explorer, Ctrl+V here")
    
    def _show_placeholder_no_dnd(self):
        """Show placeholder text when drag-drop is not available."""
        if not self.image_files:
            self.image_listbox.config(fg='gray')
            self.image_listbox.delete(0, tk.END)
            self.image_listbox.insert(tk.END, "  Ctrl+C file(s) in File Explorer")
            self.image_listbox.insert(tk.END, "  then Ctrl+V here to add")
            self.image_listbox.insert(tk.END, "  (Install 'windnd' for drag & drop)")
    
    def _get_mode_description(self) -> str:
        """Get description text for current text type selection."""
        from ocr_handler import OCR_TEXT_TYPES
        
        # Get current text type key from selection
        selected_label = self.text_type_var.get() if hasattr(self, 'text_type_var') else None
        text_type_key = self.text_type
        
        # Find key from label if we have a selection
        if selected_label:
            for key, info in OCR_TEXT_TYPES.items():
                if info['label'] == selected_label:
                    text_type_key = key
                    break
        
        if text_type_key == "printed":
            return "📄 Local OCR (Tesseract) — Free & private"
        else:  # handwriting
            if self.api_key and self.api_key != "not-required":
                return f"✍️ Cloud OCR ({self.provider}) — Best for handwriting\nUsing: {self.model}"
            else:
                return f"⚠️ Handwriting mode requires an AI provider API key\nConfigure in Settings → API Keys"
    
    def _on_text_type_change(self, event=None):
        """Handle text type dropdown change."""
        # Update the mode description label
        self.mode_label.config(text=self._get_mode_description())
        
        # Update internal text_type based on selection
        from ocr_handler import OCR_TEXT_TYPES
        selected_label = self.text_type_var.get()
        for key, info in OCR_TEXT_TYPES.items():
            if info['label'] == selected_label:
                self.text_type = key
                break
        
        # Show/hide context frame based on text type
        if self.text_type == "handwriting":
            self.context_frame.pack(fill=tk.X, pady=(0, 10), before=self.mode_label.master)
        else:
            self.context_frame.pack_forget()
    
    def _update_ui_state(self):
        """Update button states and page counter based on current state."""
        num_pages = len(self.image_files)
        has_selection = len(self.image_listbox.curselection()) > 0
        
        # Update page counter
        if num_pages == 0:
            self.page_count_label.config(text="")
        elif num_pages == 1:
            self.page_count_label.config(text="1 page")
        else:
            self.page_count_label.config(text=f"{num_pages} pages")
        
        # Update button states
        if self.processing:
            self.add_btn.config(state=tk.DISABLED)
            self.remove_btn.config(state=tk.DISABLED)
            self.clear_btn.config(state=tk.DISABLED)
            self.move_up_btn.config(state=tk.DISABLED)
            self.move_down_btn.config(state=tk.DISABLED)
            self.process_btn.config(state=tk.DISABLED)
            self.cancel_btn.config(text="⏹ Stop")
        else:
            self.add_btn.config(state=tk.NORMAL)
            self.remove_btn.config(state=tk.NORMAL if has_selection else tk.DISABLED)
            self.clear_btn.config(state=tk.NORMAL if num_pages > 0 else tk.DISABLED)
            self.move_up_btn.config(state=tk.NORMAL if has_selection else tk.DISABLED)
            self.move_down_btn.config(state=tk.NORMAL if has_selection else tk.DISABLED)
            self.process_btn.config(state=tk.NORMAL if num_pages > 0 else tk.DISABLED)
            self.cancel_btn.config(text="Cancel")
    
    def _on_selection_change(self, event=None):
        """Handle listbox selection change."""
        self._update_ui_state()
    
    def _refresh_listbox(self):
        """Refresh the listbox display."""
        self.image_listbox.config(fg='black')
        self.image_listbox.delete(0, tk.END)
        
        if not self.image_files:
            self._show_placeholder()
        else:
            for i, filepath in enumerate(self.image_files, start=1):
                filename = os.path.basename(filepath)
                self.image_listbox.insert(tk.END, f"Page {i}: {filename}")
        
        self._update_ui_state()
    
    def _paste_files(self, event=None):
        """Handle Ctrl+V paste of files from clipboard."""
        try:
            # Try to get files from Windows clipboard using win32clipboard
            try:
                import win32clipboard
                import win32con
                
                win32clipboard.OpenClipboard()
                try:
                    # Check if clipboard has files (CF_HDROP format)
                    if win32clipboard.IsClipboardFormatAvailable(win32con.CF_HDROP):
                        files = win32clipboard.GetClipboardData(win32con.CF_HDROP)
                        if files:
                            valid_files = []
                            for f in files:
                                if os.path.isfile(f):
                                    ext = os.path.splitext(f)[1].lower()
                                    if ext in self.ALL_EXTENSIONS:
                                        valid_files.append(f)
                            
                            if valid_files:
                                self._handle_dropped_files(valid_files)
                            else:
                                messagebox.showinfo(
                                    "No Valid Files",
                                    "No supported image or PDF files found in clipboard.",
                                    parent=self.dialog
                                )
                            return
                finally:
                    win32clipboard.CloseClipboard()
            except ImportError:
                pass  # win32clipboard not available, try text fallback
            
            # Fallback: try to get text from clipboard (might be file paths)
            try:
                clipboard_text = self.dialog.clipboard_get()
                if clipboard_text:
                    # Check if it looks like file path(s)
                    lines = clipboard_text.strip().split('\n')
                    valid_files = []
                    for line in lines:
                        line = line.strip().strip('"')
                        if os.path.isfile(line):
                            ext = os.path.splitext(line)[1].lower()
                            if ext in self.ALL_EXTENSIONS:
                                valid_files.append(line)
                    
                    if valid_files:
                        self._handle_dropped_files(valid_files)
                        return
            except tk.TclError:
                pass  # Clipboard empty or not text
            
            # If we get here, no valid files were found
            messagebox.showinfo(
                "Paste Files",
                "To paste files:\n\n"
                "1. Open File Explorer\n"
                "2. Select file(s) and press Ctrl+C\n"
                "3. Come back here and press Ctrl+V",
                parent=self.dialog
            )
            
        except Exception as e:
            messagebox.showerror("Paste Error", f"Could not paste files: {e}", parent=self.dialog)
    
    def _open_file_explorer(self):
        """Open Windows File Explorer for the user to drag files from."""
        import subprocess
        import os
        
        # Try to open a sensible default location
        # Check for common image/document locations
        possible_paths = [
            os.path.expanduser("~/Pictures"),
            os.path.expanduser("~/Documents"),
            os.path.expanduser("~/Desktop"),
            os.path.expanduser("~")
        ]
        
        # Use the first path that exists
        open_path = None
        for path in possible_paths:
            if os.path.exists(path):
                open_path = path
                break
        
        try:
            if open_path:
                # Open Explorer to the specific folder
                subprocess.Popen(['explorer', open_path])
            else:
                # Just open Explorer
                subprocess.Popen(['explorer'])
        except Exception as e:
            messagebox.showerror("Error", f"Could not open File Explorer: {e}", parent=self.dialog)
    
    def _add_images(self):
        """Open file dialog to add images or PDFs."""
        filetypes = [
            ("All supported", " ".join(f"*{ext}" for ext in self.ALL_EXTENSIONS)),
            ("Image files", " ".join(f"*{ext}" for ext in self.IMAGE_EXTENSIONS)),
            ("PDF files", "*.pdf"),
            ("PNG files", "*.png"),
            ("JPEG files", "*.jpg *.jpeg"),
            ("TIFF files", "*.tif *.tiff"),
            ("All files", "*.*")
        ]
        
        files = filedialog.askopenfilenames(
            parent=self.dialog,
            title="Select file(s) then click 'Open'",
            filetypes=filetypes
        )
        
        if files:
            self._add_files_to_list(list(files))
    
    def _natural_sort_key(self, filename: str):
        """Generate key for natural sorting (page1, page2, page10 instead of page1, page10, page2)."""
        return [int(c) if c.isdigit() else c.lower() for c in re.split(r'(\d+)', filename)]
    
    def _remove_selected(self):
        """Remove selected images from the list."""
        selection = list(self.image_listbox.curselection())
        if not selection:
            return
        
        # Don't remove placeholder
        if not self.image_files:
            return
        
        # Remove in reverse order to maintain indices
        for index in reversed(selection):
            if index < len(self.image_files):
                del self.image_files[index]
        
        self._refresh_listbox()
        
        if self.image_files:
            self.status_label.config(text=f"Ready to process {len(self.image_files)} page(s)")
        else:
            self.status_label.config(text="")
    
    def _clear_all(self):
        """Clear all images."""
        if not self.image_files:
            return
        
        if messagebox.askyesno("Clear All", "Remove all pages from the list?", parent=self.dialog):
            self.image_files.clear()
            self._refresh_listbox()
            self.status_label.config(text="Drag images here or click Add Images")
    
    def _move_up(self):
        """Move selected item up in the list."""
        selection = self.image_listbox.curselection()
        if not selection or selection[0] == 0 or not self.image_files:
            return
        
        index = selection[0]
        self.image_files[index], self.image_files[index - 1] = \
            self.image_files[index - 1], self.image_files[index]
        
        self._refresh_listbox()
        self.image_listbox.selection_set(index - 1)
        self.image_listbox.see(index - 1)
    
    def _move_down(self):
        """Move selected item down in the list."""
        selection = self.image_listbox.curselection()
        if not selection or selection[0] >= len(self.image_files) - 1 or not self.image_files:
            return
        
        index = selection[0]
        self.image_files[index], self.image_files[index + 1] = \
            self.image_files[index + 1], self.image_files[index]
        
        self._refresh_listbox()
        self.image_listbox.selection_set(index + 1)
        self.image_listbox.see(index + 1)
    
    def _process_ocr(self):
        """Start OCR processing of all images."""
        if not self.image_files:
            messagebox.showwarning("No Images", "Please add images first.", parent=self.dialog)
            return
        
        self.processing = True
        self.cancelled = False
        self._update_ui_state()
        
        # Show progress bar
        self.progress_bar.pack(fill=tk.X, pady=(5, 0))
        self.progress_var.set(0)
        
        # Run OCR in thread
        import threading
        thread = threading.Thread(target=self._process_ocr_thread, daemon=True)
        thread.start()
    
    def _process_ocr_thread(self):
        """OCR processing thread - delegates to ocr_handler."""
        from ocr_handler import process_multiple_images_ocr
        
        # Map text_type to ocr_mode for the processor
        # printed -> local_first (Tesseract)
        # handwriting -> cloud_direct (Vision AI)
        if self.text_type == "handwriting":
            effective_ocr_mode = "cloud_direct"
        else:
            effective_ocr_mode = "local_first"
        
        # Get context hint for handwriting mode
        context_hint = self.context_var.get().strip() if self.text_type == "handwriting" else ""
        
        def progress_callback(page_num, total_pages, message):
            """Update UI with progress."""
            progress = ((page_num - 1) / total_pages) * 100
            self.dialog.after(0, lambda: self.progress_var.set(progress))
            self.dialog.after(0, lambda m=message, p=page_num, t=total_pages: 
                self.status_label.config(text=f"Page {p}/{t}: {m}"))
        
        def cancel_check():
            """Check if user cancelled."""
            return self.cancelled
        
        try:
            success, result = process_multiple_images_ocr(
                image_files=self.image_files,
                ocr_mode=effective_ocr_mode,
                language=self.ocr_language,
                quality=self.ocr_quality,
                confidence_threshold=self.confidence_threshold,
                provider=self.provider,
                model=self.model,
                api_key=self.api_key,
                progress_callback=progress_callback,
                cancel_check=cancel_check,
                text_type=self.text_type,
                context_hint=context_hint
            )
            
            # Update progress to 100%
            self.dialog.after(0, lambda: self.progress_var.set(100))
            
            if success:
                # Check if we used local OCR and confidence is below threshold
                if self.text_type == "printed" and result:
                    avg_confidence = self._calculate_average_confidence(result)
                    
                    if avg_confidence is not None and avg_confidence < self.confidence_threshold:
                        # Confidence is low - prompt user to try cloud AI
                        self.dialog.after(0, lambda: self._prompt_for_cloud_retry(
                            result, avg_confidence, context_hint
                        ))
                        return
                
                self.ocr_result = result
                self.dialog.after(0, lambda: self._handle_success(result))
            else:
                if "cancelled" in str(result).lower():
                    self.dialog.after(0, self._handle_cancelled)
                else:
                    self.dialog.after(0, lambda: self._handle_error(result))
                    
        except Exception as e:
            import traceback
            error_details = traceback.format_exc()
            print(f"OCR processing error:\n{error_details}")
            self.dialog.after(0, lambda: self._handle_error(str(e)))
    
    def _calculate_average_confidence(self, entries):
        """Calculate average confidence score from OCR results."""
        confidences = []
        for entry in entries:
            if isinstance(entry, dict) and 'confidence' in entry:
                conf = entry.get('confidence')
                if conf is not None:
                    confidences.append(conf)
        
        if confidences:
            return sum(confidences) / len(confidences)
        return None
    
    def _prompt_for_cloud_retry(self, local_results, avg_confidence, context_hint):
        """Prompt user to retry with cloud AI due to low confidence."""
        self.processing = False
        self._update_ui_state()
        
        # Check if cloud AI is available
        if not self.api_key or self.api_key == "not-required":
            # No API key - just show the local results with a warning
            messagebox.showwarning(
                "Low Confidence OCR",
                f"Local OCR confidence is low ({avg_confidence:.0f}%).\n\n"
                f"Your threshold is set to {self.confidence_threshold}%.\n\n"
                "Online AI could produce better results, but no API key is configured.\n"
                "Configure an API key in Settings to enable this option.",
                parent=self.dialog
            )
            self.ocr_result = local_results
            self._handle_success(local_results)
            return
        
        # Ask user if they want to retry with cloud AI
        response = messagebox.askyesnocancel(
            "Low Confidence - Try Online AI?",
            f"Local OCR confidence is low ({avg_confidence:.0f}%).\n"
            f"Your threshold is set to {self.confidence_threshold}%.\n\n"
            "(You can adjust your confidence threshold in Settings → OCR Settings → OCR Processing Mode)\n\n"
            f"Would you like to retry using online AI ({self.provider})?\n\n"
            "• Yes = Retry with online AI (uses API credits)\n"
            "• No = Keep local OCR results\n"
            "• Cancel = Go back to edit settings",
            parent=self.dialog
        )
        
        if response is True:
            # User wants to retry with cloud AI
            self._retry_with_cloud_ai(context_hint)
        elif response is False:
            # User wants to keep local results
            self.ocr_result = local_results
            self._handle_success(local_results)
        else:
            # User cancelled - reset UI
            self.progress_bar.pack_forget()
            self.status_label.config(text=f"Ready to process {len(self.image_files)} page(s)")
    
    def _retry_with_cloud_ai(self, context_hint):
        """Retry OCR processing with cloud AI."""
        self.processing = True
        self._update_ui_state()
        
        # Show progress bar
        self.progress_bar.pack(fill=tk.X, pady=(5, 0))
        self.progress_var.set(0)
        self.status_label.config(text="Retrying with online AI...")
        
        # Run in thread
        import threading
        thread = threading.Thread(
            target=self._cloud_ocr_thread,
            args=(context_hint,),
            daemon=True
        )
        thread.start()
    
    def _cloud_ocr_thread(self, context_hint):
        """Cloud OCR processing thread."""
        from ocr_handler import process_multiple_images_ocr
        
        def progress_callback(page_num, total_pages, message):
            progress = ((page_num - 1) / total_pages) * 100
            self.dialog.after(0, lambda: self.progress_var.set(progress))
            self.dialog.after(0, lambda m=message, p=page_num, t=total_pages: 
                self.status_label.config(text=f"Page {p}/{t}: {m} (Online AI)"))
        
        def cancel_check():
            return self.cancelled
        
        try:
            success, result = process_multiple_images_ocr(
                image_files=self.image_files,
                ocr_mode="cloud_direct",  # Force cloud AI
                language=self.ocr_language,
                quality=self.ocr_quality,
                confidence_threshold=self.confidence_threshold,
                provider=self.provider,
                model=self.model,
                api_key=self.api_key,
                progress_callback=progress_callback,
                cancel_check=cancel_check,
                text_type="handwriting",  # Use handwriting mode for cloud
                context_hint=context_hint
            )
            
            self.dialog.after(0, lambda: self.progress_var.set(100))
            
            if success:
                self.ocr_result = result
                self.dialog.after(0, lambda: self._handle_success(result))
            else:
                self.dialog.after(0, lambda: self._handle_error(result))
                
        except Exception as e:
            self.dialog.after(0, lambda: self._handle_error(str(e)))
    
    def _handle_success(self, entries):
        """Handle successful OCR completion - show preview with edit options."""
        self.processing = False
        self.progress_bar.pack_forget()
        self._update_ui_state()
        
        num_pages = len(self.image_files)
        num_segments = len(entries)
        
        # Calculate and display confidence if available
        avg_confidence = self._calculate_average_confidence(entries)
        confidence_info = ""
        if avg_confidence is not None:
            confidence_info = f" (confidence: {avg_confidence:.0f}%)"
        
        # Store entries for later
        self.ocr_entries = entries
        
        # Combine entries into text for preview
        self.ocr_text = '\n\n'.join(entry['text'] for entry in entries)
        
        self.status_label.config(
            text=f"✅ Complete! Extracted {num_segments} text segments from {num_pages} page(s){confidence_info}"
        )
        
        # Show the result preview dialog
        self._show_result_preview()
    
    def _show_result_preview(self):
        """
        Show OCR result preview with options to edit or accept.
        This replaces the main dialog content with a preview/edit interface.
        The preview window is independent so user can minimize the main app.
        """
        # Clear the current dialog content
        for widget in self.dialog.winfo_children():
            widget.destroy()
        
        # Make dialog independent (not tied to parent window)
        # This allows user to minimize main app while keeping preview open
        self.dialog.transient('')  # Remove transient binding
        self.dialog.grab_release()  # Release modal grab
        
        # Resize dialog for preview - allow to be made short for stacking
        self.dialog.geometry("700x500")
        self.dialog.minsize(400, 250)  # Allow narrow/short for comparison with PDF
        self.dialog.title("📋 OCR Result Preview")
        
        main_frame = ttk.Frame(self.dialog, padding=15)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Header
        header_frame = ttk.Frame(main_frame)
        header_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(
            header_frame,
            text="📋 OCR Result Preview",
            font=('Arial', 14, 'bold')
        ).pack(side=tk.LEFT)
        
        num_pages = len(self.image_files)
        ttk.Label(
            header_frame,
            text=f"{num_pages} page(s) processed",
            font=('Arial', 10),
            foreground='gray'
        ).pack(side=tk.RIGHT)
        
        # Instructions
        ttk.Label(
            main_frame,
            text="Review the extracted text below. You can edit it manually, use voice commands, or accept as-is.",
            font=('Arial', 9),
            foreground='gray',
            wraplength=650
        ).pack(fill=tk.X, pady=(0, 10))
        
        # Button frame - pack BEFORE editor so it stays at bottom
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(side=tk.BOTTOM, fill=tk.X, pady=(10, 0))
        
        # Left side - editing options
        edit_btns = ttk.Frame(btn_frame)
        edit_btns.pack(side=tk.LEFT)
        
        ttk.Button(
            edit_btns,
            text="🎤 Edit with Voice",
            command=self._edit_with_voice,
            width=16
        ).pack(side=tk.LEFT, padx=(0, 5))
        
        ttk.Label(
            edit_btns,
            text="Speak corrections",
            font=('Arial', 8),
            foreground='gray'
        ).pack(side=tk.LEFT, padx=(0, 15))
        
        # Middle - Copy and Save As buttons
        middle_btns = ttk.Frame(btn_frame)
        middle_btns.pack(side=tk.LEFT, padx=(20, 0))
        
        ttk.Button(
            middle_btns,
            text="📋 Copy",
            command=self._copy_to_clipboard,
            width=10
        ).pack(side=tk.LEFT, padx=(0, 5))
        
        # Save As dropdown menu
        self.save_as_btn = tk.Menubutton(
            middle_btns,
            text="💾 Save As ▼",
            relief=tk.RAISED,
            font=('Arial', 9),
            width=12
        )
        self.save_as_btn.pack(side=tk.LEFT, padx=(0, 5))
        
        save_menu = tk.Menu(self.save_as_btn, tearoff=0)
        self.save_as_btn.config(menu=save_menu)
        
        save_menu.add_command(
            label="📄 Save as .txt",
            command=lambda: self._save_as_format('txt')
        )
        save_menu.add_command(
            label="📝 Save as .docx",
            command=lambda: self._save_as_format('docx')
        )
        save_menu.add_command(
            label="📋 Save as .rtf",
            command=lambda: self._save_as_format('rtf')
        )
        save_menu.add_command(
            label="📕 Save as .pdf",
            command=lambda: self._save_as_format('pdf')
        )
        
        # Right side - main actions
        action_btns = ttk.Frame(btn_frame)
        action_btns.pack(side=tk.RIGHT)
        
        ttk.Button(
            action_btns,
            text="Cancel",
            command=self._on_close,
            width=10
        ).pack(side=tk.LEFT, padx=(0, 5))
        
        ttk.Button(
            action_btns,
            text="✓ Accept & Save to Library",
            command=self._accept_result,
            width=22
        ).pack(side=tk.LEFT)
        
        # Text preview/editor - pack AFTER buttons so it fills remaining space
        editor_frame = ttk.LabelFrame(main_frame, text="Extracted Text", padding=5)
        editor_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        
        text_container = ttk.Frame(editor_frame)
        text_container.pack(fill=tk.BOTH, expand=True)
        
        self.preview_text = tk.Text(
            text_container,
            wrap=tk.WORD,
            font=('Georgia', 11),
            undo=True,
            padx=10,
            pady=10,
            bg='#FFFEF5'
        )
        self.preview_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        scrollbar = ttk.Scrollbar(text_container, orient=tk.VERTICAL, command=self.preview_text.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.preview_text.config(yscrollcommand=scrollbar.set)
        
        # Insert the OCR text
        self.preview_text.insert('1.0', self.ocr_text)
    
    def _copy_to_clipboard(self):
        """Copy the current text to clipboard."""
        text = self.preview_text.get('1.0', 'end-1c')
        if text:
            self.dialog.clipboard_clear()
            self.dialog.clipboard_append(text)
            from tkinter import messagebox
            messagebox.showinfo("Copied", "Text copied to clipboard!", parent=self.dialog)
        else:
            from tkinter import messagebox
            messagebox.showwarning("Empty", "No text to copy.", parent=self.dialog)
    
    def _save_as_format(self, export_format: str):
        """
        Save the current text to a file in the specified format.
        
        Args:
            export_format: One of 'txt', 'docx', 'rtf', 'pdf'
        """
        from tkinter import filedialog, messagebox
        import os
        import re
        import datetime
        
        text = self.preview_text.get('1.0', 'end-1c').strip()
        if not text:
            messagebox.showwarning("Empty", "No text to save.", parent=self.dialog)
            return
        
        # Generate default filename from first image file
        default_name = "OCR_Transcription"
        if self.image_files:
            base_name = os.path.splitext(os.path.basename(self.image_files[0]))[0]
            # Clean the name
            clean_name = re.sub(r'[<>:"/\\|?*]', '-', base_name)
            clean_name = ''.join(c for c in clean_name if ord(c) >= 32)
            if clean_name:
                default_name = clean_name[:50]
        
        # File type configurations
        filetypes = {
            'txt': [("Text files", "*.txt"), ("All files", "*.*")],
            'docx': [("Word documents", "*.docx"), ("All files", "*.*")],
            'rtf': [("RTF documents", "*.rtf"), ("All files", "*.*")],
            'pdf': [("PDF files", "*.pdf"), ("All files", "*.*")]
        }
        
        filepath = filedialog.asksaveasfilename(
            title="Save OCR Text As",
            defaultextension=f".{export_format}",
            initialfile=f"{default_name}.{export_format}",
            filetypes=filetypes.get(export_format, [("All files", "*.*")]),
            parent=self.dialog
        )
        
        if not filepath:
            return  # User cancelled
        
        try:
            export_date = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            source_info = f"OCR from {len(self.image_files)} image(s)"
            
            if export_format == 'txt':
                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write("=" * 60 + "\n")
                    f.write(f"OCR Transcription\n")
                    f.write("=" * 60 + "\n\n")
                    f.write(f"Source: {source_info}\n")
                    f.write(f"Exported: {export_date}\n\n")
                    f.write("=" * 60 + "\n\n")
                    f.write(text)
                messagebox.showinfo("Saved", f"Text saved to:\n{filepath}", parent=self.dialog)
            
            elif export_format == 'docx':
                try:
                    from docx import Document
                    from docx.shared import Pt
                    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
                    
                    doc = Document()
                    
                    # Title
                    title = doc.add_heading('OCR Transcription', 0)
                    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    
                    # Metadata
                    doc.add_heading('Document Information', level=2)
                    meta = doc.add_paragraph()
                    meta.add_run("Source: ").bold = True
                    meta.add_run(f"{source_info}\n")
                    meta.add_run("Exported: ").bold = True
                    meta.add_run(f"{export_date}\n")
                    
                    doc.add_paragraph()  # Spacing
                    
                    # Content
                    doc.add_heading('Content', level=2)
                    for para in text.split('\n\n'):
                        if para.strip():
                            doc.add_paragraph(para.strip())
                    
                    doc.save(filepath)
                    messagebox.showinfo("Saved", f"Document saved to:\n{filepath}", parent=self.dialog)
                except ImportError:
                    messagebox.showerror("Error", "python-docx not installed.\n\nInstall with: pip install python-docx", parent=self.dialog)
            
            elif export_format == 'rtf':
                # Create RTF directly
                rtf_content = []
                rtf_content.append(r'{\rtf1\ansi\deff0')
                rtf_content.append(r'{\fonttbl{\f0 Times New Roman;}}')
                rtf_content.append(r'\f0\fs24')
                
                # Title
                rtf_content.append(r'{\b\fs32 OCR Transcription}\par')
                rtf_content.append(r'\par')
                
                # Metadata
                rtf_content.append(r'{\b Source: }' + source_info + r'\par')
                rtf_content.append(r'{\b Exported: }' + export_date + r'\par')
                rtf_content.append(r'\par')
                
                # Content
                for line in text.split('\n'):
                    if line.strip():
                        safe_line = line.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')
                        rtf_content.append(safe_line + r'\par')
                
                rtf_content.append(r'}')
                
                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write('\n'.join(rtf_content))
                messagebox.showinfo("Saved", f"RTF saved to:\n{filepath}", parent=self.dialog)
            
            elif export_format == 'pdf':
                try:
                    from reportlab.lib.pagesizes import letter
                    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                    from reportlab.lib.units import inch
                    from reportlab.lib.enums import TA_CENTER
                    
                    pdf_doc = SimpleDocTemplate(filepath, pagesize=letter)
                    styles = getSampleStyleSheet()
                    story = []
                    
                    # Title
                    title_style = ParagraphStyle(
                        'Title',
                        parent=styles['Heading1'],
                        fontSize=20,
                        alignment=TA_CENTER
                    )
                    story.append(Paragraph('OCR Transcription', title_style))
                    story.append(Spacer(1, 0.3 * inch))
                    
                    # Metadata
                    meta_style = styles['Normal']
                    story.append(Paragraph(f"<b>Source:</b> {source_info}", meta_style))
                    story.append(Paragraph(f"<b>Exported:</b> {export_date}", meta_style))
                    story.append(Spacer(1, 0.3 * inch))
                    
                    # Content
                    for para in text.split('\n\n'):
                        if para.strip():
                            safe_para = para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                            story.append(Paragraph(safe_para, styles['Normal']))
                            story.append(Spacer(1, 0.1 * inch))
                    
                    pdf_doc.build(story)
                    messagebox.showinfo("Saved", f"PDF saved to:\n{filepath}", parent=self.dialog)
                except ImportError:
                    messagebox.showerror("Error", "reportlab not installed.\n\nInstall with: pip install reportlab", parent=self.dialog)
        
        except Exception as e:
            messagebox.showerror("Save Error", f"Failed to save:\n{str(e)}", parent=self.dialog)
    
    def _edit_with_voice(self):
        """Open voice editing dialog for the OCR result."""
        try:
            from voice_edit_dialog import open_voice_edit_dialog
            
            # Get current text from preview (may have been manually edited)
            current_text = self.preview_text.get('1.0', 'end-1c')
            
            # Open voice edit dialog
            edited_text = open_voice_edit_dialog(
                self.dialog,
                self.app,
                current_text,
                title="Voice Edit OCR Result"
            )
            
            if edited_text is not None:
                # Update preview with edited text
                self.preview_text.delete('1.0', tk.END)
                self.preview_text.insert('1.0', edited_text)
                self.ocr_text = edited_text
                
        except ImportError as e:
            from tkinter import messagebox
            messagebox.showerror(
                "Module Not Found",
                f"Could not load voice editing module:\n{str(e)}",
                parent=self.dialog
            )
    
    def _accept_result(self):
        """Accept the current text and pass to app."""
        # Get final text from preview (may have been edited)
        final_text = self.preview_text.get('1.0', 'end-1c')
        
        # Convert back to entries format
        # Each paragraph becomes an entry
        paragraphs = [p.strip() for p in final_text.split('\n\n') if p.strip()]
        
        entries = []
        for i, para in enumerate(paragraphs, start=1):
            entries.append({
                'start': 1,  # Page number not tracked after editing
                'text': para,
                'location': f'Segment {i}'
            })
        
        # Clean up temp files
        self._cleanup_temp_files()
        
        # Close dialog and pass result to app
        self.dialog.destroy()
        self.app._handle_multi_image_ocr_result(entries, self.image_files)

    def _handle_error(self, error: str):
        """Handle OCR error."""
        self.processing = False
        self.progress_bar.pack_forget()
        self._update_ui_state()
        
        error_short = error[:100] + "..." if len(error) > 100 else error
        self.status_label.config(text=f"❌ Error: {error_short}")
        messagebox.showerror("OCR Error", error, parent=self.dialog)
    
    def _handle_cancelled(self):
        """Handle user cancellation."""
        self.processing = False
        self.progress_bar.pack_forget()
        self._update_ui_state()
        
        self.status_label.config(text="⚠️ Processing cancelled")
    
    def _on_close(self):
        """Handle dialog close or cancel button."""
        if self.processing:
            # Set cancelled flag - thread will check this
            self.cancelled = True
            self.status_label.config(text="Stopping...")
            return
        
        # Clean up temp files
        self._cleanup_temp_files()
        self.dialog.destroy()
    
    def _cleanup_temp_files(self):
        """Remove temporary image files created from PDFs."""
        for temp_path in self.temp_files:
            try:
                if os.path.exists(temp_path):
                    os.unlink(temp_path)
            except Exception:
                pass  # Best effort cleanup
        self.temp_files.clear()
