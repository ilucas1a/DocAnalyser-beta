"""
thread_viewer.py - Conversation Thread Viewer Window
Displays conversation threads with follow-up capability

Extracted from Main.py to reduce file size and enable follow-up from thread window.

v2.0 - Added multi-source document support with:
  - Independent collapsible sections for each source document
  - Lazy loading (content only rendered when expanded)
  - Configurable character warning threshold
  - Default collapsed when > 2 sources
"""

import tkinter as tk
from tkinter import ttk, scrolledtext, messagebox, filedialog
import datetime
import os
import threading
import re
import webbrowser
from typing import Optional, Callable, Dict, List, Any

# Import from our modules
from document_library import get_document_by_id
from utils import safe_filename, format_display_date
from document_export import export_conversation_thread, get_file_extension_and_types, export_document

# Import help system
try:
    from context_help import add_help, HELP_TEXTS
except ImportError:
    # Fallback if context_help not available
    def add_help(*args, **kwargs): pass
    HELP_TEXTS = {}

# === CONFIGURATION CONSTANTS ===
DEFAULT_CHAR_WARNING_THRESHOLD = 150000  # Characters - warn if expanding would exceed this
DEFAULT_COLLAPSE_THRESHOLD = 2  # Sources - default collapsed if more than this many

def get_clean_filename(text: str, max_length: int = 50) -> str:
    """Wrapper for safe_filename for backward compatibility"""
    return safe_filename(text, max_length)

class ThreadViewerWindow:
    """
    Unified Viewer Window - displays both source documents and conversation threads.
    
    This window serves as the single viewer for all document content, replacing
    the previous separate Full Text and Thread Viewer windows.
    
    Two Modes:
    - Source Mode: Displays the source document(s) in prose/transcript format
    - Conversation Mode: Displays the conversation thread in chat format
                         with optional collapsible source section(s)
    
    Features:
    - Toggle between Source and Conversation modes
    - Collapsible source document section(s) in both modes
    - Multiple source documents displayed as separate collapsible sections
    - Lazy loading for performance with large documents
    - Follow-up question capability
    - Copy, Save, and Export functionality
    - New Conversation (Same Source) for starting fresh
    - All content is editable
    """
    
    def __init__(
        self,
        parent: tk.Tk,
        current_thread: List[Dict],
        thread_message_count: int,
        current_document_id: Optional[str],
        current_document_text: str,
        current_document_source: str,
        model_var: tk.StringVar,
        provider_var: tk.StringVar,
        api_key_var: tk.StringVar,
        config: Dict,
        # Callbacks
        on_followup_complete: Optional[Callable[[str, str], None]] = None,
        on_clear_thread: Optional[Callable[[], None]] = None,
        refresh_library: Optional[Callable[[], None]] = None,
        get_ai_handler: Optional[Callable] = None,
        build_threaded_messages: Optional[Callable[[str], List[Dict]]] = None,
        add_message_to_thread: Optional[Callable[[str, str], None]] = None,
        font_size: int = 10,
        # New parameters for "New Conversation (Same Source)" feature
        document_class: str = "source",
        source_document_id: Optional[str] = None,
        on_start_new_conversation: Optional[Callable[[str], bool]] = None,
        # Unified Viewer callback for mode changes
        on_mode_change: Optional[Callable[[str], None]] = None,
        # Chunking support for initial prompts from Source Mode
        process_with_chunking: Optional[Callable] = None,
        current_entries: Optional[List] = None,
        current_document_type: str = "text",
        # Initial mode selection
        initial_mode: Optional[str] = None,
        # NEW: Multi-source document support
        source_documents: Optional[List[Dict]] = None,
        # NEW: Reference to main app for context synchronization
        app: Optional[object] = None,
    ):
        """
        Initialize the Thread Viewer Window.
        
        Args:
            parent: Parent Tk window
            current_thread: List of message dictionaries with 'role' and 'content'
            thread_message_count: Number of exchanges in the thread
            current_document_id: ID of the current document
            current_document_text: Text content of the current document (for backward compatibility)
            current_document_source: Source identifier of the current document
            model_var: StringVar for selected AI model
            provider_var: StringVar for selected AI provider
            api_key_var: StringVar for API key
            config: Configuration dictionary
            on_followup_complete: Callback when follow-up is processed (question, response)
            on_clear_thread: Callback to clear the thread
            refresh_library: Callback to refresh library display
            get_ai_handler: Function to get AI handler instance
            build_threaded_messages: Function to build threaded message list
            add_message_to_thread: Function to add message to thread
            font_size: Base font size for text display (default 10)
            process_with_chunking: Callback to process prompts with chunking (for initial prompts)
            current_entries: Document entries for chunking
            current_document_type: Type of document (text/audio_transcription)
            source_documents: List of source document dicts with 'title', 'text', 'source' keys
                             If None, falls back to single document from current_document_text
        """
        self.parent = parent
        self.current_thread = current_thread
        self.thread_message_count = thread_message_count
        self.current_document_id = current_document_id
        self.current_document_source = current_document_source
        self.model_var = model_var
        self.provider_var = provider_var
        self.api_key_var = api_key_var
        self.config = config
        # Note: font_size is read from config dynamically via _get_font_size()
        
        # === MULTI-SOURCE DOCUMENT SUPPORT ===
        # If source_documents provided, use it; otherwise create from single document
        if source_documents is not None and len(source_documents) > 0:
            self.source_documents = source_documents
        elif current_document_text:
            # Backward compatibility: wrap single document in list
            self.source_documents = [{
                'title': current_document_source or 'Document',
                'text': current_document_text,
                'source': current_document_source or '',
                'char_count': len(current_document_text) if current_document_text else 0
            }]
        else:
            self.source_documents = []
        
        # For backward compatibility, also keep current_document_text as combined text
        if self.source_documents:
            self.current_document_text = "\n\n".join(
                f"=== {doc.get('title', f'Document {i+1}')} ===\n{doc.get('text', '')}"
                for i, doc in enumerate(self.source_documents)
            ) if len(self.source_documents) > 1 else self.source_documents[0].get('text', '')
        else:
            self.current_document_text = current_document_text
        
        # Source document collapse state tracking
        # Dict mapping source index -> bool (True = expanded, False = collapsed)
        self.source_expanded_state = {}
        
        # Track which sources have been rendered (for lazy loading)
        # NOTE: Must be initialized BEFORE _init_source_collapse_state() which uses it
        self.source_content_rendered = {}
        
        self._init_source_collapse_state()
        
        # Callbacks
        self.on_followup_complete = on_followup_complete
        self.on_clear_thread = on_clear_thread
        self.refresh_library = refresh_library
        self.get_ai_handler = get_ai_handler
        self.build_threaded_messages = build_threaded_messages
        self.add_message_to_thread = add_message_to_thread
        
        # New Conversation (Same Source) feature
        self.document_class = document_class
        self.source_document_id = source_document_id
        self.on_start_new_conversation = on_start_new_conversation
        
        # Chunking support for initial prompts
        self.process_with_chunking = process_with_chunking
        self.current_entries = current_entries
        self.current_document_type = current_document_type
        
        # Processing state
        self.is_processing = False
        
        # Unified Viewer mode tracking
        # Mode is 'source' (viewing source document) or 'conversation' (viewing thread)
        # Use initial_mode if provided, otherwise start in conversation mode if there's a thread
        if initial_mode:
            self.current_mode = initial_mode
        else:
            self.current_mode = 'conversation' if (self.current_thread and len(self.current_thread) > 0) else 'source'
        
        # Legacy: source_section_visible controls ALL sources visibility in conversation mode
        self.source_section_visible = True
        
        # Exchange collapse state tracking
        # Dict mapping exchange index -> bool (True = expanded, False = collapsed)
        # By default, last 2 exchanges are expanded, older ones collapsed
        self.exchange_expanded_state = {}
        self._init_exchange_collapse_state()
        
        # Callback for mode changes (to update main UI button label)
        self.on_mode_change = on_mode_change
        
        # Reference to main app for context synchronization
        self.app = app
        
        # Create window
        self._create_window()
    
    def _init_source_collapse_state(self):
        """
        Initialize collapse state for source documents.
        
        Logic:
        - If > DEFAULT_COLLAPSE_THRESHOLD (2) sources: all start collapsed
        - Otherwise: all start expanded
        """
        num_sources = len(self.source_documents)
        collapse_threshold = self.config.get('viewer_collapse_threshold', DEFAULT_COLLAPSE_THRESHOLD)
        
        # Default state: collapsed if many sources, expanded otherwise
        default_expanded = num_sources <= collapse_threshold
        
        for i in range(num_sources):
            self.source_expanded_state[i] = default_expanded
            self.source_content_rendered[i] = False  # Lazy loading - not yet rendered
    
    def _init_exchange_collapse_state(self):
        """Initialize collapse state for exchanges - last 2 expanded, older collapsed"""
        num_exchanges = self._count_exchanges()
        for i in range(num_exchanges):
            # Expand last 2 exchanges, collapse older ones
            self.exchange_expanded_state[i] = (i >= num_exchanges - 2)
    
    def _count_exchanges(self) -> int:
        """Count the number of user-assistant exchange pairs in the thread"""
        if not self.current_thread:
            return 0
        # Count user messages (each user message starts an exchange)
        return sum(1 for msg in self.current_thread if msg.get('role') == 'user')
    
    def _get_font_size(self) -> int:
        """Get current font size from config (allows dynamic updates)."""
        return self.config.get('font_size', 10)
    
    def _get_char_warning_threshold(self) -> int:
        """Get the character warning threshold from config."""
        return self.config.get('viewer_char_warning_threshold', DEFAULT_CHAR_WARNING_THRESHOLD)
    
    def _calculate_total_expanded_chars(self) -> int:
        """Calculate total characters that would be displayed if all sources were expanded."""
        total = 0
        for doc in self.source_documents:
            total += doc.get('char_count', len(doc.get('text', '')))
        return total
    
    def _calculate_currently_expanded_chars(self) -> int:
        """Calculate total characters currently expanded (visible)."""
        total = 0
        for i, doc in enumerate(self.source_documents):
            if self.source_expanded_state.get(i, False):
                total += doc.get('char_count', len(doc.get('text', '')))
        return total
    
    def _check_expansion_warning(self, expanding_index: int = None) -> bool:
        """
        Check if expanding would exceed the character warning threshold.
        
        Args:
            expanding_index: Index of source being expanded (None if expanding all)
            
        Returns:
            True if expansion should proceed, False if user cancelled
        """
        threshold = self._get_char_warning_threshold()
        
        # Calculate what the new total would be
        if expanding_index is not None:
            # Single source being expanded
            current_chars = self._calculate_currently_expanded_chars()
            new_doc_chars = self.source_documents[expanding_index].get(
                'char_count', 
                len(self.source_documents[expanding_index].get('text', ''))
            )
            projected_total = current_chars + new_doc_chars
        else:
            # Expanding all
            projected_total = self._calculate_total_expanded_chars()
        
        if projected_total > threshold:
            # Show warning
            result = messagebox.askyesno(
                "Large Content Warning",
                f"Expanding this content would display approximately {projected_total:,} characters.\n\n"
                f"This may slow down the viewer on some systems.\n\n"
                f"(Current threshold: {threshold:,} characters - configurable in Settings)\n\n"
                "Continue anyway?",
                parent=self.window
            )
            return result
        
        return True  # Under threshold, proceed
    
    def _update_window_title(self):
        """Update window title based on current mode"""
        num_sources = len(self.source_documents)
        if num_sources > 1:
            source_desc = f"{num_sources} sources"
        else:
            source_desc = self.current_document_source or 'Unknown'
        
        if self.current_mode == 'source':
            title = f"📄 Source Document - {source_desc}"
        else:
            title = f"💬 Conversation - {source_desc}"
        self.window.title(title)
    
    def _create_window(self):
        """Create and configure the unified viewer window"""
        self.window = tk.Toplevel(self.parent)
        self._update_window_title()  # Set title based on mode
        self.window.geometry("700x780+0+0")  # Wider and taller, positioned at top-left of screen
        self.window.minsize(500, 600)  # Reasonable minimum - ensure buttons are visible
        self.window.configure(bg='#dcdad5')  # Match main window background
        
        # Note: transient() removed - it prevents minimize/maximize buttons on Windows
        # The window will now have standard OS window controls (minimize, maximize, close)
        
        # Get document info
        self._load_document_info()
        
        # Build UI
        self._create_header()
        self._create_find_replace_bar()  # New find-replace bar
        self._create_document_info()
        self._create_thread_display()
        self._create_followup_section()
        self._create_button_bar()
        
        # Bind keyboard shortcuts
        # Note: Ctrl+C, Ctrl+V, and Ctrl+Z work natively (no binding needed)
        self.window.bind('<Escape>', lambda e: self._close_window())
        self.window.bind('<Control-Shift-s>', lambda e: self._save_thread(None))  # Save to disk
        self.window.bind('<Control-Return>', lambda e: self._submit_followup())
        self.window.bind('<Control-h>', lambda e: self._focus_find_field())  # Focus find field
        self.window.bind('<Control-r>', lambda e: self._focus_find_field())  # Alternative shortcut
        
        # Handle window close - save and clear thread
        self.window.protocol('WM_DELETE_WINDOW', self._close_window)
    
    def _close_window(self):
        """Close window - auto-save edits and keep content intact in main app"""
        # Auto-save any edits based on current mode (happens silently)
        try:
            if self.current_mode == 'conversation':
                self._save_edits_to_thread()
            else:
                self._save_source_edits()
        except Exception as e:
            print(f"⚠️ Error saving on close: {e}")
            # Continue closing anyway - don't trap the user
        
        # Notify main app that viewer is closing (so button state updates)
        # Pass 'closed' as a special mode to signal viewer closure
        if self.on_mode_change:
            self.on_mode_change('closed')
        
        # Just close the window - don't clear the thread
        # The thread remains in the main app so View Thread button stays enabled
        # Thread will be saved when: a new document loads, app closes, or "New Conversation" is clicked
        self.window.destroy()
    
    def _load_document_info(self):
        """Load document metadata"""
        # For multi-source, use first document or combined info
        if len(self.source_documents) == 1:
            self.doc_title = self.source_documents[0].get('title', 'Unknown Document')
            self.source_info = self.source_documents[0].get('source', 'N/A')
        elif len(self.source_documents) > 1:
            self.doc_title = f"{len(self.source_documents)} Source Documents"
            self.source_info = "Multiple sources"
        else:
            self.doc_title = getattr(self, 'current_document_source', 'Unknown Document')
            self.source_info = "N/A"
        
        self.fetched_date = "N/A"
        self.published_date = "N/A"  # Original publication date
        
        if self.current_document_id:
            doc = get_document_by_id(self.current_document_id)
            if doc:
                if len(self.source_documents) <= 1:
                    self.doc_title = doc.get('title', self.doc_title)
                    self.source_info = doc.get('source', 'N/A')
                
                # Format the fetched/imported date
                raw_fetched = doc.get('fetched', '')
                if raw_fetched:
                    self.fetched_date = format_display_date(raw_fetched)
                
                # Get published_date from metadata and format it
                metadata = doc.get('metadata', {})
                if isinstance(metadata, dict):
                    raw_pub_date = metadata.get('published_date', '')
                    if raw_pub_date:
                        self.published_date = format_display_date(raw_pub_date)
    
    def _on_focus_in(self, event=None):
        """
        Handle window focus event.
        Auto-refresh if the thread has been updated externally (e.g., from main window).
        """
        # Only refresh if we're not currently processing
        if self.is_processing:
            return
        
        # Check if the thread has grown since we last displayed it
        current_length = len(self.current_thread)
        if current_length != self._last_thread_length:
            # Thread was updated externally - refresh the display
            self._last_thread_length = current_length
            # Update message count (thread has 2 messages per exchange)
            self.thread_message_count = current_length // 2
            self._refresh_thread_display()
    
    def _create_header(self):
        """Create header frame with title, message count, branch selector, and quick actions"""
        # === ROW 1: Title, info, and action buttons ===
        header_frame = ttk.Frame(self.window, padding=(10, 10, 10, 2))
        header_frame.pack(fill=tk.X)
        
        # Dynamic heading based on mode
        heading_text = "📄 Source Document" if self.current_mode == 'source' else "💬 Conversation Thread"
        if len(self.source_documents) > 1:
            heading_text += f" ({len(self.source_documents)} sources)"
        
        self.heading_label = ttk.Label(
            header_frame, 
            text=heading_text,
            font=('Arial', 14, 'bold')
        )
        self.heading_label.pack(side=tk.LEFT)
        
        # Show message count (for conversation mode) or document info (for source mode)
        if self.current_mode == 'conversation':
            info_text = f"{self.thread_message_count} exchange{'s' if self.thread_message_count != 1 else ''}"
        else:
            # For source mode, show document length indicator
            total_chars = self._calculate_total_expanded_chars()
            if total_chars > 10000:
                info_text = f"{total_chars:,} characters (long document)"
            elif total_chars > 0:
                info_text = f"{total_chars:,} characters"
            else:
                info_text = "No content"
        
        self.message_count_label = ttk.Label(
            header_frame, 
            text=info_text,
            font=('Arial', 10), 
            foreground='blue'
        )
        self.message_count_label.pack(side=tk.LEFT, padx=10)
        
        # === Quick action buttons (right side of row 1) ===
        controls_frame = ttk.Frame(header_frame)
        controls_frame.pack(side=tk.RIGHT)
        
        # Hamburger menu button (contains Copy Link and other options)
        menu_btn = ttk.Button(
            controls_frame,
            text="☰",
            command=self._show_menu,
            width=3
        )
        menu_btn.pack(side=tk.LEFT, padx=2)
        
        # Undo Edits button (conveys that text is editable, provides undo functionality)
        undo_btn = ttk.Button(
            header_frame,
            text="Undo Edits",
            command=self._undo_edit,
            width=12
        )
        undo_btn.pack(side=tk.RIGHT, padx=(0, 10))
        if HELP_TEXTS:
            add_help(undo_btn, **HELP_TEXTS.get("thread_undo_button",
                {"title": "Undo Edits", 
                 "description": "Undo your last edit. Press multiple times to undo further. Ctrl+Z also works. Use Ctrl+Y to redo. Edits are auto-saved when you close this window."}))
        
        # === ROW 2: Branch Selector (on its own line for better layout) ===
        # Use a wrapper frame that's always packed to maintain position
        self.branch_row_wrapper = ttk.Frame(self.window)
        self.branch_row_wrapper.pack(fill=tk.X)
        
        self.branch_selector_frame = ttk.Frame(self.branch_row_wrapper, padding=(10, 2, 10, 5))
        # Don't pack yet - _populate_branch_selector will handle visibility
        
        ttk.Label(
            self.branch_selector_frame,
            text="Branch:",
            font=('Arial', 9)
        ).pack(side=tk.LEFT, padx=(0, 5))
        
        self.branch_var = tk.StringVar()
        self.branch_combo = ttk.Combobox(
            self.branch_selector_frame,
            textvariable=self.branch_var,
            state='readonly',
            width=40  # Wider now that it has its own row
        )
        self.branch_combo.pack(side=tk.LEFT)
        self.branch_combo.bind('<<ComboboxSelected>>', self._on_branch_selected)
        
        # Initialize branch list storage
        self._branch_list = []
        
        # Initialize branch list and visibility
        self._populate_branch_selector()
    
    def _undo_edit(self):
        """Undo the last edit in the text widget"""
        try:
            self.thread_text.edit_undo()
        except tk.TclError:
            # Nothing to undo - silently ignore
            pass
    
    # === BRANCH SELECTOR METHODS ===
    
    def _get_source_document_id(self) -> str:
        """
        Get the source document ID for the current context.
        If viewing a response document, returns its parent source ID.
        If viewing a source document, returns the current document ID.
        """
        
        if not self.current_document_id:
            return None
        
        from document_library import get_document_by_id, is_source_document
        
        # Check if current document is a source
        is_src = is_source_document(self.current_document_id)
        
        if is_src:
            return self.current_document_id
        
        # Otherwise, get parent from metadata
        doc = get_document_by_id(self.current_document_id)
        if doc:
            doc_title = doc.get('title', '?')[:30]
            metadata = doc.get('metadata', {})
            parent_id = metadata.get('parent_document_id') or metadata.get('original_document_id')
            if parent_id:
                return parent_id
        else:
            pass
        
        return None
    
    def _populate_branch_selector(self):
        """
        Populate the branch selector dropdown with available conversation branches.
        Shows/hides the selector based on whether branches exist.
        """
        from document_library import get_response_branches_for_source, get_document_by_id
        
        
        # Get source document ID
        source_id = self._get_source_document_id()
        
        if not source_id:
            # No source document - hide branch selector
            self.branch_selector_frame.pack_forget()
            self._branch_list = []
            return
        
        # Get all response branches for this source
        branches = get_response_branches_for_source(source_id)
        
        # Store branch info for later use
        self._branch_list = branches if branches else []
        
        if not branches:
            # No branches yet - hide selector (will show after first response)
            self.branch_selector_frame.pack_forget()
            return
        
        # Build display names for dropdown
        branch_names = []
        current_branch_name = None
        
        for branch in branches:
            # Extract display name from title (remove [Response] prefix if present)
            title = branch.get('title', 'Unknown')
            if title.startswith('[Response]'):
                display_name = title[10:].strip()  # Remove '[Response]' prefix
                # Further trim if it starts with prompt name
                if ':' in display_name:
                    display_name = display_name.split(':', 1)[1].strip()
            else:
                display_name = title
            
            # Truncate if too long
            if len(display_name) > 30:
                display_name = display_name[:27] + "..."
            
            # Add processing indicator if branch is still being processed
            if branch.get('is_processing', False):
                display_name += " ⏳"
            
            branch_names.append(display_name)
            
            # Check if this is the current branch
            if branch.get('doc_id') == self.current_document_id:
                current_branch_name = display_name
        
        # Add "+ New Branch..." option
        branch_names.append("➕ New Branch...")
        
        # Update combobox
        self.branch_combo['values'] = branch_names
        
        # Set current selection
        if current_branch_name:
            self.branch_var.set(current_branch_name)
        elif branches:
            # Default to first branch if current not found
            self.branch_var.set(branch_names[0])
        else:
            pass
        
        
        # Show the selector (in case it was hidden)
        # Pack as its own row below the header
        self.branch_selector_frame.pack(fill=tk.X)
        
        # Update visibility based on mode
        self._update_branch_selector_visibility()
    
    def _update_branch_selector_visibility(self):
        """Show/hide branch selector based on available branches"""
        if not hasattr(self, 'branch_selector_frame'):
            return
        
        # Show branch selector in BOTH modes when there are branches
        # This allows users to see and access conversations from either mode
        if hasattr(self, '_branch_list') and self._branch_list:
            self.branch_selector_frame.pack(fill=tk.X)
        else:
            self.branch_selector_frame.pack_forget()
    
    def _on_branch_selected(self, event=None):
        """
        Handle branch selection from dropdown.
        Switches to the selected conversation branch.
        """
        selected = self.branch_var.get()
        
        if not selected:
            return
        
        # Check if "+ New Branch..." was selected
        if selected in ("+ New Branch...", "➕ New Branch..."):
            self._create_new_branch_from_selector()
            return
        
        # Find the branch ID for the selected name
        if not hasattr(self, '_branch_list') or not self._branch_list:
            return
        
        
        selected_branch_id = None
        for i, branch in enumerate(self._branch_list):
            title = branch.get('title', '')
            branch_id = branch.get('doc_id', '')  # Note: 'doc_id' not 'id'
            
            # Match using same logic as _populate_branch_selector
            if title.startswith('[Response]'):
                display_name = title[10:].strip()
                if ':' in display_name:
                    display_name = display_name.split(':', 1)[1].strip()
            else:
                display_name = title
            
            if len(display_name) > 30:
                display_name = display_name[:27] + "..."
            
            # Add processing indicator to match what _populate_branch_selector shows
            if branch.get('is_processing', False):
                display_name += " ⏳"
            
            is_match = display_name == selected
            match_marker = "✅ MATCH" if is_match else ""
            
            if is_match:
                selected_branch_id = branch_id
                break
        
        if not selected_branch_id:
            # Restore the previous selection
            self._populate_branch_selector()
            return
        
        # Don't switch if already on this branch AND in conversation mode
        # In source mode, we should switch to show the branch's conversation
        if selected_branch_id == self.current_document_id and self.current_mode == 'conversation':
            return
        
        # If in source mode viewing the same branch, just switch to conversation mode
        if selected_branch_id == self.current_document_id and self.current_mode == 'source':
            self.switch_mode('conversation')
            return
        
        
        # Save current edits before switching
        try:
            if self.current_mode == 'conversation':
                self._save_edits_to_thread()
        except ValueError:
            pass
        
        # Load the selected branch
        self._load_branch(selected_branch_id)
    
    def _load_branch(self, branch_doc_id: str):
        """
        Load a conversation branch by document ID.
        Updates the viewer to show the selected branch's content.
        """
        from document_library import get_document_by_id, load_thread_from_document, load_document_entries, save_thread_to_document
        
        
        # =================================================================
        # CRITICAL FIX: Save current thread to its document BEFORE clearing
        # This prevents data loss when switching branches, since current_thread
        # is a shared reference with Main.py
        # =================================================================
        if self.current_document_id and self.current_document_id != branch_doc_id:
            if self.current_thread and len(self.current_thread) > 0:
                try:
                    metadata = {
                        "model": self.model_var.get(),
                        "provider": self.provider_var.get(),
                        "last_updated": datetime.datetime.now().isoformat(),
                        "message_count": self.thread_message_count
                    }
                    save_thread_to_document(self.current_document_id, self.current_thread, metadata)
                except Exception as e:
                    pass
        
        # Get the branch document
        branch_doc = get_document_by_id(branch_doc_id)
        if not branch_doc:
            return
        
        
        # Load thread from this branch
        thread_data, thread_metadata = load_thread_from_document(branch_doc_id)
        
        
        # Update viewer state
        self.current_document_id = branch_doc_id
        self.doc_title = branch_doc.get('title', 'Unknown')
        
        # Update thread
        self.current_thread.clear()
        if thread_data:
            self.current_thread.extend(thread_data)
        self.thread_message_count = len([m for m in self.current_thread if m.get('role') == 'user'])
        
        # Reset exchange collapse state for new thread
        self.exchange_expanded_state.clear()
        self._init_exchange_collapse_state()
        
        # Update window title
        self._update_window_title()
        
        # Switch to conversation mode and refresh display
        self.current_mode = 'conversation'
        self._refresh_thread_display()
        self._update_mode_buttons()
        
        # === CRITICAL: Load source document entries for processing ===
        # Get parent source document ID from metadata
        metadata = branch_doc.get('metadata', {})
        source_doc_id = metadata.get('parent_document_id') or metadata.get('original_document_id')
        
        if source_doc_id:
            source_entries = load_document_entries(source_doc_id)
            if source_entries:
                self.current_entries = source_entries
                
                # Also load source documents for display
                source_doc = get_document_by_id(source_doc_id)
                if source_doc:
                    from utils import entries_to_text_with_speakers, entries_to_text
                    # Determine if audio transcription
                    doc_type = source_doc.get('type', 'text')
                    is_audio = doc_type == 'audio_transcription'
                    
                    if is_audio:
                        source_text = entries_to_text_with_speakers(
                            source_entries,
                            timestamp_interval=self.config.get('timestamp_interval', 'every_segment')
                        )
                    else:
                        source_text = entries_to_text(source_entries)
                    
                    # Update source documents for display
                    self.source_documents = [{
                        'title': source_doc.get('title', 'Source Document'),
                        'text': source_text,
                        'source': source_doc.get('source', ''),
                        'char_count': len(source_text) if source_text else 0
                    }]
                    self.current_document_text = source_text
                    self.current_document_source = source_doc.get('source', '')
            else:
                pass
        else:
            pass
        
        # Update main app context if available
        if hasattr(self, 'app') and self.app:
            self.app.current_document_id = branch_doc_id
            self.app.current_document_class = 'response'
            if hasattr(self.app, 'current_thread'):
                self.app.current_thread.clear()
                if thread_data:
                    self.app.current_thread.extend(thread_data)
            if hasattr(self.app, 'thread_message_count'):
                self.app.thread_message_count = self.thread_message_count
            
            # CRITICAL: Update entries in main app for processing
            if hasattr(self, 'current_entries') and self.current_entries:
                if hasattr(self.app, 'current_entries'):
                    self.app.current_entries = self.current_entries
                    
            # Also update source document reference
            if source_doc_id:
                if hasattr(self.app, 'source_document_id'):
                    self.app.source_document_id = source_doc_id
                if hasattr(self.app, 'current_document_metadata'):
                    self.app.current_document_metadata = metadata
                # CRITICAL: Update document text in main app for processing
                if hasattr(self, 'current_document_text') and self.current_document_text:
                    if hasattr(self.app, 'current_document_text'):
                        self.app.current_document_text = self.current_document_text
                    
            if hasattr(self.app, 'update_button_states'):
                self.app.update_button_states()
        
        if hasattr(self, 'app') and self.app:
            pass
    
    def _create_new_branch_from_selector(self):
        """
        Handle "+ New Branch..." selection from dropdown.
        Shows a dialog to create a new conversation branch.
        """
        from tkinter import simpledialog
        
        # Prompt for branch name
        branch_name = simpledialog.askstring(
            "New Conversation Branch",
            "Enter a name for the new conversation branch:",
            parent=self.window
        )
        
        if not branch_name or not branch_name.strip():
            # Cancelled or empty - restore previous selection
            self._populate_branch_selector()
            return
        
        branch_name = branch_name.strip()
        
        # Get source document ID
        source_id = self._get_source_document_id()
        if not source_id:
            self._populate_branch_selector()
            return
        
        # Create the new branch document
        self._create_new_branch_and_switch(branch_name, source_id)
    
    def _create_new_branch_and_switch(self, branch_name: str, source_doc_id: str):
        """
        Create a new empty conversation branch and switch to it.
        """
        import datetime
        from document_library import add_document_to_library, get_document_by_id
        
        
        # Get source document info
        source_doc = get_document_by_id(source_doc_id)
        if not source_doc:
            return
        
        source_title = source_doc.get('title', 'Unknown')
        source_url = source_doc.get('source', '')
        
        # Create title for the response document
        response_title = f"[Response] {branch_name}"
        
        # Create metadata linking back to source
        # NOTE: pre_created=True ensures get_response_branches_for_source includes
        # this branch even before any exchanges are added (0-exchange filter).
        # manually_created=True distinguishes user-created branches from auto-created
        # ones so they don't show the processing hourglass icon.
        response_metadata = {
            "original_document_id": source_doc_id,
            "parent_document_id": source_doc_id,
            "source_title": source_title,
            "branch_name": branch_name,
            "created": datetime.datetime.now().isoformat(),
            "editable": True,
            "pre_created": True,
            "manually_created": True
        }
        
        # Create the response document
        # CRITICAL: Include branch_name in source to generate UNIQUE doc_id for each branch
        # Without this, all branches for the same source get the same ID and overwrite each other!
        new_doc_id = add_document_to_library(
            doc_type="conversation_thread",
            source=f"Conversation about: {source_url or source_title} - Branch: {branch_name}",
            title=response_title,
            entries=[{"text": f"Conversation branch: {branch_name}", "location": "Header"}],
            metadata=response_metadata,
            document_class="response"
        )
        
        if not new_doc_id:
            self._populate_branch_selector()
            return
        
        
        # Switch to the new branch
        self._load_branch(new_doc_id)
        
        if hasattr(self, 'app') and self.app:
            pass
        
        # Refresh the branch selector to include the new branch
        self._populate_branch_selector()
    
    # === END BRANCH SELECTOR METHODS ===
    
    def _copy_link(self):
        """Copy a reference link to clipboard"""
        if self.current_document_source:
            self.window.clipboard_clear()
            self.window.clipboard_append(self.current_document_source)
            # Brief visual feedback could be added here
    
    def _show_menu(self):
        """Show a dropdown menu with additional options"""
        menu = tk.Menu(self.window, tearoff=0)
        menu.add_command(label="Copy Source Link", command=self._copy_link)
        menu.add_separator()
        menu.add_command(label="Copy All", command=self._copy_thread)
        menu.add_command(label="Undo Edits", command=self._undo_edit)
        menu.add_separator()
        menu.add_command(label="Expand All", command=lambda: self._set_all_expanded(True))
        menu.add_command(label="Collapse All", command=lambda: self._set_all_expanded(False))
        menu.add_separator()
        menu.add_command(label="Close", command=self._close_window)
        
        # Position menu below the button
        try:
            menu.tk_popup(self.window.winfo_pointerx(), self.window.winfo_pointery())
        finally:
            menu.grab_release()
    
    def _set_all_expanded(self, expanded: bool):
        """Set all exchanges AND all sources to expanded or collapsed state"""
        # Check warning threshold if expanding
        if expanded and not self._check_expansion_warning():
            return  # User cancelled
        
        num_exchanges = self._count_exchanges()
        
        # Set all exchanges
        for i in range(num_exchanges):
            self.exchange_expanded_state[i] = expanded
        
        # Set ALL source sections
        for i in range(len(self.source_documents)):
            self.source_expanded_state[i] = expanded
        
        # Update legacy flag
        self.source_section_visible = expanded
        
        # Update source toggle button if it exists
        if hasattr(self, 'source_toggle_btn'):
            if expanded:
                self.source_toggle_btn.config(text="Hide Sources")
            else:
                self.source_toggle_btn.config(text="Show Sources")
        
        # Update expand/collapse button
        if hasattr(self, 'exchanges_toggle_btn'):
            if expanded:
                self.exchanges_toggle_btn.config(text="Collapse All")
            else:
                self.exchanges_toggle_btn.config(text="Expand All")
        
        # Save any edits before refreshing (otherwise they'll be lost)
        self._save_edits_before_refresh()
        
        self._refresh_thread_display()
    
    def _create_find_replace_bar(self):
        """Create a compact find-replace bar for fixing names/places"""
        # Container frame with subtle background
        self.find_replace_frame = ttk.Frame(self.window, padding=(10, 5, 10, 5))
        self.find_replace_frame.pack(fill=tk.X)
        
        # Find field
        ttk.Label(
            self.find_replace_frame,
            text="Find:",
            font=('Arial', 9)
        ).pack(side=tk.LEFT, padx=(0, 5))
        
        self.find_var = tk.StringVar()
        self.find_entry = ttk.Entry(
            self.find_replace_frame,
            textvariable=self.find_var,
            width=20
        )
        self.find_entry.pack(side=tk.LEFT, padx=(0, 10))
        
        # Bind Enter in find field to find next
        self.find_entry.bind('<Return>', lambda e: self._find_next())
        
        # Replace field
        ttk.Label(
            self.find_replace_frame,
            text="Replace:",
            font=('Arial', 9)
        ).pack(side=tk.LEFT, padx=(0, 5))
        
        self.replace_var = tk.StringVar()
        self.replace_entry = ttk.Entry(
            self.find_replace_frame,
            textvariable=self.replace_var,
            width=20
        )
        self.replace_entry.pack(side=tk.LEFT, padx=(0, 10))
        
        # Bind Enter in replace field to replace next
        self.replace_entry.bind('<Return>', lambda e: self._replace_next())
        
        # Buttons
        find_btn = ttk.Button(
            self.find_replace_frame,
            text="Find",
            command=self._find_next,
            width=6
        )
        find_btn.pack(side=tk.LEFT, padx=2)
        
        replace_btn = ttk.Button(
            self.find_replace_frame,
            text="Replace",
            command=self._replace_next,
            width=8
        )
        replace_btn.pack(side=tk.LEFT, padx=2)
        
        replace_all_btn = ttk.Button(
            self.find_replace_frame,
            text="Replace All",
            command=self._replace_all,
            width=10
        )
        replace_all_btn.pack(side=tk.LEFT, padx=2)
        
        # Match count label
        self.match_count_label = ttk.Label(
            self.find_replace_frame,
            text="",
            font=('Arial', 9),
            foreground='gray'
        )
        self.match_count_label.pack(side=tk.LEFT, padx=(10, 0))
        
        # Add help to the find-replace components
        if HELP_TEXTS:
            add_help(self.find_entry, **HELP_TEXTS.get("thread_find_field", 
                {"title": "Find", "description": "Enter text to find. Press Enter to find next."}))
            add_help(self.replace_entry, **HELP_TEXTS.get("thread_replace_field", 
                {"title": "Replace", "description": "Enter replacement text. Press Enter to replace."}))
            add_help(find_btn, **HELP_TEXTS.get("thread_find_button", 
                {"title": "Find", "description": "Find the next occurrence of the search text."}))
            add_help(replace_btn, **HELP_TEXTS.get("thread_replace_button", 
                {"title": "Replace", "description": "Replace the current selection and find next."}))
            add_help(replace_all_btn, **HELP_TEXTS.get("thread_replace_all_button", 
                {"title": "Replace All", "description": "Replace all occurrences throughout the conversation."}))
        
        # Track current search position
        self._search_start = "1.0"
    
    def _focus_find_field(self):
        """Focus the find entry field"""
        self.find_entry.focus_set()
        self.find_entry.select_range(0, tk.END)
        return "break"
    
    def _find_next(self):
        """Find the next occurrence of the search text"""
        search_text = self.find_var.get()
        if not search_text:
            self.match_count_label.config(text="")
            return
        
        # Remove any existing highlight
        self.thread_text.tag_remove("search_highlight", "1.0", tk.END)
        
        # Configure highlight tag
        self.thread_text.tag_config("search_highlight", background="yellow", foreground="black")
        
        # Search from current position
        pos = self.thread_text.search(
            search_text, 
            self._search_start, 
            stopindex=tk.END,
            nocase=True  # Case-insensitive search
        )
        
        if not pos:
            # Wrap around to beginning
            pos = self.thread_text.search(
                search_text, 
                "1.0", 
                stopindex=self._search_start,
                nocase=True
            )
        
        if pos:
            # Calculate end position
            end_pos = f"{pos}+{len(search_text)}c"
            
            # Highlight and select the found text
            self.thread_text.tag_add("search_highlight", pos, end_pos)
            self.thread_text.tag_add(tk.SEL, pos, end_pos)
            self.thread_text.mark_set(tk.INSERT, end_pos)
            self.thread_text.see(pos)
            
            # Update search start for next find
            self._search_start = end_pos
            
            # Count total matches
            self._update_match_count(search_text)
        else:
            self.match_count_label.config(text="Not found", foreground='red')
            self._search_start = "1.0"
    
    def _update_match_count(self, search_text):
        """Count and display the number of matches"""
        count = 0
        start = "1.0"
        while True:
            pos = self.thread_text.search(
                search_text, 
                start, 
                stopindex=tk.END,
                nocase=True
            )
            if not pos:
                break
            count += 1
            start = f"{pos}+1c"
        
        if count > 0:
            self.match_count_label.config(text=f"{count} match{'es' if count != 1 else ''}", foreground='gray')
        else:
            self.match_count_label.config(text="Not found", foreground='red')
    
    def _replace_next(self):
        """Replace the current selection and find next"""
        search_text = self.find_var.get()
        replace_text = self.replace_var.get()
        
        if not search_text:
            return
        
        # Check if there's a selection that matches the search text
        try:
            selected = self.thread_text.get(tk.SEL_FIRST, tk.SEL_LAST)
            if selected.lower() == search_text.lower():
                # Replace the selection
                self.thread_text.delete(tk.SEL_FIRST, tk.SEL_LAST)
                self.thread_text.insert(tk.INSERT, replace_text)
        except tk.TclError:
            # No selection - just find next
            pass
        
        # Find next occurrence
        self._find_next()
    
    def _replace_all(self):
        """Replace all occurrences of the search text"""
        search_text = self.find_var.get()
        replace_text = self.replace_var.get()
        
        if not search_text:
            return
        
        # Count replacements
        count = 0
        
        # Start from the beginning
        self.thread_text.mark_set(tk.INSERT, "1.0")
        
        while True:
            # Find next occurrence
            pos = self.thread_text.search(
                search_text, 
                tk.INSERT, 
                stopindex=tk.END,
                nocase=True
            )
            
            if not pos:
                break
            
            # Calculate end position
            end_pos = f"{pos}+{len(search_text)}c"
            
            # Delete and insert replacement
            self.thread_text.delete(pos, end_pos)
            self.thread_text.insert(pos, replace_text)
            
            # Move cursor past the replacement
            self.thread_text.mark_set(tk.INSERT, f"{pos}+{len(replace_text)}c")
            
            count += 1
        
        # Remove any highlights
        self.thread_text.tag_remove("search_highlight", "1.0", tk.END)
        
        # Update status
        if count > 0:
            self.match_count_label.config(
                text=f"Replaced {count} occurrence{'s' if count != 1 else ''}", 
                foreground='green'
            )
        else:
            self.match_count_label.config(text="Not found", foreground='red')
        
        # Reset search position
        self._search_start = "1.0"
    
    def _create_document_info(self):
        """Create document information section"""
        doc_info_frame = ttk.Frame(self.window, padding=(10, 0, 10, 5))
        doc_info_frame.pack(fill=tk.X)
        
        ttk.Label(
            doc_info_frame, 
            text="SOURCE DOCUMENT INFORMATION:",
            font=('Arial', 9, 'bold'), 
            foreground='#2c3e50'
        ).pack(anchor=tk.W, pady=(0, 5))
        
        # For multi-source, show count
        if len(self.source_documents) > 1:
            ttk.Label(
                doc_info_frame, 
                text=f"  📚 Sources: {len(self.source_documents)} documents loaded",
                font=('Arial', 9), 
                foreground='gray'
            ).pack(anchor=tk.W)
        else:
            ttk.Label(
                doc_info_frame, 
                text=f"  📄 Title: {self.doc_title}",
                font=('Arial', 9), 
                foreground='gray'
            ).pack(anchor=tk.W)
            
            ttk.Label(
                doc_info_frame, 
                text=f"  🔗 Source: {self.source_info}",
                font=('Arial', 9), 
                foreground='gray'
            ).pack(anchor=tk.W)
        
        # Show published date if available
        if hasattr(self, 'published_date') and self.published_date and self.published_date != 'N/A':
            ttk.Label(
                doc_info_frame, 
                text=f"  📅 Published: {self.published_date}",
                font=('Arial', 9), 
                foreground='gray'
            ).pack(anchor=tk.W)
        
        ttk.Label(
            doc_info_frame, 
            text=f"  📅 Imported: {self.fetched_date}",
            font=('Arial', 9), 
            foreground='gray'
        ).pack(anchor=tk.W)
        
        # AI Provider and Model info
        provider = self.provider_var.get() if self.provider_var else "N/A"
        model = self.model_var.get() if self.model_var else "N/A"
        
        ttk.Label(
            doc_info_frame, 
            text=f"  🤖 AI: {provider} / {model}",
            font=('Arial', 9), 
            foreground='gray'
        ).pack(anchor=tk.W)
    
    def _create_thread_display(self):
        """Create the main thread content display"""
        content_frame = ttk.Frame(self.window, padding=10)
        content_frame.pack(fill=tk.BOTH, expand=True)
        
        # Create scrolled text widget - white background indicates editable
        # undo=True enables built-in undo/redo (Ctrl+Z works automatically)
        self.thread_text = scrolledtext.ScrolledText(
            content_frame, 
            wrap=tk.WORD,
            font=('Arial', self._get_font_size()),
            bg='white',  # White background to indicate editable
            height=16,
            undo=True,  # Enable undo/redo
            maxundo=-1  # Unlimited undo steps
        )
        self.thread_text.pack(fill=tk.BOTH, expand=True)
        
        # Bind mouse wheel to only scroll this widget (prevent propagation to parent)
        def on_mousewheel(event):
            """Handle mouse wheel - scroll only this widget"""
            # Windows uses event.delta, Linux uses Button-4/Button-5
            if event.delta:
                self.thread_text.yview_scroll(int(-1 * (event.delta / 120)), "units")
            return "break"  # Prevent event propagation
        
        def on_mousewheel_linux(event):
            """Handle mouse wheel on Linux"""
            if event.num == 4:
                self.thread_text.yview_scroll(-3, "units")
            elif event.num == 5:
                self.thread_text.yview_scroll(3, "units")
            return "break"
        
        # Bind for Windows/Mac
        self.thread_text.bind("<MouseWheel>", on_mousewheel)
        # Bind for Linux
        self.thread_text.bind("<Button-4>", on_mousewheel_linux)
        self.thread_text.bind("<Button-5>", on_mousewheel_linux)
        
        # Add help to thread display area
        if HELP_TEXTS:
            add_help(self.thread_text, **HELP_TEXTS.get("thread_display_area", 
                {"title": "Conversation Thread", "description": "Your conversation history - editable"}))
        
        # Note: Tag configurations are done in _refresh_thread_display to support font size changes
        
        # Populate with current thread content
        self._refresh_thread_display()
    
    def _refresh_thread_display(self):
        """Refresh the display based on current mode (source or conversation)"""
        self.thread_text.config(state=tk.NORMAL)
        self.thread_text.delete('1.0', tk.END)
        
        # Ensure heading and title match current mode
        self._update_heading()
        self._update_window_title()
        
        font_size = self._get_font_size()
        
        # Update the widget's base font
        self.thread_text.config(font=('Arial', font_size))
        
        # Configure tags (used by both modes)
        self.thread_text.tag_config("normal", font=('Arial', font_size))
        self.thread_text.tag_config("bold", font=('Arial', font_size, 'bold'))
        self.thread_text.tag_config("italic", font=('Arial', font_size, 'italic'))
        self.thread_text.tag_config("header", font=('Arial', font_size, 'bold'), foreground='#2c3e50')
        self.thread_text.tag_config("bullet", font=('Arial', font_size), lmargin1=20, lmargin2=35)
        self.thread_text.tag_config("user", foreground='#2E4053', font=('Arial', font_size, 'bold'))
        self.thread_text.tag_config("assistant", foreground='#16537E', font=('Arial', font_size, 'bold'))
        self.thread_text.tag_config("timestamp", foreground='#7F8C8D', font=('Arial', max(8, font_size - 2)))
        self.thread_text.tag_config("divider", foreground='#BDC3C7')
        self.thread_text.tag_config("processing", foreground='#E67E22', font=('Arial', font_size, 'italic'))
        self.thread_text.tag_config("success", foreground='#27AE60', font=('Arial', font_size, 'bold'))
        self.thread_text.tag_config("error", foreground='#E74C3C', font=('Arial', font_size, 'bold'))
        self.thread_text.tag_config("source_header", font=('Arial', font_size, 'bold'), foreground='#1a5276', background='#d4e6f1')
        self.thread_text.tag_config("source_text", font=('Arial', font_size), foreground='#2c3e50')
        self.thread_text.tag_config("collapsed_indicator", font=('Arial', font_size, 'italic'), foreground='#7f8c8d')
        # Exchange header style - clickable section headers for collapsible exchanges
        self.thread_text.tag_config("exchange_header", font=('Arial', font_size, 'bold'), foreground='#1a5276', background='#e8f4f8')
        
        # Dispatch to appropriate display method based on mode
        if self.current_mode == 'source':
            self._display_source_mode()
        else:
            self._display_conversation_mode()
        
        # Add clickable hyperlinks
        self._make_links_clickable()
        
        # Scroll position depends on mode
        if self.current_mode == 'source':
            self.thread_text.see('1.0')  # Scroll to top for source
        else:
            # Scroll to top of last exchange so user can see the last Q&A
            num_exchanges = self._count_exchanges()
            if num_exchanges > 0:
                last_exchange_idx = num_exchanges - 1
                self.window.after(10, lambda: self._scroll_to_exchange(last_exchange_idx))
            else:
                self.thread_text.see('1.0')  # No exchanges, scroll to top
        
        # Update info label based on mode
        self._update_info_label()
        
        # Reset modification flag - content was just loaded, not edited
        self.thread_text.edit_modified(False)
    
    def _display_source_mode(self):
        """Display source document(s) in prose format with collapsible sections"""
        if not self.source_documents:
            self.thread_text.insert(tk.END, "No source document loaded.\n", "normal")
            return
        
        # Single document: display directly (no header needed)
        if len(self.source_documents) == 1:
            source_text = self.source_documents[0].get('text', '')
            if source_text:
                self.thread_text.insert(tk.END, source_text, "source_text")
            else:
                self.thread_text.insert(tk.END, "No content in source document.\n", "normal")
            return
        
        # Multiple documents: display as collapsible sections
        for i, doc in enumerate(self.source_documents):
            self._insert_source_header_multi(i)
            if self.source_expanded_state.get(i, False):
                self._insert_source_content_multi(i)
    
    def _display_conversation_mode(self):
        """Display conversation thread with collapsible source(s) and exchanges"""
        # Show collapsible source section(s)
        if self.source_documents:
            if len(self.source_documents) == 1:
                # Single source: use original behavior
                self._insert_source_header()
                if self.source_section_visible:
                    self._insert_source_content()
            else:
                # Multiple sources: show each as separate collapsible section
                for i, doc in enumerate(self.source_documents):
                    self._insert_source_header_multi(i)
                    if self.source_expanded_state.get(i, False):
                        self._insert_source_content_multi(i, truncate=True)
                
                # Add separator after all sources
                self.thread_text.insert(tk.END, "\n" + "═" * 60 + "\n\n", "divider")
        
        # Check for empty conversation
        if not self.current_thread or len(self.current_thread) == 0:
            self.thread_text.insert(tk.END, "No conversation yet.\n\n", "normal")
            self.thread_text.insert(tk.END, "Use the follow-up field below to start a conversation.\n", "normal")
            return
        
        # Group messages into exchanges (user + assistant pairs)
        exchanges = self._group_messages_into_exchanges()
        
        # Update collapse state for any new exchanges
        self._update_exchange_collapse_state(len(exchanges))
        
        current_time = datetime.datetime.now()
        
        for exchange_idx, exchange in enumerate(exchanges):
            is_expanded = self.exchange_expanded_state.get(exchange_idx, True)
            
            # Calculate approximate timestamp for this exchange
            approx_time = current_time - datetime.timedelta(
                minutes=(len(exchanges) - exchange_idx) * 5
            )
            timestamp_str = approx_time.strftime("%H:%M")
            
            # Get preview of user question for collapsed view
            user_msg = exchange.get('user', {})
            user_content = user_msg.get('content', 'Question')
            question_preview = user_content[:80] + "..." if len(user_content) > 80 else user_content
            question_preview = question_preview.replace('\n', ' ')
            
            if is_expanded:
                # Expanded exchange header
                self._insert_exchange_header(exchange_idx, timestamp_str, question_preview, expanded=True)
                
                # User message
                self.thread_text.insert(tk.END, f"🧑 YOU ", "user")
                self.thread_text.insert(tk.END, f"[{timestamp_str}]\n", "timestamp")
                self.thread_text.insert(tk.END, f"{user_content}\n", "normal")
                
                # Assistant message (if present)
                assistant_msg = exchange.get('assistant', {})
                if assistant_msg:
                    assistant_content = assistant_msg.get('content', '')
                    msg_provider = assistant_msg.get('provider', '')
                    msg_model = assistant_msg.get('model', '')
                    
                    if msg_provider or msg_model:
                        ai_label = f"🤖 {msg_provider}" if msg_provider else "🤖 AI"
                        if msg_model and msg_model != msg_provider:
                            ai_label += f" ({msg_model})"
                    else:
                        ai_label = "🤖 AI"
                    
                    self.thread_text.insert(tk.END, f"\n{ai_label} ", "assistant")
                    self.thread_text.insert(tk.END, f"[{timestamp_str}]\n", "timestamp")
                    
                    # Render AI response content with markdown formatting
                    self._render_markdown_content(assistant_content)
                
                # End of exchange
                self.thread_text.insert(tk.END, "\n" + "─" * 60 + "\n", "divider")
            else:
                # Collapsed exchange - just show header
                self._insert_exchange_header(exchange_idx, timestamp_str, question_preview, expanded=False)
    
    def _insert_source_header_multi(self, index: int):
        """Insert a clickable header for a specific source document (multi-source mode)"""
        is_expanded = self.source_expanded_state.get(index, False)
        indicator = "▼" if is_expanded else "▶"
        
        doc = self.source_documents[index]
        title = doc.get('title', f'Document {index + 1}')
        char_count = doc.get('char_count', len(doc.get('text', '')))
        
        # Create header text
        header_text = f"{indicator} SOURCE DOCUMENT {index + 1}: {title}"
        
        # Insert header with tag for click handling
        tag_name = f"source_header_{index}"
        
        self.thread_text.insert(tk.END, header_text, (tag_name, "exchange_header"))
        
        # Show character count
        self.thread_text.insert(tk.END, f" ({char_count:,} chars)", "timestamp")
        
        if not is_expanded:
            # Show hint when collapsed
            self.thread_text.insert(tk.END, " - click to expand", "collapsed_indicator")
        
        self.thread_text.insert(tk.END, "\n", "")
        
        # Configure tag for click handling
        self.thread_text.tag_bind(tag_name, "<Button-1>", lambda e, idx=index: self._toggle_source_document(idx))
        self.thread_text.tag_bind(tag_name, "<Enter>", lambda e: self.thread_text.config(cursor="hand2"))
        self.thread_text.tag_bind(tag_name, "<Leave>", lambda e: self.thread_text.config(cursor=""))
    
    def _insert_source_content_multi(self, index: int, truncate: bool = False):
        """
        Insert the content for a specific source document (multi-source mode).
        
        Args:
            index: Index of the source document
            truncate: If True, truncate long content (for conversation mode)
        """
        self.thread_text.insert(tk.END, "─" * 60 + "\n", "divider")
        
        doc = self.source_documents[index]
        source_text = doc.get('text', '')
        
        if truncate:
            max_display = 5000  # Characters to show in conversation mode
            if len(source_text) > max_display:
                truncated = source_text[:max_display]
                self.thread_text.insert(tk.END, truncated, "source_text")
                remaining = len(source_text) - max_display
                self.thread_text.insert(
                    tk.END, 
                    f"\n\n... [{remaining:,} more characters - switch to Source Mode to view full document]\n", 
                    "collapsed_indicator"
                )
            else:
                self.thread_text.insert(tk.END, source_text, "source_text")
        else:
            # Full content in source mode
            self.thread_text.insert(tk.END, source_text, "source_text")
        
        # End of source section
        self.thread_text.insert(tk.END, "\n" + "─" * 60 + "\n\n", "divider")
        
        # Mark as rendered (for lazy loading tracking)
        self.source_content_rendered[index] = True
    
    def _toggle_source_document(self, index: int):
        """Toggle the expanded/collapsed state of a specific source document"""
        current_state = self.source_expanded_state.get(index, False)
        new_state = not current_state
        
        # If expanding, check warning threshold
        if new_state and not self._check_expansion_warning(index):
            return  # User cancelled
        
        self.source_expanded_state[index] = new_state
        
        # Save any edits before refreshing (otherwise they'll be lost)
        self._save_edits_before_refresh()
        
        self._refresh_thread_display()
    
    def _group_messages_into_exchanges(self) -> list:
        """Group thread messages into user-assistant exchange pairs"""
        exchanges = []
        current_exchange = {}
        
        for msg in self.current_thread:
            role = msg.get('role', 'unknown')
            
            if role == 'user':
                # Start new exchange
                if current_exchange:
                    exchanges.append(current_exchange)
                current_exchange = {'user': msg}
            elif role == 'assistant':
                # Add to current exchange
                current_exchange['assistant'] = msg
            elif role == 'system':
                # Skip system messages in display
                continue
        
        # Don't forget the last exchange
        if current_exchange:
            exchanges.append(current_exchange)
        
        return exchanges
    
    def _update_exchange_collapse_state(self, num_exchanges: int):
        """Update collapse state when new exchanges are added"""
        for i in range(num_exchanges):
            if i not in self.exchange_expanded_state:
                # New exchange - expand last 2, collapse older
                self.exchange_expanded_state[i] = (i >= num_exchanges - 2)
    
    def _insert_exchange_header(self, exchange_idx: int, timestamp: str, question_preview: str, expanded: bool):
        """Insert a clickable header for an exchange"""
        indicator = "▼" if expanded else "▶"
        
        # Create header text with exchange number
        header_text = f"{indicator} Exchange {exchange_idx + 1}"
        
        # Insert header with tag for click handling
        tag_name = f"exchange_header_{exchange_idx}"
        
        self.thread_text.insert(tk.END, header_text, (tag_name, "exchange_header"))
        self.thread_text.insert(tk.END, f" [{timestamp}]", "timestamp")
        
        if not expanded:
            # Show preview when collapsed
            self.thread_text.insert(tk.END, f" - {question_preview}", "collapsed_indicator")
        
        self.thread_text.insert(tk.END, "\n", "")
        
        # Configure tag for click handling
        self.thread_text.tag_bind(tag_name, "<Button-1>", lambda e, idx=exchange_idx: self._toggle_exchange(idx))
        self.thread_text.tag_bind(tag_name, "<Enter>", lambda e: self.thread_text.config(cursor="hand2"))
        self.thread_text.tag_bind(tag_name, "<Leave>", lambda e: self.thread_text.config(cursor=""))
    
    def _toggle_exchange(self, exchange_idx: int):
        """Toggle the expanded/collapsed state of an exchange"""
        current_state = self.exchange_expanded_state.get(exchange_idx, True)
        new_state = not current_state
        self.exchange_expanded_state[exchange_idx] = new_state
        
        # Save any edits before refreshing (otherwise they'll be lost)
        self._save_edits_before_refresh()
        
        self._refresh_thread_display()
        
        # If expanding, scroll to the beginning of this exchange
        if new_state:
            self.window.after(10, lambda: self._scroll_to_exchange(exchange_idx))
    
    def _scroll_to_exchange(self, exchange_idx: int):
        """Scroll to show the beginning of the specified exchange"""
        try:
            # Search for the exchange header tag
            tag_name = f"exchange_header_{exchange_idx}"
            
            # Find the range of this tag
            tag_ranges = self.thread_text.tag_ranges(tag_name)
            if tag_ranges:
                # Scroll to the start of the tag
                start_index = tag_ranges[0]
                self.thread_text.see(start_index)
                # Position at top of visible area
                self.thread_text.yview_moveto(0)
                # Now scroll to show the exchange at the top
                self.thread_text.see(start_index)
        except Exception as e:
            pass
    
    def _get_exchange_at_cursor(self) -> int:
        """
        Determine which exchange the cursor is currently in.
        
        Returns:
            Exchange index (0-based), or -1 if cursor is not in any exchange
        """
        try:
            # Get current cursor position
            cursor_pos = self.thread_text.index(tk.INSERT)
            
            # Get all exchange header tags and their positions
            exchanges = self._group_messages_into_exchanges()
            num_exchanges = len(exchanges)
            
            if num_exchanges == 0:
                return -1
            
            # Find which exchange header the cursor is after
            # Exchange headers have tags like "exchange_header_0", "exchange_header_1", etc.
            cursor_line = int(cursor_pos.split('.')[0])
            
            last_exchange_before_cursor = -1
            
            for i in range(num_exchanges):
                tag_name = f"exchange_header_{i}"
                tag_ranges = self.thread_text.tag_ranges(tag_name)
                
                if tag_ranges:
                    # Get the line number of this exchange header
                    header_start = str(tag_ranges[0])
                    header_line = int(header_start.split('.')[0])
                    
                    if header_line <= cursor_line:
                        last_exchange_before_cursor = i
                    else:
                        # We've passed the cursor position
                        break
            
            return last_exchange_before_cursor
            
        except Exception as e:
            return -1
    
    def _delete_current_exchange(self):
        """Delete the exchange where the cursor is currently positioned"""
        # Only works in conversation mode
        if self.current_mode != 'conversation':
            self._set_status("⚠️ Switch to Conversation mode to delete exchanges")
            return
        
        exchanges = self._group_messages_into_exchanges()
        num_exchanges = len(exchanges)
        
        if num_exchanges == 0:
            self._set_status("⚠️ No exchanges to delete")
            return
        
        # Find which exchange the cursor is in
        exchange_idx = self._get_exchange_at_cursor()
        
        if exchange_idx < 0:
            self._set_status("⚠️ Click inside an exchange first, then delete")
            return
        
        # Get preview of the exchange for confirmation
        exchange = exchanges[exchange_idx]
        user_content = exchange.get('user', {}).get('content', '')
        preview = user_content[:60] + "..." if len(user_content) > 60 else user_content
        preview = preview.replace('\n', ' ')
        
        # Confirm deletion
        confirm = messagebox.askyesno(
            "Delete Exchange?",
            f"Delete Exchange {exchange_idx + 1} of {num_exchanges}?\n\n"
            f"Question: \"{preview}\"\n\n"
            f"This will remove both the question and the AI response."
        )
        
        if not confirm:
            return
        
        # Find the messages to remove from current_thread
        # Each exchange has a user message and optionally an assistant message
        user_msg = exchange.get('user')
        assistant_msg = exchange.get('assistant')
        
        # Remove from current_thread
        if user_msg and user_msg in self.current_thread:
            self.current_thread.remove(user_msg)
        if assistant_msg and assistant_msg in self.current_thread:
            self.current_thread.remove(assistant_msg)
        
        # Update message count
        self.thread_message_count = len([m for m in self.current_thread if m.get('role') == 'user'])
        
        # Update exchange collapse state (shift indices down)
        new_collapse_state = {}
        for idx, expanded in self.exchange_expanded_state.items():
            if idx < exchange_idx:
                new_collapse_state[idx] = expanded
            elif idx > exchange_idx:
                new_collapse_state[idx - 1] = expanded
            # idx == exchange_idx is skipped (deleted)
        self.exchange_expanded_state = new_collapse_state
        
        # Save to document library
        if self.current_document_id:
            from document_library import save_thread_to_document
            metadata = {
                "model": self.model_var.get(),
                "provider": self.provider_var.get(),
                "last_updated": datetime.datetime.now().isoformat(),
                "message_count": self.thread_message_count,
                "edited": True
            }
            save_thread_to_document(self.current_document_id, self.current_thread, metadata)
        
        # Refresh display
        self._refresh_thread_display()
        
        # Update button visibility if no exchanges left
        self._update_delete_button_visibility()
        
        remaining = len(self._group_messages_into_exchanges())
        self._set_status(f"✅ Exchange deleted ({remaining} remaining)")
    
    def _update_delete_button_visibility(self):
        """Show/hide the delete exchange button based on mode and exchange count"""
        if hasattr(self, 'delete_exchange_btn'):
            BTN_PAD = 2
            if self.current_mode == 'conversation' and self._count_exchanges() > 0:
                # Check if already packed
                if not self.delete_exchange_btn.winfo_ismapped():
                    # Pack after save_as_btn to maintain correct order
                    self.delete_exchange_btn.pack(side=tk.LEFT, padx=BTN_PAD, after=self.save_as_btn)
            else:
                self.delete_exchange_btn.pack_forget()
    
    def _insert_source_header(self):
        """Insert a clickable header for the source document section (single-source mode)"""
        indicator = "▼" if self.source_section_visible else "▶"
        
        # Create header text
        header_text = f"{indicator} SOURCE DOCUMENT"
        
        # Insert header with tag for click handling
        tag_name = "source_header_clickable"
        
        self.thread_text.insert(tk.END, header_text, (tag_name, "exchange_header"))  # Reuse exchange_header style
        
        if not self.source_section_visible:
            # Show hint when collapsed
            self.thread_text.insert(tk.END, " (click to expand)", "collapsed_indicator")
        
        self.thread_text.insert(tk.END, "\n", "")
        
        # Configure tag for click handling
        self.thread_text.tag_bind(tag_name, "<Button-1>", lambda e: self._toggle_source_section())
        self.thread_text.tag_bind(tag_name, "<Enter>", lambda e: self.thread_text.config(cursor="hand2"))
        self.thread_text.tag_bind(tag_name, "<Leave>", lambda e: self.thread_text.config(cursor=""))
    
    def _insert_source_content(self):
        """Insert the source document content (single-source mode, called when source is expanded)"""
        self.thread_text.insert(tk.END, "─" * 60 + "\n", "divider")
        
        # Source content (truncated if very long for conversation view)
        source_text = self.source_documents[0].get('text', '') if self.source_documents else ''
        max_source_display = 5000  # Characters to show in conversation mode
        
        if len(source_text) > max_source_display:
            truncated = source_text[:max_source_display]
            self.thread_text.insert(tk.END, truncated, "source_text")
            remaining = len(source_text) - max_source_display
            self.thread_text.insert(tk.END, f"\n\n... [{remaining:,} more characters - switch to Source Mode to view full document]\n", "collapsed_indicator")
        else:
            self.thread_text.insert(tk.END, source_text, "source_text")
        
        # End of source section
        self.thread_text.insert(tk.END, "\n" + "═" * 60 + "\n\n", "divider")
    
    def _toggle_source_section(self):
        """Toggle the expanded/collapsed state of the source document section (single-source mode)"""
        # If expanding, check warning threshold
        if not self.source_section_visible:
            if not self._check_expansion_warning(0):
                return  # User cancelled
        
        self.source_section_visible = not self.source_section_visible
        
        # Save any edits before refreshing (otherwise they'll be lost)
        self._save_edits_before_refresh()
        
        self._refresh_thread_display()
        self._update_mode_buttons()
        
        # Update the button text if it exists
        if hasattr(self, 'source_toggle_btn'):
            if self.source_section_visible:
                self.source_toggle_btn.config(text="Hide Source")
            else:
                self.source_toggle_btn.config(text="Show Source")
    
    def _update_info_label(self):
        """Update the info label based on current mode"""
        if self.current_mode == 'conversation':
            info_text = f"{self.thread_message_count} exchange{'s' if self.thread_message_count != 1 else ''}"
        else:
            # For source mode, show document length
            total_chars = self._calculate_total_expanded_chars()
            if len(self.source_documents) > 1:
                info_text = f"{len(self.source_documents)} sources, {total_chars:,} total chars"
            elif total_chars > 10000:
                info_text = f"{total_chars:,} characters (long document)"
            elif total_chars > 0:
                info_text = f"{total_chars:,} characters"
            else:
                info_text = "No content"
        
        self.message_count_label.config(text=info_text)
    
    def switch_mode(self, new_mode: str):
        """
        Switch between 'source' and 'conversation' modes.
        
        Args:
            new_mode: Either 'source' or 'conversation'
        """
        
        if new_mode not in ('source', 'conversation'):
            return
        
        if new_mode == self.current_mode:
            return  # Already in this mode
        
        # Save any edits before switching (happens silently)
        try:
            if self.current_mode == 'conversation':
                self._save_edits_to_thread()
            elif self.current_mode == 'source':
                self._save_source_edits()
        except Exception as e:
            print(f"⚠️ Error saving on mode switch: {e}")
            # Continue switching anyway
        
        self.current_mode = new_mode
        
        # Update UI
        self._update_window_title()
        self._update_heading()
        self._refresh_thread_display()
        self._update_mode_buttons()
        self._update_branch_selector_visibility()
        
        # Notify main app of mode change (for button label update)
        if self.on_mode_change:
            self.on_mode_change(new_mode)
    
    def _toggle_mode(self):
        """Toggle between source and conversation modes (button callback)"""
        if self.current_mode == 'source':
            self.switch_mode('conversation')
        else:
            self.switch_mode('source')
    
    def toggle_source_visibility(self):
        """Toggle the visibility of ALL source sections (button command)"""
        if self.current_mode == 'conversation':
            if len(self.source_documents) == 1:
                # Single source: use original behavior
                self._toggle_source_section()
            else:
                # Multiple sources: toggle all
                # Determine current state (if most are expanded, collapse all; otherwise expand all)
                expanded_count = sum(1 for i in range(len(self.source_documents)) 
                                   if self.source_expanded_state.get(i, False))
                new_state = expanded_count <= len(self.source_documents) / 2
                
                # If expanding, check warning
                if new_state and not self._check_expansion_warning():
                    return
                
                for i in range(len(self.source_documents)):
                    self.source_expanded_state[i] = new_state
                
                self.source_section_visible = new_state
                
                # Save any edits before refreshing (otherwise they'll be lost)
                self._save_edits_before_refresh()
                
                self._refresh_thread_display()
                self._update_mode_buttons()
    
    def _toggle_all_exchanges(self):
        """Toggle all exchanges AND all sources between expanded and collapsed"""
        num_exchanges = self._count_exchanges()
        num_sources = len(self.source_documents)
        
        # Check if most content is expanded or collapsed
        expanded_exchange_count = sum(1 for i in range(num_exchanges) if self.exchange_expanded_state.get(i, True))
        expanded_source_count = sum(1 for i in range(num_sources) if self.source_expanded_state.get(i, False))
        
        total_items = num_exchanges + num_sources
        total_expanded = expanded_exchange_count + expanded_source_count
        
        most_expanded = total_expanded > total_items / 2 if total_items > 0 else True
        
        # Toggle all to opposite state
        new_state = not most_expanded
        
        # If expanding, check warning
        if new_state and not self._check_expansion_warning():
            return
        
        # Toggle exchanges
        for i in range(num_exchanges):
            self.exchange_expanded_state[i] = new_state
        
        # Toggle all source sections
        for i in range(num_sources):
            self.source_expanded_state[i] = new_state
        
        # Update legacy flag
        self.source_section_visible = new_state
        
        # Update source toggle button if it exists
        if hasattr(self, 'source_toggle_btn'):
            btn_text = "Hide Sources" if new_state else "Show Sources"
            if num_sources == 1:
                btn_text = "Hide Source" if new_state else "Show Source"
            self.source_toggle_btn.config(text=btn_text)
        
        # Update button text
        if hasattr(self, 'exchanges_toggle_btn'):
            if new_state:
                self.exchanges_toggle_btn.config(text="Collapse All")
            else:
                self.exchanges_toggle_btn.config(text="Expand All")
        
        # Save any edits before refreshing (otherwise they'll be lost)
        self._save_edits_before_refresh()
        
        # Refresh display
        self._refresh_thread_display()
    
    def _update_heading(self):
        """Update the heading label based on current mode"""
        if self.current_mode == 'source':
            heading_text = "📄 Source Document"
        else:
            heading_text = "💬 Conversation Thread"
        
        if len(self.source_documents) > 1:
            heading_text += f" ({len(self.source_documents)} sources)"
        
        self.heading_label.config(text=heading_text)
    
    def _update_mode_buttons(self):
        """Update mode toggle buttons based on current state"""
        if not hasattr(self, 'mode_toggle_btn'):
            return
        
        num_sources = len(self.source_documents)
        BTN_PAD = 2
        
        if self.current_mode == 'source':
            # In source mode - show button to go to conversation (if conversation exists)
            has_conversation = self.current_thread and len(self.current_thread) > 0
            if has_conversation:
                self.mode_toggle_btn.config(text="View Thread", state=tk.NORMAL)
            else:
                self.mode_toggle_btn.config(text="No Thread", state=tk.DISABLED)
            # Hide source visibility toggle in source mode
            if hasattr(self, 'source_toggle_btn'):
                self.source_toggle_btn.pack_forget()
            # Hide exchanges toggle in source mode
            if hasattr(self, 'exchanges_toggle_btn'):
                self.exchanges_toggle_btn.pack_forget()
        else:
            # In conversation mode - show button to go to source
            self.mode_toggle_btn.config(text="View Source", state=tk.NORMAL)
            
            # Track what was last packed for proper ordering
            last_packed = self.mode_toggle_btn
            
            # Show source visibility toggle
            if hasattr(self, 'source_toggle_btn') and self.source_documents:
                # Pack if not already visible
                if not self.source_toggle_btn.winfo_ismapped():
                    self.source_toggle_btn.pack(side=tk.LEFT, padx=BTN_PAD, after=last_packed)
                # Update last_packed regardless of whether we just packed it
                last_packed = self.source_toggle_btn
                
                # Determine button text based on source visibility
                if num_sources == 1:
                    btn_text = "Hide Source" if self.source_section_visible else "Show Source"
                else:
                    expanded_count = sum(1 for i in range(num_sources) if self.source_expanded_state.get(i, False))
                    if expanded_count > num_sources / 2:
                        btn_text = "Hide Sources"
                    else:
                        btn_text = "Show Sources"
                self.source_toggle_btn.config(text=btn_text)
            
            # Show exchanges toggle if there are exchanges
            if hasattr(self, 'exchanges_toggle_btn'):
                num_exchanges = self._count_exchanges()
                if num_exchanges > 0:
                    # Pack if not already visible
                    if not self.exchanges_toggle_btn.winfo_ismapped():
                        self.exchanges_toggle_btn.pack(side=tk.LEFT, padx=BTN_PAD, after=last_packed)
                    # Update button text based on current state
                    expanded_count = sum(1 for i in range(num_exchanges) if self.exchange_expanded_state.get(i, True))
                    if expanded_count > num_exchanges / 2:
                        self.exchanges_toggle_btn.config(text="Collapse All")
                    else:
                        self.exchanges_toggle_btn.config(text="Expand All")
                else:
                    self.exchanges_toggle_btn.pack_forget()
        
        # Update delete exchange button visibility
        self._update_delete_button_visibility()
    
    def _save_source_edits(self):
        """Save edits made to source document (when in source mode)"""
        if self.current_mode != 'source':
            return
        
        try:
            # Get current text from widget
            edited_text = self.thread_text.get('1.0', tk.END).rstrip()
            
            # For single source, update directly
            if len(self.source_documents) == 1:
                original_text = self.source_documents[0].get('text', '')
                if edited_text != original_text:
                    self.source_documents[0]['text'] = edited_text
                    self.source_documents[0]['char_count'] = len(edited_text)
                    # Update legacy field
                    self.current_document_text = edited_text
                    
                    # === FIX: Actually persist the changes to the Documents Library ===
                    if self.current_document_id:
                        from document_library import load_document_entries, get_document_by_id, DATA_DIR, save_json_atomic, load_library, save_library
                        import os
                        
                        # Get the document
                        doc = get_document_by_id(self.current_document_id)
                        if doc:
                            # Load current entries
                            entries = load_document_entries(self.current_document_id)
                            
                            if entries and len(entries) > 0:
                                # Update the text in the first entry (main content)
                                entries[0]['text'] = edited_text
                                entries[0]['char_count'] = len(edited_text)
                                
                                # Save entries directly to file (bypass editable check for source docs)
                                # Users should be able to correct transcription errors!
                                entries_file = os.path.join(DATA_DIR, f"doc_{self.current_document_id}_entries.json")
                                save_json_atomic(entries_file, entries)
                                
                                # Update the library's last_edited timestamp
                                library = load_library()
                                for idx, lib_doc in enumerate(library["documents"]):
                                    if lib_doc.get("id") == self.current_document_id:
                                        if "metadata" not in library["documents"][idx]:
                                            library["documents"][idx]["metadata"] = {}
                                        library["documents"][idx]["metadata"]["last_edited"] = datetime.datetime.now().isoformat()
                                        break
                                save_library(library)
                                
                                print(f"✅ Source document edits saved to library: {self.current_document_id}")
                                self._set_status("✅ Source edits saved")
                            else:
                                print(f"⚠️ No entries found for document {self.current_document_id}")
                    
            # For multiple sources, we'd need more sophisticated parsing
            # For now, don't save edits to multi-source (would need to track which section was edited)
            
        except Exception as e:
            print(f"⚠️ Error saving source edits: {e}")
            import traceback
            traceback.print_exc()

    # === REMAINING METHODS (unchanged from original) ===
    # The following methods are carried over from the original implementation
    # with no changes needed for multi-source support:
    
    def _render_markdown_content(self, content: str):
        """
        Render markdown-formatted content into the thread text widget.
        Supports: **bold**, *italic*, # headers, and bullets
        """
        lines = content.split('\n')
        
        for line in lines:
            # Headers: # Header or ## Header
            if line.strip().startswith('#'):
                text = line.lstrip('# ').strip()
                self.thread_text.insert(tk.END, text + '\n', 'header')
                continue
            
            # Bullets: - item or * item
            if line.strip().startswith(('- ', '* ')):
                text = '• ' + line.strip()[2:]
                self.thread_text.insert(tk.END, text + '\n', 'bullet')
                continue
            
            # Process bold and italic inline
            self._render_inline_markdown(line)
            self.thread_text.insert(tk.END, '\n', 'normal')
    
    def _render_inline_markdown(self, line: str):
        """
        Render inline markdown (bold, italic) in a line of text.
        """
        # Pattern: **bold** or *italic*
        combined_pattern = r'(\*\*(.*?)\*\*|\*(.*?)\*)'
        
        current_pos = 0
        matches = list(re.finditer(combined_pattern, line))
        
        if matches:
            for match in matches:
                # Insert text before the match
                if match.start() > current_pos:
                    self.thread_text.insert(tk.END, line[current_pos:match.start()], 'normal')
                
                # Insert the formatted text
                if match.group(0).startswith('**'):
                    # Bold
                    self.thread_text.insert(tk.END, match.group(2), 'bold')
                else:
                    # Italic
                    self.thread_text.insert(tk.END, match.group(3), 'italic')
                
                current_pos = match.end()
            
            # Insert remaining text
            if current_pos < len(line):
                self.thread_text.insert(tk.END, line[current_pos:], 'normal')
        else:
            # No markdown formatting, just insert the line
            self.thread_text.insert(tk.END, line, 'normal')
    
    def _make_links_clickable(self):
        """
        Find all URLs in the thread text and make them clickable.
        Handles both plain URLs and Markdown-style links [text](url).
        """
        # Configure hyperlink tag
        self.thread_text.tag_config("hyperlink", foreground="blue", underline=True)
        
        # Get all content
        content = self.thread_text.get('1.0', tk.END)
        
        # Pattern to find Markdown-style links: [text](url)
        markdown_pattern = r'\[([^\]]+)\]\((https?://[^\)]+)\)'
        # Pattern to find plain URLs (not already in markdown format)
        plain_url_pattern = r'(?<!\()(?<!\]\()(https?://[^\s\)\]\>"]+)'
        
        # Store URL positions and cleaned URLs
        self.url_locations = []
        
        # Find Markdown-style links
        for match in re.finditer(markdown_pattern, content):
            link_text = match.group(1)
            raw_url = match.group(2)
            # Remove spaces from URL (AI sometimes adds them)
            clean_url = raw_url.replace(' ', '')
            start_pos = f"1.0+{match.start()}c"
            end_pos = f"1.0+{match.end()}c"
            
            # Tag the entire link
            self.thread_text.tag_add("hyperlink", start_pos, end_pos)
            # Store for click handling
            self.url_locations.append((start_pos, end_pos, clean_url))
        
        # Find plain URLs (that aren't part of markdown links)
        for match in re.finditer(plain_url_pattern, content):
            url = match.group(0)
            # Clean up URL - remove trailing punctuation that's likely not part of the URL
            while url and url[-1] in '.,;:!?\'"':
                url = url[:-1]
            
            start_pos = f"1.0+{match.start()}c"
            end_pos = f"1.0+{match.start() + len(url)}c"
            
            # Check if this URL is already tagged (part of a markdown link)
            already_tagged = False
            for existing_start, existing_end, _ in self.url_locations:
                if (self.thread_text.compare(start_pos, ">=", existing_start) and 
                    self.thread_text.compare(end_pos, "<=", existing_end)):
                    already_tagged = True
                    break
            
            if not already_tagged:
                # Tag the URL
                self.thread_text.tag_add("hyperlink", start_pos, end_pos)
                # Store for click handling
                self.url_locations.append((start_pos, end_pos, url))
        
        # Bind click event
        self.thread_text.tag_bind("hyperlink", "<Button-1>", self._on_link_click)
        
        # Change cursor when hovering over links
        self.thread_text.tag_bind("hyperlink", "<Enter>",
                                   lambda e: self.thread_text.config(cursor="hand2"))
        self.thread_text.tag_bind("hyperlink", "<Leave>",
                                   lambda e: self.thread_text.config(cursor=""))
    
    def _on_link_click(self, event):
        """Handle click on a hyperlink."""
        # Get the index of the click
        index = self.thread_text.index(f"@{event.x},{event.y}")
        
        # Check which URL was clicked
        for start, end, url in self.url_locations:
            # Check if click is within this URL's range
            if (self.thread_text.compare(index, ">=", start) and 
                self.thread_text.compare(index, "<=", end)):
                # Open the URL in the default browser
                try:
                    webbrowser.open(url)
                except Exception as e:
                    messagebox.showerror("Error", f"Could not open URL:\n{url}\n\nError: {e}")
                return "break"
    
    def _create_followup_section(self):
        """Create the follow-up question input section"""
        followup_frame = ttk.LabelFrame(self.window, text="Ask a Follow-up Question", padding=10)
        followup_frame.pack(fill=tk.X, padx=10, pady=(0, 5))
        
        # Input area with submit button
        input_frame = ttk.Frame(followup_frame)
        input_frame.pack(fill=tk.X)
        
        # Submit button FIRST (pack on right) so it gets priority space
        self.submit_btn = ttk.Button(
            input_frame, 
            text="Submit",
            command=self._submit_followup,
            width=10
        )
        self.submit_btn.pack(side=tk.RIGHT, padx=(10, 0))
        
        # Add help to submit button
        if HELP_TEXTS:
            add_help(self.submit_btn, **HELP_TEXTS.get("thread_submit_button", 
                {"title": "Submit", "description": "Send your follow-up question"}))
        
        # Input text area (pack after submit so it fills remaining space)
        self.followup_input = scrolledtext.ScrolledText(
            input_frame, 
            wrap=tk.WORD, 
            height=2,
            font=('Arial', 10)
        )
        self.followup_input.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        # Bind mouse wheel to only scroll this widget (prevent propagation to parent)
        def on_mousewheel(event):
            if event.delta:
                self.followup_input.yview_scroll(int(-1 * (event.delta / 120)), "units")
            return "break"
        
        def on_mousewheel_linux(event):
            if event.num == 4:
                self.followup_input.yview_scroll(-3, "units")
            elif event.num == 5:
                self.followup_input.yview_scroll(3, "units")
            return "break"
        
        self.followup_input.bind("<MouseWheel>", on_mousewheel)
        self.followup_input.bind("<Button-4>", on_mousewheel_linux)
        self.followup_input.bind("<Button-5>", on_mousewheel_linux)
        
        # Add help to followup input
        if HELP_TEXTS:
            add_help(self.followup_input, **HELP_TEXTS.get("thread_followup_input", 
                {"title": "Follow-up Question", "description": "Type your follow-up question here"}))
        
        # Hint label
        ttk.Label(
            followup_frame,
            text="💡 Ctrl+Enter to submit | Ctrl+Z undo | Ctrl+Y redo | Ctrl+H find/replace",
            font=('Arial', 8),
            foreground='gray'
        ).pack(anchor=tk.W, pady=(5, 0))
    
    def _create_button_bar(self):
        """Create the button bar at the bottom - split into two rows for clarity"""
        # Uniform button width for ALL buttons
        BTN_WIDTH = 15  # Standard width for all buttons
        BTN_PAD = 2  # Horizontal padding between buttons
        
        # Container for both rows - padx=10 to match followup_frame alignment
        # Extra bottom padding to ensure buttons don't get cut off
        button_container = ttk.Frame(self.window, padding=(10, 5, 10, 15))
        button_container.pack(fill=tk.X, padx=10, pady=(0, 5))
        
        # === ROW 1: Mode and View Controls ===
        self.row1 = ttk.Frame(button_container)
        self.row1.pack(fill=tk.X, pady=(0, 5))
        
        # Left side frame for mode buttons (using pack for auto-alignment)
        self.row1_left = ttk.Frame(self.row1)
        self.row1_left.pack(side=tk.LEFT)
        
        # Mode toggle button (always visible)
        initial_mode_text = "View Source" if self.current_mode == 'conversation' else "View Thread"
        self.mode_toggle_btn = ttk.Button(
            self.row1_left,
            text=initial_mode_text,
            command=self._toggle_mode,
            width=BTN_WIDTH
        )
        self.mode_toggle_btn.pack(side=tk.LEFT, padx=(0, BTN_PAD))
        if HELP_TEXTS:
            add_help(self.mode_toggle_btn, **HELP_TEXTS.get("unified_viewer_mode_toggle",
                {"title": "Toggle View Mode", 
                 "description": "Switch between viewing the source document and the conversation thread"}))
        
        # Source visibility toggle (only in conversation mode)
        num_sources = len(self.source_documents)
        source_btn_text = "Hide Source" if self.source_section_visible else "Show Source"
        if num_sources > 1:
            source_btn_text = "Hide Sources" if self.source_section_visible else "Show Sources"
        
        self.source_toggle_btn = ttk.Button(
            self.row1_left,
            text=source_btn_text,
            command=self.toggle_source_visibility,
            width=BTN_WIDTH
        )
        if self.current_mode == 'conversation' and self.source_documents:
            self.source_toggle_btn.pack(side=tk.LEFT, padx=BTN_PAD)
        if HELP_TEXTS:
            add_help(self.source_toggle_btn, **HELP_TEXTS.get("unified_viewer_source_toggle",
                {"title": "Show/Hide Source(s)", 
                 "description": "Toggle visibility of source document section(s) in conversation view"}))
        
        # Expand/Collapse all exchanges button (only in conversation mode with exchanges)
        self.exchanges_toggle_btn = ttk.Button(
            self.row1_left,
            text="Expand All",
            command=self._toggle_all_exchanges,
            width=BTN_WIDTH
        )
        if self.current_mode == 'conversation' and self._count_exchanges() > 0:
            self.exchanges_toggle_btn.pack(side=tk.LEFT, padx=BTN_PAD)
        if HELP_TEXTS:
            add_help(self.exchanges_toggle_btn, **HELP_TEXTS.get("unified_viewer_exchanges_toggle",
                {"title": "Expand/Collapse All", 
                 "description": "Toggle visibility of all conversation exchanges and source sections at once"}))
        
        # Status label
        self.status_label = ttk.Label(
            self.row1_left,
            text="",
            font=('Arial', 9),
            foreground='green'
        )
        self.status_label.pack(side=tk.LEFT, padx=(10, 0))
        
        # Close button on RIGHT side (aligned with Submit button above)
        self.close_btn = ttk.Button(
            self.row1, 
            text="Close", 
            command=self._close_window, 
            width=10  # Same width as Submit button
        )
        self.close_btn.pack(side=tk.RIGHT, padx=(10, 0))
        if HELP_TEXTS:
            add_help(self.close_btn, **HELP_TEXTS.get("thread_close_button", 
                {"title": "Close", "description": "Close this window"}))
        
        # === ROW 2: Content Actions (using pack for auto-alignment) ===
        self.row2 = ttk.Frame(button_container)
        self.row2.pack(fill=tk.X, anchor=tk.W)
        
        # Left side frame for action buttons
        self.row2_left = ttk.Frame(self.row2)
        self.row2_left.pack(side=tk.LEFT)
        
        # Copy button
        self.copy_btn = ttk.Button(
            self.row2_left, 
            text="Copy",
            command=self._show_copy_dialog,
            width=BTN_WIDTH
        )
        self.copy_btn.pack(side=tk.LEFT, padx=(0, BTN_PAD))
        if HELP_TEXTS:
            add_help(self.copy_btn, **HELP_TEXTS.get("thread_copy_button", 
                {"title": "Copy", "description": "Copy content to clipboard"}))
        
        # Save As button
        self.save_as_btn = ttk.Button(
            self.row2_left,
            text="Save As",
            command=self._show_save_as_dialog,
            width=BTN_WIDTH
        )
        self.save_as_btn.pack(side=tk.LEFT, padx=BTN_PAD)
        if HELP_TEXTS:
            add_help(self.save_as_btn, **HELP_TEXTS.get("thread_save_disk_menu", 
                {"title": "Save to Disk", "description": "Save thread to a file"}))
        
        # Delete Exchange button (only show in conversation mode with exchanges)
        self.delete_exchange_btn = ttk.Button(
            self.row2_left,
            text="Delete Exchange",
            command=self._delete_current_exchange,
            width=BTN_WIDTH
        )
        if self.current_mode == 'conversation' and self._count_exchanges() > 0:
            self.delete_exchange_btn.pack(side=tk.LEFT, padx=BTN_PAD)
        if HELP_TEXTS:
            add_help(self.delete_exchange_btn, **HELP_TEXTS.get("thread_delete_exchange_button",
                {"title": "Delete Exchange", 
                 "description": "Delete the exchange (question + answer) where your cursor is positioned. Click inside an exchange first, then click this button."}))
        
        # New Branch button (only show for response/product documents with a source)
        self.new_branch_btn = None
        if self._can_start_new_conversation():
            self.new_branch_btn = ttk.Button(
                self.row2_left,
                text="New Branch",
                command=self._start_new_conversation,
                width=BTN_WIDTH
            )
            self.new_branch_btn.pack(side=tk.LEFT, padx=BTN_PAD)
            if HELP_TEXTS:
                add_help(self.new_branch_btn, **HELP_TEXTS.get("thread_new_branch_button",
                    {"title": "New Branch", 
                     "description": "Start a new conversation branch using the same source document. The new branch appears in the Branch dropdown at the top."}))
    
    def _set_status(self, message: str, duration_ms: int = 3000):
        """
        Set a temporary status message that fades after duration.
        
        Args:
            message: The status message to display
            duration_ms: How long to show the message (default 3 seconds)
        """
        self.status_label.config(text=message)
        # Clear the message after duration
        if duration_ms > 0:
            self.window.after(duration_ms, lambda: self.status_label.config(text=""))
    
    def _submit_followup(self):
        """Submit a follow-up question or initial prompt (with chunking if needed)"""
        question = self.followup_input.get('1.0', tk.END).strip()
        
        if not question:
            messagebox.showerror("Error", "Please enter a question.")
            return
        
        # LM Studio doesn't require an API key
        is_local_ai = self.provider_var.get() == "LM Studio (Local)"
        if not is_local_ai and (not self.model_var.get() or not self.api_key_var.get()):
            messagebox.showerror("Error", "Please select a model and enter an API key.")
            return
        
        if self.is_processing:
            messagebox.showinfo("Please Wait", "Already processing a question. Please wait.")
            return
        
        # ============================================================
        # CHECK IF WE SHOULD SHOW BRANCH PICKER
        # Show branch picker if:
        # 1. Current document IS a source document, OR
        # 2. We're in "source" mode viewing a Response's parent source
        # ============================================================
        from document_library import is_source_document, get_response_branches_for_source, get_document_by_id
        
        # DEBUG: Print document info
        
        doc = get_document_by_id(self.current_document_id) if self.current_document_id else None
        
        # Determine if we should show the branch picker
        should_show_branch_picker = False
        source_doc_id_for_branches = None
        
        if doc:
            is_src = is_source_document(self.current_document_id)
            
            if is_src:
                # Case 1: Directly viewing a source document
                should_show_branch_picker = True
                source_doc_id_for_branches = self.current_document_id
            else:
                # Case 2 & 3: Viewing a Response document (in ANY mode)
                # Check if this Response has a parent source document
                metadata = doc.get('metadata', {})
                parent_id = metadata.get('parent_document_id') or metadata.get('original_document_id')
                if parent_id:
                    should_show_branch_picker = True
                    source_doc_id_for_branches = parent_id
                else:
                    pass
        else:
            pass
        
        if should_show_branch_picker and source_doc_id_for_branches:
            # Show branch picker for source document
            self._handle_source_document_followup(question, source_doc_id_for_branches)
            return
        else:
            pass
        
        # ============================================================
        # DETECT INITIAL PROMPT SCENARIO
        # If in Source Mode with no conversation AND we have chunking callback,
        # use chunking to process the initial prompt properly
        # ============================================================
        is_initial_prompt = (
            self.current_mode == 'source' and 
            (not self.current_thread or len(self.current_thread) == 0) and
            self.process_with_chunking is not None and
            self.current_entries is not None and
            len(self.current_entries) > 0
        )
        
        if is_initial_prompt:
            # Use chunking callback for initial prompt
            self._submit_initial_prompt_with_chunking(question)
            return
        
        # ============================================================
        # REGULAR FOLLOW-UP FLOW (existing code)
        # ============================================================
        # Check we have the required callbacks
        if not self.get_ai_handler or not self.build_threaded_messages or not self.add_message_to_thread:
            messagebox.showerror(
                "Error", 
                "Follow-up from thread viewer requires AI handler callbacks.\n"
                "Please use the Follow-up button from the main window."
            )
            return
        
        # Start processing
        self.is_processing = True
        self.submit_btn.config(state=tk.DISABLED, text="⏳ Processing...")
        self.followup_input.config(state=tk.DISABLED)
        
        # Show processing indicator in thread display
        self.thread_text.insert(tk.END, "\n" + "─" * 80 + "\n", "divider")
        self.thread_text.insert(tk.END, f"\n🧑 YOU [Now]\n", "user")
        self.thread_text.insert(tk.END, f"{question}\n", "")
        self.thread_text.insert(tk.END, "\n🤖 AI [Processing...]\n", "processing")
        self.thread_text.see(tk.END)
        
        # Clear input
        self.followup_input.config(state=tk.NORMAL)
        self.followup_input.delete('1.0', tk.END)
        self.followup_input.config(state=tk.DISABLED)
        
        # Process in background thread
        process_thread = threading.Thread(
            target=self._process_followup_thread,
            args=(question,),
            daemon=True
        )
        process_thread.start()
    
    def _handle_source_document_followup(self, question: str, source_doc_id: str = None):
        """
        Handle follow-up when viewing a source document or response document.
        Shows branch picker dialog and redirects to appropriate response document.
        
        Supports multi-select: user can choose to save to multiple branches.
        After processing, shows navigation dialog to choose where to go.
        
        Args:
            question: The question to process
            source_doc_id: ID of the source document (defaults to self.current_document_id)
        """
        
        from document_library import get_response_branches_for_source, get_document_by_id, is_source_document
        
        # Use provided source_doc_id or fall back to current document
        if source_doc_id is None:
            source_doc_id = self.current_document_id
        
        
        try:
            from branch_picker_dialog import show_branch_picker
        except ImportError as e:
            messagebox.showerror("Import Error", f"Could not import branch_picker_dialog:\n{e}")
            return
        
        # Get source document info
        source_doc = get_document_by_id(source_doc_id)
        if not source_doc:
            messagebox.showerror("Error", "Could not find source document.")
            return
        
        source_title = source_doc.get('title', 'Unknown')
        
        # Get existing response branches
        branches = get_response_branches_for_source(source_doc_id)
        
        # Determine current branch ID (for pre-selection)
        # If we're in a response document, use that as current branch
        current_branch_id = None
        if self.current_document_id and not is_source_document(self.current_document_id):
            current_branch_id = self.current_document_id
        
        # Show branch picker dialog
        result = show_branch_picker(
            parent=self.window,
            source_document_id=source_doc_id,
            source_title=source_title,
            existing_branches=branches,
            current_branch_id=current_branch_id,
            current_mode=self.current_mode,
            action_description="save this response"
        )
        
        if result.get('cancelled', True):
            # User cancelled
            return
        
        selected_existing = result.get('existing_branches', [])
        new_branch_names = result.get('new_branches', [])
        stay_in_current_view = result.get('stay_in_current_view', False)
        
        
        # Store info for handling after AI processing
        self._pending_multi_save = {
            'question': question,
            'source_doc_id': source_doc_id,
            'source_title': source_title,
            'existing_branches': selected_existing,
            'new_branch_names': new_branch_names,
            'stay_in_current_view': stay_in_current_view,
            'original_mode': self.current_mode,
            'original_document_id': self.current_document_id,
            'primary_branch_id': None  # Will be set below
        }
        
        # Determine primary branch (where we'll process the question)
        if new_branch_names:
            # Create a new branch as primary
            branch_name = new_branch_names[0]
            self._create_new_branch_and_process(question, branch_name, source_doc_id)
        elif selected_existing:
            # Use first selected existing branch as primary
            primary_id = selected_existing[0]
            self._pending_multi_save['primary_branch_id'] = primary_id
            self._switch_to_branch_and_process(primary_id, question)
        else:
            # No selection (shouldn't happen due to dialog validation)
            messagebox.showwarning("No Selection", "Please select a branch or create a new one.")
            return
    
    def _submit_initial_prompt_with_chunking(self, prompt: str):
        """
        Process an initial prompt using chunking (for large documents in Source Mode).
        This delegates to the main app's chunking logic via callback.
        """
        # Start processing
        self.is_processing = True
        self.submit_btn.config(state=tk.DISABLED, text="⏳ Processing...")
        self.followup_input.config(state=tk.DISABLED)
        
        # Show processing indicator - but in Source Mode format
        self.thread_text.config(state=tk.NORMAL)
        self.thread_text.insert(tk.END, "\n\n" + "═" * 60 + "\n", "divider")
        self.thread_text.insert(tk.END, "⚙️ Processing with AI (chunking enabled for large documents)...\n", "processing")
        self.thread_text.see(tk.END)
        
        # Clear input
        self.followup_input.config(state=tk.NORMAL)
        self.followup_input.delete('1.0', tk.END)
        self.followup_input.config(state=tk.DISABLED)
        
        # Status callback - update the viewer's display
        # Track the processing status line position
        self._processing_line_start = None
        
        def status_callback(status: str):
            def update():
                try:
                    self.thread_text.config(state=tk.NORMAL)
                    
                    # More robust: Search for and replace the processing status line
                    content = self.thread_text.get("1.0", tk.END)
                    lines = content.split("\n")
                    
                    # Find the line starting with ⚙️ or ⏳
                    for i, line in enumerate(lines):
                        if line.startswith("⚙️") or line.startswith("⏳"):
                            # Calculate position: line number is i+1 (1-indexed)
                            line_start = f"{i+1}.0"
                            line_end = f"{i+1}.end"
                            self.thread_text.delete(line_start, line_end)
                            self.thread_text.insert(line_start, status, "processing")
                            break
                    else:
                        # If no processing line found, just append
                        self.thread_text.insert(tk.END, f"{status}\n", "processing")
                    
                    self.thread_text.see(tk.END)
                    self.thread_text.config(state=tk.DISABLED)
                except Exception as e:
                    import traceback
                    traceback.print_exc()
            self.window.after(0, update)
        
        # Completion callback - with debug output
        def complete_callback(success: bool, result: str):
            try:
                # Check if window still exists
                if self.window and self.window.winfo_exists():
                    self.window.after(0, lambda: self._safe_handle_initial_prompt_result(prompt, success, result))
                else:
                    pass
            except Exception as e:
                import traceback
                traceback.print_exc()
        
        # Call the chunking processor (runs in background thread)
        self.process_with_chunking(prompt, status_callback, complete_callback)
    
    def _safe_handle_initial_prompt_result(self, prompt: str, success: bool, result: str):
        """Wrapper to catch and report errors from _handle_initial_prompt_result"""
        try:
            self._handle_initial_prompt_result(prompt, success, result)
        except Exception as e:
            import traceback
            traceback.print_exc()
            # Try to show error to user
            try:
                messagebox.showerror("Processing Error", f"Error updating display:\n{e}")
            except:
                pass
    
    def _handle_initial_prompt_result(self, prompt: str, success: bool, result: str):
        """Handle result from initial prompt processing with chunking"""
        
        self.is_processing = False
        self.submit_btn.config(state=tk.NORMAL, text="Submit")
        self.followup_input.config(state=tk.NORMAL)
        
        if success:
            # Show completion message in the viewer
            try:
                self.thread_text.config(state=tk.NORMAL)
                
                # More robust: Find and replace the processing line with completion message
                content = self.thread_text.get("1.0", tk.END)
                lines = content.split("\n")
                
                replaced = False
                for i, line in enumerate(lines):
                    if line.startswith("⚙️") or line.startswith("⏳"):
                        line_start = f"{i+1}.0"
                        line_end = f"{i+1}.end"
                        self.thread_text.delete(line_start, line_end)
                        self.thread_text.insert(line_start, "✅ Processing complete!", "success")
                        replaced = True
                        break
                
                if not replaced:
                    # If no processing line found, just append
                    self.thread_text.insert(tk.END, "\n✅ Processing complete!\n", "success")
                
                self.thread_text.see(tk.END)
                self.thread_text.config(state=tk.DISABLED)
            except Exception as e:
                import traceback
                traceback.print_exc()
            
            # Notify main app first (this may update the thread)
            if self.on_followup_complete:
                self.on_followup_complete(prompt, result)
            
            # === CRITICAL: Reload thread from document ===
            # After creating a new branch, self.current_thread was cleared.
            # We need to reload from the document to get the saved conversation.
            try:
                from document_library import load_thread_from_document
                # FIXED: Properly unpack the tuple (thread_list, metadata)
                reloaded_thread, thread_metadata = load_thread_from_document(self.current_document_id)
                if reloaded_thread and isinstance(reloaded_thread, list):
                    # Update the shared thread reference
                    self.current_thread.clear()
                    self.current_thread.extend(reloaded_thread)
                    self.thread_message_count = len([m for m in reloaded_thread if isinstance(m, dict) and m.get('role') == 'user'])
                else:
                    pass
            except Exception as e:
                import traceback
                traceback.print_exc()
            
            # === MULTI-SAVE: Copy exchange to additional branches ===
            if hasattr(self, '_pending_multi_save') and self._pending_multi_save:
                self._copy_exchange_to_additional_branches(prompt, result)
            
            # === HANDLE NAVIGATION PREFERENCE ===
            # Check if user wanted to stay in current view (e.g., keep reading source)
            stay_in_current = False
            original_mode = 'conversation'  # Default if no preference stored
            original_doc_id = self.current_document_id
            
            if hasattr(self, '_pending_multi_save') and self._pending_multi_save:
                stay_in_current = self._pending_multi_save.get('stay_in_current_view', False)
                original_mode = self._pending_multi_save.get('original_mode', 'source')
                original_doc_id = self._pending_multi_save.get('original_document_id', self.current_document_id)
            
            if stay_in_current:
                # User wants to stay in current view (e.g., continue reading source)
                # RESTORE the original mode (branch creation may have changed it)
                self.current_mode = original_mode
                
                # If they were viewing the source document, switch back to it
                if original_mode == 'source' and original_doc_id:
                    source_doc_id = self._pending_multi_save.get('source_doc_id')
                    if source_doc_id:
                        # Load the source document back into view
                        from document_library import get_document_by_id
                        source_doc = get_document_by_id(source_doc_id)
                        if source_doc:
                            # Restore source document context
                            self.current_document_id = source_doc_id
                            # Update window title to reflect source
                            source_title = source_doc.get('title', 'Source Document')
                            self.window.title(f"Source Document - {source_title[:50]}")
                
                # Refresh the display in the correct mode
                self.window.after(500, lambda: self._refresh_thread_display())
                # Show a brief notification that the response was saved
                self.window.after(600, lambda: self._show_saved_notification())
            else:
                # User wants to GO TO the conversation to see the response
                # Set mode to conversation (may already be set, but ensure it)
                self.current_mode = 'conversation'
                
                # CRITICAL: Always refresh the display, even if mode was already 'conversation'
                # This is needed because _create_new_branch_and_process sets the mode but 
                # doesn't refresh the display (it's busy processing the AI request)
                def refresh_to_conversation():
                    self._update_window_title()
                    self._update_heading()
                    self._refresh_thread_display()
                    self._update_mode_buttons()
                    self._update_branch_selector_visibility()
                    # Notify main app of mode change
                    if self.on_mode_change:
                        self.on_mode_change('conversation')
                
                self.window.after(500, refresh_to_conversation)
            
            # Refresh branch selector (new branch may have been created)
            self.window.after(600, lambda: self._delayed_populate_branch_selector("POST-PROCESSING"))
            
            # Clear the pending multi-save info
            if hasattr(self, '_pending_multi_save'):
                self._pending_multi_save = None
            
        else:
            # Show error in the viewer
            try:
                self.thread_text.config(state=tk.NORMAL)
                self.thread_text.delete("end-3l", "end-1l")
                self.thread_text.insert(tk.END, "❌ Processing failed - see error message\n", "error")
                self.thread_text.see(tk.END)
                self.thread_text.config(state=tk.DISABLED)
            except:
                pass
            
            # Show error dialog
            messagebox.showerror(
                "AI Processing Error",
                f"Failed to process prompt:\n\n{result}"
            )
            
            # Refresh the display to remove processing indicator
            self._refresh_thread_display()
    
    def _delayed_populate_branch_selector(self, caller: str = "unknown"):
        """Wrapper to tag delayed branch selector calls for debugging"""
        self._populate_branch_selector()
    
    def _do_switch_mode_with_debug(self, mode: str):
        """Debug wrapper for switch_mode"""
        try:
            self.switch_mode(mode)
        except Exception as e:
            import traceback
            traceback.print_exc()
    
    def _show_saved_notification(self):
        """
        Show a brief, non-intrusive notification that the response was saved.
        Used when user chooses to stay in current view after asking a question.
        """
        try:
            # Create a temporary notification label
            notification = tk.Label(
                self.window,
                text="✅ Response saved to conversation",
                font=('Arial', 10, 'bold'),
                fg='white',
                bg='#27AE60',  # Green background
                padx=15,
                pady=8
            )
            
            # Position at top of window
            notification.place(relx=0.5, y=10, anchor='n')
            
            # Auto-hide after 2.5 seconds
            def hide_notification():
                try:
                    notification.destroy()
                except:
                    pass
            
            self.window.after(2500, hide_notification)
            
        except Exception as e:
            pass
    
    def _copy_exchange_to_additional_branches(self, question: str, ai_response: str):
        """
        Copy the exchange (question + AI response) to additional selected branches.
        This implements the multi-save feature where one response can be saved to multiple conversations.
        
        Args:
            question: The user's question
            ai_response: The AI's response
        """
        if not hasattr(self, '_pending_multi_save') or not self._pending_multi_save:
            return
        
        
        from document_library import (
            get_document_by_id, save_thread_to_document, load_thread_from_document,
            add_document_to_library
        )
        import datetime
        
        pending = self._pending_multi_save
        primary_branch_id = self.current_document_id  # The branch we just processed to
        selected_existing = pending.get('existing_branches', [])
        new_branch_names = pending.get('new_branches', [])
        source_doc_id = pending.get('source_doc_id')
        source_title = pending.get('source_title', 'Unknown')
        
        
        # Build the exchange to copy
        exchange = [
            {'role': 'user', 'content': question},
            {'role': 'assistant', 'content': ai_response}
        ]
        
        # Copy to additional EXISTING branches (skip the primary)
        for branch_id in selected_existing:
            if branch_id == primary_branch_id:
                continue
            
            try:
                
                # Load existing thread
                existing_thread, metadata = load_thread_from_document(branch_id)
                if existing_thread is None:
                    existing_thread = []
                
                # Append the new exchange
                existing_thread.extend(exchange)
                
                # Save back
                save_thread_to_document(branch_id, existing_thread)
                
            except Exception as e:
                pass
        
        # Create additional NEW branches (skip the first one if it was the primary)
        # The first new branch was created as primary, so we need to create any additional ones
        branches_created_as_primary = 1 if new_branch_names else 0
        additional_new_branches = new_branch_names[branches_created_as_primary:] if len(new_branch_names) > branches_created_as_primary else []
        
        for branch_name in additional_new_branches:
            try:
                
                # Generate name if not provided
                if not branch_name:
                    # Auto-generate from question
                    branch_name = question[:30].strip()
                    if len(question) > 30:
                        branch_name += "..."
                
                response_title = f"[Response] {branch_name}"
                
                # Create metadata for new branch
                response_metadata = {
                    'original_document_id': source_doc_id,
                    'parent_document_id': source_doc_id,
                    'source_title': source_title,
                    'branch_name': branch_name,
                    'created': datetime.datetime.now().isoformat(),
                    'editable': True,
                    'pre_created': False  # Already has content
                }
                
                # Get source URL for the source field
                source_doc = get_document_by_id(source_doc_id)
                source_url = source_doc.get('source', '') if source_doc else ''
                
                # Create the document
                new_doc_id = add_document_to_library(
                    doc_type="conversation_thread",
                    source=f"Conversation about: {source_url or source_title} - Branch: {branch_name}",
                    title=response_title,
                    entries=[{"text": f"Conversation branch: {branch_name}", "location": "Header"}],
                    metadata=response_metadata,
                    document_class="response"
                )
                
                if new_doc_id:
                    # Save the exchange to it
                    save_thread_to_document(new_doc_id, exchange)
                else:
                    pass
                    
            except Exception as e:
                pass
        
    
    def _create_new_branch_and_process(self, question: str, branch_name: str = None, source_doc_id: str = None):
        """
        Create a new response branch and process the initial prompt.
        
        This creates the response document FIRST with the user's chosen name,
        then switches context and processes the question.
        
        Args:
            question: The prompt/question to process
            branch_name: Optional name for the new branch (used for title)
            source_doc_id: ID of the source document to link to (defaults to current_document_id)
        """
        import datetime
        from document_library import (
            add_document_to_library, get_document_by_id, load_document_entries,
            save_thread_to_document
        )
        
        # Use provided source_doc_id or fall back to current document
        if source_doc_id is None:
            source_doc_id = self.current_document_id
        
        # Get source document info
        source_doc = get_document_by_id(source_doc_id)
        if not source_doc:
            messagebox.showerror("Error", "Could not find source document.")
            return
        
        source_title = source_doc.get('title', 'Unknown')
        source_url = source_doc.get('source', '')
        source_id = source_doc_id  # Use the provided/determined source ID
        
        # Generate branch name if not provided
        if not branch_name:
            # Use first part of the question as branch name
            branch_name = question[:50] + "..." if len(question) > 50 else question
            branch_name = branch_name.replace('\n', ' ')
        
        # Create title for the response document
        response_title = f"[Response] {branch_name}"
        
        # Create metadata linking back to source
        response_metadata = {
            "original_document_id": source_id,
            "parent_document_id": source_id,
            "source_title": source_title,
            "branch_name": branch_name,
            "created": datetime.datetime.now().isoformat(),
            "editable": True,
            "pre_created": True  # Flag so Main.py knows not to create another
        }
        
        # Create the response document
        # CRITICAL: Include branch_name in source to generate UNIQUE doc_id for each branch
        # Without this, all branches for the same source get the same ID and overwrite each other!
        new_doc_id = add_document_to_library(
            doc_type="conversation_thread",
            source=f"Conversation about: {source_url or source_title} - Branch: {branch_name}",
            title=response_title,
            entries=[{"text": f"Conversation branch: {branch_name}", "location": "Header"}],
            metadata=response_metadata,
            document_class="response"
        )
        
        if not new_doc_id:
            messagebox.showerror("Error", "Failed to create response document.")
            return
        
        
        # VERIFY the new document's metadata
        try:
            verify_new = get_document_by_id(new_doc_id)
            if verify_new:
                new_meta = verify_new.get('metadata', {})
            else:
                pass
        except Exception as ve:
            pass
        
        # =================================================================
        # CRITICAL FIX: Save current thread to its document BEFORE clearing
        # This prevents data loss when creating a new branch while another
        # branch has an unsaved conversation
        # =================================================================
        if self.current_document_id and self.current_document_id != new_doc_id:
            if self.current_thread and len(self.current_thread) > 0:
                try:
                    save_metadata = {
                        "model": self.model_var.get(),
                        "provider": self.provider_var.get(),
                        "last_updated": datetime.datetime.now().isoformat(),
                        "message_count": self.thread_message_count
                    }
                    save_thread_to_document(self.current_document_id, self.current_thread, save_metadata)
                    
                    # VERIFY the save worked
                    try:
                        from document_library import get_document_by_id
                        verify_doc = get_document_by_id(self.current_document_id)
                        if verify_doc:
                            saved_thread = verify_doc.get('conversation_thread', [])
                            saved_exchanges = len([m for m in saved_thread if m.get('role') == 'user'])
                        else:
                            pass
                    except Exception as ve:
                        pass
                except Exception as e:
                    pass
            else:
                pass
        else:
            pass
        
        # === CRITICAL: Update Thread Viewer's context ===
        self.current_document_id = new_doc_id
        self.current_thread.clear()
        self.thread_message_count = 0
        
        # Update window title
        self.window.title(f"Conversation - {branch_name}")
        self.doc_title = response_title
        
        # Switch to conversation mode
        self.current_mode = 'conversation'
        
        # Refresh the branch selector to show the new branch immediately
        self._populate_branch_selector()
        
        # === CRITICAL: Update Main App's context via callback ===
        # This ensures the main app also knows we switched documents
        if hasattr(self, 'app') and self.app:
            self.app.current_document_id = new_doc_id
            self.app.current_document_class = 'response'
            # Also update metadata so save function sees it
            if hasattr(self.app, 'current_document_metadata'):
                self.app.current_document_metadata = response_metadata.copy()
            else:
                # Create the attribute if it doesn't exist
                self.app.current_document_metadata = response_metadata.copy()
        
        # Make sure we have the source document's entries for processing
        if not self.current_entries or len(self.current_entries) == 0:
            source_entries = load_document_entries(source_id)
            if source_entries:
                self.current_entries = source_entries
        
        # Now process the question
        # Use chunking for large documents, regular follow-up for small ones
        if self.process_with_chunking and self.current_entries and len(self.current_entries) > 0:
            self._submit_initial_prompt_with_chunking(question)
        else:
            self._submit_followup_direct(question)
    
    def _switch_to_branch_and_process(self, branch_doc_id: str, question: str):
        """
        Switch to an existing response branch and add the follow-up.
        
        Args:
            branch_doc_id: Document ID of the response branch
            question: The prompt/question to process
        """
        from document_library import get_document_by_id, load_thread_from_document, load_document_entries, save_thread_to_document
        import datetime
        
        # Get the branch document
        branch_doc = get_document_by_id(branch_doc_id)
        if not branch_doc:
            messagebox.showerror("Error", "Could not find the selected conversation branch.")
            return
        
        # =================================================================
        # CRITICAL FIX: Check if we're already on this branch
        # If so, DON'T reload from document - keep the in-memory thread
        # (which may have unsaved exchanges)
        # =================================================================
        already_on_this_branch = (self.current_document_id == branch_doc_id)
        
        if already_on_this_branch:
            # Already on this branch - just process the question with current thread
            print(f"📌 Already on branch {branch_doc_id}, keeping in-memory thread ({len(self.current_thread)} messages)")
            
            # Make sure we're in conversation mode
            self.current_mode = 'conversation'
            
            # Refresh the branch selector
            self._populate_branch_selector()
            
            # Process the question directly
            self._submit_followup_direct(question)
            return
        
        # =================================================================
        # SWITCHING TO A DIFFERENT BRANCH
        # Save current thread first, then load the target branch's thread
        # =================================================================
        print(f"🔀 Switching from branch {self.current_document_id} to {branch_doc_id}")
        
        # Save current thread to its document BEFORE switching
        if self.current_document_id and self.current_thread and len(self.current_thread) > 0:
            try:
                save_metadata = {
                    "model": self.model_var.get(),
                    "provider": self.provider_var.get(),
                    "last_updated": datetime.datetime.now().isoformat(),
                    "message_count": self.thread_message_count
                }
                save_thread_to_document(self.current_document_id, self.current_thread, save_metadata)
                print(f"💾 Saved {self.thread_message_count} exchanges to old branch {self.current_document_id}")
            except Exception as e:
                print(f"⚠️ Failed to save old thread: {e}")
        
        # Load existing thread from the target branch
        existing_thread, thread_metadata = load_thread_from_document(branch_doc_id)
        
        # Update our local context to point to the new branch
        self.current_document_id = branch_doc_id
        
        # Load/update thread from the target branch
        if existing_thread:
            self.current_thread.clear()
            self.current_thread.extend(existing_thread)
            self.thread_message_count = len([m for m in existing_thread if m.get('role') == 'user'])
            print(f"📂 Loaded {self.thread_message_count} exchanges from new branch")
        else:
            self.current_thread.clear()
            self.thread_message_count = 0
            print(f"📂 New branch has no existing thread")
        
        # Get the new document info
        self.doc_title = branch_doc.get('title', 'Unknown')
        self.window.title(f"Conversation - {branch_doc.get('source', '')[:60]}")
        
        # Switch to conversation mode since we're adding to a conversation
        self.current_mode = 'conversation'
        
        # Refresh the branch selector to show the new branch immediately
        self._populate_branch_selector()
        
        # Now submit the follow-up using the regular flow
        self._submit_followup_direct(question)
    
    def _submit_followup_direct(self, question: str):
        """
        Submit a follow-up question directly using the standard flow.
        Used when we've already set up the context (switched documents, etc.)
        
        Args:
            question: The question to submit
        """
        # Check we have the required callbacks
        if not self.get_ai_handler or not self.build_threaded_messages or not self.add_message_to_thread:
            messagebox.showerror(
                "Error", 
                "Follow-up requires AI handler callbacks.\n"
                "Please try from the main window."
            )
            return
        
        # Start processing
        self.is_processing = True
        self.submit_btn.config(state=tk.DISABLED, text="⏳ Processing...")
        self.followup_input.config(state=tk.DISABLED)
        
        # Refresh display to show current state
        self._refresh_thread_display()
        
        # Show processing indicator
        self.thread_text.config(state=tk.NORMAL)
        self.thread_text.insert(tk.END, "\n" + "─" * 80 + "\n", "divider")
        self.thread_text.insert(tk.END, f"\n🧑 YOU [Now]\n", "user")
        self.thread_text.insert(tk.END, f"{question}\n", "")
        self.thread_text.insert(tk.END, "\n🤖 AI [Processing...]\n", "processing")
        self.thread_text.see(tk.END)
        
        # Clear input (it should already be clear, but just in case)
        self.followup_input.config(state=tk.NORMAL)
        self.followup_input.delete('1.0', tk.END)
        self.followup_input.config(state=tk.DISABLED)
        
        # Process in background thread
        process_thread = threading.Thread(
            target=self._process_followup_thread,
            args=(question,),
            daemon=True
        )
        process_thread.start()
    
    def _process_followup_thread(self, question: str):
        """Process follow-up question in background thread"""
        try:
            # Build threaded messages
            messages = self.build_threaded_messages(question)
            
            # Call AI provider
            ai_handler = self.get_ai_handler()
            success, result = ai_handler.call_ai_provider(
                provider=self.provider_var.get(),
                model=self.model_var.get(),
                messages=messages,
                api_key=self.api_key_var.get()
            )
            
            # Update on main thread
            self.window.after(0, self._handle_followup_result, question, success, result)
            
        except Exception as e:
            self.window.after(0, self._handle_followup_result, question, False, str(e))
    
    def _handle_followup_result(self, question: str, success: bool, result: str):
        """Handle follow-up result on main thread"""
        self.is_processing = False
        self.submit_btn.config(state=tk.NORMAL, text="Submit")
        self.followup_input.config(state=tk.NORMAL)
        
        if success:
            # Add to thread via callback (this updates the main app's thread state)
            self.add_message_to_thread("user", question)
            self.add_message_to_thread("assistant", result)
            
            # Update our local reference and count
            self.thread_message_count += 1
            
            # =================================================================
            # CRITICAL FIX: Save thread to document IMMEDIATELY after adding
            # This prevents data loss if branch picker dialog appears again
            # =================================================================
            if self.current_document_id:
                try:
                    from document_library import save_thread_to_document
                    import datetime
                    save_metadata = {
                        "model": self.model_var.get(),
                        "provider": self.provider_var.get(),
                        "last_updated": datetime.datetime.now().isoformat(),
                        "message_count": self.thread_message_count
                    }
                    save_thread_to_document(self.current_document_id, self.current_thread, save_metadata)
                    print(f"💾 Auto-saved {self.thread_message_count} exchanges to document {self.current_document_id}")
                except Exception as e:
                    print(f"⚠️ Failed to auto-save thread: {e}")
            
            # Notify main app
            if self.on_followup_complete:
                self.on_followup_complete(question, result)
            
            # Refresh the display
            self._refresh_thread_display()
            
            # Switch to conversation mode to show the response (if not already there)
            if self.current_mode != 'conversation':
                self.window.after(500, lambda: self._do_switch_mode_with_debug('conversation'))
            
            # Refresh branch selector (in case context changed)
            self._populate_branch_selector()
            
        else:
            # Show error
            messagebox.showerror(
                "AI Processing Error",
                f"Failed to process follow-up question:\n\n{result}"
            )
            
            # Remove the "Processing..." indicator
            self._refresh_thread_display()
    
    def _save_edits_before_refresh(self):
        """
        Save any pending edits before refreshing the display.
        This prevents edits from being lost when the display is rebuilt.
        Called before _refresh_thread_display() in functions that might lose edits.
        """
        try:
            # Only save if text has been modified
            is_modified = self.thread_text.edit_modified()
            if not is_modified:
                return  # No edits to save
            
            if self.current_mode == 'conversation':
                # Check if we have conversation markers (i.e., displaying conversation, not source)
                widget_text = self.thread_text.get('1.0', 'end-1c')
                has_conversation = '🧑 YOU' in widget_text or '👤 YOU' in widget_text or '🤖 ' in widget_text
                
                if has_conversation and self.current_thread:
                    print("💾 Auto-saving edits before refresh...")
                    self._save_edits_to_thread()
            elif self.current_mode == 'source':
                self._save_source_edits()
        except Exception as e:
            print(f"⚠️ _save_edits_before_refresh error: {e}")
    
    def _save_edits_to_thread(self):
        """
        Parse the edited text from the display and save changes back to the thread.
        This allows users to fix names, places, and other AI mistakes.
        
        SMART MERGE: For expanded exchanges, parse edits from the widget.
        For collapsed exchanges, preserve original content from self.current_thread.
        This ensures edits are never lost, regardless of collapse state.
        """
        # Skip if there's no conversation to save
        if not self.current_thread or len(self.current_thread) == 0:
            return  # Nothing to save - no error needed
        
        # Skip if we're in source mode (source edits are saved separately)
        if self.current_mode == 'source':
            self._save_source_edits()
            return
        
        # Check if widget actually contains conversation markers
        try:
            widget_text = self.thread_text.get('1.0', 'end-1c')
            has_user_marker = '🧑 YOU' in widget_text or '👤 YOU' in widget_text
            has_ai_marker = '🤖 ' in widget_text
            
            if not has_user_marker and not has_ai_marker:
                # Widget doesn't contain conversation markers - just preserve original thread
                if self.current_document_id and self.current_thread:
                    from document_library import save_thread_to_document
                    metadata = {
                        "model": self.model_var.get(),
                        "provider": self.provider_var.get(),
                        "last_updated": datetime.datetime.now().isoformat(),
                        "message_count": self.thread_message_count
                    }
                    save_thread_to_document(self.current_document_id, self.current_thread, metadata)
                return
        except:
            pass  # If check fails, continue with normal flow
        
        # Check if modifications were made
        try:
            is_modified = self.thread_text.edit_modified()
            if not is_modified:
                # No edits - still save to ensure thread is persisted
                if self.current_document_id and self.current_thread:
                    from document_library import save_thread_to_document
                    metadata = {
                        "model": self.model_var.get(),
                        "provider": self.provider_var.get(),
                        "last_updated": datetime.datetime.now().isoformat(),
                        "message_count": self.thread_message_count
                    }
                    save_thread_to_document(self.current_document_id, self.current_thread, metadata)
                return
        except Exception as e:
            print(f"⚠️ edit_modified() check failed: {e} - continuing with save")
        
        # ================================================================
        # SMART MERGE: Parse expanded exchanges, preserve collapsed ones
        # ================================================================
        
        # Group original thread into exchanges for reference
        original_exchanges = self._group_messages_into_exchanges()
        num_exchanges = len(original_exchanges)
        
        # Parse the widget text to find exchange sections
        edited_text = self.thread_text.get('1.0', tk.END)
        lines = edited_text.split('\n')
        
        # Build a map of exchange_index -> parsed content (for expanded exchanges only)
        parsed_exchanges = {}  # {index: {'user': content, 'assistant': content}}
        
        current_exchange_idx = -1
        current_role = None
        current_content_lines = []
        exchange_is_expanded = {}  # Track which exchanges are expanded in the widget
        
        import re
        exchange_header_pattern = re.compile(r'^([▼▶])\s*Exchange\s+(\d+)')
        
        for line_idx, line in enumerate(lines):
            line_num = line_idx + 1
            
            # Check for exchange header (▼ Exchange 1 or ▶ Exchange 1)
            header_match = exchange_header_pattern.match(line.strip())
            if header_match:
                # Save previous content if any
                if current_exchange_idx >= 0 and current_role and current_content_lines:
                    content_text = self._reconstruct_markdown_content(current_content_lines)
                    if content_text:
                        if current_exchange_idx not in parsed_exchanges:
                            parsed_exchanges[current_exchange_idx] = {}
                        parsed_exchanges[current_exchange_idx][current_role] = content_text
                
                # Start new exchange
                indicator = header_match.group(1)
                exchange_num = int(header_match.group(2))
                current_exchange_idx = exchange_num - 1  # Convert to 0-based index
                exchange_is_expanded[current_exchange_idx] = (indicator == '▼')
                current_role = None
                current_content_lines = []
                continue
            
            # Check for role markers
            is_user_marker = (
                line.startswith('🧑 YOU') or 
                line.startswith('👤 YOU') or
                line.startswith('YOU [') or
                line.startswith('🧑‍💻 YOU')
            )
            is_ai_marker = (
                line.startswith('🤖 AI') or 
                (line.startswith('🤖 ') and '[' in line) or
                line.startswith('AI [') or
                line.startswith('🤖AI')
            )
            
            if is_user_marker:
                # Save previous content
                if current_exchange_idx >= 0 and current_role and current_content_lines:
                    content_text = self._reconstruct_markdown_content(current_content_lines)
                    if content_text:
                        if current_exchange_idx not in parsed_exchanges:
                            parsed_exchanges[current_exchange_idx] = {}
                        parsed_exchanges[current_exchange_idx][current_role] = content_text
                current_role = 'user'
                current_content_lines = []
            elif is_ai_marker:
                # Save previous content
                if current_exchange_idx >= 0 and current_role and current_content_lines:
                    content_text = self._reconstruct_markdown_content(current_content_lines)
                    if content_text:
                        if current_exchange_idx not in parsed_exchanges:
                            parsed_exchanges[current_exchange_idx] = {}
                        parsed_exchanges[current_exchange_idx][current_role] = content_text
                current_role = 'assistant'
                current_content_lines = []
            elif line.startswith('─' * 10):  # Divider line
                # Save previous content before divider
                if current_exchange_idx >= 0 and current_role and current_content_lines:
                    content_text = self._reconstruct_markdown_content(current_content_lines)
                    if content_text:
                        if current_exchange_idx not in parsed_exchanges:
                            parsed_exchanges[current_exchange_idx] = {}
                        parsed_exchanges[current_exchange_idx][current_role] = content_text
                current_role = None
                current_content_lines = []
            elif current_role and current_exchange_idx >= 0:
                # Skip timestamp lines and certain markers
                if not (line.strip().startswith('[') and line.strip().endswith(']') and len(line.strip()) <= 12):
                    if not (line.strip().startswith('▼') or line.strip().startswith('▶')):
                        current_content_lines.append((line_num, line))
        
        # Don't forget the last content
        if current_exchange_idx >= 0 and current_role and current_content_lines:
            content_text = self._reconstruct_markdown_content(current_content_lines)
            if content_text:
                if current_exchange_idx not in parsed_exchanges:
                    parsed_exchanges[current_exchange_idx] = {}
                parsed_exchanges[current_exchange_idx][current_role] = content_text
        
        # ================================================================
        # MERGE: Combine parsed (expanded) and original (collapsed) content
        # ================================================================
        new_thread = []
        
        for idx in range(num_exchanges):
            is_expanded = exchange_is_expanded.get(idx, self.exchange_expanded_state.get(idx, True))
            
            if is_expanded and idx in parsed_exchanges:
                # Use parsed content from widget (user's edits)
                parsed = parsed_exchanges[idx]
                if 'user' in parsed:
                    new_thread.append({'role': 'user', 'content': parsed['user']})
                elif idx < len(original_exchanges) and 'user' in original_exchanges[idx]:
                    # Fallback to original if parsing failed for user
                    new_thread.append({'role': 'user', 'content': original_exchanges[idx]['user'].get('content', '')})
                
                if 'assistant' in parsed:
                    new_thread.append({'role': 'assistant', 'content': parsed['assistant']})
                elif idx < len(original_exchanges) and 'assistant' in original_exchanges[idx]:
                    # Fallback to original if parsing failed for assistant
                    new_thread.append({'role': 'assistant', 'content': original_exchanges[idx]['assistant'].get('content', '')})
            else:
                # Collapsed or not in parsed - use original content
                if idx < len(original_exchanges):
                    orig = original_exchanges[idx]
                    if 'user' in orig:
                        new_thread.append({'role': 'user', 'content': orig['user'].get('content', '')})
                    if 'assistant' in orig:
                        new_thread.append({'role': 'assistant', 'content': orig['assistant'].get('content', '')})
        
        # Validate we got something
        if not new_thread:
            # Parsing completely failed - keep original
            print("⚠️ Parse failed completely - preserving original thread")
            if self.current_document_id:
                from document_library import save_thread_to_document
                metadata = {
                    "model": self.model_var.get(),
                    "provider": self.provider_var.get(),
                    "last_updated": datetime.datetime.now().isoformat(),
                    "message_count": self.thread_message_count
                }
                save_thread_to_document(self.current_document_id, self.current_thread, metadata)
            return
        
        # Update the thread
        self.current_thread.clear()
        self.current_thread.extend(new_thread)
        
        # Update message count
        self.thread_message_count = len([m for m in new_thread if m.get('role') == 'user'])
        
        # Save to document library
        if self.current_document_id:
            from document_library import save_thread_to_document
            metadata = {
                "model": self.model_var.get(),
                "provider": self.provider_var.get(),
                "last_updated": datetime.datetime.now().isoformat(),
                "message_count": self.thread_message_count,
                "edited": True
            }
            save_thread_to_document(self.current_document_id, self.current_thread, metadata)
            print(f"✅ Edits auto-saved ({len(new_thread)} messages, {len(parsed_exchanges)} exchanges parsed)")

    def _reconstruct_markdown_content(self, content_lines: list) -> str:
        """
        Reconstruct markdown from content lines by reading formatting tags.
        
        Args:
            content_lines: List of (line_number, text) tuples
            
        Returns:
            Text with markdown formatting (## for headers, ** for bold, etc.)
        """
        result_lines = []
        
        for line_num, line_text in content_lines:
            if not line_text.strip():
                result_lines.append('')
                continue
            
            # Check block-level formatting at the start of the line
            line_start_pos = f"{line_num}.0"
            try:
                tags_at_start = self.thread_text.tag_names(line_start_pos)
            except:
                tags_at_start = ()
            
            is_header = 'header' in tags_at_start
            is_bullet = 'bullet' in tags_at_start or line_text.strip().startswith('•')
            
            # Handle header lines
            if is_header:
                # Remove any existing ## prefix to avoid doubling
                clean_text = line_text.strip()
                if clean_text.startswith('## '):
                    clean_text = clean_text[3:]
                elif clean_text.startswith('### '):
                    clean_text = clean_text[4:]
                result_lines.append(f"## {clean_text}")
                continue
            
            # Handle bullet lines
            if is_bullet:
                clean_text = line_text.strip()
                if clean_text.startswith('•'):
                    clean_text = clean_text[1:].strip()
                elif clean_text.startswith('- '):
                    clean_text = clean_text[2:].strip()
                elif clean_text.startswith('* '):
                    clean_text = clean_text[2:].strip()
                # Reconstruct inline formatting within the bullet
                formatted_text = self._reconstruct_inline_markdown(line_num, clean_text, len(line_text) - len(line_text.lstrip()))
                result_lines.append(f"- {formatted_text}")
                continue
            
            # Regular line - check for inline formatting (bold, italic)
            formatted_text = self._reconstruct_inline_markdown(line_num, line_text, 0)
            result_lines.append(formatted_text)
        
        return '\n'.join(result_lines).strip()

    def _reconstruct_inline_markdown(self, line_num: int, text: str, col_offset: int) -> str:
        """
        Reconstruct inline markdown (**bold**, *italic*) by reading tags from the Text widget.
        
        Args:
            line_num: The line number in the Text widget
            text: The text content
            col_offset: Column offset (for indented lines)
            
        Returns:
            Text with markdown markers for bold/italic
        """
        if not text:
            return ""
        
        result = []
        current_bold = False
        current_italic = False
        current_segment = []
        
        for i, char in enumerate(text):
            col = col_offset + i
            pos = f"{line_num}.{col}"
            
            try:
                tags = self.thread_text.tag_names(pos)
            except:
                tags = ()
            
            # Check for bold (but not header/user/assistant which are whole-line bold)
            is_bold = 'bold' in tags and 'header' not in tags and 'user' not in tags and 'assistant' not in tags
            is_italic = 'italic' in tags
            
            # If formatting changed, flush current segment
            if is_bold != current_bold or is_italic != current_italic:
                if current_segment:
                    segment_text = ''.join(current_segment)
                    if current_bold and current_italic:
                        result.append(f"***{segment_text}***")
                    elif current_bold:
                        result.append(f"**{segment_text}**")
                    elif current_italic:
                        result.append(f"*{segment_text}*")
                    else:
                        result.append(segment_text)
                    current_segment = []
                
                current_bold = is_bold
                current_italic = is_italic
            
            current_segment.append(char)
        
        # Flush final segment
        if current_segment:
            segment_text = ''.join(current_segment)
            if current_bold and current_italic:
                result.append(f"***{segment_text}***")
            elif current_bold:
                result.append(f"**{segment_text}**")
            elif current_italic:
                result.append(f"*{segment_text}*")
            else:
                result.append(segment_text)
        
        return ''.join(result)
    
    def _copy_source_only(self, plain_text=False):
        """
        Copy just the source document to clipboard.
        
        Args:
            plain_text: If True, copy as plain text. If False, copy as formatted HTML.
        """
        import sys
        
        if not self.current_document_text:
            self._set_status("⚠️ No source document available")
            return
        
        try:
            if plain_text:
                # Build plain text version with metadata header
                lines = []
                lines.append("SOURCE DOCUMENT INFORMATION:")
                lines.append(f"Title: {self.doc_title}")
                lines.append(f"Source: {self.source_info}")
                if hasattr(self, 'published_date') and self.published_date and self.published_date != 'N/A':
                    lines.append(f"Published: {self.published_date}")
                lines.append(f"Imported: {self.fetched_date}")
                lines.append("")
                lines.append("=" * 60)
                lines.append("")
                lines.append(self.current_document_text)
                
                content = "\n".join(lines)
                self.window.clipboard_clear()
                self.window.clipboard_append(content)
                self._set_status("✅ Source copied to clipboard (plain text)")
            else:
                # Build formatted HTML version
                html_parts = []
                
                # Add metadata header
                html_parts.append(f'<h1 style="color: #2C3E50; font-size: 16pt; text-align: center;">{self._escape_html(self.doc_title)}</h1>')
                html_parts.append('<hr style="border: 1px solid #ccc;">')
                html_parts.append('<p style="font-size: 10pt; color: #555;">')
                html_parts.append(f'<b>Source:</b> {self._escape_html(self.source_info)}<br>')
                if hasattr(self, 'published_date') and self.published_date and self.published_date != 'N/A':
                    html_parts.append(f'<b>Published:</b> {self.published_date}<br>')
                html_parts.append(f'<b>Imported:</b> {self.fetched_date}')
                html_parts.append('</p>')
                html_parts.append('<hr style="border: 1px solid #ccc;">')
                html_parts.append('<br>')
                
                # Add source document text - preserve paragraphs
                source_paragraphs = self.current_document_text.split('\n\n')
                for para in source_paragraphs:
                    if para.strip():
                        html_parts.append(f'<p style="margin: 8pt 0;">{self._escape_html(para.strip())}</p>')
                
                # Wrap in HTML document
                html_body = '\n'.join(html_parts)
                html_doc = f'''<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
body {{ font-family: Calibri, Arial, sans-serif; font-size: 11pt; line-height: 1.5; }}
</style>
</head>
<body>
{html_body}
</body>
</html>'''
                
                # Copy to clipboard based on platform
                if sys.platform == 'win32':
                    success = self._copy_html_to_clipboard_windows(html_doc)
                else:
                    # Fallback for non-Windows
                    self._copy_source_only(plain_text=True)
                    return
                
                if success:
                    self._set_status("✅ Source copied (formatted)! Paste into Word/Outlook.")
                else:
                    # Fallback to plain text
                    self._copy_source_only(plain_text=True)
                    
        except Exception as e:
            import traceback
            traceback.print_exc()
            messagebox.showerror("Copy Error", f"Failed to copy source:\n{str(e)}")
            # Fallback to plain text
            try:
                if not plain_text:
                    self._copy_source_only(plain_text=True)
            except:
                pass

    def _fix_numbered_lists(self, text: str) -> str:
        """
        Fix markdown-style numbered lists where all items start with '1.'
        Convert them to properly sequential numbers (1, 2, 3, etc.)
        
        Handles the common AI pattern where numbered items have explanation
        paragraphs between them:
            1. **First point**
            The speaker argues that...
            
            1. **Second point**
            Meta-awareness is described as...
        """
        import re
        
        lines = text.split('\n')
        result = []
        list_counter = 0
        
        for line in lines:
            # Check if line starts with a number followed by period and space
            match = re.match(r'^(\s*)(\d+)\.\s+(.*)', line)
            
            if match:
                indent = match.group(1)
                content = match.group(3)
                list_counter += 1
                result.append(f"{indent}{list_counter}. {content}")
            else:
                # Reset counter only on structural breaks: headings, HRs, bullet lists
                stripped = line.strip()
                if (stripped.startswith('## ') or stripped.startswith('### ') or 
                    stripped.startswith('# ') or stripped == '---' or
                    stripped.startswith('- ') or stripped.startswith('* ') or
                    stripped.startswith('===')):
                    list_counter = 0
                # Empty lines and regular paragraphs do NOT reset counter
                # (AI often puts explanations between numbered items)
                result.append(line)
        
        return '\n'.join(result)

    def _copy_thread(self):
        """Copy entire thread to clipboard with metadata header (plain text)"""
        # Build metadata header
        metadata = []
        metadata.append("SOURCE DOCUMENT INFORMATION:")
        metadata.append(f"Title: {self.doc_title}")
        metadata.append(f"Source: {self.source_info}")
        if hasattr(self, 'published_date') and self.published_date and self.published_date != 'N/A':
            metadata.append(f"Published: {self.published_date}")
        metadata.append(f"Imported: {self.fetched_date}")
        metadata.append("")
        metadata.append("=" * 60)
        metadata.append("CONVERSATION THREAD")
        metadata.append("=" * 60)
        metadata.append("")
        
        # Combine metadata with thread content
        header = "\n".join(metadata)
        thread_content = self.thread_text.get('1.0', tk.END)
        
        # Fix numbered lists (convert "1. 1. 1." to "1. 2. 3.")
        thread_content = self._fix_numbered_lists(thread_content)
        
        full_content = header + thread_content
        
        self.window.clipboard_clear()
        self.window.clipboard_append(full_content)
        self._set_status("✅ Thread copied to clipboard (plain text)")

    def _copy_thread_formatted(self):
        """
        Copy thread to clipboard as formatted HTML.
        If user made edits, saves them first. Otherwise uses original markdown.
        """
        import sys
        
        try:
            # Save edits if modifications were made (the save function checks edit_modified)
            # If no edits, this returns early and preserves original markdown in self.current_thread
            try:
                if self.current_mode == 'conversation' and self.current_thread:
                    self._save_edits_to_thread()
            except ValueError:
                # User cancelled save dialog
                return
            except Exception as e:
                print(f"⚠️ Error saving edits: {e}")
            
            # Now build HTML from the saved thread (which has markdown)
            html_content = self._thread_to_html()
            
            if not html_content:
                self._set_status("⚠️ No content to copy")
                return
            
            # Copy to clipboard based on platform
            if sys.platform == 'win32':
                success = self._copy_html_to_clipboard_windows(html_content)
            else:
                # Fallback for non-Windows: copy as plain text
                plain_text = self.thread_text.get('1.0', tk.END)
                plain_text = self._fix_numbered_lists(plain_text)
                self.window.clipboard_clear()
                self.window.clipboard_append(plain_text)
                self._set_status("ℹ️ Formatted copy is Windows-only. Plain text copied.")
                return
            
            if success:
                self._set_status("✅ Formatted content copied! Paste into Word/Outlook.")
            else:
                # Fallback to plain text
                plain_text = self.thread_text.get('1.0', tk.END)
                plain_text = self._fix_numbered_lists(plain_text)
                self.window.clipboard_clear()
                self.window.clipboard_append(plain_text)
                self._set_status("⚠️ HTML copy failed. Plain text copied instead.")
                
        except Exception as e:
            import traceback
            traceback.print_exc()
            messagebox.showerror("Copy Error", f"Failed to copy formatted content:\n{str(e)}")
            # Fallback to plain text
            try:
                plain_text = self.thread_text.get('1.0', tk.END)
                plain_text = self._fix_numbered_lists(plain_text)
                self.window.clipboard_clear()
                self.window.clipboard_append(plain_text)
                self._set_status("⚠️ Error occurred. Plain text copied instead.")
            except:
                pass

    def _copy_expanded_only(self):
        """
        Copy only the expanded exchanges to clipboard as plain text.
        Collapsed exchanges are completely omitted (no headers, no indicators).
        Perfect for sharing specific exchanges without surrounding context.
        """
        exchanges = self._group_messages_into_exchanges()
        
        if not exchanges:
            self._set_status("⚠️ No exchanges to copy")
            return
        
        # Find which exchanges are expanded
        expanded_indices = [i for i in range(len(exchanges)) 
                          if self.exchange_expanded_state.get(i, True)]
        
        if not expanded_indices:
            self._set_status("⚠️ No exchanges are expanded. Expand the exchanges you want to copy.")
            return
        
        # Build content from expanded exchanges only
        content_parts = []
        
        # Add brief header
        content_parts.append(f"From: {self.doc_title}")
        content_parts.append(f"Source: {self.source_info}")
        content_parts.append("")
        content_parts.append("=" * 60)
        content_parts.append("")
        
        for idx in expanded_indices:
            exchange = exchanges[idx]
            
            # User message
            user_msg = exchange.get('user', {})
            user_content = user_msg.get('content', '')
            content_parts.append("YOU:")
            content_parts.append(user_content)
            content_parts.append("")
            
            # Assistant message
            assistant_msg = exchange.get('assistant', {})
            if assistant_msg:
                assistant_content = assistant_msg.get('content', '')
                provider = assistant_msg.get('provider', 'AI')
                content_parts.append(f"{provider}:")
                content_parts.append(assistant_content)
                content_parts.append("")
            
            # Separator between exchanges (if more than one expanded)
            if idx != expanded_indices[-1]:
                content_parts.append("-" * 40)
                content_parts.append("")
        
        full_content = "\n".join(content_parts)
        
        # Fix numbered lists (convert "1. 1. 1." to "1. 2. 3.")
        full_content = self._fix_numbered_lists(full_content)
        
        self.window.clipboard_clear()
        self.window.clipboard_append(full_content)
        
        count = len(expanded_indices)
        self._set_status(f"✅ Copied {count} expanded exchange{'s' if count != 1 else ''} (plain text)")

    def _copy_expanded_only_formatted(self):
        """
        Copy only the expanded exchanges to clipboard as formatted HTML.
        Collapsed exchanges are completely omitted.
        Uses the same HTML generation as _thread_to_html for consistency.
        """
        import sys
        
        try:
            exchanges = self._group_messages_into_exchanges()
            
            if not exchanges:
                self._set_status("⚠️ No exchanges to copy")
                return
            
            # Find which exchanges are expanded
            expanded_indices = []
            for i in range(len(exchanges)):
                is_expanded = self.exchange_expanded_state.get(i, True)
                if is_expanded:
                    expanded_indices.append(i)
            
            if not expanded_indices:
                self._set_status("⚠️ No exchanges are expanded. Expand the exchanges you want to copy.")
                return
            
            # Build a filtered thread containing only expanded exchanges
            filtered_thread = []
            for idx in expanded_indices:
                exchange = exchanges[idx]
                if 'user' in exchange:
                    filtered_thread.append(exchange['user'])
                if 'assistant' in exchange:
                    filtered_thread.append(exchange['assistant'])
            
            # Now build HTML using same approach as _thread_to_html
            html_parts = []
            
            # Add metadata header (simplified for partial export)
            html_parts.append(f'<h1 style="color: #2C3E50; font-size: 16pt; text-align: center;">{self._escape_html(self.doc_title)}</h1>')
            html_parts.append('<hr style="border: 1px solid #ccc;">')
            html_parts.append('<p style="font-size: 10pt; color: #555;">')
            html_parts.append(f'<b>Source:</b> {self._escape_html(self.source_info)}<br>')
            if hasattr(self, 'published_date') and self.published_date and self.published_date != 'N/A':
                html_parts.append(f'<b>Published:</b> {self.published_date}<br>')
            html_parts.append(f'<b>Imported:</b> {self.fetched_date}')
            html_parts.append('</p>')
            html_parts.append('<hr style="border: 1px solid #ccc;">')
            html_parts.append('<br>')
            
            # Process each message in the filtered thread (same as _thread_to_html)
            for msg in filtered_thread:
                role = msg.get('role', '')
                content = msg.get('content', '')
                timestamp = msg.get('timestamp', '')
                
                if role == 'user':
                    time_str = f" [{timestamp}]" if timestamp else ""
                    html_parts.append(f'<p style="color: #2E4053; font-weight: bold; margin-top: 15pt;">🧑 YOU{time_str}</p>')
                    # Convert user content (usually plain text)
                    html_parts.append(f'<p style="margin: 6pt 0;">{self._escape_html(content)}</p>')
                    
                elif role == 'assistant':
                    provider = msg.get('provider', 'AI')
                    model = msg.get('model', '')
                    time_str = f" [{timestamp}]" if timestamp else ""
                    
                    if model and model != provider:
                        label = f"🤖 {provider} ({model}){time_str}"
                    else:
                        label = f"🤖 {provider}{time_str}"
                    
                    html_parts.append(f'<p style="color: #16537E; font-weight: bold; margin-top: 15pt;">{self._escape_html(label)}</p>')
                    
                    # Convert assistant content (has markdown) - use same method as _thread_to_html
                    content_html = self._markdown_to_html_content(content)
                    html_parts.append(content_html)
                    
                    # Add divider
                    html_parts.append('<hr style="border: 1px solid #ddd; margin: 15pt 0;">')
            
            # Wrap in HTML document (same structure as _thread_to_html)
            html_body = '\n'.join(html_parts)
            html_doc = f'''<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
body {{ font-family: Calibri, Arial, sans-serif; font-size: 11pt; line-height: 1.5; }}
h1, h2, h3 {{ font-weight: bold; }}
b, strong {{ font-weight: bold; }}
i, em {{ font-style: italic; }}
ul, ol {{ margin: 6pt 0; padding-left: 25pt; }}
li {{ margin: 4pt 0; }}
</style>
</head>
<body>
{html_body}
</body>
</html>'''
            
            # Copy to clipboard based on platform
            if sys.platform == 'win32':
                success = self._copy_html_to_clipboard_windows(html_doc)
            else:
                # Fallback for non-Windows
                self._copy_expanded_only()  # Use plain text version
                return
            
            if success:
                count = len(expanded_indices)
                self._set_status(f"✅ Copied {count} expanded exchange{'s' if count != 1 else ''} (formatted)")
            else:
                # Fallback to plain text
                self._copy_expanded_only()
                
        except Exception as e:
            import traceback
            traceback.print_exc()
            messagebox.showerror("Copy Error", f"Failed to copy formatted content:\n{str(e)}")
            # Fallback to plain text
            try:
                self._copy_expanded_only()
            except:
                pass

    def _copy_complete(self, plain_text=False):
        """
        Copy the source document AND the conversation thread to clipboard.
        
        Args:
            plain_text: If True, copy as plain text. If False, copy as formatted HTML.
        """
        import sys
        
        try:
            if not self.current_document_text:
                self._set_status("⚠️ No source document available")
                return
            
            if plain_text:
                # Build plain text version
                lines = []
                
                # Add metadata header
                lines.append("=" * 60)
                lines.append(f"TITLE: {self.doc_title}")
                lines.append(f"SOURCE: {self.source_info}")
                if hasattr(self, 'published_date') and self.published_date and self.published_date != 'N/A':
                    lines.append(f"PUBLISHED: {self.published_date}")
                lines.append(f"IMPORTED: {self.fetched_date}")
                lines.append("=" * 60)
                lines.append("")
                
                # Add source document
                lines.append("─" * 40)
                lines.append("SOURCE DOCUMENT")
                lines.append("─" * 40)
                lines.append("")
                lines.append(self.current_document_text)
                lines.append("")
                
                # Add conversation thread
                lines.append("─" * 40)
                lines.append("CONVERSATION THREAD")
                lines.append("─" * 40)
                lines.append("")
                
                # Add each message
                for msg in self.current_thread:
                    role = msg.get('role', '')
                    content = msg.get('content', '')
                    timestamp = msg.get('timestamp', '')
                    
                    if role == 'user':
                        time_str = f" [{timestamp}]" if timestamp else ""
                        lines.append(f"🧑 YOU{time_str}")
                        lines.append(content)
                        lines.append("")
                    elif role == 'assistant':
                        provider = msg.get('provider', 'AI')
                        model = msg.get('model', '')
                        time_str = f" [{timestamp}]" if timestamp else ""
                        
                        if model and model != provider:
                            label = f"🤖 {provider} ({model}){time_str}"
                        else:
                            label = f"🤖 {provider}{time_str}"
                        
                        lines.append(label)
                        lines.append(content)
                        lines.append("-" * 30)
                        lines.append("")
                
                # Copy to clipboard
                text = "\n".join(lines)
                # Fix numbered lists (convert "1. 1. 1." to "1. 2. 3.")
                text = self._fix_numbered_lists(text)
                self.window.clipboard_clear()
                self.window.clipboard_append(text)
                self._set_status("✅ Complete content copied (plain text)")
                
            else:
                # Build formatted HTML version
                html_parts = []
                
                # Add metadata header
                html_parts.append(f'<h1 style="color: #2C3E50; font-size: 16pt; text-align: center;">{self._escape_html(self.doc_title)}</h1>')
                html_parts.append('<hr style="border: 1px solid #ccc;">')
                html_parts.append('<p style="font-size: 10pt; color: #555;">')
                html_parts.append(f'<b>Source:</b> {self._escape_html(self.source_info)}<br>')
                if hasattr(self, 'published_date') and self.published_date and self.published_date != 'N/A':
                    html_parts.append(f'<b>Published:</b> {self.published_date}<br>')
                html_parts.append(f'<b>Imported:</b> {self.fetched_date}')
                html_parts.append('</p>')
                html_parts.append('<hr style="border: 1px solid #ccc;">')
                html_parts.append('<br>')
                
                # Add source document section
                html_parts.append('<h2 style="color: #34495E; font-size: 14pt; background-color: #EBF5FB; padding: 8pt; border-left: 4px solid #3498DB;">📄 Source Document</h2>')
                # Convert source document text - escape HTML and preserve paragraphs
                source_paragraphs = self.current_document_text.split('\n\n')
                for para in source_paragraphs:
                    if para.strip():
                        html_parts.append(f'<p style="margin: 8pt 0;">{self._escape_html(para.strip())}</p>')
                
                html_parts.append('<br>')
                html_parts.append('<hr style="border: 2px solid #3498DB; margin: 20pt 0;">')
                html_parts.append('<br>')
                
                # Add conversation thread section
                html_parts.append('<h2 style="color: #34495E; font-size: 14pt; background-color: #E8F8F5; padding: 8pt; border-left: 4px solid #1ABC9C;">💬 Conversation Thread</h2>')
                html_parts.append('<br>')
                
                # Add each message
                for msg in self.current_thread:
                    role = msg.get('role', '')
                    content = msg.get('content', '')
                    timestamp = msg.get('timestamp', '')
                    
                    if role == 'user':
                        time_str = f" [{timestamp}]" if timestamp else ""
                        html_parts.append(f'<p style="color: #2E4053; font-weight: bold; margin-top: 15pt;">🧑 YOU{time_str}</p>')
                        html_parts.append(f'<p style="margin: 6pt 0;">{self._escape_html(content)}</p>')
                        
                    elif role == 'assistant':
                        provider = msg.get('provider', 'AI')
                        model = msg.get('model', '')
                        time_str = f" [{timestamp}]" if timestamp else ""
                        
                        if model and model != provider:
                            label = f"🤖 {provider} ({model}){time_str}"
                        else:
                            label = f"🤖 {provider}{time_str}"
                        
                        html_parts.append(f'<p style="color: #16537E; font-weight: bold; margin-top: 15pt;">{self._escape_html(label)}</p>')
                        
                        # Convert assistant content (has markdown)
                        content_html = self._markdown_to_html_content(content)
                        html_parts.append(content_html)
                        
                        # Add divider
                        html_parts.append('<hr style="border: 1px solid #ddd; margin: 15pt 0;">')
                
                # Wrap in HTML document
                html_body = '\n'.join(html_parts)
                html_doc = f'''<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
body {{ font-family: Calibri, Arial, sans-serif; font-size: 11pt; line-height: 1.5; }}
h1, h2, h3 {{ font-weight: bold; }}
b, strong {{ font-weight: bold; }}
i, em {{ font-style: italic; }}
ul, ol {{ margin: 6pt 0; padding-left: 25pt; }}
li {{ margin: 4pt 0; }}
</style>
</head>
<body>
{html_body}
</body>
</html>'''
                
                # Copy to clipboard based on platform
                if sys.platform == 'win32':
                    success = self._copy_html_to_clipboard_windows(html_doc)
                else:
                    # Fallback for non-Windows
                    self._copy_complete(plain_text=True)
                    return
                
                if success:
                    self._set_status("✅ Complete content copied (formatted)! Paste into Word/Outlook.")
                else:
                    # Fallback to plain text
                    self._copy_complete(plain_text=True)
                    
        except Exception as e:
            import traceback
            traceback.print_exc()
            messagebox.showerror("Copy Error", f"Failed to copy content:\n{str(e)}")
            # Fallback to plain text
            try:
                if not plain_text:
                    self._copy_complete(plain_text=True)
            except:
                pass

    def _thread_to_html(self) -> str:
        """
        Convert the current thread (with markdown) to HTML.
        """
        html_parts = []
        
        # Add metadata header
        html_parts.append(f'<h1 style="color: #2C3E50; font-size: 16pt; text-align: center;">{self._escape_html(self.doc_title)}</h1>')
        html_parts.append('<hr style="border: 1px solid #ccc;">')
        html_parts.append('<p style="font-size: 10pt; color: #555;">')
        html_parts.append(f'<b>Source:</b> {self._escape_html(self.source_info)}<br>')
        if hasattr(self, 'published_date') and self.published_date and self.published_date != 'N/A':
            html_parts.append(f'<b>Published:</b> {self.published_date}<br>')
        html_parts.append(f'<b>Imported:</b> {self.fetched_date}')
        html_parts.append('</p>')
        html_parts.append('<hr style="border: 1px solid #ccc;">')
        html_parts.append('<br>')
        
        # Process each message in the thread
        for msg in self.current_thread:
            role = msg.get('role', '')
            content = msg.get('content', '')
            timestamp = msg.get('timestamp', '')
            
            if role == 'user':
                time_str = f" [{timestamp}]" if timestamp else ""
                html_parts.append(f'<p style="color: #2E4053; font-weight: bold; margin-top: 15pt;">🧑 YOU{time_str}</p>')
                # Convert user content (usually plain text)
                html_parts.append(f'<p style="margin: 6pt 0;">{self._escape_html(content)}</p>')
                
            elif role == 'assistant':
                provider = msg.get('provider', 'AI')
                model = msg.get('model', '')
                time_str = f" [{timestamp}]" if timestamp else ""
                
                if model and model != provider:
                    label = f"🤖 {provider} ({model}){time_str}"
                else:
                    label = f"🤖 {provider}{time_str}"
                
                html_parts.append(f'<p style="color: #16537E; font-weight: bold; margin-top: 15pt;">{self._escape_html(label)}</p>')
                
                # Convert assistant content (has markdown)
                content_html = self._markdown_to_html_content(content)
                html_parts.append(content_html)
                
                # Add divider
                html_parts.append('<hr style="border: 1px solid #ddd; margin: 15pt 0;">')
        
        # Wrap in HTML document
        html_body = '\n'.join(html_parts)
        html_doc = f'''<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
body {{ font-family: Calibri, Arial, sans-serif; font-size: 11pt; line-height: 1.5; }}
h1, h2, h3 {{ font-weight: bold; }}
b, strong {{ font-weight: bold; }}
i, em {{ font-style: italic; }}
ul, ol {{ margin: 6pt 0; padding-left: 25pt; }}
li {{ margin: 4pt 0; }}
</style>
</head>
<body>
{html_body}
</body>
</html>'''
        
        return html_doc

    def _markdown_to_html_content(self, markdown_text: str) -> str:
        """
        Convert markdown text to HTML.
        Handles: **bold**, *italic*, ## headings, - bullets, numbered lists
        """
        if not markdown_text:
            return ""
        
        lines = markdown_text.split('\n')
        html_parts = []
        in_list = False
        list_type = None
        list_counter = 0  # Track numbered list counter for Gmail-compatible explicit numbering
        
        for line in lines:
            stripped = line.strip()
            
            # Empty line - close bullet lists but preserve numbered list state
            # (AI output often has empty lines between numbered items)
            if not stripped:
                if in_list and list_type == 'ul':
                    html_parts.append('</ul>')
                    in_list = False
                    list_type = None
                # Don't reset ol_paragraphs counter on empty lines - only on actual content change
                continue
            
            # Horizontal rule
            if stripped == '---':
                if in_list:
                    html_parts.append(f'</{list_type}>')
                    in_list = False
                list_type = None
                list_counter = 0
                html_parts.append('<hr style="border: 1px solid #ddd;">')
                continue
            
            # Heading 2: ## Title
            if stripped.startswith('## '):
                if in_list:
                    html_parts.append(f'</{list_type}>')
                    in_list = False
                list_type = None
                list_counter = 0
                text = self._convert_inline_markdown(stripped[3:])
                html_parts.append(f'<h2 style="color: #2C3E50; font-size: 13pt; margin: 12pt 0 6pt 0;">{text}</h2>')
                continue
            
            # Heading 3: ### Title
            if stripped.startswith('### '):
                if in_list:
                    html_parts.append(f'</{list_type}>')
                    in_list = False
                list_type = None
                list_counter = 0
                text = self._convert_inline_markdown(stripped[4:])
                html_parts.append(f'<h3 style="color: #34495E; font-size: 12pt; margin: 10pt 0 4pt 0;">{text}</h3>')
                continue
            
            # Bullet list: - item or * item
            if stripped.startswith('- ') or stripped.startswith('* '):
                if not in_list or list_type != 'ul':
                    if in_list:
                        html_parts.append(f'</{list_type}>')
                    html_parts.append('<ul>')
                    in_list = True
                    list_type = 'ul'
                text = self._convert_inline_markdown(stripped[2:])
                html_parts.append(f'<li>{text}</li>')
                continue
            
            # Numbered list: 1. item
            # Use explicit numbers in paragraphs for better Gmail compatibility
            if re.match(r'^\d+\.\s+', stripped):
                # Close any bullet list if open
                if in_list and list_type == 'ul':
                    html_parts.append('</ul>')
                    in_list = False
                    list_type = None
                # Track numbered list counter (but don't use <ol> tags)
                if list_type != 'ol_paragraphs':
                    list_type = 'ol_paragraphs'
                    list_counter = 0
                list_counter += 1
                text = self._convert_inline_markdown(re.sub(r'^\d+\.\s+', '', stripped))
                # Use paragraph with explicit number instead of <li> for Gmail compatibility
                html_parts.append(f'<p style="margin: 3pt 0 3pt 20pt;">{list_counter}. {text}</p>')
                continue
            
            # Block quote: > text
            if stripped.startswith('> '):
                if in_list:
                    html_parts.append(f'</{list_type}>')
                    in_list = False
                list_type = None
                list_counter = 0
                text = self._convert_inline_markdown(stripped[2:])
                html_parts.append(f'<blockquote style="margin: 6pt 0 6pt 20pt; padding-left: 10pt; border-left: 3px solid #ccc; color: #555; font-style: italic;">{text}</blockquote>')
                continue
            
            # Regular paragraph
            if in_list:
                html_parts.append(f'</{list_type}>')
                in_list = False
            # Only reset numbered list counter on structural breaks (headings, HR, etc.)
            # Regular paragraphs between numbered items should NOT reset the counter
            # because AI often puts explanation paragraphs between numbered points
            if list_type == 'ul':
                list_type = None
            # Keep list_type and list_counter alive for ol_paragraphs
            
            text = self._convert_inline_markdown(stripped)
            html_parts.append(f'<p style="margin: 6pt 0;">{text}</p>')
        
        # Close any remaining list
        if in_list:
            html_parts.append(f'</{list_type}>')
        
        return '\n'.join(html_parts)

    def _convert_inline_markdown(self, text: str) -> str:
        """Convert inline markdown (**bold**, *italic*) to HTML."""
        if not text:
            return ""
        
        # Escape HTML first
        text = self._escape_html(text)
        
        # Convert **bold** to <b>bold</b>
        text = re.sub(r'\*\*([^*]+)\*\*', r'<b>\1</b>', text)
        
        # Convert *italic* to <i>italic</i>
        text = re.sub(r'(?<!\*)\*([^*]+)\*(?!\*)', r'<i>\1</i>', text)
        
        return text

    def _text_widget_to_html(self) -> str:
        """
        Convert the Text widget content to HTML, preserving visual formatting.
        Reads the actual tags applied to text ranges in the widget.
        """
        # Get all content
        content = self.thread_text.get('1.0', 'end-1c')
        
        # Build HTML by scanning through the text and checking tags at each position
        html_parts = []
        
        # Add metadata header
        html_parts.append(f'<h1 style="color: #2C3E50; font-size: 16pt; text-align: center;">{self._escape_html(self.doc_title)}</h1>')
        html_parts.append('<hr style="border: 1px solid #ccc;">')
        html_parts.append('<p style="font-size: 10pt; color: #555;">')
        html_parts.append(f'<b>Source:</b> {self._escape_html(self.source_info)}<br>')
        if hasattr(self, 'published_date') and self.published_date and self.published_date != 'N/A':
            html_parts.append(f'<b>Published:</b> {self.published_date}<br>')
        html_parts.append(f'<b>Imported:</b> {self.fetched_date}')
        html_parts.append('</p>')
        html_parts.append('<hr style="border: 1px solid #ccc;">')
        html_parts.append('<br>')
        
        # Process text line by line to handle block-level formatting
        lines = content.split('\n')
        line_num = 1
        
        for line in lines:
            if not line:
                html_parts.append('<p>&nbsp;</p>')
                line_num += 1
                continue
            
            # Get the start index for this line
            line_start = f"{line_num}.0"
            
            # Check what tags are at the start of the line for block-level formatting
            tags_at_start = self.thread_text.tag_names(line_start)
            
            # Determine block-level style
            is_header = 'header' in tags_at_start
            is_user = 'user' in tags_at_start
            is_assistant = 'assistant' in tags_at_start
            is_divider = 'divider' in tags_at_start
            is_bullet = 'bullet' in tags_at_start
            is_exchange_header = 'exchange_header' in tags_at_start
            is_source_header = 'source_header' in tags_at_start
            
            # Handle divider lines
            if is_divider or line.strip().startswith('─' * 10):
                html_parts.append('<hr style="border: 1px solid #ddd; margin: 10pt 0;">')
                line_num += 1
                continue
            
            # Handle exchange/source headers (clickable in viewer)
            if is_exchange_header or is_source_header:
                bg_color = '#d4e6f1' if is_source_header else '#e8f4f8'
                html_parts.append(f'<p style="background-color: {bg_color}; padding: 5pt; font-weight: bold; color: #1a5276;">{self._escape_html(line)}</p>')
                line_num += 1
                continue
            
            # Handle user/assistant markers
            if is_user:
                html_parts.append(f'<p style="color: #2E4053; font-weight: bold; margin-top: 12pt;">{self._escape_html(line)}</p>')
                line_num += 1
                continue
            
            if is_assistant:
                html_parts.append(f'<p style="color: #16537E; font-weight: bold; margin-top: 12pt;">{self._escape_html(line)}</p>')
                line_num += 1
                continue
            
            # Handle headers
            if is_header:
                html_parts.append(f'<h2 style="color: #2c3e50; font-size: 13pt; margin: 12pt 0 6pt 0;">{self._escape_html(line)}</h2>')
                line_num += 1
                continue
            
            # Handle bullets
            if is_bullet or line.strip().startswith('•'):
                bullet_text = line.strip()
                if bullet_text.startswith('•'):
                    bullet_text = bullet_text[1:].strip()
                html_parts.append(f'<li style="margin: 3pt 0;">{self._process_inline_formatting(line_num, bullet_text)}</li>')
                line_num += 1
                continue
            
            # Regular paragraph - process inline formatting
            formatted_line = self._process_inline_formatting(line_num, line)
            html_parts.append(f'<p style="margin: 6pt 0;">{formatted_line}</p>')
            line_num += 1
        
        # Wrap in HTML document
        html_body = '\n'.join(html_parts)
        html_doc = f'''<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
body {{ font-family: Calibri, Arial, sans-serif; font-size: 11pt; line-height: 1.4; }}
h1, h2, h3 {{ font-weight: bold; }}
b, strong {{ font-weight: bold; }}
i, em {{ font-style: italic; }}
li {{ margin-left: 20pt; }}
</style>
</head>
<body>
{html_body}
</body>
</html>'''
        
        return html_doc

    def _process_inline_formatting(self, line_num: int, line_text: str) -> str:
        """
        Process inline formatting (bold, italic) for a line by checking tags at each character.
        Returns HTML string with <b> and <i> tags.
        """
        if not line_text:
            return ""
        
        result = []
        current_bold = False
        current_italic = False
        current_text = []
        
        for col, char in enumerate(line_text):
            pos = f"{line_num}.{col}"
            try:
                tags = self.thread_text.tag_names(pos)
            except:
                tags = ()
            
            is_bold = 'bold' in tags or 'header' in tags or 'user' in tags or 'assistant' in tags
            is_italic = 'italic' in tags
            
            # Check if formatting changed
            if is_bold != current_bold or is_italic != current_italic:
                # Flush current text with previous formatting
                if current_text:
                    text = self._escape_html(''.join(current_text))
                    if current_bold and current_italic:
                        result.append(f'<b><i>{text}</i></b>')
                    elif current_bold:
                        result.append(f'<b>{text}</b>')
                    elif current_italic:
                        result.append(f'<i>{text}</i>')
                    else:
                        result.append(text)
                    current_text = []
                
                current_bold = is_bold
                current_italic = is_italic
            
            current_text.append(char)
        
        # Flush remaining text
        if current_text:
            text = self._escape_html(''.join(current_text))
            if current_bold and current_italic:
                result.append(f'<b><i>{text}</i></b>')
            elif current_bold:
                result.append(f'<b>{text}</b>')
            elif current_italic:
                result.append(f'<i>{text}</i>')
            else:
                result.append(text)
        
        return ''.join(result)

    def _escape_html(self, text: str) -> str:
        """Escape HTML special characters."""
        if not text:
            return ""
        return (text
                .replace('&', '&amp;')
                .replace('<', '&lt;')
                .replace('>', '&gt;')
                .replace('"', '&quot;'))

    def _get_thread_content_with_markdown(self) -> str:
        """
        Get thread content, reconstructing markdown from the stored thread data.
        This ensures we have the original markdown formatting, not the display text.
        """
        lines = []
        
        # Add metadata header
        lines.append(f"**{self.doc_title}**")
        lines.append("")
        lines.append(f"**Source:** {self.source_info}")
        if hasattr(self, 'published_date') and self.published_date and self.published_date != 'N/A':
            lines.append(f"**Published:** {self.published_date}")
        lines.append(f"**Imported:** {self.fetched_date}")
        lines.append("")
        lines.append("---")
        lines.append("")
        
        # Add conversation exchanges from stored thread (preserves original markdown)
        for msg in self.current_thread:
            role = msg.get('role', '')
            content = msg.get('content', '')
            timestamp = msg.get('timestamp', '')
            
            if role == 'user':
                time_str = f" [{timestamp}]" if timestamp else ""
                lines.append(f"**🧑 YOU{time_str}**")
                lines.append("")
                lines.append(content)
                lines.append("")
            elif role == 'assistant':
                provider = msg.get('provider', 'AI')
                model = msg.get('model', '')
                time_str = f" [{timestamp}]" if timestamp else ""
                
                if model and model != provider:
                    label = f"**🤖 {provider} ({model}){time_str}**"
                else:
                    label = f"**🤖 {provider}{time_str}**"
                
                lines.append(label)
                lines.append("")
                lines.append(content)
                lines.append("")
                lines.append("---")
                lines.append("")
        
        return "\n".join(lines)

    def _markdown_to_html(self, markdown_text: str) -> str:
        """
        Convert markdown text to HTML for clipboard.
        Handles: **bold**, *italic*, ## headings, - bullets, numbered lists, --- hr
        """
        lines = markdown_text.split('\n')
        html_lines = []
        in_list = False
        list_type = None  # 'ul' or 'ol'
        
        for line in lines:
            stripped = line.strip()
            
            # Empty line - close any open list
            if not stripped:
                if in_list:
                    html_lines.append(f'</{list_type}>')
                    in_list = False
                    list_type = None
                html_lines.append('<p>&nbsp;</p>')
                continue
            
            # Horizontal rule
            if stripped == '---':
                if in_list:
                    html_lines.append(f'</{list_type}>')
                    in_list = False
                    list_type = None
                html_lines.append('<hr style="border: 1px solid #ccc;">')
                continue
            
            # Heading 2: ## Title
            if stripped.startswith('## '):
                if in_list:
                    html_lines.append(f'</{list_type}>')
                    in_list = False
                    list_type = None
                text = self._inline_markdown_to_html(stripped[3:])
                html_lines.append(f'<h2 style="color: #2C3E50; font-size: 16pt; margin: 12pt 0 6pt 0;">{text}</h2>')
                continue
            
            # Heading 3: ### Title
            if stripped.startswith('### '):
                if in_list:
                    html_lines.append(f'</{list_type}>')
                    in_list = False
                    list_type = None
                text = self._inline_markdown_to_html(stripped[4:])
                html_lines.append(f'<h3 style="color: #34495E; font-size: 13pt; margin: 10pt 0 4pt 0;">{text}</h3>')
                continue
            
            # Bullet list: - item or * item
            if stripped.startswith('- ') or stripped.startswith('* '):
                if not in_list or list_type != 'ul':
                    if in_list:
                        html_lines.append(f'</{list_type}>')
                    html_lines.append('<ul style="margin: 6pt 0; padding-left: 20pt;">')
                    in_list = True
                    list_type = 'ul'
                text = self._inline_markdown_to_html(stripped[2:])
                html_lines.append(f'<li style="margin: 3pt 0;">{text}</li>')
                continue
            
            # Numbered list: 1. item
            if re.match(r'^\d+\.\s+', stripped):
                if not in_list or list_type != 'ol':
                    if in_list:
                        html_lines.append(f'</{list_type}>')
                    html_lines.append('<ol style="margin: 6pt 0; padding-left: 20pt;">')
                    in_list = True
                    list_type = 'ol'
                text = self._inline_markdown_to_html(re.sub(r'^\d+\.\s+', '', stripped))
                html_lines.append(f'<li style="margin: 3pt 0;">{text}</li>')
                continue
            
            # Block quote: > text
            if stripped.startswith('> '):
                if in_list:
                    html_lines.append(f'</{list_type}>')
                    in_list = False
                    list_type = None
                text = self._inline_markdown_to_html(stripped[2:])
                html_lines.append(f'<blockquote style="margin: 6pt 0 6pt 20pt; padding-left: 10pt; border-left: 3px solid #ccc; color: #555; font-style: italic;">{text}</blockquote>')
                continue
            
            # Regular paragraph - close any open list first
            if in_list:
                html_lines.append(f'</{list_type}>')
                in_list = False
                list_type = None
            
            text = self._inline_markdown_to_html(stripped)
            html_lines.append(f'<p style="margin: 6pt 0;">{text}</p>')
        
        # Close any remaining open list
        if in_list:
            html_lines.append(f'</{list_type}>')
        
        # Wrap in basic HTML structure
        html_body = '\n'.join(html_lines)
        html_doc = f'''<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
body {{ font-family: Calibri, Arial, sans-serif; font-size: 11pt; line-height: 1.4; }}
h2 {{ font-weight: bold; }}
h3 {{ font-weight: bold; }}
b, strong {{ font-weight: bold; }}
i, em {{ font-style: italic; }}
</style>
</head>
<body>
{html_body}
</body>
</html>'''
        
        return html_doc

    def _inline_markdown_to_html(self, text: str) -> str:
        """Convert inline markdown (**bold**, *italic*) to HTML tags."""
        # Escape HTML special characters first
        text = text.replace('&', '&amp;')
        text = text.replace('<', '&lt;')
        text = text.replace('>', '&gt;')
        
        # Convert **bold** to <b>bold</b>
        text = re.sub(r'\*\*([^*]+)\*\*', r'<b>\1</b>', text)
        
        # Convert *italic* to <i>italic</i> (but not ** which is bold)
        text = re.sub(r'(?<!\*)\*([^*]+)\*(?!\*)', r'<i>\1</i>', text)
        
        return text

    def _copy_html_to_clipboard_windows(self, html_content: str) -> bool:
        """
        Copy HTML content to Windows clipboard in CF_HTML format.
        This allows pasting with formatting into Word, Outlook, etc.
        """
        # First try pywin32 if available (most reliable)
        try:
            import win32clipboard
            import win32con
            
            
            # Build CF_HTML data
            cf_html = self._build_cf_html(html_content)
            
            # Register HTML format
            CF_HTML = win32clipboard.RegisterClipboardFormat("HTML Format")
            
            win32clipboard.OpenClipboard()
            try:
                win32clipboard.EmptyClipboard()
                
                # Set HTML format - pywin32 expects bytes for custom formats
                cf_html_bytes = cf_html.encode('utf-8')
                win32clipboard.SetClipboardData(CF_HTML, cf_html_bytes)
                
                # Also set plain text as fallback
                plain_text = self._html_to_plain_text(html_content)
                win32clipboard.SetClipboardText(plain_text, win32con.CF_UNICODETEXT)
                
                return True
            finally:
                win32clipboard.CloseClipboard()
                
        except ImportError:
            pass
        except Exception as e:
            import traceback
            traceback.print_exc()
        
        # Fallback: use ctypes directly
        try:
            import ctypes
            
            user32 = ctypes.windll.user32
            kernel32 = ctypes.windll.kernel32
            
            # Register HTML clipboard format
            CF_HTML = user32.RegisterClipboardFormatW("HTML Format")
            CF_UNICODETEXT = 13
            GMEM_MOVEABLE = 0x0002
            
            
            # Build CF_HTML data
            cf_html = self._build_cf_html(html_content)
            cf_html_bytes = cf_html.encode('utf-8') + b'\x00'
            
            
            # Open clipboard
            if not user32.OpenClipboard(None):
                return False
            
            try:
                user32.EmptyClipboard()
                
                # Allocate and set HTML data
                h_html = kernel32.GlobalAlloc(GMEM_MOVEABLE, len(cf_html_bytes))
                if h_html:
                    p_html = kernel32.GlobalLock(h_html)
                    if p_html:
                        ctypes.memmove(p_html, cf_html_bytes, len(cf_html_bytes))
                        kernel32.GlobalUnlock(h_html)
                        result = user32.SetClipboardData(CF_HTML, h_html)
                        if not result:
                            kernel32.GlobalFree(h_html)
                    else:
                        pass
                else:
                    pass
                
                # Also set plain text
                plain_text = self._html_to_plain_text(html_content)
                plain_bytes = (plain_text + '\x00').encode('utf-16-le')
                
                h_text = kernel32.GlobalAlloc(GMEM_MOVEABLE, len(plain_bytes))
                if h_text:
                    p_text = kernel32.GlobalLock(h_text)
                    if p_text:
                        ctypes.memmove(p_text, plain_bytes, len(plain_bytes))
                        kernel32.GlobalUnlock(h_text)
                        user32.SetClipboardData(CF_UNICODETEXT, h_text)
                
                return True
                
            finally:
                user32.CloseClipboard()
                
        except Exception as e:
            import traceback
            traceback.print_exc()
            return False

    def _build_cf_html(self, html_content: str) -> str:
        """
        Build the CF_HTML clipboard format string.
        
        CF_HTML format requires a header with byte positions:
        Version:0.9
        StartHTML:XXXXXXXX
        EndHTML:XXXXXXXX
        StartFragment:XXXXXXXX
        EndFragment:XXXXXXXX
        <html>...<!--StartFragment-->CONTENT<!--EndFragment-->...</html>
        """
        # Markers
        START_FRAG = "<!--StartFragment-->"
        END_FRAG = "<!--EndFragment-->"
        
        # Find or create body section
        body_start = html_content.find('<body')
        if body_start == -1:
            # No body tag, wrap content
            html_with_markers = f"<html><body>{START_FRAG}{html_content}{END_FRAG}</body></html>"
        else:
            # Find end of body tag
            body_tag_end = html_content.find('>', body_start) + 1
            body_close = html_content.find('</body>')
            if body_close == -1:
                body_close = len(html_content)
            
            html_with_markers = (
                html_content[:body_tag_end] +
                START_FRAG +
                html_content[body_tag_end:body_close] +
                END_FRAG +
                html_content[body_close:]
            )
        
        # Header template - MUST use exact format with \r\n line endings
        # The numbers will be formatted as 8-digit zero-padded integers
        header_template = (
            "Version:0.9\r\n"
            "StartHTML:{:08d}\r\n"
            "EndHTML:{:08d}\r\n"
            "StartFragment:{:08d}\r\n"
            "EndFragment:{:08d}\r\n"
        )
        
        # Calculate header length (with placeholder values)
        dummy_header = header_template.format(0, 0, 0, 0)
        header_len = len(dummy_header.encode('utf-8'))
        
        # Calculate byte positions (must be byte positions, not character positions)
        html_bytes = html_with_markers.encode('utf-8')
        
        start_html = header_len
        end_html = header_len + len(html_bytes)
        
        # Find fragment markers in the byte string
        start_frag_marker_bytes = START_FRAG.encode('utf-8')
        end_frag_marker_bytes = END_FRAG.encode('utf-8')
        
        start_frag_pos = html_bytes.find(start_frag_marker_bytes)
        end_frag_pos = html_bytes.find(end_frag_marker_bytes)
        
        # StartFragment points to AFTER the marker, EndFragment points to BEFORE the marker
        start_fragment = header_len + start_frag_pos + len(start_frag_marker_bytes)
        end_fragment = header_len + end_frag_pos
        
        # Build final header with actual positions
        header = header_template.format(start_html, end_html, start_fragment, end_fragment)
        
        result = header + html_with_markers
        
        # Debug output
        
        return result

    def _html_to_plain_text(self, html: str) -> str:
        """Convert HTML to plain text as fallback."""
        # Remove HTML tags
        text = re.sub(r'<style[^>]*>.*?</style>', '', html, flags=re.DOTALL)
        text = re.sub(r'<script[^>]*>.*?</script>', '', text, flags=re.DOTALL)
        text = re.sub(r'<[^>]+>', '', text)
        # Decode HTML entities
        text = text.replace('&nbsp;', ' ')
        text = text.replace('&amp;', '&')
        text = text.replace('&lt;', '<')
        text = text.replace('&gt;', '>')
        text = text.replace('&quot;', '"')
        # Clean up whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)
        return text.strip()

    def _get_selected_text(self) -> str:
        """Get the currently selected text from the thread display, or None if no selection"""
        try:
            selected = self.thread_text.get(tk.SEL_FIRST, tk.SEL_LAST)
            return selected if selected.strip() else None
        except tk.TclError:
            # No selection
            return None
    
    def _format_selected_text(self, text: str) -> str:
        """
        Clean up and format selected text:
        - Fix spacing issues
        - Normalize indentation
        - Clean up line breaks
        """
        if not text:
            return ""
        
        lines = text.split('\n')
        formatted_lines = []
        
        for line in lines:
            # Strip trailing whitespace but preserve intentional indentation
            line = line.rstrip()
            # Normalize multiple spaces to single space (except at start)
            if line:
                indent = len(line) - len(line.lstrip())
                content = ' '.join(line.split())
                line = ' ' * indent + content if indent else content
            formatted_lines.append(line)
        
        # Join lines and clean up multiple blank lines
        result = '\n'.join(formatted_lines)
        # Replace 3+ consecutive newlines with 2
        while '\n\n\n' in result:
            result = result.replace('\n\n\n', '\n\n')
        
        return result.strip()
    
    def _copy_selection_to_clipboard(self):
        """Copy selected text to clipboard"""
        selected = self._get_selected_text()
        
        if not selected:
            self._set_status("⚠️ Please select text first")
            return
        
        formatted = self._format_selected_text(selected)
        
        # Fix numbered lists (convert "1. 1. 1." to "1. 2. 3.")
        formatted = self._fix_numbered_lists(formatted)
        
        self.window.clipboard_clear()
        self.window.clipboard_append(formatted)
        
        # Show brief confirmation
        word_count = len(formatted.split())
        self._set_status(f"✅ Selection copied ({word_count} words)")

    def _copy_selection_formatted(self):
        """Copy selected text to clipboard with HTML formatting preserved."""
        import sys
        
        # Check if there's a selection
        try:
            sel_start = self.thread_text.index(tk.SEL_FIRST)
            sel_end = self.thread_text.index(tk.SEL_LAST)
        except tk.TclError:
            self._set_status("⚠️ Please select text first")
            return
        
        # Get selected text and reconstruct markdown from visual formatting
        selected_text = self.thread_text.get(sel_start, sel_end)
        
        # Build list of (line_num, text) for the selection
        start_line, start_col = map(int, sel_start.split('.'))
        
        content_lines = []
        lines = selected_text.split('\n')
        current_line = start_line
        
        for i, line in enumerate(lines):
            col_offset = start_col if i == 0 else 0
            content_lines.append((current_line, col_offset, line))
            current_line += 1
        
        # Reconstruct markdown from visual formatting
        markdown_text = self._reconstruct_selection_markdown(content_lines)
        
        # Convert markdown to HTML
        html_content = self._selection_markdown_to_html(markdown_text)
        
        # Copy to clipboard
        if sys.platform == 'win32':
            success = self._copy_html_to_clipboard_windows(html_content)
        else:
            fixed_text = self._fix_numbered_lists(selected_text)
            self.window.clipboard_clear()
            self.window.clipboard_append(fixed_text)
            self._set_status("ℹ️ Formatted copy is Windows-only. Plain text copied.")
            return
        
        if success:
            word_count = len(selected_text.split())
            self._set_status(f"✅ Selection copied formatted ({word_count} words)")
        else:
            fixed_text = self._fix_numbered_lists(selected_text)
            self.window.clipboard_clear()
            self.window.clipboard_append(fixed_text)
            self._set_status("⚠️ HTML copy failed. Plain text copied.")

    def _reconstruct_selection_markdown(self, content_lines: list) -> str:
        """
        Reconstruct markdown from selected content lines.
        
        Args:
            content_lines: List of (line_number, col_offset, text) tuples
        """
        result_lines = []
        
        for line_num, col_offset, line_text in content_lines:
            if not line_text.strip():
                result_lines.append('')
                continue
            
            # Check block-level formatting
            line_start_pos = f"{line_num}.{col_offset}"
            try:
                tags_at_start = self.thread_text.tag_names(line_start_pos)
            except:
                tags_at_start = ()
            
            is_header = 'header' in tags_at_start
            is_bullet = 'bullet' in tags_at_start or line_text.strip().startswith('•')
            is_user = 'user' in tags_at_start
            is_assistant = 'assistant' in tags_at_start
            
            # Handle user/assistant markers - keep as-is
            if is_user or is_assistant:
                result_lines.append(line_text)
                continue
            
            # Handle dividers
            if line_text.strip().startswith('─' * 10):
                result_lines.append('---')
                continue
            
            # Handle headers
            if is_header:
                clean_text = line_text.strip()
                if not clean_text.startswith('## '):
                    clean_text = f"## {clean_text}"
                result_lines.append(clean_text)
                continue
            
            # Handle bullets
            if is_bullet:
                clean_text = line_text.strip()
                if clean_text.startswith('•'):
                    clean_text = clean_text[1:].strip()
                if not clean_text.startswith('- '):
                    clean_text = f"- {clean_text}"
                result_lines.append(clean_text)
                continue
            
            # Regular line - check for inline bold/italic
            formatted = self._reconstruct_line_inline_markdown(line_num, col_offset, line_text)
            result_lines.append(formatted)
        
        return '\n'.join(result_lines)

    def _reconstruct_line_inline_markdown(self, line_num: int, col_offset: int, text: str) -> str:
        """Reconstruct inline markdown for a single line."""
        if not text:
            return ""
        
        result = []
        current_bold = False
        current_segment = []
        
        for i, char in enumerate(text):
            col = col_offset + i
            pos = f"{line_num}.{col}"
            
            try:
                tags = self.thread_text.tag_names(pos)
            except:
                tags = ()
            
            # Check for bold (excluding header/user/assistant which are whole-line bold)
            is_bold = 'bold' in tags and 'header' not in tags and 'user' not in tags and 'assistant' not in tags
            
            if is_bold != current_bold:
                if current_segment:
                    segment_text = ''.join(current_segment)
                    if current_bold:
                        result.append(f"**{segment_text}**")
                    else:
                        result.append(segment_text)
                    current_segment = []
                current_bold = is_bold
            
            current_segment.append(char)
        
        # Flush final segment
        if current_segment:
            segment_text = ''.join(current_segment)
            if current_bold:
                result.append(f"**{segment_text}**")
            else:
                result.append(segment_text)
        
        return ''.join(result)

    def _selection_markdown_to_html(self, markdown_text: str) -> str:
        """Convert selection markdown to HTML document."""
        content_html = self._markdown_to_html_content(markdown_text)
        
        html_doc = f'''<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
body {{ font-family: Calibri, Arial, sans-serif; font-size: 11pt; line-height: 1.5; }}
h2 {{ font-weight: bold; color: #2C3E50; font-size: 13pt; }}
b, strong {{ font-weight: bold; }}
ul, ol {{ margin: 6pt 0; padding-left: 25pt; }}
li {{ margin: 4pt 0; }}
</style>
</head>
<body>
{content_html}
</body>
</html>'''
        
        return html_doc

    def _selection_to_html(self, start_index: str, end_index: str) -> str:
        """
        Convert selected text range to HTML, preserving visual formatting.
        """
        # Get the selected text
        selected_text = self.thread_text.get(start_index, end_index)
        
        html_parts = []
        
        # Parse start position
        start_line, start_col = map(int, start_index.split('.'))
        end_line, end_col = map(int, end_index.split('.'))
        
        # Process each line in the selection
        lines = selected_text.split('\n')
        current_line = start_line
        
        for i, line in enumerate(lines):
            if not line:
                html_parts.append('<p>&nbsp;</p>')
                current_line += 1
                continue
            
            # Determine the column offset for this line
            if i == 0:
                col_offset = start_col
            else:
                col_offset = 0
            
            line_start = f"{current_line}.{col_offset}"
            
            # Check tags at start of line for block-level formatting
            try:
                tags_at_start = self.thread_text.tag_names(line_start)
            except:
                tags_at_start = ()
            
            is_header = 'header' in tags_at_start
            is_user = 'user' in tags_at_start
            is_assistant = 'assistant' in tags_at_start
            is_divider = 'divider' in tags_at_start
            is_bullet = 'bullet' in tags_at_start
            
            # Handle dividers
            if is_divider or line.strip().startswith('─' * 10):
                html_parts.append('<hr style="border: 1px solid #ddd;">')
                current_line += 1
                continue
            
            # Handle user/assistant markers
            if is_user:
                html_parts.append(f'<p style="color: #2E4053; font-weight: bold;">{self._escape_html(line)}</p>')
                current_line += 1
                continue
            
            if is_assistant:
                html_parts.append(f'<p style="color: #16537E; font-weight: bold;">{self._escape_html(line)}</p>')
                current_line += 1
                continue
            
            # Handle headers
            if is_header:
                html_parts.append(f'<h2 style="color: #2c3e50; font-size: 13pt;">{self._escape_html(line)}</h2>')
                current_line += 1
                continue
            
            # Handle bullets
            if is_bullet or line.strip().startswith('•'):
                bullet_text = line.strip()
                if bullet_text.startswith('•'):
                    bullet_text = bullet_text[1:].strip()
                formatted = self._process_inline_formatting_range(current_line, col_offset, line)
                html_parts.append(f'<li style="margin: 3pt 0;">{formatted}</li>')
                current_line += 1
                continue
            
            # Regular paragraph with inline formatting
            formatted = self._process_inline_formatting_range(current_line, col_offset, line)
            html_parts.append(f'<p style="margin: 6pt 0;">{formatted}</p>')
            current_line += 1
        
        # Wrap in HTML
        html_body = '\n'.join(html_parts)
        html_doc = f'''<!DOCTYPE html>
<html>
<head><meta charset="utf-8"></head>
<body style="font-family: Calibri, Arial, sans-serif; font-size: 11pt;">
{html_body}
</body>
</html>'''
        
        return html_doc

    def _process_inline_formatting_range(self, line_num: int, col_offset: int, line_text: str) -> str:
        """
        Process inline formatting for a specific range of text.
        """
        if not line_text:
            return ""
        
        result = []
        current_bold = False
        current_italic = False
        current_text = []
        
        for i, char in enumerate(line_text):
            col = col_offset + i
            pos = f"{line_num}.{col}"
            try:
                tags = self.thread_text.tag_names(pos)
            except:
                tags = ()
            
            is_bold = 'bold' in tags or 'header' in tags
            is_italic = 'italic' in tags
            
            if is_bold != current_bold or is_italic != current_italic:
                if current_text:
                    text = self._escape_html(''.join(current_text))
                    if current_bold and current_italic:
                        result.append(f'<b><i>{text}</i></b>')
                    elif current_bold:
                        result.append(f'<b>{text}</b>')
                    elif current_italic:
                        result.append(f'<i>{text}</i>')
                    else:
                        result.append(text)
                    current_text = []
                
                current_bold = is_bold
                current_italic = is_italic
            
            current_text.append(char)
        
        if current_text:
            text = self._escape_html(''.join(current_text))
            if current_bold and current_italic:
                result.append(f'<b><i>{text}</i></b>')
            elif current_bold:
                result.append(f'<b>{text}</b>')
            elif current_italic:
                result.append(f'<i>{text}</i>')
            else:
                result.append(text)
        
        return ''.join(result)
    
    def _save_selection(self, format_ext: str):
        """Save selected text to file in specified format with proper formatting"""
        selected = self._get_selected_text()
        
        if not selected:
            self._set_status("⚠️ Please select text first")
            return
        
        formatted = self._format_selected_text(selected)
        
        # Import the enhanced formatter
        from doc_formatter import save_formatted_document
        from utils import safe_filename
        from tkinter import filedialog
        
        # Prepare filename - use first few words of selection
        preview = ' '.join(formatted.split()[:5])
        if len(preview) > 30:
            preview = preview[:30]
        clean_name = safe_filename(f"Selection - {preview}")
        
        # Set up file dialog based on format
        format_map = {
            '.txt': ('txt', [("Text files", "*.txt"), ("All files", "*.*")]),
            '.docx': ('docx', [("Word documents", "*.docx"), ("All files", "*.*")]),
            '.rtf': ('rtf', [("RTF documents", "*.rtf"), ("All files", "*.*")]),
            '.pdf': ('pdf', [("PDF files", "*.pdf"), ("All files", "*.*")])
        }
        
        export_format, filetypes = format_map.get(format_ext, ('txt', [("All files", "*.*")]))
        
        filepath = filedialog.asksaveasfilename(
            parent=self.window,
            title="Save Selection As",
            defaultextension=format_ext,
            initialfile=f"{clean_name}{format_ext}",
            filetypes=filetypes
        )
        
        if not filepath:
            return  # User cancelled
        
        # Use enhanced formatter to save with proper formatting
        success = save_formatted_document(
            filepath=filepath,
            content_text=formatted,
            title=f"Selection from: {self.doc_title}",
            source=self.source_info,
            imported_date=self.fetched_date,
            doc_class="selection",
            export_format=export_format,
            published_date=getattr(self, 'published_date', None)
        )
        
        if success:
            filename = filepath.split('/')[-1].split('\\')[-1]  # Get just filename
            self._set_status(f"✅ Selection saved to {filename}")
        else:
            messagebox.showerror("Error", "Failed to save selection. Check console for details.")

    def _show_copy_dialog(self):
        """
        Show a unified dialog to choose what content to copy to clipboard.
        Includes: Source Only, All Exchanges, Expanded Only, Complete (Source+Thread), Selection
        """
        # Create dialog window
        dialog = tk.Toplevel(self.window)
        dialog.title("Copy")
        dialog.transient(self.window)
        dialog.grab_set()
        
        # Center on parent window
        dialog.geometry("450x480")
        dialog_x = self.window.winfo_x() + (self.window.winfo_width() - 450) // 2
        dialog_y = self.window.winfo_y() + (self.window.winfo_height() - 480) // 2
        dialog.geometry(f"+{dialog_x}+{dialog_y}")
        
        # Main frame with padding
        main_frame = ttk.Frame(dialog, padding=15)
        main_frame.pack(fill='both', expand=True)
        
        # Title
        title_label = ttk.Label(main_frame, text="What do you want to copy?", 
                                font=('Segoe UI', 11, 'bold'))
        title_label.pack(anchor='w', pady=(0, 15))
        
        # Check if source document is available
        has_source = bool(self.current_document_text)
        
        # Check if thread exists
        has_thread = bool(self.current_thread and len(self.current_thread) > 0)
        
        # Default selection based on current mode
        if self.current_mode == 'source':
            default_choice = "source" if has_source else "thread"
        else:
            default_choice = "thread"
        
        # Variable to track selection
        copy_choice = tk.StringVar(value=default_choice)
        
        # Count expanded exchanges for the label
        exchanges = self._group_messages_into_exchanges()
        expanded_count = sum(1 for i in range(len(exchanges)) 
                           if self.exchange_expanded_state.get(i, True))
        total_count = len(exchanges)
        
        # Check if there's a text selection
        has_selection = False
        selection_word_count = 0
        try:
            selection = self.thread_text.get(tk.SEL_FIRST, tk.SEL_LAST)
            if selection.strip():
                has_selection = True
                selection_word_count = len(selection.split())
        except tk.TclError:
            pass
        
        # Option 1: Source Only
        opt1_frame = ttk.Frame(main_frame)
        opt1_frame.pack(fill='x', pady=5)
        
        source_rb = ttk.Radiobutton(opt1_frame, text="Source Only", variable=copy_choice, 
                                     value="source")
        source_rb.pack(anchor='w')
        
        if has_source:
            ttk.Label(opt1_frame, text="Just the original source document", 
                     foreground='gray').pack(anchor='w', padx=(20, 0))
        else:
            ttk.Label(opt1_frame, text="(Not available - no source document loaded)", 
                     foreground='red').pack(anchor='w', padx=(20, 0))
            source_rb.config(state='disabled')
        
        # Option 2: All Exchanges (Thread)
        opt2_frame = ttk.Frame(main_frame)
        opt2_frame.pack(fill='x', pady=5)
        
        thread_rb = ttk.Radiobutton(opt2_frame, text="All Exchanges", variable=copy_choice, 
                                     value="thread")
        thread_rb.pack(anchor='w')
        
        if has_thread:
            ttk.Label(opt2_frame, text=f"Copy all {total_count} exchange(s) to clipboard", 
                     foreground='gray').pack(anchor='w', padx=(20, 0))
        else:
            ttk.Label(opt2_frame, text="(Not available - no conversation yet)", 
                     foreground='red').pack(anchor='w', padx=(20, 0))
            thread_rb.config(state='disabled')
        
        # Option 3: Expanded Only
        opt3_frame = ttk.Frame(main_frame)
        opt3_frame.pack(fill='x', pady=5)
        
        expanded_text = f"Expanded Only ({expanded_count} of {total_count} exchanges)"
        expanded_rb = ttk.Radiobutton(opt3_frame, text=expanded_text, variable=copy_choice, 
                                       value="expanded")
        expanded_rb.pack(anchor='w')
        
        if has_thread and total_count > 0:
            ttk.Label(opt3_frame, text="Only exchanges you've expanded (collapsed ones are omitted)", 
                     foreground='gray').pack(anchor='w', padx=(20, 0))
        else:
            ttk.Label(opt3_frame, text="(Not available - no exchanges)", 
                     foreground='red').pack(anchor='w', padx=(20, 0))
            expanded_rb.config(state='disabled')
        
        # Option 4: Complete (Source + Thread)
        opt4_frame = ttk.Frame(main_frame)
        opt4_frame.pack(fill='x', pady=5)
        
        complete_rb = ttk.Radiobutton(opt4_frame, text="Complete: Source + Thread", 
                                       variable=copy_choice, value="complete")
        complete_rb.pack(anchor='w')
        
        if has_source and has_thread:
            ttk.Label(opt4_frame, text="Source document AND all conversation exchanges", 
                     foreground='gray').pack(anchor='w', padx=(20, 0))
        else:
            ttk.Label(opt4_frame, text="(Requires both source and thread)", 
                     foreground='red').pack(anchor='w', padx=(20, 0))
            complete_rb.config(state='disabled')
        
        # Option 5: Selection
        opt5_frame = ttk.Frame(main_frame)
        opt5_frame.pack(fill='x', pady=5)
        
        if has_selection:
            selection_text = f"Selection ({selection_word_count} words)"
            selection_rb = ttk.Radiobutton(opt5_frame, text=selection_text, 
                                           variable=copy_choice, value="selection")
            selection_rb.pack(anchor='w')
            ttk.Label(opt5_frame, text="Copy the currently selected text", 
                     foreground='gray').pack(anchor='w', padx=(20, 0))
        else:
            selection_rb = ttk.Radiobutton(opt5_frame, text="Selection", 
                                           variable=copy_choice, value="selection",
                                           state='disabled')
            selection_rb.pack(anchor='w')
            ttk.Label(opt5_frame, text="(No text selected - select text first to enable)", 
                     foreground='red').pack(anchor='w', padx=(20, 0))
        
        # Separator
        ttk.Separator(main_frame, orient='horizontal').pack(fill='x', pady=15)
        
        # Format selection
        format_frame = ttk.Frame(main_frame)
        format_frame.pack(fill='x', pady=5)
        
        ttk.Label(format_frame, text="Format:").pack(side='left')
        
        format_var = tk.StringVar(value="formatted")
        format_combo = ttk.Combobox(format_frame, textvariable=format_var, 
                                     values=["Plain Text", "Formatted (for Word/Email)"],
                                     state='readonly', width=22)
        format_combo.pack(side='left', padx=(10, 0))
        format_combo.current(1)  # Default to formatted
        
        # Buttons
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill='x', pady=(20, 0))
        
        def on_copy():
            choice = copy_choice.get()
            fmt = format_var.get()
            is_plain = "Plain" in fmt
            dialog.destroy()
            
            if choice == "source":
                self._copy_source_only(plain_text=is_plain)
            elif choice == "thread":
                if is_plain:
                    self._copy_thread()
                else:
                    self._copy_thread_formatted()
            elif choice == "expanded":
                if is_plain:
                    self._copy_expanded_only()
                else:
                    self._copy_expanded_only_formatted()
            elif choice == "complete":
                self._copy_complete(plain_text=is_plain)
            elif choice == "selection":
                if is_plain:
                    self._copy_selection_to_clipboard()
                else:
                    self._copy_selection_formatted()
        
        def on_cancel():
            dialog.destroy()
        
        ttk.Button(button_frame, text="Copy", command=on_copy, width=10).pack(side='right', padx=5)
        ttk.Button(button_frame, text="Cancel", command=on_cancel, width=10).pack(side='right')
        
        # Handle Enter and Escape
        dialog.bind('<Return>', lambda e: on_copy())
        dialog.bind('<Escape>', lambda e: on_cancel())
        
        # Focus on dialog
        dialog.focus_set()
        dialog.wait_window()

    def _show_save_as_dialog(self):
        """
        Show a dialog to choose what content to save, then proceed to file save dialog.
        """
        # Create dialog window
        dialog = tk.Toplevel(self.window)
        dialog.title("Save As")
        dialog.transient(self.window)
        dialog.grab_set()
        
        # Center on parent window
        dialog.geometry("450x460")
        dialog_x = self.window.winfo_x() + (self.window.winfo_width() - 450) // 2
        dialog_y = self.window.winfo_y() + (self.window.winfo_height() - 460) // 2
        dialog.geometry(f"+{dialog_x}+{dialog_y}")
        
        # Main frame with padding
        main_frame = ttk.Frame(dialog, padding=15)
        main_frame.pack(fill='both', expand=True)
        
        # Title
        title_label = ttk.Label(main_frame, text="What do you want to save?", 
                                font=('Segoe UI', 11, 'bold'))
        title_label.pack(anchor='w', pady=(0, 15))
        
        # Check if source document is available
        has_source = bool(self.current_document_text)
        
        # Check if thread exists
        has_thread = bool(self.current_thread and len(self.current_thread) > 0)
        
        # Default selection based on current mode
        if self.current_mode == 'source':
            default_choice = "source" if has_source else "thread"
        else:
            default_choice = "thread"
        
        # Variable to track selection
        save_choice = tk.StringVar(value=default_choice)
        
        # Count expanded exchanges for the label
        exchanges = self._group_messages_into_exchanges()
        expanded_count = sum(1 for i in range(len(exchanges)) 
                           if self.exchange_expanded_state.get(i, True))
        total_count = len(exchanges)
        
        # Option 1: Source Only
        opt1_frame = ttk.Frame(main_frame)
        opt1_frame.pack(fill='x', pady=5)
        
        source_rb = ttk.Radiobutton(opt1_frame, text="Source Only", variable=save_choice, 
                                     value="source")
        source_rb.pack(anchor='w')
        
        if has_source:
            ttk.Label(opt1_frame, text="Just the original source document", 
                     foreground='gray').pack(anchor='w', padx=(20, 0))
        else:
            ttk.Label(opt1_frame, text="(Not available - no source document loaded)", 
                     foreground='red').pack(anchor='w', padx=(20, 0))
            source_rb.config(state='disabled')
        
        # Option 2: Thread
        opt2_frame = ttk.Frame(main_frame)
        opt2_frame.pack(fill='x', pady=5)
        
        thread_rb = ttk.Radiobutton(opt2_frame, text="Thread", variable=save_choice, 
                                     value="thread")
        thread_rb.pack(anchor='w')
        
        if has_thread:
            ttk.Label(opt2_frame, text="Your questions and AI responses only", 
                     foreground='gray').pack(anchor='w', padx=(20, 0))
        else:
            ttk.Label(opt2_frame, text="(Not available - no conversation yet)", 
                     foreground='red').pack(anchor='w', padx=(20, 0))
            thread_rb.config(state='disabled')
        
        # Option 3: Expanded Only
        opt3_frame = ttk.Frame(main_frame)
        opt3_frame.pack(fill='x', pady=5)
        
        expanded_text = f"Expanded Only ({expanded_count} of {total_count} exchanges)"
        expanded_rb = ttk.Radiobutton(opt3_frame, text=expanded_text, variable=save_choice, 
                                       value="expanded")
        expanded_rb.pack(anchor='w')
        
        if has_thread and total_count > 0:
            ttk.Label(opt3_frame, text="Only exchanges you've expanded (collapsed ones are omitted)", 
                     foreground='gray').pack(anchor='w', padx=(20, 0))
        else:
            ttk.Label(opt3_frame, text="(Not available - no exchanges)", 
                     foreground='red').pack(anchor='w', padx=(20, 0))
            expanded_rb.config(state='disabled')
        
        # Option 4: Complete (disabled if no source)
        opt4_frame = ttk.Frame(main_frame)
        opt4_frame.pack(fill='x', pady=5)
        
        complete_rb = ttk.Radiobutton(opt4_frame, text="Complete: Source + Thread", 
                                       variable=save_choice, value="complete")
        complete_rb.pack(anchor='w')
        
        if has_source and has_thread:
            ttk.Label(opt4_frame, text="Original source document AND all conversation exchanges", 
                     foreground='gray').pack(anchor='w', padx=(20, 0))
        else:
            ttk.Label(opt4_frame, text="(Requires both source and thread)", 
                     foreground='red').pack(anchor='w', padx=(20, 0))
            complete_rb.config(state='disabled')
        
        # Separator
        ttk.Separator(main_frame, orient='horizontal').pack(fill='x', pady=15)
        
        # Format selection
        format_frame = ttk.Frame(main_frame)
        format_frame.pack(fill='x', pady=5)
        
        ttk.Label(format_frame, text="Format:").pack(side='left')
        
        format_var = tk.StringVar(value=".docx")
        format_combo = ttk.Combobox(format_frame, textvariable=format_var, 
                                     values=[".docx", ".txt", ".rtf", ".pdf"],
                                     state='readonly', width=10)
        format_combo.pack(side='left', padx=(10, 0))
        
        # Buttons
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill='x', pady=(20, 0))
        
        def on_save():
            choice = save_choice.get()
            fmt = format_var.get()
            dialog.destroy()
            
            if choice == "source":
                self._save_source_only(fmt)
            elif choice == "thread":
                self._save_thread(fmt)
            elif choice == "expanded":
                self._save_expanded_only(fmt)
            elif choice == "complete":
                self._save_complete(fmt)
        
        def on_cancel():
            dialog.destroy()
        
        ttk.Button(button_frame, text="Save", command=on_save, width=10).pack(side='right', padx=5)
        ttk.Button(button_frame, text="Cancel", command=on_cancel, width=10).pack(side='right')
        
        # Handle Enter and Escape
        dialog.bind('<Return>', lambda e: on_save())
        dialog.bind('<Escape>', lambda e: on_cancel())
        
        # Focus on dialog
        dialog.focus_set()
        dialog.wait_window()

    def _save_source_only(self, format_ext=None):
        """Save just the source document to a file."""
        if not self.current_document_text:
            messagebox.showwarning("No Source", "No source document available to save.")
            return
        
        # Save any pending source edits first
        try:
            self._save_source_edits()
        except ValueError:
            return
        except Exception:
            pass
        
        # Set up file dialog based on requested format
        if format_ext:
            ext, filetypes = get_file_extension_and_types(format_ext.lstrip('.'))
            default_ext = ext
        else:
            default_ext = ".txt"
            filetypes = [
                ("Text files", "*.txt"),
                ("Word Document", "*.docx"),
                ("RTF files", "*.rtf"),
                ("PDF files", "*.pdf"),
                ("All files", "*.*")
            ]
        
        # Get clean filename from document title
        clean_title = get_clean_filename(self.doc_title)
        
        file_path = filedialog.asksaveasfilename(
            defaultextension=default_ext,
            filetypes=filetypes,
            initialfile=f"{clean_title}_source{format_ext or '.txt'}"
        )
        
        if not file_path:
            return
        
        try:
            ext = os.path.splitext(file_path)[1].lower().lstrip('.')
            if ext not in ['txt', 'docx', 'rtf', 'pdf']:
                ext = 'txt'
            
            # Save using document_export's export_document function
            metadata = {
                'title': self.doc_title,
                'source': self.source_info,
                'published_date': getattr(self, 'published_date', None) if getattr(self, 'published_date', 'N/A') != 'N/A' else None,
                'imported_date': self.fetched_date,
                'doc_class': 'source'
            }
            success, msg = export_document(file_path, self.current_document_text, ext, metadata, show_messages=False)
            if not success:
                raise Exception(msg)
            
            self._set_status(f"✅ Source saved to {os.path.basename(file_path)}")
            
        except Exception as e:
            messagebox.showerror("Save Error", f"Failed to save source:\n{str(e)}")

    def _save_thread(self, format_ext=None):
        """Save content to file - source document when in source mode, thread when in conversation mode"""
        # First, save any pending edits to ensure we export the latest content
        try:
            if self.current_mode == 'conversation' and self.current_thread:
                self._save_edits_to_thread()
            elif self.current_mode == 'source':
                self._save_source_edits()
        except ValueError:
            # User cancelled the save edits dialog
            return
        except Exception as e:
            pass
            # Log but continue - edits may not need saving
        
        # Determine what we're saving based on current mode
        is_source_mode = self.current_mode == 'source'
        
        # Set up file dialog based on requested format
        if format_ext:
            ext, filetypes = get_file_extension_and_types(format_ext.lstrip('.'))
            default_ext = ext
        else:
            # Show all formats if no specific format requested
            default_ext = ".txt"
            filetypes = [
                ("Text files", "*.txt"),
                ("Word Document", "*.docx"),
                ("RTF files", "*.rtf"),
                ("PDF files", "*.pdf"),
                ("All files", "*.*")
            ]
        
        # Get clean filename from document title
        clean_title = get_clean_filename(self.doc_title)
        
        # Different filename suffix based on mode
        suffix = "" if is_source_mode else "_thread"
        
        file_path = filedialog.asksaveasfilename(
            defaultextension=default_ext,
            filetypes=filetypes,
            initialfile=f"{clean_title}{suffix}{format_ext or '.txt'}"
        )
        
        if not file_path:
            return
        
        try:
            # Get file extension and determine format
            ext = os.path.splitext(file_path)[1].lower().lstrip('.')
            if ext not in ['txt', 'docx', 'rtf', 'pdf']:
                ext = 'txt'
            
            if is_source_mode:
                # === SOURCE MODE: Save the source document text ===
                from doc_formatter import save_formatted_document
                
                success = save_formatted_document(
                    filepath=file_path,
                    content_text=self.current_document_text or "",
                    title=self.doc_title,
                    source=self.source_info,
                    imported_date=self.fetched_date,
                    doc_class="source",
                    export_format=ext,
                    published_date=getattr(self, 'published_date', None)
                )
                
                if success:
                    filename = os.path.basename(file_path)
                    self._set_status(f"✅ Source document saved to {filename}")
                else:
                    messagebox.showerror("Error", "Failed to save source document. Check console for details.")
            else:
                # === CONVERSATION MODE: Save the conversation thread ===
                # Build thread metadata
                thread_metadata = {
                    'doc_title': self.doc_title,
                    'source_info': self.source_info,
                    'published_date': getattr(self, 'published_date', None),
                    'fetched_date': self.fetched_date,
                    'provider': self.provider_var.get() if self.provider_var else 'N/A',
                    'model': self.model_var.get() if self.model_var else 'N/A',
                    'message_count': self.thread_message_count
                }
                
                # Use consolidated export function
                success, result = export_conversation_thread(
                    filepath=file_path,
                    format=ext,
                    thread_messages=self.current_thread,
                    thread_metadata=thread_metadata,
                    show_messages=True
                )
                
                if success:
                    filename = os.path.basename(file_path)
                    self._set_status(f"✅ Conversation thread saved to {filename}")
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save:\n{str(e)}")
            import traceback
            traceback.print_exc()

    def _save_expanded_only(self, format_ext='.txt'):
        """
        Save only the expanded exchanges to a file.
        Collapsed exchanges are completely omitted.
        
        Args:
            format_ext: File extension ('.txt', '.docx', '.rtf', or '.pdf')
        """
        exchanges = self._group_messages_into_exchanges()
        
        if not exchanges:
            messagebox.showwarning("No Content", "No exchanges to save.")
            return
        
        # Find which exchanges are expanded
        expanded_indices = [i for i in range(len(exchanges)) 
                          if self.exchange_expanded_state.get(i, True)]
        
        if not expanded_indices:
            messagebox.showwarning("No Expanded Exchanges", 
                "No exchanges are currently expanded.\n\n"
                "Please expand the exchange(s) you want to save, then try again.")
            return
        
        # Build content from expanded exchanges only
        expanded_thread = []
        for idx in expanded_indices:
            exchange = exchanges[idx]
            if 'user' in exchange:
                expanded_thread.append(exchange['user'])
            if 'assistant' in exchange:
                expanded_thread.append(exchange['assistant'])
        
        # Set up file dialog - support all 4 formats
        ext = format_ext.lstrip('.')
        format_filetypes = {
            'docx': [("Word Document", "*.docx"), ("All files", "*.*")],
            'txt': [("Text files", "*.txt"), ("All files", "*.*")],
            'rtf': [("RTF Document", "*.rtf"), ("All files", "*.*")],
            'pdf': [("PDF Document", "*.pdf"), ("All files", "*.*")]
        }
        filetypes = format_filetypes.get(ext, [("All files", "*.*")])
        
        # Get clean filename
        clean_title = get_clean_filename(self.doc_title)
        count = len(expanded_indices)
        suffix = f"_exchange{'s' if count != 1 else ''}_{count}"
        
        file_path = filedialog.asksaveasfilename(
            defaultextension=format_ext,
            filetypes=filetypes,
            initialfile=f"{clean_title}{suffix}{format_ext}"
        )
        
        if not file_path:
            return
        
        try:
            # Build metadata
            thread_metadata = {
                'doc_title': self.doc_title,
                'source_info': self.source_info,
                'published_date': getattr(self, 'published_date', None),
                'fetched_date': self.fetched_date,
                'model': self.model_var.get(),
                'provider': self.provider_var.get(),
                'message_count': count,
                'note': f"Exported {count} of {len(exchanges)} exchange{'s' if len(exchanges) != 1 else ''}"
            }
            
            # Use consolidated export function
            success, result = export_conversation_thread(
                filepath=file_path,
                format=ext,
                thread_messages=expanded_thread,
                thread_metadata=thread_metadata,
                show_messages=True
            )
            
            if success:
                filename = os.path.basename(file_path)
                self._set_status(f"✅ Saved {count} expanded exchange{'s' if count != 1 else ''} to {filename}")
            else:
                messagebox.showerror("Error", f"Failed to save: {result}")
                
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save:\n{str(e)}")
            import traceback
            traceback.print_exc()

    def _save_complete(self, format_ext='.txt'):
        """
        Save both the source document AND the conversation thread to a single file.
        This provides complete context for sharing.
        
        Args:
            format_ext: File extension ('.txt', '.docx', '.rtf', or '.pdf')
        """
        # Check we have both source and conversation
        if not self.current_document_text:
            messagebox.showwarning("No Source", 
                "No source document available.\n\n"
                "Use the regular Save As options to save just the conversation.")
            return
        
        if not self.current_thread or len(self.current_thread) == 0:
            messagebox.showwarning("No Conversation", 
                "No conversation exchanges available.\n\n"
                "Use the regular Save As options to save just the source document.")
            return
        
        # Set up file dialog - support all 4 formats
        ext = format_ext.lstrip('.')
        format_filetypes = {
            'docx': [("Word Document", "*.docx"), ("All files", "*.*")],
            'txt': [("Text files", "*.txt"), ("All files", "*.*")],
            'rtf': [("RTF Document", "*.rtf"), ("All files", "*.*")],
            'pdf': [("PDF Document", "*.pdf"), ("All files", "*.*")]
        }
        filetypes = format_filetypes.get(ext, [("All files", "*.*")])
        
        # Get clean filename
        clean_title = get_clean_filename(self.doc_title)
        
        file_path = filedialog.asksaveasfilename(
            defaultextension=format_ext,
            filetypes=filetypes,
            initialfile=f"{clean_title}_complete{format_ext}"
        )
        
        if not file_path:
            return
        
        try:
            if ext == 'docx':
                self._save_complete_docx(file_path)
            elif ext == 'pdf':
                self._save_complete_pdf(file_path)
            elif ext == 'rtf':
                self._save_complete_rtf(file_path)
            else:
                self._save_complete_txt(file_path)
                
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save:\n{str(e)}")
            import traceback
            traceback.print_exc()

    def _save_complete_txt(self, file_path: str):
        """Save complete document (source + thread) as plain text."""
        lines = []
        
        # === HEADER ===
        lines.append("=" * 70)
        lines.append("COMPLETE DOCUMENT EXPORT")
        lines.append("=" * 70)
        lines.append("")
        lines.append(f"Title: {self.doc_title}")
        lines.append(f"Source: {self.source_info}")
        if hasattr(self, 'published_date') and self.published_date and self.published_date != 'N/A':
            lines.append(f"Published: {self.published_date}")
        lines.append(f"Imported: {self.fetched_date}")
        lines.append(f"Exported: {datetime.datetime.now().strftime('%d-%b-%Y %H:%M')}")
        lines.append("")
        
        # === SOURCE DOCUMENT ===
        lines.append("=" * 70)
        lines.append("SOURCE DOCUMENT")
        lines.append("=" * 70)
        lines.append("")
        lines.append(self.current_document_text or "(No source text available)")
        lines.append("")
        
        # === CONVERSATION THREAD ===
        lines.append("=" * 70)
        lines.append("CONVERSATION THREAD")
        lines.append("=" * 70)
        lines.append("")
        
        exchanges = self._group_messages_into_exchanges()
        for i, exchange in enumerate(exchanges):
            lines.append(f"--- Exchange {i + 1} ---")
            lines.append("")
            
            # User message
            user_msg = exchange.get('user', {})
            if user_msg:
                timestamp = user_msg.get('timestamp', '')
                time_str = f" [{timestamp}]" if timestamp else ""
                lines.append(f"YOU{time_str}:")
                lines.append(user_msg.get('content', ''))
                lines.append("")
            
            # Assistant message
            assistant_msg = exchange.get('assistant', {})
            if assistant_msg:
                provider = assistant_msg.get('provider', 'AI')
                model = assistant_msg.get('model', '')
                timestamp = assistant_msg.get('timestamp', '')
                time_str = f" [{timestamp}]" if timestamp else ""
                
                if model and model != provider:
                    label = f"{provider} ({model}){time_str}:"
                else:
                    label = f"{provider}{time_str}:"
                
                lines.append(label)
                lines.append(assistant_msg.get('content', ''))
                lines.append("")
            
            lines.append("")
        
        # Write to file
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(lines))
        
        filename = os.path.basename(file_path)
        self._set_status(f"✅ Complete document saved to {filename}")

    def _save_complete_rtf(self, file_path: str):
        """Save complete document (source + thread) as RTF."""
        # RTF is essentially text with formatting codes
        # We'll create a simple RTF with basic formatting
        
        def rtf_escape(text: str) -> str:
            """Escape special RTF characters."""
            if not text:
                return ""
            # Escape backslash, curly braces, and handle unicode
            text = text.replace('\\', '\\\\')
            text = text.replace('{', '\\{')
            text = text.replace('}', '\\}')
            # Handle line breaks
            text = text.replace('\n', '\\par\n')
            return text
        
        rtf_lines = []
        
        # RTF header
        rtf_lines.append('{\\rtf1\\ansi\\deff0')
        rtf_lines.append('{\\fonttbl{\\f0 Calibri;}{\\f1 Arial;}}')
        rtf_lines.append('\\f0\\fs22')  # Default font and size
        
        # === TITLE ===
        rtf_lines.append('\\pard\\qc\\b\\fs32')  # Centered, bold, larger
        rtf_lines.append(rtf_escape(self.doc_title))
        rtf_lines.append('\\par\\b0\\fs22\\ql')  # Reset to normal, left-aligned
        rtf_lines.append('\\par')
        
        # === METADATA ===
        rtf_lines.append('\\pard\\fs20\\cf1')  # Smaller, gray text
        rtf_lines.append(f'\\b Source:\\b0  {rtf_escape(self.source_info)}\\par')
        if hasattr(self, 'published_date') and self.published_date and self.published_date != 'N/A':
            rtf_lines.append(f'\\b Published:\\b0  {self.published_date}\\par')
        rtf_lines.append(f'\\b Imported:\\b0  {self.fetched_date}\\par')
        rtf_lines.append(f'\\b Exported:\\b0  {datetime.datetime.now().strftime("%d-%b-%Y %H:%M")}\\par')
        rtf_lines.append('\\fs22\\cf0\\par')  # Reset font size and color
        
        # === SOURCE DOCUMENT SECTION ===
        rtf_lines.append('\\pard\\brdrb\\brdrs\\brdrw10\\par')  # Horizontal line
        rtf_lines.append('\\par\\b\\fs28 SOURCE DOCUMENT\\b0\\fs22\\par')
        rtf_lines.append('\\brdrb\\brdrs\\brdrw10\\par\\par')
        
        # Add source text
        source_text = self.current_document_text or "(No source text available)"
        for para_text in source_text.split('\n\n'):
            if para_text.strip():
                rtf_lines.append(rtf_escape(para_text.strip()))
                rtf_lines.append('\\par\\par')
        
        # === CONVERSATION THREAD SECTION ===
        rtf_lines.append('\\par\\brdrb\\brdrs\\brdrw10\\par')
        rtf_lines.append('\\par\\b\\fs28 CONVERSATION THREAD\\b0\\fs22\\par')
        rtf_lines.append('\\brdrb\\brdrs\\brdrw10\\par\\par')
        
        # Add exchanges
        exchanges = self._group_messages_into_exchanges()
        for i, exchange in enumerate(exchanges):
            # Exchange header
            rtf_lines.append(f'\\b Exchange {i + 1}\\b0\\par')
            rtf_lines.append('\\par')
            
            # User message
            user_msg = exchange.get('user', {})
            if user_msg:
                timestamp = user_msg.get('timestamp', '')
                time_str = f" [{timestamp}]" if timestamp else ""
                rtf_lines.append(f'\\b\\cf2 YOU{time_str}:\\cf0\\b0\\par')
                rtf_lines.append(rtf_escape(user_msg.get('content', '')))
                rtf_lines.append('\\par\\par')
            
            # Assistant message
            assistant_msg = exchange.get('assistant', {})
            if assistant_msg:
                provider = assistant_msg.get('provider', 'AI')
                model = assistant_msg.get('model', '')
                timestamp = assistant_msg.get('timestamp', '')
                time_str = f" [{timestamp}]" if timestamp else ""
                
                if model and model != provider:
                    label = f"{provider} ({model}){time_str}:"
                else:
                    label = f"{provider}{time_str}:"
                
                rtf_lines.append(f'\\b\\cf1 {rtf_escape(label)}\\cf0\\b0\\par')
                
                # Add AI response content
                content = assistant_msg.get('content', '')
                for para_text in content.split('\n\n'):
                    if para_text.strip():
                        rtf_lines.append(rtf_escape(para_text.strip()))
                        rtf_lines.append('\\par\\par')
            
            # Add separator between exchanges
            if i < len(exchanges) - 1:
                rtf_lines.append('\\par\\pard\\qc\\emdash\\emdash\\emdash\\emdash\\emdash\\ql\\par\\par')
        
        # Close RTF
        rtf_lines.append('}')
        
        # Write to file
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(rtf_lines))
        
        filename = os.path.basename(file_path)
        self._set_status(f"✅ Complete document saved to {filename}")

    def _save_complete_docx(self, file_path: str):
        """Save complete document (source + thread) as Word document."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            messagebox.showerror("Missing Module", 
                "python-docx is required for Word export.\n\n"
                "Install with: pip install python-docx")
            return
        
        doc = Document()
        
        # === TITLE ===
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(self.doc_title)
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # === METADATA ===
        doc.add_paragraph()
        meta_para = doc.add_paragraph()
        meta_para.add_run("Source: ").bold = True
        meta_para.add_run(self.source_info)
        
        if hasattr(self, 'published_date') and self.published_date and self.published_date != 'N/A':
            meta_para = doc.add_paragraph()
            meta_para.add_run("Published: ").bold = True
            meta_para.add_run(str(self.published_date))
        
        meta_para = doc.add_paragraph()
        meta_para.add_run("Imported: ").bold = True
        meta_para.add_run(self.fetched_date)
        
        meta_para = doc.add_paragraph()
        meta_para.add_run("Exported: ").bold = True
        meta_para.add_run(datetime.datetime.now().strftime('%d-%b-%Y %H:%M'))
        
        # === SOURCE DOCUMENT SECTION ===
        doc.add_paragraph()
        doc.add_paragraph("_" * 50)
        
        section_heading = doc.add_paragraph()
        heading_run = section_heading.add_run("SOURCE DOCUMENT")
        heading_run.bold = True
        heading_run.font.size = Pt(14)
        
        doc.add_paragraph("_" * 50)
        doc.add_paragraph()
        
        # Add source text (split into paragraphs)
        source_text = self.current_document_text or "(No source text available)"
        for para_text in source_text.split('\n\n'):
            if para_text.strip():
                doc.add_paragraph(para_text.strip())
        
        # === CONVERSATION THREAD SECTION ===
        doc.add_paragraph()
        doc.add_paragraph("_" * 50)
        
        section_heading = doc.add_paragraph()
        heading_run = section_heading.add_run("CONVERSATION THREAD")
        heading_run.bold = True
        heading_run.font.size = Pt(14)
        
        doc.add_paragraph("_" * 50)
        doc.add_paragraph()
        
        # Add exchanges
        exchanges = self._group_messages_into_exchanges()
        for i, exchange in enumerate(exchanges):
            # Exchange header
            ex_header = doc.add_paragraph()
            ex_header.add_run(f"Exchange {i + 1}").bold = True
            
            # User message
            user_msg = exchange.get('user', {})
            if user_msg:
                timestamp = user_msg.get('timestamp', '')
                time_str = f" [{timestamp}]" if timestamp else ""
                
                user_para = doc.add_paragraph()
                user_label = user_para.add_run(f"YOU{time_str}:")
                user_label.bold = True
                
                doc.add_paragraph(user_msg.get('content', ''))
            
            # Assistant message
            assistant_msg = exchange.get('assistant', {})
            if assistant_msg:
                provider = assistant_msg.get('provider', 'AI')
                model = assistant_msg.get('model', '')
                timestamp = assistant_msg.get('timestamp', '')
                time_str = f" [{timestamp}]" if timestamp else ""
                
                if model and model != provider:
                    label = f"{provider} ({model}){time_str}:"
                else:
                    label = f"{provider}{time_str}:"
                
                ai_para = doc.add_paragraph()
                ai_label = ai_para.add_run(label)
                ai_label.bold = True
                
                # Add AI response content (handle markdown-ish formatting)
                content = assistant_msg.get('content', '')
                for para_text in content.split('\n\n'):
                    if para_text.strip():
                        doc.add_paragraph(para_text.strip())
            
            # Add separator between exchanges
            if i < len(exchanges) - 1:
                doc.add_paragraph()
                sep = doc.add_paragraph()
                sep.add_run("─" * 30)
                doc.add_paragraph()
        
        # Save the document
        doc.save(file_path)
        
        filename = os.path.basename(file_path)
        self._set_status(f"✅ Complete document saved to {filename}")

    def _save_complete_pdf(self, file_path: str):
        """Save complete document (source + thread) as PDF."""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.lib.enums import TA_CENTER
        except ImportError:
            messagebox.showerror("Missing Module", 
                "reportlab is required for PDF export.\n\n"
                "Install with: pip install reportlab")
            return
        
        # Create the PDF document
        doc = SimpleDocTemplate(
            file_path,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        # Get styles
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            alignment=TA_CENTER,
            spaceAfter=12
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceBefore=20,
            spaceAfter=10
        )
        
        meta_style = ParagraphStyle(
            'CustomMeta',
            parent=styles['Normal'],
            fontSize=10,
            textColor='gray'
        )
        
        normal_style = styles['Normal']
        
        user_style = ParagraphStyle(
            'UserStyle',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            textColor='#2E4053',
            spaceBefore=12
        )
        
        ai_style = ParagraphStyle(
            'AIStyle',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            textColor='#16537E',
            spaceBefore=12
        )
        
        # Build the document content
        story = []
        
        # === TITLE ===
        story.append(Paragraph(self._escape_html(self.doc_title), title_style))
        story.append(Spacer(1, 12))
        
        # === METADATA ===
        story.append(Paragraph(f"<b>Source:</b> {self._escape_html(self.source_info)}", meta_style))
        if hasattr(self, 'published_date') and self.published_date and self.published_date != 'N/A':
            story.append(Paragraph(f"<b>Published:</b> {self.published_date}", meta_style))
        story.append(Paragraph(f"<b>Imported:</b> {self.fetched_date}", meta_style))
        story.append(Paragraph(f"<b>Exported:</b> {datetime.datetime.now().strftime('%d-%b-%Y %H:%M')}", meta_style))
        story.append(Spacer(1, 20))
        
        # === SOURCE DOCUMENT SECTION ===
        story.append(Paragraph("─" * 50, normal_style))
        story.append(Paragraph("SOURCE DOCUMENT", heading_style))
        story.append(Paragraph("─" * 50, normal_style))
        story.append(Spacer(1, 12))
        
        # Add source text (split into paragraphs, escape HTML)
        source_text = self.current_document_text or "(No source text available)"
        for para_text in source_text.split('\n\n'):
            if para_text.strip():
                # Escape any HTML-like characters
                safe_text = self._escape_html(para_text.strip())
                story.append(Paragraph(safe_text, normal_style))
                story.append(Spacer(1, 6))
        
        # === CONVERSATION THREAD SECTION ===
        story.append(Spacer(1, 20))
        story.append(Paragraph("─" * 50, normal_style))
        story.append(Paragraph("CONVERSATION THREAD", heading_style))
        story.append(Paragraph("─" * 50, normal_style))
        story.append(Spacer(1, 12))
        
        # Add exchanges
        exchanges = self._group_messages_into_exchanges()
        for i, exchange in enumerate(exchanges):
            # Exchange header
            story.append(Paragraph(f"<b>Exchange {i + 1}</b>", normal_style))
            story.append(Spacer(1, 6))
            
            # User message
            user_msg = exchange.get('user', {})
            if user_msg:
                timestamp = user_msg.get('timestamp', '')
                time_str = f" [{timestamp}]" if timestamp else ""
                story.append(Paragraph(f"YOU{time_str}:", user_style))
                
                content = user_msg.get('content', '')
                safe_content = self._escape_html(content)
                story.append(Paragraph(safe_content, normal_style))
                story.append(Spacer(1, 6))
            
            # Assistant message
            assistant_msg = exchange.get('assistant', {})
            if assistant_msg:
                provider = assistant_msg.get('provider', 'AI')
                model = assistant_msg.get('model', '')
                timestamp = assistant_msg.get('timestamp', '')
                time_str = f" [{timestamp}]" if timestamp else ""
                
                if model and model != provider:
                    label = f"{provider} ({model}){time_str}:"
                else:
                    label = f"{provider}{time_str}:"
                
                story.append(Paragraph(label, ai_style))
                
                # Add AI response content (split into paragraphs)
                content = assistant_msg.get('content', '')
                for para_text in content.split('\n\n'):
                    if para_text.strip():
                        safe_text = self._escape_html(para_text.strip())
                        story.append(Paragraph(safe_text, normal_style))
                        story.append(Spacer(1, 6))
            
            # Add separator between exchanges
            if i < len(exchanges) - 1:
                story.append(Spacer(1, 12))
                story.append(Paragraph("─" * 30, normal_style))
                story.append(Spacer(1, 12))
        
        # Build the PDF
        doc.build(story)
        
        filename = os.path.basename(file_path)
        self._set_status(f"✅ Complete document saved to {filename}")

    def _can_start_new_conversation(self) -> bool:
        """
        Check if the "New Conversation (Same Source)" button should be shown.
        
        Returns True if:
        - This is a response/product document (not a source document)
        - There is a source_document_id (the original source is known)
        - The callback is available
        """
        # Must be a response/product document
        if self.document_class not in ['response', 'product', 'processed_output']:
            return False
        
        # Must have a source document ID
        if not self.source_document_id:
            return False
        
        # Must have the callback
        if not self.on_start_new_conversation:
            return False
        
        return True
    
    def _start_new_conversation(self):
        """
        Start a new conversation using the original source document.
        
        This closes the Thread Viewer and loads the source document in the main UI,
        clearing the current conversation so the user can start fresh.
        """
        # Check for unsaved edits in the thread display
        # (The thread text is editable, so user may have made changes)
        
        # Confirm with user
        result = messagebox.askyesnocancel(
            "New Conversation (Same Source)",
            "Start a new conversation using the original source document?\n\n"
            "• The current conversation will be saved\n"
            "• The original source document will be loaded\n"
            "• You can then ask new questions about it\n\n"
            "Continue?"
        )
        
        if result is None:  # Cancel
            return
        
        if result:  # Yes
            # Save current thread first (via clear_thread callback which saves)
            if self.on_clear_thread:
                self.on_clear_thread()
            
            # Call the callback to load the source document
            if self.on_start_new_conversation:
                success = self.on_start_new_conversation(self.source_document_id)
                if success:
                    # Close this window
                    self.window.destroy()
                # If not successful, the callback shows an error message
                # and we keep the window open

def show_thread_viewer(
    parent: tk.Tk,
    current_thread: List[Dict],
    thread_message_count: int,
    current_document_id: Optional[str],
    current_document_text: str,
    current_document_source: str,
    model_var: tk.StringVar,
    provider_var: tk.StringVar,
    api_key_var: tk.StringVar,
    config: Dict,
    on_followup_complete: Optional[Callable[[str, str], None]] = None,
    on_clear_thread: Optional[Callable[[], None]] = None,
    refresh_library: Optional[Callable[[], None]] = None,
    get_ai_handler: Optional[Callable] = None,
    build_threaded_messages: Optional[Callable[[str], List[Dict]]] = None,
    add_message_to_thread: Optional[Callable[[str, str], None]] = None,
    attachment_manager: Optional[Any] = None,
    font_size: int = 10,
    # New parameters for "New Conversation (Same Source)" feature
    document_class: str = "source",
    source_document_id: Optional[str] = None,
    on_start_new_conversation: Optional[Callable[[str], bool]] = None,
    # Unified Viewer callback
    on_mode_change: Optional[Callable[[str], None]] = None,
    # Chunking support for initial prompts from Source Mode
    process_with_chunking: Optional[Callable] = None,
    current_entries: Optional[List] = None,
    current_document_type: str = "text",
    # Initial mode selection
    initial_mode: Optional[str] = None,
    # NEW: Multi-source document support
    source_documents: Optional[List[Dict]] = None,
    # NEW: Reference to main app for context synchronization
    app: Optional[object] = None,
) -> ThreadViewerWindow:
    """
    Convenience function to show the unified viewer window.
    
    The unified viewer can display either:
    - Source documents (in prose format)
    - Conversation threads (in chat format)
    
    Args:
        source_documents: Optional list of source document dicts, each containing:
            - 'title': Display title for the document
            - 'text': Full text content
            - 'source': Source URL or file path
            - 'char_count': (optional) Character count for display
        
        If source_documents is not provided, falls back to current_document_text.
    
    Returns the ThreadViewerWindow instance.
    """
    # Check for content: document OR attachments OR standalone conversation
    has_document = bool(current_document_text) or (source_documents and len(source_documents) > 0)
    has_attachments = attachment_manager and attachment_manager.get_attachment_count() > 0
    has_conversation = thread_message_count > 0 and current_thread and len(current_thread) > 0
    
    # Unified viewer can show source document OR conversation
    # Only show error if there's absolutely no content
    if not has_document and not has_attachments and not has_conversation:
        messagebox.showinfo("No Content", "Please load a document or add attachments first.")
        return None
    
    # For attachments-only or standalone mode, create a placeholder document text
    effective_document_text = current_document_text
    effective_document_source = current_document_source
    effective_source_documents = source_documents
    
    if not has_document and has_attachments:
        # Create summary for attachments-only mode
        att_count = attachment_manager.get_attachment_count()
        att_names = [att.get('filename', 'Unknown') for att in attachment_manager.attachments]
        
        # Create source_documents list from attachments
        effective_source_documents = []
        for att in attachment_manager.attachments:
            effective_source_documents.append({
                'title': att.get('filename', 'Unknown'),
                'text': att.get('text', ''),
                'source': att.get('path', att.get('source', '')),
                'char_count': len(att.get('text', ''))
            })
        
        effective_document_text = f"[Attachments-only mode: {att_count} document(s)]\n\nDocuments:\n" + "\n".join(f"  - {name}" for name in att_names)
        effective_document_source = f"Multiple attachments ({att_count} documents)"
        
    elif not has_document and not has_attachments and has_conversation:
        # Standalone conversation mode - no source document
        effective_document_text = "[Standalone conversation - no source document]"
        effective_document_source = "Standalone conversation"
    
    return ThreadViewerWindow(
        parent=parent,
        current_thread=current_thread,
        thread_message_count=thread_message_count,
        current_document_id=current_document_id,
        current_document_text=effective_document_text,
        current_document_source=effective_document_source,
        model_var=model_var,
        provider_var=provider_var,
        api_key_var=api_key_var,
        config=config,
        on_followup_complete=on_followup_complete,
        on_clear_thread=on_clear_thread,
        refresh_library=refresh_library,
        get_ai_handler=get_ai_handler,
        build_threaded_messages=build_threaded_messages,
        add_message_to_thread=add_message_to_thread,
        font_size=font_size,
        # New parameters for "New Conversation (Same Source)" feature
        document_class=document_class,
        source_document_id=source_document_id,
        on_start_new_conversation=on_start_new_conversation,
        on_mode_change=on_mode_change,
        # Chunking support
        process_with_chunking=process_with_chunking,
        current_entries=current_entries,
        current_document_type=current_document_type,
        # Initial mode
        initial_mode=initial_mode,
        # Multi-source support
        source_documents=effective_source_documents,
        # App reference for context sync
        app=app,
    )
