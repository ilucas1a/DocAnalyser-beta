"""
doc_formatter.py

Enhanced document formatting for DocAnalyser exports.
Handles proper formatting of conversation threads and AI responses
with support for markdown-style formatting in DOCX, RTF, PDF, and TXT.

Usage:
    from doc_formatter import save_formatted_document
"""

import os
import re
import datetime
from tkinter import filedialog, messagebox


def parse_markdown_text(text: str) -> list:
    """
    Parse text with markdown-style formatting into structured elements.
    
    Returns a list of dicts with 'type' and 'content' keys:
    - type: 'heading2', 'heading3', 'bullet', 'numbered', 'blockquote', 'paragraph'
    - content: the text content (may include inline formatting markers)
    """
    elements = []
    lines = text.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Skip empty lines
        if not stripped:
            i += 1
            continue
        
        # Heading 2: ## Title
        if stripped.startswith('## '):
            elements.append({'type': 'heading2', 'content': stripped[3:]})
            i += 1
            continue
        
        # Heading 3: ### Title
        if stripped.startswith('### '):
            elements.append({'type': 'heading3', 'content': stripped[4:]})
            i += 1
            continue
        
        # Horizontal rule: ---
        if stripped == '---':
            elements.append({'type': 'hr', 'content': ''})
            i += 1
            continue
        
        # Blockquote: > text
        if stripped.startswith('> '):
            elements.append({'type': 'blockquote', 'content': stripped[2:]})
            i += 1
            continue
        
        # Numbered list: 1. item or 2. item etc
        if re.match(r'^\d+\.\s+', stripped):
            content = re.sub(r'^\d+\.\s+', '', stripped)
            elements.append({'type': 'numbered', 'content': content})
            i += 1
            continue
        
        # Bullet list: - item or * item
        if stripped.startswith('- ') or stripped.startswith('* '):
            content = stripped[2:]
            # Check for nested items (indented)
            indent_level = len(line) - len(line.lstrip())
            elements.append({'type': 'bullet', 'content': content, 'indent': indent_level})
            i += 1
            continue
        
        # Regular paragraph
        elements.append({'type': 'paragraph', 'content': stripped})
        i += 1
    
    return elements


def add_formatted_paragraph(doc, text: str, style='Normal', is_italic=False):
    """
    Add a paragraph with inline markdown formatting (**bold**, *italic*).
    
    Args:
        doc: python-docx Document object
        text: Text that may contain **bold** or *italic* markers
        style: Paragraph style name
        is_italic: If True, make the entire paragraph italic
    """
    para = doc.add_paragraph(style=style)
    
    if is_italic:
        # Entire paragraph is italic
        run = para.add_run(text)
        run.italic = True
        return para
    
    # Parse inline formatting: **bold** and *italic*
    # Pattern matches **bold**, *italic*, or regular text
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|[^*]+)'
    parts = re.findall(pattern, text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            # Italic text
            run = para.add_run(part[1:-1])
            run.italic = True
        else:
            # Regular text
            para.add_run(part)
    
    return para


def save_formatted_docx(
    filepath: str,
    content_text: str,
    title: str = "Document",
    source: str = "Unknown",
    imported_date: str = "Unknown",
    doc_class: str = "selection",
    is_conversation: bool = False,
    published_date: str = None
) -> bool:
    """
    Save content as a properly formatted DOCX file with metadata header.
    
    Handles:
    - Document metadata header (title, source, dates)
    - Markdown headings (## and ###)
    - Bold (**text**) and italic (*text*)
    - Bullet and numbered lists
    - Block quotes
    - Conversation formatting (user questions in italic)
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.enum.style import WD_STYLE_TYPE
        
        doc = Document()
        
        # Set up styles
        styles = doc.styles
        
        # Modify Normal style
        normal_style = styles['Normal']
        normal_style.font.name = 'Calibri'
        normal_style.font.size = Pt(11)
        
        # === ADD METADATA HEADER SECTION ===
        export_date = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        # Title as main heading
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Separator line
        doc.add_paragraph('─' * 60).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Document Information section
        info_heading = doc.add_paragraph()
        info_run = info_heading.add_run('Document Information')
        info_run.bold = True
        info_run.font.size = Pt(12)
        info_run.font.color.rgb = RGBColor(0x34, 0x49, 0x5E)
        
        # Metadata fields
        def add_metadata_field(label: str, value: str):
            """Helper to add a metadata field with bold label"""
            para = doc.add_paragraph()
            label_run = para.add_run(f'{label}: ')
            label_run.bold = True
            label_run.font.size = Pt(10)
            value_run = para.add_run(value)
            value_run.font.size = Pt(10)
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.left_indent = Inches(0.25)
        
        # Add metadata fields
        if source and source not in ['Unknown', 'N/A', '']:
            add_metadata_field('Source', source)
        
        if published_date and published_date not in ['N/A', 'Unknown', '', None]:
            add_metadata_field('Published', published_date)
        
        if imported_date and imported_date not in ['Unknown', 'N/A', '']:
            add_metadata_field('Imported', imported_date)
        
        add_metadata_field('Exported', export_date)
        
        if doc_class and doc_class not in ['selection', '']:
            add_metadata_field('Type', doc_class.replace('_', ' ').title())
        
        # Separator before content
        doc.add_paragraph()
        separator = doc.add_paragraph('─' * 60)
        separator.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph()
        
        # === END METADATA HEADER SECTION ===
        
        # Process the content
        lines = content_text.split('\n')
        i = 0
        in_user_section = False
        
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            
            # Skip empty lines but add spacing
            if not stripped:
                doc.add_paragraph()
                i += 1
                continue
            
            # Conversation markers
            if stripped.startswith('🧑 YOU') or stripped.startswith('YOU ['):
                para = doc.add_paragraph()
                run = para.add_run(stripped)
                run.bold = True
                run.font.color.rgb = RGBColor(0x2E, 0x40, 0x53)  # Dark blue-gray
                in_user_section = True
                i += 1
                continue
            
            if stripped.startswith('🤖 AI') or stripped.startswith('AI ['):
                para = doc.add_paragraph()
                run = para.add_run(stripped)
                run.bold = True
                run.font.color.rgb = RGBColor(0x16, 0x53, 0x7E)  # Blue
                in_user_section = False
                i += 1
                continue
            
            # Heading 2: ## Title
            if stripped.startswith('## '):
                heading_text = stripped[3:]
                # Remove any ** markers from heading
                heading_text = heading_text.replace('**', '')
                para = doc.add_paragraph()
                run = para.add_run(heading_text)
                run.bold = True
                run.font.size = Pt(13)
                run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
                doc.add_paragraph()  # Add space after heading
                i += 1
                continue
            
            # Heading 3: ### Title
            if stripped.startswith('### '):
                heading_text = stripped[4:]
                heading_text = heading_text.replace('**', '')
                para = doc.add_paragraph()
                run = para.add_run(heading_text)
                run.bold = True
                run.font.size = Pt(12)
                i += 1
                continue
            
            # Horizontal rule
            if stripped == '---':
                para = doc.add_paragraph('─' * 50)
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                i += 1
                continue
            
            # Block quote
            if stripped.startswith('> '):
                quote_text = stripped[2:]
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Inches(0.5)
                run = para.add_run(quote_text)
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                i += 1
                continue
            
            # Numbered list
            if re.match(r'^\d+\.\s+', stripped):
                content = re.sub(r'^\d+\.\s+', '', stripped)
                add_formatted_paragraph(doc, content, style='List Number')
                i += 1
                continue
            
            # Bullet list (including nested)
            if stripped.startswith('- ') or stripped.startswith('* '):
                content = stripped[2:]
                indent = len(line) - len(line.lstrip())
                para = add_formatted_paragraph(doc, content, style='List Bullet')
                if indent > 0:
                    para.paragraph_format.left_indent = Inches(0.25 * (indent // 2 + 1))
                i += 1
                continue
            
            # Regular paragraph
            # If we're in a user section, make it italic
            add_formatted_paragraph(doc, stripped, is_italic=in_user_section)
            i += 1
        
        doc.save(filepath)
        return True
        
    except ImportError:
        print("❌ python-docx library not installed")
        return False
    except Exception as e:
        print(f"❌ Error saving DOCX: {str(e)}")
        import traceback
        traceback.print_exc()
        return False


def save_formatted_txt(
    filepath: str,
    content_text: str,
    title: str = "Document",
    source: str = "Unknown",
    imported_date: str = "Unknown",
    doc_class: str = "selection",
    published_date: str = None
) -> bool:
    """Save content as formatted plain text with metadata header."""
    try:
        export_date = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write("=" * 70 + "\n")
            f.write(f"{title}\n")
            f.write("=" * 70 + "\n\n")
            
            # Document Information section
            f.write("DOCUMENT INFORMATION:\n")
            if source and source not in ['Unknown', 'N/A', '']:
                f.write(f"  Source: {source}\n")
            if published_date and published_date not in ['N/A', 'Unknown', '', None]:
                f.write(f"  Published: {published_date}\n")
            if imported_date and imported_date not in ['Unknown', 'N/A', '']:
                f.write(f"  Imported: {imported_date}\n")
            f.write(f"  Exported: {export_date}\n")
            if doc_class and doc_class not in ['selection', '']:
                f.write(f"  Type: {doc_class.replace('_', ' ').title()}\n")
            
            f.write("\n" + "-" * 70 + "\n\n")
            
            # Clean up markdown for plain text
            clean_text = content_text
            # Convert **bold** to BOLD
            clean_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', clean_text)
            # Convert *italic* to _italic_
            clean_text = re.sub(r'\*([^*]+)\*', r'\1', clean_text)
            # Convert ## headings to uppercase
            clean_text = re.sub(r'^## (.+)$', lambda m: '\n' + m.group(1).upper() + '\n' + '-' * len(m.group(1)), clean_text, flags=re.MULTILINE)
            # Convert ### headings
            clean_text = re.sub(r'^### (.+)$', lambda m: '\n' + m.group(1) + '\n', clean_text, flags=re.MULTILINE)
            
            f.write(clean_text)
        
        return True
    except Exception as e:
        print(f"❌ Error saving TXT: {str(e)}")
        return False


def save_formatted_pdf(
    filepath: str,
    content_text: str,
    title: str = "Document",
    source: str = "Unknown",
    imported_date: str = "Unknown",
    doc_class: str = "selection",
    published_date: str = None
) -> bool:
    """Save content as formatted PDF with metadata header."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
        from reportlab.lib.units import inch
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
        from reportlab.lib.colors import HexColor
        
        pdf_doc = SimpleDocTemplate(filepath, pagesize=letter,
                                     leftMargin=inch, rightMargin=inch,
                                     topMargin=inch, bottomMargin=inch)
        styles = getSampleStyleSheet()
        story = []
        
        export_date = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        # === METADATA HEADER SECTION ===
        # Title style
        title_style = ParagraphStyle(
            'TitleStyle',
            parent=styles['Heading1'],
            fontSize=16,
            textColor=HexColor('#2C3E50'),
            alignment=TA_CENTER,
            spaceAfter=10
        )
        
        # Info style for metadata
        info_style = ParagraphStyle(
            'InfoStyle',
            parent=styles['Normal'],
            fontSize=10,
            leftIndent=20,
            spaceAfter=2
        )
        
        info_heading_style = ParagraphStyle(
            'InfoHeadingStyle',
            parent=styles['Normal'],
            fontSize=12,
            fontName='Helvetica-Bold',
            textColor=HexColor('#34495E'),
            spaceAfter=8,
            spaceBefore=12
        )
        
        # Add title
        story.append(Paragraph(title.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'), title_style))
        story.append(Paragraph('─' * 60, styles['Normal']))
        story.append(Spacer(1, 10))
        
        # Document Information heading
        story.append(Paragraph('Document Information', info_heading_style))
        
        # Metadata fields
        if source and source not in ['Unknown', 'N/A', '']:
            safe_source = source.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            story.append(Paragraph(f'<b>Source:</b> {safe_source}', info_style))
        
        if published_date and published_date not in ['N/A', 'Unknown', '', None]:
            story.append(Paragraph(f'<b>Published:</b> {published_date}', info_style))
        
        if imported_date and imported_date not in ['Unknown', 'N/A', '']:
            story.append(Paragraph(f'<b>Imported:</b> {imported_date}', info_style))
        
        story.append(Paragraph(f'<b>Exported:</b> {export_date}', info_style))
        
        if doc_class and doc_class not in ['selection', '']:
            story.append(Paragraph(f"<b>Type:</b> {doc_class.replace('_', ' ').title()}", info_style))
        
        # Separator
        story.append(Spacer(1, 15))
        story.append(Paragraph('─' * 60, styles['Normal']))
        story.append(Spacer(1, 15))
        
        # === END METADATA HEADER SECTION ===
        
        # Custom styles
        heading2_style = ParagraphStyle(
            'Heading2Custom',
            parent=styles['Heading2'],
            fontSize=13,
            textColor=HexColor('#2C3E50'),
            spaceAfter=10,
            spaceBefore=15
        )
        
        heading3_style = ParagraphStyle(
            'Heading3Custom',
            parent=styles['Heading3'],
            fontSize=12,
            spaceAfter=8,
            spaceBefore=12
        )
        
        user_style = ParagraphStyle(
            'UserStyle',
            parent=styles['Normal'],
            fontSize=10,
            textColor=HexColor('#2E4053'),
            fontName='Helvetica-Bold'
        )
        
        ai_style = ParagraphStyle(
            'AIStyle',
            parent=styles['Normal'],
            fontSize=10,
            textColor=HexColor('#16537E'),
            fontName='Helvetica-Bold'
        )
        
        italic_style = ParagraphStyle(
            'ItalicStyle',
            parent=styles['Normal'],
            fontSize=10,
            fontName='Helvetica-Oblique'
        )
        
        quote_style = ParagraphStyle(
            'QuoteStyle',
            parent=styles['Normal'],
            fontSize=10,
            leftIndent=30,
            fontName='Helvetica-Oblique',
            textColor=HexColor('#555555')
        )
        
        # Process content
        lines = content_text.split('\n')
        in_user_section = False
        
        for line in lines:
            stripped = line.strip()
            
            if not stripped:
                story.append(Spacer(1, 6))
                continue
            
            # Escape special characters
            safe_line = stripped.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            
            # Conversation markers
            if stripped.startswith('🧑 YOU') or stripped.startswith('YOU ['):
                story.append(Paragraph(safe_line, user_style))
                in_user_section = True
                continue
            
            if stripped.startswith('🤖 AI') or stripped.startswith('AI ['):
                story.append(Paragraph(safe_line, ai_style))
                in_user_section = False
                continue
            
            # Headings
            if stripped.startswith('## '):
                heading_text = stripped[3:].replace('**', '')
                safe_heading = heading_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(safe_heading, heading2_style))
                continue
            
            if stripped.startswith('### '):
                heading_text = stripped[4:].replace('**', '')
                safe_heading = heading_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(safe_heading, heading3_style))
                continue
            
            # Horizontal rule
            if stripped == '---':
                story.append(Paragraph('─' * 50, styles['Normal']))
                continue
            
            # Block quote
            if stripped.startswith('> '):
                quote_text = stripped[2:]
                safe_quote = quote_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(safe_quote, quote_style))
                continue
            
            # Convert markdown bold/italic to HTML-style for reportlab
            formatted = safe_line
            formatted = re.sub(r'\*\*([^*]+)\*\*', r'<b>\1</b>', formatted)
            formatted = re.sub(r'\*([^*]+)\*', r'<i>\1</i>', formatted)
            
            # Use italic style for user section
            if in_user_section:
                story.append(Paragraph(formatted, italic_style))
            else:
                story.append(Paragraph(formatted, styles['Normal']))
        
        pdf_doc.build(story)
        return True
        
    except ImportError:
        print("❌ reportlab library not installed")
        return False
    except Exception as e:
        print(f"❌ Error saving PDF: {str(e)}")
        import traceback
        traceback.print_exc()
        return False


def save_formatted_document(
    filepath: str,
    content_text: str,
    title: str = "Document",
    source: str = "Unknown",
    imported_date: str = "Unknown",
    doc_class: str = "selection",
    export_format: str = "docx",
    published_date: str = None
) -> bool:
    """
    Save content with proper formatting based on export format.
    
    Args:
        filepath: Full path to save file
        content_text: The content to save (may contain markdown formatting)
        title: Document title
        source: Source URL or path
        imported_date: When document was imported
        doc_class: Document class
        export_format: One of 'txt', 'docx', 'rtf', 'pdf'
        published_date: Original publication date (optional)
    
    Returns:
        True if successful, False otherwise
    """
    if export_format == 'txt':
        return save_formatted_txt(filepath, content_text, title, source, imported_date, doc_class, published_date)
    elif export_format == 'docx':
        return save_formatted_docx(filepath, content_text, title, source, imported_date, doc_class, published_date=published_date)
    elif export_format == 'rtf':
        # Save as DOCX, user can convert to RTF
        docx_path = filepath.replace('.rtf', '.docx')
        success = save_formatted_docx(docx_path, content_text, title, source, imported_date, doc_class, published_date=published_date)
        if success:
            messagebox.showinfo(
                "RTF Export",
                f"Document saved as DOCX at:\n{docx_path}\n\n"
                "To convert to RTF, open in Word and Save As RTF."
            )
        return success
    elif export_format == 'pdf':
        return save_formatted_pdf(filepath, content_text, title, source, imported_date, doc_class, published_date)
    else:
        print(f"❌ Unknown format: {export_format}")
        return False
