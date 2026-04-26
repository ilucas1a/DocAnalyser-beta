"""
test_docx_links.py — Generate a sample .docx exercising every link
type, then print a checklist of what to verify by eye in Word.

Run from a terminal in the project directory:
    python test_docx_links.py

Produces test_docx_links_output.docx in the same folder.  Open it in
Word (or upload to Drive and open as a Google Doc) and check that
each link works as described.

This script exercises the SAME functions DocAnalyser uses for real
exports — _add_markdown_content_to_docx and the helpers — so a green
result here means the digest export and individual summary export
will work correctly too.
"""

import os
import sys

PROJECT_DIR = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, PROJECT_DIR)


SAMPLE_CONTENT = """## Key Points {#key-points}

- **Strategy:** Russia views the Iran conflict as a multipolar realignment moment. [Detail](#point-1)
- **Diplomacy:** Sino-Russian coordination has accelerated since the strikes began. [Detail](#point-2)
- **Energy:** Watch [oil-market signals](https://www.eia.gov/petroleum/) over the next week. [Detail](#point-3)

See [Sources](#sources) at the bottom for the full list of underlying interviews.

## Detail

### Point 1 — Russia's strategic framing {#point-1}

Russia is treating the Iran-US escalation as evidence that the unipolar moment has ended. The MoD statement on 25 April was unusually explicit. Key implication: expect deeper *Russia-China-Iran* coordination on sanctions evasion and energy markets. [Back](#key-points)

### Point 2 — Sino-Russian coordination {#point-2}

The Beijing-Moscow alignment has shifted from rhetorical to operational, with practical steps on payment systems and dual-use exports. **This matters** because it shortens the timeline for a viable BRICS+ payments architecture. [Back](#key-points)

### Point 3 — Energy market signals {#point-3}

Brent has stayed remarkably calm so far — under $80 — but watch the Strait of Hormuz traffic data published at [marinetraffic.com](https://www.marinetraffic.com). A sustained tanker drop-off would be the canary. [Back](#key-points)

## Sources {#sources}

1. Alexander Mercouris — [Russian MoD Iran Hail Joint Alliance](https://www.youtube.com/watch?v=example1) (25 Apr 2026)
2. Glenn Diesen — [Michael Hudson: Iran War Ignites Global Financial Armageddon](https://www.youtube.com/watch?v=example2) (25 Apr 2026)
3. Lt Col Daniel Davis — [Deep Dive Intel Briefing: What We Learned This Week](https://www.youtube.com/watch?v=example3) (26 Apr 2026)

[Back to Key Points](#key-points)
"""


def main():
    try:
        from docx import Document
    except ImportError:
        print("ERROR: python-docx is not installed.")
        print("       pip install python-docx")
        return

    # Import the functions DocAnalyser actually uses for real exports.
    from document_export import _add_markdown_content_to_docx
    from docx_helpers import add_external_hyperlink

    out_path = os.path.join(PROJECT_DIR, "test_docx_links_output.docx")

    doc = Document()
    doc.add_heading("DocAnalyser — Hyperlink test document", level=0)

    # Simulate a metadata header with a URL field.  Same path the
    # document export takes when saving an individual summary.
    doc.add_heading("Document Information", level=2)
    meta = doc.add_paragraph()
    meta.add_run("Title: ").bold = True
    meta.add_run("Test digest — 26 Apr 2026\n")
    meta.add_run("Source: ").bold = True
    test_url = "https://www.youtube.com/watch?v=dQw4w9WgXcQ"
    add_external_hyperlink(meta, test_url, test_url)
    meta.add_run("\n")
    meta.add_run("Type: ").bold = True
    meta.add_run("test\n")

    doc.add_paragraph()
    doc.add_heading("Content", level=2)

    # Render the sample content using the same path the digest uses.
    _add_markdown_content_to_docx(doc, SAMPLE_CONTENT)

    doc.save(out_path)

    # Walk the document XML to count what was actually emitted.
    from docx.oxml.ns import qn
    n_internal   = 0
    n_external   = 0
    n_bookmarks  = 0
    bookmark_names = []
    for hyper in doc.element.iter(qn("w:hyperlink")):
        if hyper.get(qn("w:anchor")):
            n_internal += 1
        elif hyper.get(qn("r:id")):
            n_external += 1
    for bm in doc.element.iter(qn("w:bookmarkStart")):
        n_bookmarks += 1
        name = bm.get(qn("w:name"))
        if name:
            bookmark_names.append(name)

    print(f"Wrote: {out_path}")
    print()
    print("Counts in the generated docx:")
    print(f"  External hyperlinks (URLs):     {n_external}    (expect 6)")
    print(f"  Internal hyperlinks (#anchors): {n_internal}    (expect 8)")
    print(f"  Bookmarks:                      {n_bookmarks}    (expect 8)")
    print()
    print(f"  Bookmark names: {sorted(bookmark_names)}")
    print()
    print("Open the docx in Word and check:")
    print("  1. The Source URL at the top is blue/underlined and Ctrl+Click follows it")
    print("     (a YouTube URL — should open the video).")
    print("  2. The two inline external links (EIA, marinetraffic) are clickable.")
    print("  3. Each [Detail] link in Key Points jumps to the matching ### heading.")
    print("  4. Each [Back] link in the detail sections jumps back to its")
    print("     originating Key Points bullet (NOT to the top of Key Points).")
    print("  5. The [Sources] inline link jumps to the ## Sources heading.")
    print("  6. The three YouTube links in Sources are clickable.")
    print("  7. The numbered list 1./2./3. in Sources renders as a real numbered")
    print("     list (auto-numbered, not literal '1. ' text).")
    print("  8. {#point-1} markers do NOT appear as visible text anywhere.")
    print()
    print("Then upload to Google Drive, open as a Google Doc (right-click →")
    print("Open with → Google Docs), and re-check the same eight items.")


if __name__ == "__main__":
    main()
