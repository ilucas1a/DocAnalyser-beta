"""
transcript_word_toolkit.py
===========================
Exports DocAnalyser audio transcript entries to a Word .docx file.

Each transcript paragraph is written as a single Word paragraph:

    [MM:SS]  [Speaker name]:  paragraph text…

The [MM:SS] is plain styled text (8pt, grey) — not a hyperlink and not a
macro button.  The user reads the timestamp in Word, types it into the
companion player's Jump field, and the player seeks instantly.  This
avoids Word security warnings entirely.

The timestamp format is also the key used by word_editor_panel.py to
parse the document back into DocAnalyser entries when the user clicks
"Save edits to DocAnalyser".

Usage:
    from transcript_word_toolkit import export_transcript_to_word
    ok, path = export_transcript_to_word(filepath, entries, title, audio_path)

Author: DocAnalyser Development Team
"""

from __future__ import annotations

import datetime
import os
import logging
from typing import List, Dict, Optional, Tuple

logger = logging.getLogger(__name__)


# ── Formatting constants ───────────────────────────────────────────────────────
_TS_FONT_PT   = 8       # timestamp text size in points
_TS_COLOR     = "999999"  # timestamp colour (grey)
_BODY_FONT_PT = 11      # paragraph body text size
_SPK_FONT_PT  = 11      # speaker label size


def _fmt_time(seconds: float) -> str:
    """Format seconds as MM:SS or H:MM:SS."""
    s = max(0, int(round(seconds)))
    if s >= 3600:
        return f"{s // 3600}:{(s % 3600) // 60:02d}:{s % 60:02d}"
    return f"{s // 60:02d}:{s % 60:02d}"


def _plain_run(paragraph, text: str,
               font_size_pt: Optional[int] = None,
               bold: bool = False,
               color_hex: Optional[str] = None):
    """Add a plain run to a python-docx paragraph."""
    run = paragraph.add_run(text)
    if font_size_pt is not None:
        from docx.shared import Pt
        run.font.size = Pt(font_size_pt)
    if bold:
        run.bold = True
    if color_hex:
        from docx.shared import RGBColor
        r = int(color_hex[0:2], 16)
        g = int(color_hex[2:4], 16)
        b = int(color_hex[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)
    return run


def _add_bookmark(run, bk_id: int, bk_name: str) -> None:
    """
    Wrap a single run's XML element with a named Word bookmark.

    Produces the OOXML structure:
        <w:bookmarkStart w:id="N" w:name="bk_name"/>
        <w:r> ... run content ... </w:r>
        <w:bookmarkEnd w:id="N"/>

    These bookmarks are later used by word_editor_panel.py to reliably
    locate timestamp and speaker-label runs via COM without any
    character-position arithmetic.

    Naming convention:
        TS_p_N   — primary paragraph timestamp  [MM:SS]
        SP_N     — speaker label               [Speaker]:
        TS_s_N   — secondary sentence timestamp {MM:SS}
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tag_start = OxmlElement('w:bookmarkStart')
    tag_start.set(qn('w:id'),   str(bk_id))
    tag_start.set(qn('w:name'), bk_name)

    tag_end = OxmlElement('w:bookmarkEnd')
    tag_end.set(qn('w:id'), str(bk_id))

    run._r.addprevious(tag_start)   # insert before the <w:r>
    run._r.addnext(tag_end)         # insert after  the <w:r>


# =============================================================================
# Main export function
# =============================================================================

def export_transcript_to_word(
    filepath: str,
    entries:  List[Dict],
    title:    str,
    audio_path:  Optional[str] = None,
    metadata:    Optional[Dict] = None,
    show_messages: bool = True,
) -> Tuple[bool, str]:
    """
    Export transcript entries to a .docx file.

    Each entry produces one Word paragraph in the form:
        [MM:SS]  [Speaker]:  paragraph text…

    The [MM:SS] is plain grey text — no hyperlink, no macro.
    To play audio, the user reads the timestamp and types it into the
    companion player's "Jump to" field.

    The metadata block at the top includes an "Audio file:" line that
    launch_transcript.py uses to find the audio automatically.

    Args:
        filepath:      Output .docx path.
        entries:       List of entry dicts; each needs 'text' and 'start'.
                       Optional: 'speaker'.
        title:         Document title (heading at top of document).
        audio_path:    Path to the source audio file.  Stored in the
                       metadata block so launch_transcript.py can find it.
        metadata:      Optional extra metadata (source, imported_date, etc.).
        show_messages: Show a success dialog when done (default True).

    Returns:
        (True, filepath)          on success
        (False, error_message)    on failure
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    except ImportError:
        return False, (
            "python-docx is not installed.\n"
            "Install with:  pip install python-docx"
        )

    meta = metadata or {}
    doc  = Document()

    # Current date and Windows username for the "last edited" metadata fields
    now_str  = datetime.datetime.now().strftime("%d-%b-%Y  %H:%M")
    username = (
        os.getenv("USERNAME")           # Windows
        or os.getenv("USER")            # macOS / Linux
        or "Unknown"
    )

    # ── Document-wide font default ─────────────────────────────────────────
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(_BODY_FONT_PT)

    # ── Title ──────────────────────────────────────────────────────────────
    title_para = doc.add_heading(title[:120], level=0)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # ── Document information block ─────────────────────────────────────────
    # "Audio file:" is read by launch_transcript.py to find the audio
    # automatically when the user opens this document later.
    doc.add_heading("Document Information", level=2)
    info = doc.add_paragraph()
    _plain_run(info, "Date last edited: ", bold=True)
    _plain_run(info, f"{now_str}\n")
    _plain_run(info, "By: ", bold=True)
    _plain_run(info, f"{username}\n")
    _plain_run(info, "Title: ", bold=True)
    _plain_run(info, f"{title}\n")
    if audio_path:
        _plain_run(info, "Audio file: ", bold=True)
        _plain_run(info, f"{audio_path}\n")
    if meta.get("source"):
        _plain_run(info, "Source: ", bold=True)
        _plain_run(info, f"{meta['source']}\n")
    if meta.get("imported_date"):
        _plain_run(info, "Date imported: ", bold=True)
        _plain_run(info, f"{meta['imported_date']}\n")

    # ── Usage note ─────────────────────────────────────────────────────────
    doc.add_paragraph()
    note = doc.add_paragraph()
    _plain_run(
        note,
        "Timestamps are shown as plain text (e.g. [04:23]).  "
        "To hear any point in the recording, read the timestamp and type it "
        "into the DocAnalyser Speaker Panel's 'Jump to' field.  "
        "Keep the [MM:SS] timestamps intact \u2014 they are used to sync "
        "edits back to DocAnalyser.  "
        "Tiny grey {MM:SS} markers within each paragraph are sentence-level "
        "timestamps: if you split a paragraph by pressing Enter, the second "
        "half will already carry the correct timestamp for that sentence.",
        font_size_pt=9,
        color_hex="666666",
    )

    # ── Transcript ─────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_heading("Transcript", level=2)

    # Bookmark ID counter — start at 100 to avoid any python-docx internal IDs.
    bk_id      = 100
    n_primary  = 0
    n_speaker  = 0
    n_sentence = 0

    for entry in entries:
        body = (entry.get("text") or "").strip()
        if not body:
            continue

        speaker = (entry.get("speaker") or "").strip()
        start   = float(entry.get("start", 0.0))
        ts      = _fmt_time(start)

        # One Word paragraph per transcript entry.
        # Plain text form:  [MM:SS]  [Speaker]:  body text…
        # This exact format is what word_editor_panel._PARA_RE parses back.
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after  = Pt(4)

        # [MM:SS]  — small grey, not a hyperlink
        run_ts = _plain_run(para, f"[{ts}]", font_size_pt=_TS_FONT_PT, color_hex=_TS_COLOR)
        _add_bookmark(run_ts, bk_id, f"TS_p_{n_primary}")
        bk_id += 1; n_primary += 1

        # Two non-breaking spaces as separator
        _plain_run(para, "\u00a0\u00a0")

        # [Speaker]:  — bold, normal size
        spk_label = speaker if speaker else "—"
        run_spk = _plain_run(para, f"[{spk_label}]: ", bold=True, font_size_pt=_SPK_FONT_PT)
        _add_bookmark(run_spk, bk_id, f"SP_{n_speaker}")
        bk_id += 1; n_speaker += 1

        # Body text with embedded sentence-level {MM:SS} markers.
        # Each sentence is prefixed with a tiny (7pt) light-grey {MM:SS} so
        # that if the user splits the paragraph by pressing Enter in Word,
        # the second half already starts with the correct sentence timestamp.
        # Only emitted when the entry has multiple sentences; single-sentence
        # entries skip it (the paragraph header timestamp IS the timestamp).
        sentences = entry.get("sentences", [])
        multi_sent = [s for s in sentences if (s.get("text") or "").strip()]
        if len(multi_sent) > 1:
            for sent in multi_sent:
                sent_start = float(sent.get("start", start))
                sent_ts    = _fmt_time(sent_start)
                sent_text  = sent["text"].strip()
                # {MM:SS} in 7pt light grey
                run_sent_ts = _plain_run(para, f"{{{sent_ts}}}",
                           font_size_pt=7, color_hex="cccccc")
                _add_bookmark(run_sent_ts, bk_id, f"TS_s_{n_sentence}")
                bk_id += 1; n_sentence += 1
                _plain_run(para, f"\u00a0{sent_text} ",
                           font_size_pt=_BODY_FONT_PT)
        else:
            # No sentence data or single sentence — write body text directly
            _plain_run(para, f"\u00a0{body}", font_size_pt=_BODY_FONT_PT)

    # ── Save ───────────────────────────────────────────────────────────────
    # Ensure .docx extension (not .docm — no macros needed)
    base, ext = os.path.splitext(filepath)
    if ext.lower() not in (".docx",):
        filepath = base + ".docx"

    try:
        doc.save(filepath)
    except PermissionError:
        return False, (
            f"Cannot save to:\n{filepath}\n\n"
            "The file may be open in Word.  Close it and try again."
        )

    logger.info(f"Transcript exported to Word: {filepath}")

    if show_messages:
        try:
            from tkinter import messagebox
            messagebox.showinfo(
                "Export complete",
                f"Transcript saved to:\n{filepath}\n\n"
                "Open this file in Word alongside the DocAnalyser\n"
                "companion player to edit with audio playback.",
            )
        except Exception:
            pass

    return True, filepath
