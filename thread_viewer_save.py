"""
thread_viewer_save.py - Save & Export Mixin for ThreadViewerWindow

Extracted from thread_viewer.py to improve maintainability.
Handles all save-to-file operations including format-specific exports
(TXT, RTF, DOCX, PDF) and the save-as dialog.

All methods access the parent ThreadViewerWindow's state via self.
"""

import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import os
import re
import datetime

from utils import safe_filename
from document_export import export_conversation_thread, get_file_extension_and_types, export_document


def get_clean_filename(text: str, max_length: int = 50) -> str:
    """Wrapper for safe_filename for backward compatibility"""
    return safe_filename(text, max_length)


class SaveMixin:
    """
    Mixin providing save/export operations for ThreadViewerWindow.
    
    Requires the following attributes on self:
        - thread_text: tk.Text widget
        - window: tk.Toplevel
        - current_thread: list of message dicts
        - current_document_text: str
        - source_documents: list of source doc dicts
        - doc_title, source_info, fetched_date, published_date: metadata strings
        - current_mode: 'source' or 'conversation'
        - model_var, provider_var: tk.StringVar
        - _set_status(): status display method
        - _save_source_edits(): from core ThreadViewerWindow
        - _escape_html(): from CopyMixin
    """

    def _show_save_as_dialog(self):
        """
        Show a dialog to choose what content to save, then proceed to file save dialog.
        """
        # Create dialog window
        dialog = tk.Toplevel(self.window)
        dialog.title("Save As")
        dialog.transient(self.window)
        dialog.grab_set()
        
        # Center on parent window
        dialog.geometry("450x460")
        dialog_x = self.window.winfo_x() + (self.window.winfo_width() - 450) // 2
        dialog_y = self.window.winfo_y() + (self.window.winfo_height() - 460) // 2
        dialog.geometry(f"+{dialog_x}+{dialog_y}")
        
        # Main frame with padding
        main_frame = ttk.Frame(dialog, padding=15)
        main_frame.pack(fill='both', expand=True)
        
        # Title
        title_label = ttk.Label(main_frame, text="What do you want to save?", 
                                font=('Segoe UI', 11, 'bold'))
        title_label.pack(anchor='w', pady=(0, 15))
        
        # Check if source document is available
        has_source = bool(self.current_document_text)
        
        # Check if thread exists
        has_thread = bool(self.current_thread and len(self.current_thread) > 0)
        
        # Default selection based on current mode
        if self.current_mode == 'source':
            default_choice = "source" if has_source else "thread"
        else:
            default_choice = "thread"
        
        # Variable to track selection
        save_choice = tk.StringVar(value=default_choice)
        
        # Count expanded exchanges for the label
        exchanges = self._group_messages_into_exchanges()
        expanded_count = sum(1 for i in range(len(exchanges)) 
                           if self.exchange_expanded_state.get(i, True))
        total_count = len(exchanges)
        
        # Option 1: Source Only
        opt1_frame = ttk.Frame(main_frame)
        opt1_frame.pack(fill='x', pady=5)
        
        source_rb = ttk.Radiobutton(opt1_frame, text="Source Only", variable=save_choice, 
                                     value="source")
        source_rb.pack(anchor='w')
        
        if has_source:
            ttk.Label(opt1_frame, text="Just the original source document", 
                     foreground='gray').pack(anchor='w', padx=(20, 0))
        else:
            ttk.Label(opt1_frame, text="(Not available - no source document loaded)", 
                     foreground='red').pack(anchor='w', padx=(20, 0))
            source_rb.config(state='disabled')
        
        # Option 2: Thread
        opt2_frame = ttk.Frame(main_frame)
        opt2_frame.pack(fill='x', pady=5)
        
        thread_rb = ttk.Radiobutton(opt2_frame, text="Thread", variable=save_choice, 
                                     value="thread")
        thread_rb.pack(anchor='w')
        
        if has_thread:
            ttk.Label(opt2_frame, text="Your questions and AI responses only", 
                     foreground='gray').pack(anchor='w', padx=(20, 0))
        else:
            ttk.Label(opt2_frame, text="(Not available - no conversation yet)", 
                     foreground='red').pack(anchor='w', padx=(20, 0))
            thread_rb.config(state='disabled')
        
        # Option 3: Expanded Only
        opt3_frame = ttk.Frame(main_frame)
        opt3_frame.pack(fill='x', pady=5)
        
        expanded_text = f"Expanded Only ({expanded_count} of {total_count} exchanges)"
        expanded_rb = ttk.Radiobutton(opt3_frame, text=expanded_text, variable=save_choice, 
                                       value="expanded")
        expanded_rb.pack(anchor='w')
        
        if has_thread and total_count > 0:
            ttk.Label(opt3_frame, text="Only exchanges you've expanded (collapsed ones are omitted)", 
                     foreground='gray').pack(anchor='w', padx=(20, 0))
        else:
            ttk.Label(opt3_frame, text="(Not available - no exchanges)", 
                     foreground='red').pack(anchor='w', padx=(20, 0))
            expanded_rb.config(state='disabled')
        
        # Option 4: Complete (disabled if no source)
        opt4_frame = ttk.Frame(main_frame)
        opt4_frame.pack(fill='x', pady=5)
        
        complete_rb = ttk.Radiobutton(opt4_frame, text="Complete: Source + Thread", 
                                       variable=save_choice, value="complete")
        complete_rb.pack(anchor='w')
        
        if has_source and has_thread:
            ttk.Label(opt4_frame, text="Original source document AND all conversation exchanges", 
                     foreground='gray').pack(anchor='w', padx=(20, 0))
        else:
            ttk.Label(opt4_frame, text="(Requires both source and thread)", 
                     foreground='red').pack(anchor='w', padx=(20, 0))
            complete_rb.config(state='disabled')
        
        # Separator
        ttk.Separator(main_frame, orient='horizontal').pack(fill='x', pady=15)
        
        # Format selection
        format_frame = ttk.Frame(main_frame)
        format_frame.pack(fill='x', pady=5)
        
        ttk.Label(format_frame, text="Format:").pack(side='left')
        
        format_var = tk.StringVar(value=".docx")
        format_combo = ttk.Combobox(format_frame, textvariable=format_var, 
                                     values=[".docx", ".txt", ".rtf", ".pdf"],
                                     state='readonly', width=10)
        format_combo.pack(side='left', padx=(10, 0))
        
        # Buttons
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill='x', pady=(20, 0))
        
        def on_save():
            choice = save_choice.get()
            fmt = format_var.get()
            dialog.destroy()
            
            if choice == "source":
                self._save_source_only(fmt)
            elif choice == "thread":
                self._save_thread(fmt)
            elif choice == "expanded":
                self._save_expanded_only(fmt)
            elif choice == "complete":
                self._save_complete(fmt)
        
        def on_cancel():
            dialog.destroy()
        
        ttk.Button(button_frame, text="Save", command=on_save, width=10).pack(side='right', padx=5)
        ttk.Button(button_frame, text="Cancel", command=on_cancel, width=10).pack(side='right')
        
        # Handle Enter and Escape
        dialog.bind('<Return>', lambda e: on_save())
        dialog.bind('<Escape>', lambda e: on_cancel())
        
        # Focus on dialog
        dialog.focus_set()
        dialog.wait_window()

    def _save_source_only(self, format_ext=None):
        """Save just the source document to a file."""
        if not self.current_document_text:
            messagebox.showwarning("No Source", "No source document available to save.")
            return
        
        # Save any pending source edits first
        try:
            self._save_source_edits()
        except ValueError:
            return
        except Exception:
            pass
        
        # Set up file dialog based on requested format
        if format_ext:
            ext, filetypes = get_file_extension_and_types(format_ext.lstrip('.'))
            default_ext = ext
        else:
            default_ext = ".txt"
            filetypes = [
                ("Text files", "*.txt"),
                ("Word Document", "*.docx"),
                ("RTF files", "*.rtf"),
                ("PDF files", "*.pdf"),
                ("All files", "*.*")
            ]
        
        # Get clean filename from document title
        clean_title = get_clean_filename(self.doc_title)
        
        file_path = filedialog.asksaveasfilename(
            defaultextension=default_ext,
            filetypes=filetypes,
            initialfile=f"{clean_title}_source{format_ext or '.txt'}"
        )
        
        if not file_path:
            return
        
        try:
            ext = os.path.splitext(file_path)[1].lower().lstrip('.')
            if ext not in ['txt', 'docx', 'rtf', 'pdf']:
                ext = 'txt'
            
            # Save using document_export's export_document function
            metadata = {
                'title': self.doc_title,
                'source': self.source_info,
                'published_date': getattr(self, 'published_date', None) if getattr(self, 'published_date', 'N/A') != 'N/A' else None,
                'imported_date': self.fetched_date,
                'doc_class': 'source'
            }
            success, msg = export_document(file_path, self.current_document_text, ext, metadata, show_messages=False)
            if not success:
                raise Exception(msg)
            
            self._set_status(f"✅ Source saved to {os.path.basename(file_path)}")
            
        except Exception as e:
            messagebox.showerror("Save Error", f"Failed to save source:\n{str(e)}")

    def _save_thread(self, format_ext=None):
        """Save content to file - source document when in source mode, thread when in conversation mode"""
        # First, save any pending edits to ensure we export the latest content
        try:
            if self.current_mode == 'conversation' and self.current_thread:
                self._save_edits_to_thread()
            elif self.current_mode == 'source':
                self._save_source_edits()
        except ValueError:
            # User cancelled the save edits dialog
            return
        except Exception as e:
            pass
            # Log but continue - edits may not need saving
        
        # Determine what we're saving based on current mode
        is_source_mode = self.current_mode == 'source'
        
        # Set up file dialog based on requested format
        if format_ext:
            ext, filetypes = get_file_extension_and_types(format_ext.lstrip('.'))
            default_ext = ext
        else:
            # Show all formats if no specific format requested
            default_ext = ".txt"
            filetypes = [
                ("Text files", "*.txt"),
                ("Word Document", "*.docx"),
                ("RTF files", "*.rtf"),
                ("PDF files", "*.pdf"),
                ("All files", "*.*")
            ]
        
        # Get clean filename from document title
        clean_title = get_clean_filename(self.doc_title)
        
        # Different filename suffix based on mode
        suffix = "" if is_source_mode else "_thread"
        
        file_path = filedialog.asksaveasfilename(
            defaultextension=default_ext,
            filetypes=filetypes,
            initialfile=f"{clean_title}{suffix}{format_ext or '.txt'}"
        )
        
        if not file_path:
            return
        
        try:
            # Get file extension and determine format
            ext = os.path.splitext(file_path)[1].lower().lstrip('.')
            if ext not in ['txt', 'docx', 'rtf', 'pdf']:
                ext = 'txt'
            
            if is_source_mode:
                # === SOURCE MODE: Save the source document text ===
                from doc_formatter import save_formatted_document
                
                success = save_formatted_document(
                    filepath=file_path,
                    content_text=self.current_document_text or "",
                    title=self.doc_title,
                    source=self.source_info,
                    imported_date=self.fetched_date,
                    doc_class="source",
                    export_format=ext,
                    published_date=getattr(self, 'published_date', None)
                )
                
                if success:
                    filename = os.path.basename(file_path)
                    self._set_status(f"✅ Source document saved to {filename}")
                else:
                    messagebox.showerror("Error", "Failed to save source document. Check console for details.")
            else:
                # === CONVERSATION MODE: Save the conversation thread ===
                # Build thread metadata
                thread_metadata = {
                    'doc_title': self.doc_title,
                    'source_info': self.source_info,
                    'published_date': getattr(self, 'published_date', None),
                    'fetched_date': self.fetched_date,
                    'provider': self.provider_var.get() if self.provider_var else 'N/A',
                    'model': self.model_var.get() if self.model_var else 'N/A',
                    'message_count': self.thread_message_count
                }
                
                # Use consolidated export function
                success, result = export_conversation_thread(
                    filepath=file_path,
                    format=ext,
                    thread_messages=self.current_thread,
                    thread_metadata=thread_metadata,
                    show_messages=True
                )
                
                if success:
                    filename = os.path.basename(file_path)
                    self._set_status(f"✅ Conversation thread saved to {filename}")
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save:\n{str(e)}")
            import traceback
            traceback.print_exc()

    def _save_expanded_only(self, format_ext='.txt'):
        """
        Save only the expanded exchanges to a file.
        Collapsed exchanges are completely omitted.
        
        Args:
            format_ext: File extension ('.txt', '.docx', '.rtf', or '.pdf')
        """
        exchanges = self._group_messages_into_exchanges()
        
        if not exchanges:
            messagebox.showwarning("No Content", "No exchanges to save.")
            return
        
        # Find which exchanges are expanded
        expanded_indices = [i for i in range(len(exchanges)) 
                          if self.exchange_expanded_state.get(i, True)]
        
        if not expanded_indices:
            messagebox.showwarning("No Expanded Exchanges", 
                "No exchanges are currently expanded.\n\n"
                "Please expand the exchange(s) you want to save, then try again.")
            return
        
        # Build content from expanded exchanges only
        expanded_thread = []
        for idx in expanded_indices:
            exchange = exchanges[idx]
            if 'user' in exchange:
                expanded_thread.append(exchange['user'])
            if 'assistant' in exchange:
                expanded_thread.append(exchange['assistant'])
        
        # Set up file dialog - support all 4 formats
        ext = format_ext.lstrip('.')
        format_filetypes = {
            'docx': [("Word Document", "*.docx"), ("All files", "*.*")],
            'txt': [("Text files", "*.txt"), ("All files", "*.*")],
            'rtf': [("RTF Document", "*.rtf"), ("All files", "*.*")],
            'pdf': [("PDF Document", "*.pdf"), ("All files", "*.*")]
        }
        filetypes = format_filetypes.get(ext, [("All files", "*.*")])
        
        # Get clean filename
        clean_title = get_clean_filename(self.doc_title)
        count = len(expanded_indices)
        suffix = f"_exchange{'s' if count != 1 else ''}_{count}"
        
        file_path = filedialog.asksaveasfilename(
            defaultextension=format_ext,
            filetypes=filetypes,
            initialfile=f"{clean_title}{suffix}{format_ext}"
        )
        
        if not file_path:
            return
        
        try:
            # Build metadata
            thread_metadata = {
                'doc_title': self.doc_title,
                'source_info': self.source_info,
                'published_date': getattr(self, 'published_date', None),
                'fetched_date': self.fetched_date,
                'model': self.model_var.get(),
                'provider': self.provider_var.get(),
                'message_count': count,
                'note': f"Exported {count} of {len(exchanges)} exchange{'s' if len(exchanges) != 1 else ''}"
            }
            
            # Use consolidated export function
            success, result = export_conversation_thread(
                filepath=file_path,
                format=ext,
                thread_messages=expanded_thread,
                thread_metadata=thread_metadata,
                show_messages=True
            )
            
            if success:
                filename = os.path.basename(file_path)
                self._set_status(f"✅ Saved {count} expanded exchange{'s' if count != 1 else ''} to {filename}")
            else:
                messagebox.showerror("Error", f"Failed to save: {result}")
                
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save:\n{str(e)}")
            import traceback
            traceback.print_exc()

    def _save_complete(self, format_ext='.txt'):
        """
        Save both the source document AND the conversation thread to a single file.
        This provides complete context for sharing.
        
        Args:
            format_ext: File extension ('.txt', '.docx', '.rtf', or '.pdf')
        """
        # Check we have both source and conversation
        if not self.current_document_text:
            messagebox.showwarning("No Source", 
                "No source document available.\n\n"
                "Use the regular Save As options to save just the conversation.")
            return
        
        if not self.current_thread or len(self.current_thread) == 0:
            messagebox.showwarning("No Conversation", 
                "No conversation exchanges available.\n\n"
                "Use the regular Save As options to save just the source document.")
            return
        
        # Set up file dialog - support all 4 formats
        ext = format_ext.lstrip('.')
        format_filetypes = {
            'docx': [("Word Document", "*.docx"), ("All files", "*.*")],
            'txt': [("Text files", "*.txt"), ("All files", "*.*")],
            'rtf': [("RTF Document", "*.rtf"), ("All files", "*.*")],
            'pdf': [("PDF Document", "*.pdf"), ("All files", "*.*")]
        }
        filetypes = format_filetypes.get(ext, [("All files", "*.*")])
        
        # Get clean filename
        clean_title = get_clean_filename(self.doc_title)
        
        file_path = filedialog.asksaveasfilename(
            defaultextension=format_ext,
            filetypes=filetypes,
            initialfile=f"{clean_title}_complete{format_ext}"
        )
        
        if not file_path:
            return
        
        try:
            if ext == 'docx':
                self._save_complete_docx(file_path)
            elif ext == 'pdf':
                self._save_complete_pdf(file_path)
            elif ext == 'rtf':
                self._save_complete_rtf(file_path)
            else:
                self._save_complete_txt(file_path)
                
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save:\n{str(e)}")
            import traceback
            traceback.print_exc()

    def _save_complete_txt(self, file_path: str):
        """Save complete document (source + thread) as plain text."""
        lines = []
        
        # === HEADER ===
        lines.append("=" * 70)
        lines.append("COMPLETE DOCUMENT EXPORT")
        lines.append("=" * 70)
        lines.append("")
        lines.append(f"Title: {self.doc_title}")
        lines.append(f"Source: {self.source_info}")
        if hasattr(self, 'published_date') and self.published_date and self.published_date != 'N/A':
            lines.append(f"Published: {self.published_date}")
        lines.append(f"Imported: {self.fetched_date}")
        lines.append(f"Exported: {datetime.datetime.now().strftime('%d-%b-%Y %H:%M')}")
        lines.append("")
        
        # === SOURCE DOCUMENT ===
        lines.append("=" * 70)
        lines.append("SOURCE DOCUMENT")
        lines.append("=" * 70)
        lines.append("")
        lines.append(self.current_document_text or "(No source text available)")
        lines.append("")
        
        # === CONVERSATION THREAD ===
        lines.append("=" * 70)
        lines.append("CONVERSATION THREAD")
        lines.append("=" * 70)
        lines.append("")
        
        exchanges = self._group_messages_into_exchanges()
        for i, exchange in enumerate(exchanges):
            lines.append(f"--- Exchange {i + 1} ---")
            lines.append("")
            
            # User message
            user_msg = exchange.get('user', {})
            if user_msg:
                timestamp = user_msg.get('timestamp', '')
                time_str = f" [{timestamp}]" if timestamp else ""
                lines.append(f"YOU{time_str}:")
                lines.append(user_msg.get('content', ''))
                lines.append("")
            
            # Assistant message
            assistant_msg = exchange.get('assistant', {})
            if assistant_msg:
                provider = assistant_msg.get('provider', 'AI')
                model = assistant_msg.get('model', '')
                timestamp = assistant_msg.get('timestamp', '')
                time_str = f" [{timestamp}]" if timestamp else ""
                
                if model and model != provider:
                    label = f"{provider} ({model}){time_str}:"
                else:
                    label = f"{provider}{time_str}:"
                
                lines.append(label)
                lines.append(assistant_msg.get('content', ''))
                lines.append("")
            
            lines.append("")
        
        # Write to file
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(lines))
        
        filename = os.path.basename(file_path)
        self._set_status(f"✅ Complete document saved to {filename}")

    def _save_complete_rtf(self, file_path: str):
        """Save complete document (source + thread) as RTF."""
        # RTF is essentially text with formatting codes
        # We'll create a simple RTF with basic formatting
        
        def rtf_escape(text: str) -> str:
            """Escape special RTF characters."""
            if not text:
                return ""
            # Escape backslash, curly braces, and handle unicode
            text = text.replace('\\', '\\\\')
            text = text.replace('{', '\\{')
            text = text.replace('}', '\\}')
            # Handle line breaks
            text = text.replace('\n', '\\par\n')
            return text
        
        rtf_lines = []
        
        # RTF header
        rtf_lines.append('{\\rtf1\\ansi\\deff0')
        rtf_lines.append('{\\fonttbl{\\f0 Calibri;}{\\f1 Arial;}}')
        rtf_lines.append('\\f0\\fs22')  # Default font and size
        
        # === TITLE ===
        rtf_lines.append('\\pard\\qc\\b\\fs32')  # Centered, bold, larger
        rtf_lines.append(rtf_escape(self.doc_title))
        rtf_lines.append('\\par\\b0\\fs22\\ql')  # Reset to normal, left-aligned
        rtf_lines.append('\\par')
        
        # === METADATA ===
        rtf_lines.append('\\pard\\fs20\\cf1')  # Smaller, gray text
        rtf_lines.append(f'\\b Source:\\b0  {rtf_escape(self.source_info)}\\par')
        if hasattr(self, 'published_date') and self.published_date and self.published_date != 'N/A':
            rtf_lines.append(f'\\b Published:\\b0  {self.published_date}\\par')
        rtf_lines.append(f'\\b Imported:\\b0  {self.fetched_date}\\par')
        rtf_lines.append(f'\\b Exported:\\b0  {datetime.datetime.now().strftime("%d-%b-%Y %H:%M")}\\par')
        rtf_lines.append('\\fs22\\cf0\\par')  # Reset font size and color
        
        # === SOURCE DOCUMENT SECTION ===
        rtf_lines.append('\\pard\\brdrb\\brdrs\\brdrw10\\par')  # Horizontal line
        rtf_lines.append('\\par\\b\\fs28 SOURCE DOCUMENT\\b0\\fs22\\par')
        rtf_lines.append('\\brdrb\\brdrs\\brdrw10\\par\\par')
        
        # Add source text
        source_text = self.current_document_text or "(No source text available)"
        for para_text in source_text.split('\n\n'):
            if para_text.strip():
                rtf_lines.append(rtf_escape(para_text.strip()))
                rtf_lines.append('\\par\\par')
        
        # === CONVERSATION THREAD SECTION ===
        rtf_lines.append('\\par\\brdrb\\brdrs\\brdrw10\\par')
        rtf_lines.append('\\par\\b\\fs28 CONVERSATION THREAD\\b0\\fs22\\par')
        rtf_lines.append('\\brdrb\\brdrs\\brdrw10\\par\\par')
        
        # Add exchanges
        exchanges = self._group_messages_into_exchanges()
        for i, exchange in enumerate(exchanges):
            # Exchange header
            rtf_lines.append(f'\\b Exchange {i + 1}\\b0\\par')
            rtf_lines.append('\\par')
            
            # User message
            user_msg = exchange.get('user', {})
            if user_msg:
                timestamp = user_msg.get('timestamp', '')
                time_str = f" [{timestamp}]" if timestamp else ""
                rtf_lines.append(f'\\b\\cf2 YOU{time_str}:\\cf0\\b0\\par')
                rtf_lines.append(rtf_escape(user_msg.get('content', '')))
                rtf_lines.append('\\par\\par')
            
            # Assistant message
            assistant_msg = exchange.get('assistant', {})
            if assistant_msg:
                provider = assistant_msg.get('provider', 'AI')
                model = assistant_msg.get('model', '')
                timestamp = assistant_msg.get('timestamp', '')
                time_str = f" [{timestamp}]" if timestamp else ""
                
                if model and model != provider:
                    label = f"{provider} ({model}){time_str}:"
                else:
                    label = f"{provider}{time_str}:"
                
                rtf_lines.append(f'\\b\\cf1 {rtf_escape(label)}\\cf0\\b0\\par')
                
                # Add AI response content
                content = assistant_msg.get('content', '')
                for para_text in content.split('\n\n'):
                    if para_text.strip():
                        rtf_lines.append(rtf_escape(para_text.strip()))
                        rtf_lines.append('\\par\\par')
            
            # Add separator between exchanges
            if i < len(exchanges) - 1:
                rtf_lines.append('\\par\\pard\\qc\\emdash\\emdash\\emdash\\emdash\\emdash\\ql\\par\\par')
        
        # Close RTF
        rtf_lines.append('}')
        
        # Write to file
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(rtf_lines))
        
        filename = os.path.basename(file_path)
        self._set_status(f"✅ Complete document saved to {filename}")

    def _save_complete_docx(self, file_path: str):
        """Save complete document (source + thread) as Word document."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            messagebox.showerror("Missing Module", 
                "python-docx is required for Word export.\n\n"
                "Install with: pip install python-docx")
            return
        
        doc = Document()
        
        # === TITLE ===
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(self.doc_title)
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # === METADATA ===
        doc.add_paragraph()
        meta_para = doc.add_paragraph()
        meta_para.add_run("Source: ").bold = True
        meta_para.add_run(self.source_info)
        
        if hasattr(self, 'published_date') and self.published_date and self.published_date != 'N/A':
            meta_para = doc.add_paragraph()
            meta_para.add_run("Published: ").bold = True
            meta_para.add_run(str(self.published_date))
        
        meta_para = doc.add_paragraph()
        meta_para.add_run("Imported: ").bold = True
        meta_para.add_run(self.fetched_date)
        
        meta_para = doc.add_paragraph()
        meta_para.add_run("Exported: ").bold = True
        meta_para.add_run(datetime.datetime.now().strftime('%d-%b-%Y %H:%M'))
        
        # === SOURCE DOCUMENT SECTION ===
        doc.add_paragraph()
        doc.add_paragraph("_" * 50)
        
        section_heading = doc.add_paragraph()
        heading_run = section_heading.add_run("SOURCE DOCUMENT")
        heading_run.bold = True
        heading_run.font.size = Pt(14)
        
        doc.add_paragraph("_" * 50)
        doc.add_paragraph()
        
        # Add source text (split into paragraphs)
        source_text = self.current_document_text or "(No source text available)"
        for para_text in source_text.split('\n\n'):
            if para_text.strip():
                doc.add_paragraph(para_text.strip())
        
        # === CONVERSATION THREAD SECTION ===
        doc.add_paragraph()
        doc.add_paragraph("_" * 50)
        
        section_heading = doc.add_paragraph()
        heading_run = section_heading.add_run("CONVERSATION THREAD")
        heading_run.bold = True
        heading_run.font.size = Pt(14)
        
        doc.add_paragraph("_" * 50)
        doc.add_paragraph()
        
        # Add exchanges
        exchanges = self._group_messages_into_exchanges()
        for i, exchange in enumerate(exchanges):
            # Exchange header
            ex_header = doc.add_paragraph()
            ex_header.add_run(f"Exchange {i + 1}").bold = True
            
            # User message
            user_msg = exchange.get('user', {})
            if user_msg:
                timestamp = user_msg.get('timestamp', '')
                time_str = f" [{timestamp}]" if timestamp else ""
                
                user_para = doc.add_paragraph()
                user_label = user_para.add_run(f"YOU{time_str}:")
                user_label.bold = True
                
                doc.add_paragraph(user_msg.get('content', ''))
            
            # Assistant message
            assistant_msg = exchange.get('assistant', {})
            if assistant_msg:
                provider = assistant_msg.get('provider', 'AI')
                model = assistant_msg.get('model', '')
                timestamp = assistant_msg.get('timestamp', '')
                time_str = f" [{timestamp}]" if timestamp else ""
                
                if model and model != provider:
                    label = f"{provider} ({model}){time_str}:"
                else:
                    label = f"{provider}{time_str}:"
                
                ai_para = doc.add_paragraph()
                ai_label = ai_para.add_run(label)
                ai_label.bold = True
                
                # Add AI response content (handle markdown-ish formatting)
                content = assistant_msg.get('content', '')
                for para_text in content.split('\n\n'):
                    if para_text.strip():
                        doc.add_paragraph(para_text.strip())
            
            # Add separator between exchanges
            if i < len(exchanges) - 1:
                doc.add_paragraph()
                sep = doc.add_paragraph()
                sep.add_run("─" * 30)
                doc.add_paragraph()
        
        # Save the document
        doc.save(file_path)
        
        filename = os.path.basename(file_path)
        self._set_status(f"✅ Complete document saved to {filename}")

    def _save_complete_pdf(self, file_path: str):
        """Save complete document (source + thread) as PDF."""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.lib.enums import TA_CENTER
        except ImportError:
            messagebox.showerror("Missing Module", 
                "reportlab is required for PDF export.\n\n"
                "Install with: pip install reportlab")
            return
        
        # Create the PDF document
        doc = SimpleDocTemplate(
            file_path,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        # Get styles
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            alignment=TA_CENTER,
            spaceAfter=12
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceBefore=20,
            spaceAfter=10
        )
        
        meta_style = ParagraphStyle(
            'CustomMeta',
            parent=styles['Normal'],
            fontSize=10,
            textColor='gray'
        )
        
        normal_style = styles['Normal']
        
        user_style = ParagraphStyle(
            'UserStyle',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            textColor='#2E4053',
            spaceBefore=12
        )
        
        ai_style = ParagraphStyle(
            'AIStyle',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            textColor='#16537E',
            spaceBefore=12
        )
        
        # Build the document content
        story = []
        
        # === TITLE ===
        story.append(Paragraph(self._escape_html(self.doc_title), title_style))
        story.append(Spacer(1, 12))
        
        # === METADATA ===
        story.append(Paragraph(f"<b>Source:</b> {self._escape_html(self.source_info)}", meta_style))
        if hasattr(self, 'published_date') and self.published_date and self.published_date != 'N/A':
            story.append(Paragraph(f"<b>Published:</b> {self.published_date}", meta_style))
        story.append(Paragraph(f"<b>Imported:</b> {self.fetched_date}", meta_style))
        story.append(Paragraph(f"<b>Exported:</b> {datetime.datetime.now().strftime('%d-%b-%Y %H:%M')}", meta_style))
        story.append(Spacer(1, 20))
        
        # === SOURCE DOCUMENT SECTION ===
        story.append(Paragraph("─" * 50, normal_style))
        story.append(Paragraph("SOURCE DOCUMENT", heading_style))
        story.append(Paragraph("─" * 50, normal_style))
        story.append(Spacer(1, 12))
        
        # Add source text (split into paragraphs, escape HTML)
        source_text = self.current_document_text or "(No source text available)"
        for para_text in source_text.split('\n\n'):
            if para_text.strip():
                # Escape any HTML-like characters
                safe_text = self._escape_html(para_text.strip())
                story.append(Paragraph(safe_text, normal_style))
                story.append(Spacer(1, 6))
        
        # === CONVERSATION THREAD SECTION ===
        story.append(Spacer(1, 20))
        story.append(Paragraph("─" * 50, normal_style))
        story.append(Paragraph("CONVERSATION THREAD", heading_style))
        story.append(Paragraph("─" * 50, normal_style))
        story.append(Spacer(1, 12))
        
        # Add exchanges
        exchanges = self._group_messages_into_exchanges()
        for i, exchange in enumerate(exchanges):
            # Exchange header
            story.append(Paragraph(f"<b>Exchange {i + 1}</b>", normal_style))
            story.append(Spacer(1, 6))
            
            # User message
            user_msg = exchange.get('user', {})
            if user_msg:
                timestamp = user_msg.get('timestamp', '')
                time_str = f" [{timestamp}]" if timestamp else ""
                story.append(Paragraph(f"YOU{time_str}:", user_style))
                
                content = user_msg.get('content', '')
                safe_content = self._escape_html(content)
                story.append(Paragraph(safe_content, normal_style))
                story.append(Spacer(1, 6))
            
            # Assistant message
            assistant_msg = exchange.get('assistant', {})
            if assistant_msg:
                provider = assistant_msg.get('provider', 'AI')
                model = assistant_msg.get('model', '')
                timestamp = assistant_msg.get('timestamp', '')
                time_str = f" [{timestamp}]" if timestamp else ""
                
                if model and model != provider:
                    label = f"{provider} ({model}){time_str}:"
                else:
                    label = f"{provider}{time_str}:"
                
                story.append(Paragraph(label, ai_style))
                
                # Add AI response content (split into paragraphs)
                content = assistant_msg.get('content', '')
                for para_text in content.split('\n\n'):
                    if para_text.strip():
                        safe_text = self._escape_html(para_text.strip())
                        story.append(Paragraph(safe_text, normal_style))
                        story.append(Spacer(1, 6))
            
            # Add separator between exchanges
            if i < len(exchanges) - 1:
                story.append(Spacer(1, 12))
                story.append(Paragraph("─" * 30, normal_style))
                story.append(Spacer(1, 12))
        
        # Build the PDF
        doc.build(story)
        
        filename = os.path.basename(file_path)
        self._set_status(f"✅ Complete document saved to {filename}")

