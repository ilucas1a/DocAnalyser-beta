"""
transcript_qa_splitter.py
=========================
Heuristic interviewer / interviewee splitter for Faster Whisper transcripts.

Classifies each fragment as a Question (interviewer) or Answer (interviewee)
using three signals:
  1. Fragment ends with a question mark
  2. Fragment is very short (<=6 words) — typical interviewer interjection
  3. Fragment follows a long response block (contextual)

Both panels support manual merging and a Move button to fix misclassifications.
Output: a .docx with [Q] and [A] paragraph prefixes.

Usage:
    python maintenance/transcript_qa_splitter.py
    python maintenance/transcript_qa_splitter.py "C:\\path\\to\\transcript.docx"

Output:
    transcript_qa.docx  (same folder as input, never overwrites original)

Requirements:
    pip install python-docx
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import os, re, sys
from copy import deepcopy
from typing import List, Optional, Tuple

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx not installed.  Run:  pip install python-docx")
    sys.exit(1)


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

# Fragments this short are likely interviewer interjections regardless of ?
SHORT_WORD_THRESHOLD = 6

# If a run of A-fragments exceeds this many words, the next short/? fragment
# is very likely the interviewer returning
LONG_RESPONSE_WORD_COUNT = 40

# Colours
COL = {
    'Q_frag':    '#e3f2fd',   # light blue — unprocessed Q
    'Q_merged':  '#1565c0',   # dark blue  — merged Q group
    'A_frag':    '#fff8e1',   # light amber — unprocessed A
    'A_merged':  '#e65100',   # dark orange — merged A group
    'special':   '#eeeeee',   # grey — timestamp / speaker / blank
}
FG = {
    'Q_frag':    '#000000',
    'Q_merged':  '#ffffff',
    'A_frag':    '#000000',
    'A_merged':  '#ffffff',
    'special':   '#888888',
}

TIMESTAMP_PAT = re.compile(r'^\[?\d{1,2}:\d{2}(?::\d{2})?\]?\s*$')
SPEAKER_PAT   = re.compile(
    r'^\s*(\[?(?:Speaker\s+)?[A-Z]\]?|INTERVIEWER|INTERVIEWEE|'
    r'[A-Z][A-Z\s]{0,20})\s*[:\-]\s*',
    re.IGNORECASE,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _is_special(text: str) -> bool:
    t = text.strip()
    if not t:
        return True
    if TIMESTAMP_PAT.match(t):
        return True
    if SPEAKER_PAT.match(t) and len(t) < 40:
        return True
    return False


def _word_count(text: str) -> int:
    return len(text.split())


def _copy_para_format(src, dst):
    try:
        src_pPr = src._p.find(qn('w:pPr'))
        dst_pPr = dst._p.find(qn('w:pPr'))
        if src_pPr is not None:
            new_pPr = deepcopy(src_pPr)
            if dst_pPr is not None:
                dst._p.remove(dst_pPr)
            dst._p.insert(0, new_pPr)
    except Exception:
        pass


# ---------------------------------------------------------------------------
# Data model
# ---------------------------------------------------------------------------

class Fragment:
    """One unit in either the Q or A list."""

    def __init__(self, texts: List[str], role: str,
                 kind: str = 'frag', para_ref=None):
        """
        role:  'Q' | 'A' | 'special'
        kind:  'frag' | 'merged' | 'special'
        """
        self.texts    = list(texts)
        self.role     = role
        self.kind     = kind
        self.para_ref = para_ref

    @property
    def full_text(self) -> str:
        return ' '.join(t for t in self.texts if t)

    @property
    def is_mergeable(self) -> bool:
        return self.kind in ('frag', 'merged')

    def display(self) -> str:
        txt = self.full_text
        prefix = f'[{len(self.texts)}↕] ' if self.kind == 'merged' else '  '
        short = txt if len(txt) <= 90 else txt[:87] + '...'
        return prefix + short

    def colour_key(self) -> str:
        if self.kind == 'special':
            return 'special'
        return f'{self.role}_{self.kind}'


# ---------------------------------------------------------------------------
# Classify raw fragments into Q / A
# ---------------------------------------------------------------------------

def classify_fragments(doc: Document) -> List[Fragment]:
    """
    Walk all paragraphs; return a flat list of Fragment objects
    with role Q, A, or special.
    """
    raw = []
    for para in doc.paragraphs:
        text = para.text.strip()
        raw.append((text, para))

    result: List[Fragment] = []
    words_since_q = 0   # cumulative A-words since last Q fragment

    for text, para in raw:
        if not text:
            result.append(Fragment([''], 'special', 'special', para))
            continue

        if _is_special(text):
            result.append(Fragment([text], 'special', 'special', para))
            continue

        # Classify
        has_question_mark = text.rstrip().endswith('?')
        wc = _word_count(text)
        is_short = wc <= SHORT_WORD_THRESHOLD
        follows_long_response = words_since_q >= LONG_RESPONSE_WORD_COUNT

        if has_question_mark or (is_short and follows_long_response):
            role = 'Q'
            words_since_q = 0
        else:
            role = 'A'
            words_since_q += wc

        result.append(Fragment([text], role, 'frag', para))

    return result


# ---------------------------------------------------------------------------
# Main application
# ---------------------------------------------------------------------------

class QASplitter:

    def __init__(self, root: tk.Tk):
        self.root        = root
        self.fragments:  List[Fragment] = []
        self.q_items:    List[Fragment] = []   # Q panel
        self.a_items:    List[Fragment] = []   # A panel
        self.q_undo:     List[tuple]    = []
        self.a_undo:     List[tuple]    = []
        self.source_doc: Optional[Document] = None
        self.input_path  = ''
        self.output_path = ''

        self.root.title('Transcript Q/A Splitter')
        self.root.geometry('1200x720')
        self.root.minsize(900, 500)

        self._build_ui()
        self._bind_keys()
        self.root.after(150, self._prompt_open)

    # ------------------------------------------------------------------
    # UI
    # ------------------------------------------------------------------

    def _build_ui(self):
        # Top bar
        top = ttk.Frame(self.root, padding=(10, 6))
        top.pack(fill=tk.X)
        self._title_var = tk.StringVar(value='No file loaded')
        ttk.Label(top, textvariable=self._title_var,
                  font=('Arial', 11, 'bold')).pack(side=tk.LEFT)
        ttk.Button(top, text='Open different file',
                   command=self._prompt_open).pack(side=tk.RIGHT)

        # Stats
        stats = ttk.Frame(self.root, padding=(10, 0))
        stats.pack(fill=tk.X)
        self._stats_var = tk.StringVar(value='')
        ttk.Label(stats, textvariable=self._stats_var,
                  foreground='#555', font=('Arial', 8)).pack(side=tk.LEFT)

        # Instructions
        inst = ttk.Frame(self.root, padding=(10, 2))
        inst.pack(fill=tk.X)
        ttk.Label(
            inst,
            text=('Shift+click or Ctrl+click to select multiple rows in a panel.  '
                  'Merge joins them.  Move sends selected row(s) to the other panel.  '
                  'Undo undoes the last merge in the active panel.'),
            foreground='#444', font=('Arial', 9),
        ).pack(side=tk.LEFT)

        ttk.Separator(self.root, orient='horizontal').pack(fill=tk.X)

        # ── Two-panel split ─────────────────────────────────────────
        pane = tk.PanedWindow(self.root, orient=tk.HORIZONTAL,
                               sashwidth=8, bg='#aaaaaa')
        pane.pack(fill=tk.BOTH, expand=True, padx=2, pady=2)

        self._q_frame = ttk.Frame(pane, padding=4)
        self._a_frame = ttk.Frame(pane, padding=4)
        pane.add(self._q_frame, minsize=350)
        pane.add(self._a_frame, minsize=350)

        self._q_lb, self._q_vsb = self._build_panel(
            self._q_frame,
            '❓  Interviewer / Questions  (blue)',
            '#e3f2fd',
            side='Q',
        )
        self._a_lb, self._a_vsb = self._build_panel(
            self._a_frame,
            '💬  Interviewee / Answers  (amber)',
            '#fff8e1',
            side='A',
        )

        # ── Bottom action bar ───────────────────────────────────────
        ttk.Separator(self.root, orient='horizontal').pack(fill=tk.X)
        btns = ttk.Frame(self.root, padding=(10, 8))
        btns.pack(fill=tk.X)

        self._q_merge_btn = ttk.Button(
            btns, text='⤵ Merge Q selection  (Ctrl+M)',
            command=lambda: self._merge(side='Q'),
            state=tk.DISABLED, width=26,
        )
        self._q_merge_btn.pack(side=tk.LEFT, padx=(0, 4))

        self._q_undo_btn = ttk.Button(
            btns, text='↩ Undo Q  (Ctrl+Z)',
            command=lambda: self._undo(side='Q'),
            state=tk.DISABLED, width=16,
        )
        self._q_undo_btn.pack(side=tk.LEFT, padx=4)

        self._move_to_a_btn = ttk.Button(
            btns, text='→ Move to A',
            command=lambda: self._move(from_side='Q'),
            state=tk.DISABLED, width=12,
        )
        self._move_to_a_btn.pack(side=tk.LEFT, padx=8)

        self._move_to_q_btn = ttk.Button(
            btns, text='← Move to Q',
            command=lambda: self._move(from_side='A'),
            state=tk.DISABLED, width=12,
        )
        self._move_to_q_btn.pack(side=tk.LEFT, padx=4)

        self._a_merge_btn = ttk.Button(
            btns, text='⤵ Merge A selection  (Ctrl+Shift+M)',
            command=lambda: self._merge(side='A'),
            state=tk.DISABLED, width=30,
        )
        self._a_merge_btn.pack(side=tk.LEFT, padx=8)

        self._a_undo_btn = ttk.Button(
            btns, text='↩ Undo A  (Ctrl+Shift+Z)',
            command=lambda: self._undo(side='A'),
            state=tk.DISABLED, width=20,
        )
        self._a_undo_btn.pack(side=tk.LEFT, padx=4)

        # Right side
        ttk.Button(btns, text='✖ Cancel',
                   command=self.root.destroy, width=10).pack(side=tk.RIGHT, padx=4)
        self._save_btn = ttk.Button(
            btns, text='💾 Save & Close  (Ctrl+S)',
            command=self._save_and_close,
            state=tk.DISABLED, width=22,
        )
        self._save_btn.pack(side=tk.RIGHT, padx=4)

        # Status bar
        sb = ttk.Frame(self.root, relief=tk.SUNKEN)
        sb.pack(fill=tk.X, side=tk.BOTTOM)
        self._status_var = tk.StringVar(value='Open a transcript .docx to begin.')
        ttk.Label(sb, textvariable=self._status_var,
                  anchor=tk.W, padding=(6, 2)).pack(fill=tk.X)

    def _build_panel(self, parent, label: str, bg: str, side: str):
        """Build one Q or A panel; return (listbox, scrollbar)."""
        header = ttk.Frame(parent)
        header.pack(fill=tk.X, pady=(0, 4))
        ttk.Label(header, text=label,
                  font=('Arial', 10, 'bold')).pack(side=tk.LEFT)
        count_var = tk.StringVar(value='')
        ttk.Label(header, textvariable=count_var,
                  foreground='#777', font=('Arial', 8)).pack(side=tk.RIGHT)
        if side == 'Q':
            self._q_count_var = count_var
        else:
            self._a_count_var = count_var

        frame = ttk.Frame(parent)
        frame.pack(fill=tk.BOTH, expand=True)

        vsb = ttk.Scrollbar(frame, orient=tk.VERTICAL)
        vsb.pack(side=tk.RIGHT, fill=tk.Y)
        hsb = ttk.Scrollbar(frame, orient=tk.HORIZONTAL)
        hsb.pack(side=tk.BOTTOM, fill=tk.X)

        lb = tk.Listbox(
            frame,
            yscrollcommand=vsb.set,
            xscrollcommand=hsb.set,
            selectmode=tk.EXTENDED,
            font=('Consolas', 9),
            activestyle='dotbox',
            exportselection=False,
            bg=bg,
        )
        lb.pack(fill=tk.BOTH, expand=True)
        vsb.config(command=lb.yview)
        hsb.config(command=lb.xview)

        lb.bind('<<ListboxSelect>>',
                lambda e, s=side: self._on_select(s))
        lb.bind('<Double-Button-1>',
                lambda e, s=side: self._merge(s))

        return lb, vsb

    def _bind_keys(self):
        self.root.bind('<Control-m>',       lambda e: self._merge('Q'))
        self.root.bind('<Control-M>',       lambda e: self._merge('Q'))
        self.root.bind('<Control-z>',       lambda e: self._undo('Q'))
        self.root.bind('<Control-Z>',       lambda e: self._undo('Q'))
        self.root.bind('<Control-Shift-m>', lambda e: self._merge('A'))
        self.root.bind('<Control-Shift-M>', lambda e: self._merge('A'))
        self.root.bind('<Control-Shift-z>', lambda e: self._undo('A'))
        self.root.bind('<Control-Shift-Z>', lambda e: self._undo('A'))
        self.root.bind('<Control-s>',       lambda e: self._save_and_close())
        self.root.bind('<Control-S>',       lambda e: self._save_and_close())
        self.root.bind('<Escape>',          lambda e: self.root.destroy())

    # ------------------------------------------------------------------
    # File I/O
    # ------------------------------------------------------------------

    def _prompt_open(self):
        if len(sys.argv) > 1:
            path = sys.argv[1].strip('"').strip("'")
            if os.path.exists(path):
                self._load(path)
                return
        path = filedialog.askopenfilename(
            title='Select transcript .docx file',
            filetypes=[('Word documents', '*.docx *.docm'), ('All files', '*.*')],
        )
        if path:
            self._load(path)
        else:
            self.root.destroy()

    def _load(self, path: str):
        try:
            doc = Document(path)
        except Exception as exc:
            messagebox.showerror('Cannot open file', str(exc))
            return

        self.source_doc  = doc
        self.input_path  = path
        base, _          = os.path.splitext(path)
        self.output_path = base + '_qa.docx'

        self.fragments   = classify_fragments(doc)
        self.q_items     = [f for f in self.fragments if f.role == 'Q']
        self.a_items     = [f for f in self.fragments if f.role in ('A', 'special')]

        # Build document-order lookup keyed on the underlying lxml element
        # (para._p) which is stable — python-docx Paragraph wrapper objects
        # are re-created on each .paragraphs access so id(para) is unreliable.
        self._para_order = {
            id(para._p): idx
            for idx, para in enumerate(doc.paragraphs)
        }

        self.q_undo.clear()
        self.a_undo.clear()

        fname = os.path.basename(path)
        self.root.title(f'Transcript Q/A Splitter — {fname}')
        self._title_var.set(f'Q/A Splitter — {fname}')
        self._save_btn.config(state=tk.NORMAL)
        self._refresh('Q')
        self._refresh('A')
        self._update_stats()

        n_q = len(self.q_items)
        n_a = len([f for f in self.a_items if f.role == 'A'])
        self._status(
            f'Loaded {fname}  —  {n_q} Q fragments, {n_a} A fragments.  '
            f'Review both panels and correct any misclassifications.'
        )

    # ------------------------------------------------------------------
    # List management
    # ------------------------------------------------------------------

    def _items(self, side: str) -> List[Fragment]:
        return self.q_items if side == 'Q' else self.a_items

    def _lb(self, side: str) -> tk.Listbox:
        return self._q_lb if side == 'Q' else self._a_lb

    def _refresh(self, side: str, scroll_to: Optional[int] = None):
        lb   = self._lb(side)
        items = self._items(side)
        lb.delete(0, tk.END)
        for idx, frag in enumerate(items):
            lb.insert(tk.END, frag.display())
            key = frag.colour_key()
            lb.itemconfig(idx,
                          bg=COL.get(key, '#ffffff'),
                          fg=FG.get(key, '#000000'),
                          selectbackground='#546e7a',
                          selectforeground='#ffffff')
        if scroll_to is not None:
            pos = max(0, min(scroll_to, lb.size() - 1))
            lb.see(pos)
            lb.selection_clear(0, tk.END)
            lb.selection_set(pos)

        # Update count label
        n_frag   = len([f for f in items if f.kind == 'frag'])
        n_merged = len([f for f in items if f.kind == 'merged'])
        count_var = self._q_count_var if side == 'Q' else self._a_count_var
        count_var.set(f'{n_frag} fragments  |  {n_merged} merged groups')

    def _update_stats(self):
        nq = len(self.q_items)
        na = len([f for f in self.a_items if f.role == 'A'])
        ns = len([f for f in self.a_items if f.role == 'special'])
        self._stats_var.set(
            f'Q panel: {nq} rows     '
            f'A panel: {na} rows     '
            f'Special / blank: {ns}'
        )

    # ------------------------------------------------------------------
    # Selection
    # ------------------------------------------------------------------

    def _on_select(self, side: str):
        sel = self._lb(side).curselection()
        items = self._items(side)
        mergeable = [i for i in sel if items[i].is_mergeable]

        # Enable/disable buttons
        can_merge = len(mergeable) >= 2
        if side == 'Q':
            self._q_merge_btn.config(state=tk.NORMAL if can_merge else tk.DISABLED)
            self._move_to_a_btn.config(
                state=tk.NORMAL if mergeable else tk.DISABLED)
        else:
            self._a_merge_btn.config(state=tk.NORMAL if can_merge else tk.DISABLED)
            self._move_to_q_btn.config(
                state=tk.NORMAL if mergeable else tk.DISABLED)

    # ------------------------------------------------------------------
    # Merge
    # ------------------------------------------------------------------

    def _merge(self, side: str):
        lb    = self._lb(side)
        items = self._items(side)
        sel   = lb.curselection()
        to_merge = [i for i in sorted(sel) if items[i].is_mergeable]

        if len(to_merge) < 2:
            self._status(f'Select at least 2 {side} fragments to merge.')
            return

        first, last = to_merge[0], to_merge[-1]

        # Save undo
        orig_first = deepcopy(items[first])
        absorbed   = [(i, deepcopy(items[i])) for i in to_merge[1:]]
        undo_stack = self.q_undo if side == 'Q' else self.a_undo
        undo_stack.append((first, orig_first, absorbed))

        # Merge texts into first
        all_texts = []
        for i in to_merge:
            all_texts.extend(items[i].texts)

        new_frag = Fragment(all_texts, side, 'merged', items[first].para_ref)
        items[first] = new_frag

        for i in sorted([idx for idx, _ in absorbed], reverse=True):
            items.pop(i)

        btn = self._q_undo_btn if side == 'Q' else self._a_undo_btn
        btn.config(state=tk.NORMAL)

        self._refresh(side, scroll_to=first)
        self._update_stats()
        self._status(f'Merged {len(to_merge)} {side} fragments.  Ctrl+Z to undo.')

    # ------------------------------------------------------------------
    # Undo
    # ------------------------------------------------------------------

    def _undo(self, side: str):
        undo_stack = self.q_undo if side == 'Q' else self.a_undo
        items      = self._items(side)

        if not undo_stack:
            self._status(f'Nothing to undo in {side} panel.')
            return

        first_idx, orig_first, absorbed = undo_stack.pop()
        items[first_idx] = orig_first
        for i, frag in sorted(absorbed, key=lambda x: x[0]):
            items.insert(i, frag)

        btn = self._q_undo_btn if side == 'Q' else self._a_undo_btn
        if not undo_stack:
            btn.config(state=tk.DISABLED)

        self._refresh(side, scroll_to=first_idx)
        self._update_stats()
        self._status(f'Undone last {side} merge.')

    # ------------------------------------------------------------------
    # Move between panels
    # ------------------------------------------------------------------

    def _move(self, from_side: str):
        lb        = self._lb(from_side)
        items     = self._items(from_side)
        to_side   = 'A' if from_side == 'Q' else 'Q'
        to_items  = self._items(to_side)

        sel = lb.curselection()
        to_move = [i for i in sorted(sel, reverse=True)
                   if items[i].is_mergeable]

        if not to_move:
            self._status('Select at least one fragment row to move.')
            return

        # Use _para_order (built at load time using _p XML elements as keys)
        # to find the correct insertion point in the destination panel.
        # We cannot use id(para) here because python-docx creates new wrapper
        # objects each time .paragraphs is accessed, giving different id()
        # values for the same underlying XML element.
        def _doc_pos(frag):
            if frag.para_ref is None:
                return 99999
            return self._para_order.get(id(frag.para_ref._p), 99999)

        moved = []
        for i in to_move:
            frag = items.pop(i)
            frag.role = to_side
            moved.append(frag)

        # Insert each moved fragment at its correct chronological position.
        first_inserted = len(to_items)  # fallback scroll target
        for frag in reversed(moved):    # reversed so earlier frags stay first
            pos = _doc_pos(frag)
            insert_at = len(to_items)   # default: end
            for idx, existing in enumerate(to_items):
                if _doc_pos(existing) > pos:
                    insert_at = idx
                    break
            to_items.insert(insert_at, frag)
            first_inserted = insert_at

        self._refresh(from_side)
        self._refresh(to_side, scroll_to=first_inserted)
        self._update_stats()
        self._status(
            f'Moved {len(moved)} fragment(s) from {from_side} panel to {to_side} panel.'
        )

    # ------------------------------------------------------------------
    # Save
    # ------------------------------------------------------------------

    def _save_and_close(self):
        if not (self.q_items or self.a_items):
            self.root.destroy()
            return

        n_unfused_a = len([f for f in self.a_items
                           if f.kind == 'frag' and f.role == 'A'])
        if n_unfused_a > 5:
            ans = messagebox.askyesnocancel(
                'Many unmerged A fragments',
                f'{n_unfused_a} Answer fragments have not been merged.\n\n'
                'Each will become its own paragraph.\n\n'
                'Save anyway?',
            )
            if ans is None or not ans:
                return

        try:
            out = Document()

            # Mirror source layout
            if self.source_doc:
                ss  = self.source_doc.sections[0]
                os_ = out.sections[0]
                os_.page_width    = ss.page_width
                os_.page_height   = ss.page_height
                os_.top_margin    = ss.top_margin
                os_.bottom_margin = ss.bottom_margin
                os_.left_margin   = ss.left_margin
                os_.right_margin  = ss.right_margin
                try:
                    sn = self.source_doc.styles['Normal']
                    on = out.styles['Normal']
                    if sn.font.name:  on.font.name = sn.font.name
                    if sn.font.size:  on.font.size = sn.font.size
                except Exception:
                    pass

            # Interleave Q and A items in their original document order
            # by rebuilding from self.fragments, mapping back to merged state
            # Build lookup: original para id → merged Fragment
            all_out_frags = self.q_items + self.a_items

            # Sort by original document order using the stable _para_order
            # built at load time (keyed on para._p, not id(para)).
            def sort_key(frag: Fragment):
                if frag.para_ref is None:
                    return 99999
                return self._para_order.get(id(frag.para_ref._p), 99999)

            ordered = sorted(all_out_frags, key=sort_key)

            for frag in ordered:
                if frag.kind == 'special':
                    p = out.add_paragraph(frag.full_text)
                else:
                    prefix = '[Q]  ' if frag.role == 'Q' else '[A]  '
                    p = out.add_paragraph(prefix + frag.full_text)
                if frag.para_ref:
                    _copy_para_format(frag.para_ref, p)

            out.save(self.output_path)
            messagebox.showinfo(
                'Saved',
                f'Document saved to:\n\n{self.output_path}\n\n'
                f'[Q] marks interviewer paragraphs.\n'
                f'[A] marks interviewee paragraphs.\n'
                f'Remove these prefixes in Word with Find & Replace when done.',
            )
            self.root.destroy()

        except Exception as exc:
            messagebox.showerror('Save error', str(exc))

    # ------------------------------------------------------------------
    # Status
    # ------------------------------------------------------------------

    def _status(self, msg: str):
        self._status_var.set(msg)


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main():
    root = tk.Tk()
    try:
        icon = os.path.join(
            os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
            'DocAnalyser.ico',
        )
        if os.path.exists(icon):
            root.iconbitmap(icon)
    except Exception:
        pass

    QASplitter(root)
    root.mainloop()


if __name__ == '__main__':
    main()
