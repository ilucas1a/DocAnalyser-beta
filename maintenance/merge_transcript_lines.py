"""
merge_transcript_lines.py
=========================
Merges excessively short lines in a Faster Whisper transcript Word document.

Faster Whisper splits on every audio pause, producing many 2-6 word fragments
per line. This script joins them into natural-length paragraphs while:
  - Preserving blank lines (paragraph / speaker breaks)
  - Preserving timestamp lines (e.g. [00:01:23]) unchanged
  - Preserving speaker label lines (e.g. "[A]:", "Speaker A:", "INTERVIEWER:")
  - Keeping paragraphs between --min-length and --max-length characters

Usage:
    python merge_transcript_lines.py                      # prompts for file
    python merge_transcript_lines.py transcript.docx      # processes that file
    python merge_transcript_lines.py transcript.docx --min-length 80 --max-length 400

Output:
    transcript_merged.docx  (same folder as input, never overwrites original)

Requirements:
    pip install python-docx
"""

import sys
import os
import re
import argparse
from copy import deepcopy

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx not installed.")
    print("Install it with:  pip install python-docx")
    sys.exit(1)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

TIMESTAMP_PAT = re.compile(r'^\[?\d{1,2}:\d{2}(?::\d{2})?\]?\s*$')
SPEAKER_PAT   = re.compile(
    r'^\s*(\[?(?:Speaker\s+)?[A-Z]\]?|INTERVIEWER|INTERVIEWEE|'
    r'[A-Z][A-Z\s]{0,20})\s*[:\-]\s*',
    re.IGNORECASE
)
# Sentence-ending pattern: period, !, ?, ellipsis, optionally followed by
# closing quote/bracket
SENTENCE_END_PAT = re.compile(r'[.!?…]["\')\]]*\s*$')


def is_timestamp(text: str) -> bool:
    return bool(TIMESTAMP_PAT.match(text.strip()))


def is_speaker_label(text: str) -> bool:
    t = text.strip()
    return bool(SPEAKER_PAT.match(t)) and len(t) < 40


def looks_like_sentence_end(text: str) -> bool:
    return bool(SENTENCE_END_PAT.search(text.rstrip()))


def find_best_break(text: str, target: int) -> int:
    """
    Find the best character position to break a long text near `target`.
    Tries sentence boundaries first, then word boundaries.
    Returns an index into `text` (exclusive end of first part).
    """
    # Look for a sentence-ending punctuation near the target (within ±60 chars)
    window_start = max(0, target - 60)
    window_end   = min(len(text), target + 60)
    chunk = text[window_start:window_end]

    # Find the LAST sentence boundary in the window
    best = -1
    for m in SENTENCE_END_PAT.finditer(chunk):
        best = window_start + m.end()

    if best > 0:
        return best

    # Fallback: nearest word boundary at or before target
    idx = min(target, len(text) - 1)
    while idx > 0 and text[idx] != ' ':
        idx -= 1
    return idx if idx > 0 else target


def copy_paragraph_format(src_para, dst_para):
    """Copy paragraph-level XML properties (style, spacing, indent) from src."""
    src_pPr = src_para._p.find(qn('w:pPr'))
    dst_pPr = dst_para._p.find(qn('w:pPr'))
    if src_pPr is not None:
        new_pPr = deepcopy(src_pPr)
        if dst_pPr is not None:
            dst_para._p.remove(dst_pPr)
        dst_para._p.insert(0, new_pPr)


# ---------------------------------------------------------------------------
# Core merge logic
# ---------------------------------------------------------------------------

def merge_paragraphs(doc: Document, min_length: int, max_length: int) -> Document:
    """
    Return a new Document with short lines merged and long screeds broken up.

    Rules applied in order for each incoming paragraph:
      1. Blank line  → flush accumulator; emit a blank paragraph.
      2. Timestamp   → flush; emit verbatim.
      3. Speaker label → flush; emit verbatim.
      4. Normal text → add to accumulator.
         After adding:
           a. If accumulated length >= min_length AND ends on sentence
              punctuation → flush (natural paragraph break).
           b. If accumulated length >= max_length → chop into <=max_length
              pieces at sentence / word boundaries, flush each as a paragraph.
      5. End of document → flush anything remaining.
    """
    out = Document()

    # Mirror page layout from source
    src_sec = doc.sections[0]
    out_sec = out.sections[0]
    out_sec.page_width    = src_sec.page_width
    out_sec.page_height   = src_sec.page_height
    out_sec.top_margin    = src_sec.top_margin
    out_sec.bottom_margin = src_sec.bottom_margin
    out_sec.left_margin   = src_sec.left_margin
    out_sec.right_margin  = src_sec.right_margin

    # Mirror Normal style font
    try:
        src_n = doc.styles['Normal']
        out_n = out.styles['Normal']
        if src_n.font.name:
            out_n.font.name = src_n.font.name
        if src_n.font.size:
            out_n.font.size = src_n.font.size
    except Exception:
        pass

    accumulator = []   # list of (text_str, original_para) tuples
    first_para  = None # reference for paragraph formatting

    def emit(text: str):
        """Write a single paragraph to the output document."""
        p = out.add_paragraph(text.strip())
        if first_para is not None:
            copy_paragraph_format(first_para, p)

    def chop_and_emit(text: str):
        """Break text into <=max_length chunks at natural boundaries."""
        while len(text) > max_length:
            break_at = find_best_break(text, max_length)
            emit(text[:break_at].strip())
            text = text[break_at:].strip()
        if text:
            emit(text)

    def flush(force: bool = False):
        nonlocal first_para
        if not accumulator:
            return
        joined = ' '.join(t for t, _ in accumulator)

        # Not yet ready to flush?
        if (not force
                and len(joined) < min_length
                and not looks_like_sentence_end(joined)):
            return

        # Emit — chopping if over max_length
        if len(joined) > max_length:
            chop_and_emit(joined)
        else:
            emit(joined)

        accumulator.clear()
        first_para = None

    def flush_force():
        flush(force=True)

    for para in doc.paragraphs:
        text = para.text

        # Rule 1: blank line
        if not text.strip():
            flush_force()
            out.add_paragraph('')
            continue

        # Rule 2: timestamp
        if is_timestamp(text):
            flush_force()
            p = out.add_paragraph(text)
            copy_paragraph_format(para, p)
            continue

        # Rule 3: speaker label
        if is_speaker_label(text):
            flush_force()
            p = out.add_paragraph(text)
            copy_paragraph_format(para, p)
            continue

        # Rule 4: normal text → accumulate
        if first_para is None:
            first_para = para
        accumulator.append((text.strip(), para))

        joined = ' '.join(t for t, _ in accumulator)

        # 4a: natural flush at sentence boundary
        if len(joined) >= min_length and looks_like_sentence_end(joined):
            flush(force=True)
            continue

        # 4b: forced chop when over max_length
        if len(joined) >= max_length:
            chop_and_emit(joined)
            accumulator.clear()
            first_para = None

    # Rule 5: end of document
    flush_force()

    return out


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description='Merge short transcript lines in a Word document.'
    )
    parser.add_argument(
        'input', nargs='?',
        help='Path to input .docx file (prompted if omitted)'
    )
    parser.add_argument(
        '--min-length', type=int, default=80,
        help='Minimum chars before a paragraph break is considered (default: 80)'
    )
    parser.add_argument(
        '--max-length', type=int, default=400,
        help='Maximum chars before a paragraph is forcibly broken (default: 400)'
    )
    args = parser.parse_args()

    input_path = args.input
    if not input_path:
        print("Transcript Line Merger")
        print("=" * 40)
        input_path = input(
            "Enter path to transcript .docx file\n"
            "(tip: drag the file onto this window, then press Enter):\n> "
        ).strip().strip('"').strip("'")

    if not os.path.exists(input_path):
        print(f"ERROR: File not found: {input_path}")
        sys.exit(1)

    if not input_path.lower().endswith('.docx'):
        print("ERROR: File must be a .docx file.")
        sys.exit(1)

    base, ext = os.path.splitext(input_path)
    output_path = base + '_merged' + ext

    print(f"\nInput:       {input_path}")
    print(f"Output:      {output_path}")
    print(f"Min length:  {args.min_length} chars")
    print(f"Max length:  {args.max_length} chars")
    print()

    print("Reading document...")
    doc = Document(input_path)
    total_paras = len([p for p in doc.paragraphs if p.text.strip()])
    print(f"  {total_paras} non-empty paragraphs found.")

    print("Merging and reformatting...")
    merged_doc = merge_paragraphs(doc, args.min_length, args.max_length)
    out_paras = len([p for p in merged_doc.paragraphs if p.text.strip()])
    reduction = total_paras - out_paras
    print(f"  {out_paras} paragraphs in output  "
          f"({'reduced by ' + str(reduction) if reduction > 0 else 'no change'}).")

    print("Saving...")
    merged_doc.save(output_path)
    print(f"\n✅ Done!  Saved to:\n   {output_path}")
    print(f"\nTip: If paragraphs are still too long or too short, re-run with")
    print(f"     different values, e.g.:")
    print(f"     --min-length 100 --max-length 350")


if __name__ == '__main__':
    main()
