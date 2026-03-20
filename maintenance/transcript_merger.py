"""
transcript_merger.py
====================
Interactive tool for merging short transcript fragments into coherent paragraphs.

After Faster Whisper transcribes audio it produces many short pause-group
fragments (2-6 words each). This tool shows all fragments in a scrollable
list. You select the ones that belong together and click Merge; they become
one paragraph. When you are happy with the result, Save writes a clean .docx.

Usage:
    python maintenance/transcript_merger.py
    python maintenance/transcript_merger.py "C:\\path\\to\\transcript.docx"

Output:
    transcript_merged.docx  in the same folder as the input (never overwrites
    the original).

Keyboard shortcuts:
    Ctrl+M  Merge selected
    Ctrl+Z  Undo last merge
    Ctrl+S  Save & close
    Escape  Cancel / close

Requirements:
    pip install python-docx      (already in the DocAnalyser venv)
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import os
import re
import sys
from copy import deepcopy
from typing import List, Optional, Tuple

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx not installed.  Run:  pip install python-docx")
    sys.exit(1)


# ---------------------------------------------------------------------------
# Helpers — same logic as merge_transcript_lines.py
# ---------------------------------------------------------------------------

TIMESTAMP_PAT = re.compile(r'^\[?\d{1,2}:\d{2}(?::\d{2})?\]?\s*$')
SPEAKER_PAT   = re.compile(
    r'^\s*(\[?(?:Speaker\s+)?[A-Z]\]?|INTERVIEWER|INTERVIEWEE|'
    r'[A-Z][A-Z\s]{0,20})\s*[:\-]\s*',
    re.IGNORECASE,
)


def _is_special(text: str) -> bool:
    """Return True for timestamps, speaker labels, and blanks — not mergeable."""
    t = text.strip()
    if not t:
        return True
    if TIMESTAMP_PAT.match(t):
        return True
    if SPEAKER_PAT.match(t) and len(t) < 40:
        return True
    return False


def _copy_para_format(src, dst):
    """Mirror paragraph-level XML properties (style, spacing, indent)."""
    try:
        src_pPr = src._p.find(qn('w:pPr'))
        dst_pPr = dst._p.find(qn('w:pPr'))
        if src_pPr is not None:
            new_pPr = deepcopy(src_pPr)
            if dst_pPr is not None:
                dst._p.remove(dst_pPr)
            dst._p.insert(0, new_pPr)
    except Exception:
        pass


# ---------------------------------------------------------------------------
# Data model
# ---------------------------------------------------------------------------

class Item:
    """One displayable row in the merger list."""

    def __init__(self, texts: List[str], kind: str, para_ref=None):
        self.texts    = list(texts)   # one text per original fragment
        self.kind     = kind          # 'fragment' | 'merged' | 'special' | 'blank'
        self.para_ref = para_ref      # first original paragraph (for formatting)

    @property
    def is_mergeable(self) -> bool:
        return self.kind in ('fragment', 'merged')

    @property
    def full_text(self) -> str:
        return ' '.join(t for t in self.texts if t)

    def display(self) -> str:
        if self.kind == 'blank':
            return ''
        txt = self.full_text
        if self.kind == 'merged':
            label = f'[{len(self.texts)} merged]  '
        elif self.kind == 'special':
            label = '  '
        else:
            label = '  '
        display_text = txt if len(txt) <= 95 else txt[:92] + '...'
        return label + display_text


# ---------------------------------------------------------------------------
# Main application
# ---------------------------------------------------------------------------

class TranscriptMerger:

    # Listbox background colours
    BG = {
        'fragment': '#ffffff',   # white — unprocessed
        'merged':   '#c8e6c9',   # green — merged group
        'special':  '#eeeeee',   # grey  — timestamp / speaker / blank
        'blank':    '#f9f9f9',   # near-white — blank line
    }

    def __init__(self, root: tk.Tk):
        self.root       = root
        self.items:      List[Item]  = []
        self.undo_stack: List[Tuple] = []   # (first_idx, orig_first_item, absorbed)
        self.input_path  = ''
        self.output_path = ''
        self.source_doc: Optional[Document] = None

        self.root.title('Transcript Merger')
        self.root.geometry('980x680')
        self.root.minsize(720, 500)

        self._build_ui()
        self._bind_keys()
        self.root.after(150, self._prompt_open)

    # ------------------------------------------------------------------
    # UI construction
    # ------------------------------------------------------------------

    def _build_ui(self):
        # ── Top toolbar ───────────────────────────────────────────────
        top = ttk.Frame(self.root, padding=(10, 6, 10, 4))
        top.pack(fill=tk.X)

        self._title_var = tk.StringVar(value='No file loaded')
        ttk.Label(top, textvariable=self._title_var,
                  font=('Arial', 11, 'bold')).pack(side=tk.LEFT)

        ttk.Button(top, text='Open different file',
                   command=self._prompt_open).pack(side=tk.RIGHT, padx=4)

        # ── Stats bar ─────────────────────────────────────────────────
        stats = ttk.Frame(self.root, padding=(10, 0, 10, 4))
        stats.pack(fill=tk.X)
        self._stats_var = tk.StringVar(value='')
        ttk.Label(stats, textvariable=self._stats_var,
                  foreground='#666666', font=('Arial', 8)).pack(side=tk.LEFT)

        # ── Instructions ──────────────────────────────────────────────
        inst = ttk.Frame(self.root, padding=(10, 2, 10, 4))
        inst.pack(fill=tk.X)
        ttk.Label(
            inst,
            text=(
                'Click a row to select.  '
                'Shift+click to extend selection.  '
                'Ctrl+click to add individual rows.  '
                'Then click  Merge Selected  to join them into one paragraph.'
            ),
            foreground='#444444',
            font=('Arial', 9),
        ).pack(side=tk.LEFT)

        ttk.Separator(self.root, orient='horizontal').pack(fill=tk.X)

        # ── Main split pane ───────────────────────────────────────────
        pane = tk.PanedWindow(self.root, orient=tk.HORIZONTAL,
                               sashwidth=6, relief=tk.FLAT, bg='#cccccc')
        pane.pack(fill=tk.BOTH, expand=True, padx=2, pady=2)

        # Left — fragment list
        left = ttk.Frame(pane, padding=0)
        pane.add(left, minsize=380, width=600)

        lbl_row = ttk.Frame(left)
        lbl_row.pack(fill=tk.X, padx=4, pady=(4, 2))
        ttk.Label(lbl_row, text='Transcript fragments',
                  font=('Arial', 9, 'bold')).pack(side=tk.LEFT)
        self._sel_info_var = tk.StringVar(value='')
        ttk.Label(lbl_row, textvariable=self._sel_info_var,
                  foreground='#777777', font=('Arial', 8)).pack(side=tk.RIGHT)

        lb_frame = ttk.Frame(left)
        lb_frame.pack(fill=tk.BOTH, expand=True, padx=4, pady=(0, 4))

        vsb = ttk.Scrollbar(lb_frame, orient=tk.VERTICAL)
        vsb.pack(side=tk.RIGHT, fill=tk.Y)
        hsb = ttk.Scrollbar(lb_frame, orient=tk.HORIZONTAL)
        hsb.pack(side=tk.BOTTOM, fill=tk.X)

        self._lb = tk.Listbox(
            lb_frame,
            yscrollcommand=vsb.set,
            xscrollcommand=hsb.set,
            selectmode=tk.EXTENDED,
            font=('Consolas', 9),
            activestyle='dotbox',
            exportselection=False,
            width=70,
        )
        self._lb.pack(fill=tk.BOTH, expand=True)
        vsb.config(command=self._lb.yview)
        hsb.config(command=self._lb.xview)
        self._lb.bind('<<ListboxSelect>>', self._on_select)
        self._lb.bind('<Double-Button-1>', lambda e: self._merge_selected())

        # Right — preview
        right = ttk.Frame(pane, padding=4)
        pane.add(right, minsize=260)

        ttk.Label(right, text='Merged paragraph preview',
                  font=('Arial', 9, 'bold')).pack(anchor=tk.W, pady=(4, 2))

        prev_frame = ttk.Frame(right)
        prev_frame.pack(fill=tk.BOTH, expand=True)
        pvsb = ttk.Scrollbar(prev_frame)
        pvsb.pack(side=tk.RIGHT, fill=tk.Y)
        self._preview = tk.Text(
            prev_frame,
            wrap=tk.WORD,
            font=('Arial', 10),
            state=tk.DISABLED,
            bg='#f5f5f5',
            relief='flat',
            yscrollcommand=pvsb.set,
            padx=8,
            pady=6,
        )
        self._preview.pack(fill=tk.BOTH, expand=True)
        pvsb.config(command=self._preview.yview)

        ttk.Label(right, text='Fragment count in selection:',
                  font=('Arial', 8), foreground='#888888').pack(anchor=tk.W, pady=(6, 0))
        self._count_var = tk.StringVar(value='—')
        ttk.Label(right, textvariable=self._count_var,
                  font=('Arial', 11, 'bold')).pack(anchor=tk.W)

        # ── Bottom action bar ─────────────────────────────────────────
        ttk.Separator(self.root, orient='horizontal').pack(fill=tk.X)
        btns = ttk.Frame(self.root, padding=(10, 8))
        btns.pack(fill=tk.X)

        self._merge_btn = ttk.Button(
            btns, text='⤵  Merge Selected  (Ctrl+M)',
            command=self._merge_selected, state=tk.DISABLED, width=28,
        )
        self._merge_btn.pack(side=tk.LEFT, padx=(0, 8))

        self._undo_btn = ttk.Button(
            btns, text='↩  Undo  (Ctrl+Z)',
            command=self._undo, state=tk.DISABLED, width=18,
        )
        self._undo_btn.pack(side=tk.LEFT, padx=4)

        ttk.Button(
            btns, text='✖  Cancel',
            command=self.root.destroy, width=10,
        ).pack(side=tk.RIGHT, padx=4)

        self._save_btn = ttk.Button(
            btns, text='💾  Save & Close  (Ctrl+S)',
            command=self._save_and_close, state=tk.DISABLED, width=24,
        )
        self._save_btn.pack(side=tk.RIGHT, padx=4)

        # ── Status bar ────────────────────────────────────────────────
        sb = ttk.Frame(self.root, relief=tk.SUNKEN)
        sb.pack(fill=tk.X, side=tk.BOTTOM)
        self._status_var = tk.StringVar(value='Ready — open a transcript file to begin.')
        ttk.Label(sb, textvariable=self._status_var,
                  anchor=tk.W, padding=(6, 2)).pack(fill=tk.X)

    def _bind_keys(self):
        self.root.bind('<Control-m>', lambda e: self._merge_selected())
        self.root.bind('<Control-M>', lambda e: self._merge_selected())
        self.root.bind('<Control-z>', lambda e: self._undo())
        self.root.bind('<Control-Z>', lambda e: self._undo())
        self.root.bind('<Control-s>', lambda e: self._save_and_close())
        self.root.bind('<Control-S>', lambda e: self._save_and_close())
        self.root.bind('<Escape>',    lambda e: self.root.destroy())

    # ------------------------------------------------------------------
    # File I/O
    # ------------------------------------------------------------------

    def _prompt_open(self):
        # Accept command-line argument first
        if len(sys.argv) > 1:
            path = sys.argv[1].strip('"').strip("'")
            if os.path.exists(path):
                self._load(path)
                return

        path = filedialog.askopenfilename(
            title='Select transcript .docx file',
            filetypes=[('Word documents', '*.docx *.docm'), ('All files', '*.*')],
        )
        if path:
            self._load(path)
        else:
            self.root.destroy()

    def _load(self, path: str):
        try:
            doc = Document(path)
        except Exception as exc:
            messagebox.showerror('Cannot open file', str(exc))
            return

        self.source_doc  = doc
        self.input_path  = path
        base, _          = os.path.splitext(path)
        self.output_path = base + '_merged.docx'

        self.items.clear()
        self.undo_stack.clear()

        for para in doc.paragraphs:
            raw = para.text.strip()
            if not raw:
                self.items.append(Item([''], 'blank', para))
            elif _is_special(raw):
                self.items.append(Item([raw], 'special', para))
            else:
                self.items.append(Item([raw], 'fragment', para))

        fname = os.path.basename(path)
        self._title_var.set(f'Transcript Merger — {fname}')
        self._save_btn.config(state=tk.NORMAL)
        self._refresh()
        n_frag = len([i for i in self.items if i.kind == 'fragment'])
        self._status(f'Loaded {fname}  ({n_frag} fragments to review)')

    # ------------------------------------------------------------------
    # List display
    # ------------------------------------------------------------------

    def _refresh(self, scroll_to: Optional[int] = None):
        self._lb.delete(0, tk.END)
        for idx, item in enumerate(self.items):
            self._lb.insert(tk.END, item.display())
            self._lb.itemconfig(idx, bg=self.BG.get(item.kind, '#ffffff'))

        if scroll_to is not None:
            pos = max(0, min(scroll_to, self._lb.size() - 1))
            self._lb.see(pos)
            self._lb.selection_clear(0, tk.END)
            self._lb.selection_set(pos)
            self._lb.activate(pos)

        self._update_stats()

    def _update_stats(self):
        n_total    = len(self.items)
        n_merged   = len([i for i in self.items if i.kind == 'merged'])
        n_fragment = len([i for i in self.items if i.kind == 'fragment'])
        n_special  = len([i for i in self.items if i.kind in ('special', 'blank')])
        self._stats_var.set(
            f'Total rows: {n_total}    '
            f'Merged groups: {n_merged}    '
            f'Unprocessed fragments: {n_fragment}    '
            f'Special / blank: {n_special}'
        )

    # ------------------------------------------------------------------
    # Selection handling
    # ------------------------------------------------------------------

    def _on_select(self, _event=None):
        sel = self._lb.curselection()
        mergeable = [i for i in sel if self.items[i].is_mergeable]

        # Update selection info label
        if mergeable:
            self._sel_info_var.set(f'{len(mergeable)} fragment(s) selected')
        else:
            self._sel_info_var.set('')

        # Enable/disable merge button
        can_merge = len(mergeable) >= 2
        self._merge_btn.config(
            state=tk.NORMAL if can_merge else tk.DISABLED,
            text=(f'⤵  Merge Selected ({len(mergeable)})  (Ctrl+M)'
                  if can_merge else '⤵  Merge Selected  (Ctrl+M)'),
        )

        # Update preview and count
        if mergeable:
            joined = ' '.join(
                self.items[i].full_text for i in sorted(mergeable)
            )
            self._count_var.set(str(len(mergeable)))
            self._preview.config(state=tk.NORMAL)
            self._preview.delete('1.0', tk.END)
            self._preview.insert('1.0', joined)
            self._preview.config(state=tk.DISABLED)
        else:
            self._count_var.set('—')
            self._preview.config(state=tk.NORMAL)
            self._preview.delete('1.0', tk.END)
            self._preview.config(state=tk.DISABLED)

    # ------------------------------------------------------------------
    # Merge
    # ------------------------------------------------------------------

    def _merge_selected(self):
        sel = self._lb.curselection()
        to_merge = [i for i in sorted(sel) if self.items[i].is_mergeable]

        if len(to_merge) < 2:
            self._status('Select at least 2 fragments to merge.')
            return

        # Guard: no non-mergeable items between first and last
        first, last = to_merge[0], to_merge[-1]
        blockers = [
            i for i in range(first + 1, last)
            if not self.items[i].is_mergeable
        ]
        if blockers:
            ans = messagebox.askyesno(
                'Selection spans a break',
                'Your selection includes timestamps, speaker labels, or blank lines '
                'between the first and last selected rows.\n\n'
                'Do you still want to merge just the selected fragment rows '
                '(the non-mergeable rows will be left in place)?',
                icon='warning',
            )
            if not ans:
                return

        # Save undo snapshot
        orig_first = deepcopy(self.items[first])
        absorbed   = [(i, deepcopy(self.items[i])) for i in to_merge[1:]]
        self.undo_stack.append((first, orig_first, absorbed))

        # Collect all texts into the first item
        all_texts = []
        for i in to_merge:
            all_texts.extend(self.items[i].texts)

        self.items[first] = Item(all_texts, 'merged', self.items[first].para_ref)

        # Remove absorbed items in reverse order (preserves indices)
        for i in sorted([idx for idx, _ in absorbed], reverse=True):
            self.items.pop(i)

        self._undo_btn.config(state=tk.NORMAL)
        self._refresh(scroll_to=first)
        self._status(
            f'Merged {len(to_merge)} fragments into one paragraph.  '
            f'(Ctrl+Z to undo)'
        )

    # ------------------------------------------------------------------
    # Undo
    # ------------------------------------------------------------------

    def _undo(self):
        if not self.undo_stack:
            self._status('Nothing to undo.')
            return

        first_idx, orig_first, absorbed = self.undo_stack.pop()

        # Restore first item to its pre-merge state
        self.items[first_idx] = orig_first

        # Re-insert absorbed items at their original positions
        for i, item in sorted(absorbed, key=lambda x: x[0]):
            self.items.insert(i, item)

        if not self.undo_stack:
            self._undo_btn.config(state=tk.DISABLED)

        self._refresh(scroll_to=first_idx)
        self._status('Undone last merge.')

    # ------------------------------------------------------------------
    # Save
    # ------------------------------------------------------------------

    def _save_and_close(self):
        if not self.items:
            self.root.destroy()
            return

        n_frag = len([i for i in self.items if i.kind == 'fragment'])
        if n_frag > 0:
            ans = messagebox.askyesnocancel(
                'Unprocessed fragments',
                f'{n_frag} fragment(s) have not been merged or reviewed.\n\n'
                'Each will become its own paragraph in the output.\n\n'
                'Save anyway?',
            )
            if ans is None:
                return   # Cancel — stay in tool
            if not ans:
                return   # No — go back without saving

        try:
            out = Document()

            # Mirror source page layout
            if self.source_doc:
                ss = self.source_doc.sections[0]
                os_ = out.sections[0]
                os_.page_width    = ss.page_width
                os_.page_height   = ss.page_height
                os_.top_margin    = ss.top_margin
                os_.bottom_margin = ss.bottom_margin
                os_.left_margin   = ss.left_margin
                os_.right_margin  = ss.right_margin

                # Mirror Normal style
                try:
                    sn = self.source_doc.styles['Normal']
                    on = out.styles['Normal']
                    if sn.font.name:
                        on.font.name = sn.font.name
                    if sn.font.size:
                        on.font.size = sn.font.size
                except Exception:
                    pass

            for item in self.items:
                p = out.add_paragraph(item.full_text)
                if item.para_ref:
                    _copy_para_format(item.para_ref, p)

            out.save(self.output_path)
            messagebox.showinfo(
                'Saved',
                f'Document saved to:\n\n{self.output_path}',
            )
            self.root.destroy()

        except Exception as exc:
            messagebox.showerror('Save error', str(exc))

    # ------------------------------------------------------------------
    # Status bar helper
    # ------------------------------------------------------------------

    def _status(self, msg: str):
        self._status_var.set(msg)


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main():
    root = tk.Tk()

    # Try to set app icon
    try:
        icon = os.path.join(
            os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
            'DocAnalyser.ico'
        )
        if os.path.exists(icon):
            root.iconbitmap(icon)
    except Exception:
        pass

    TranscriptMerger(root)
    root.mainloop()


if __name__ == '__main__':
    main()
